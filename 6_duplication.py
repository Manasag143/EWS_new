import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any, Optional, Union, Set, Tuple
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
from datetime import datetime
from collections import defaultdict
import difflib

# Pydantic imports
from pydantic import BaseModel, Field, validator, root_validator
from enum import Enum

# NLTK imports for rule-based deduplication
try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk.stem import PorterStemmer
    # Download required NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('punkt')
        nltk.download('stopwords')
    NLTK_AVAILABLE = True
except ImportError:
    print("Warning: NLTK not available. Using basic text processing.")
    NLTK_AVAILABLE = False

# Optional: sklearn for advanced similarity
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# ================================
# PYDANTIC MODELS
# ================================

class RiskLevel(str, Enum):
    """Enumeration for risk levels"""
    HIGH = "High"
    LOW = "Low"
    MEDIUM = "Medium"

class FinancialQuarter(str, Enum):
    """Enumeration for financial quarters"""
    Q1 = "Q1"
    Q2 = "Q2" 
    Q3 = "Q3"
    Q4 = "Q4"

class CriteriaType(str, Enum):
    """Enumeration for financial criteria types"""
    DEBT_INCREASE = "debt_increase"
    PROVISIONING = "provisioning"
    ASSET_DECLINE = "asset_decline"
    RECEIVABLE_DAYS = "receivable_days"
    PAYABLE_DAYS = "payable_days"
    DEBT_EBITDA = "debt_ebitda"
    REVENUE_DECLINE = "revenue_decline"
    ONETIME_EXPENSES = "onetime_expenses"
    MARGIN_DECLINE = "margin_decline"
    CASH_BALANCE = "cash_balance"
    SHORT_TERM_DEBT = "short_term_debt"
    MANAGEMENT_ISSUES = "management_issues"
    REGULATORY_COMPLIANCE = "regulatory_compliance"
    MARKET_COMPETITION = "market_competition"
    OPERATIONAL_DISRUPTIONS = "operational_disruptions"

class FinancialMetrics(BaseModel):
    """Model for previous year financial data"""
    debt: Optional[float] = Field(None, description="Previous debt amount in Cr")
    ebitda: Optional[float] = Field(None, description="Current quarter EBITDA in Cr")
    asset_value: Optional[float] = Field(None, description="Previous asset value in Cr")
    receivable_days: Optional[int] = Field(None, description="Previous receivable days")
    payable_days: Optional[int] = Field(None, description="Previous payable days")
    revenue: Optional[float] = Field(None, description="Previous revenue in Cr")
    profitability: Optional[float] = Field(None, description="Previous profitability in Cr")
    operating_margin: Optional[float] = Field(None, description="Previous operating margin %")
    cash_balance: Optional[float] = Field(None, description="Previous cash balance in Cr")
    current_liabilities: Optional[float] = Field(None, description="Previous current liabilities in Cr")
    
    @validator('operating_margin')
    def validate_margin_percentage(cls, v):
        if v is not None and (v < 0 or v > 100):
            raise ValueError('Operating margin must be between 0 and 100')
        return v

class CompanyInfo(BaseModel):
    """Model for company information extracted from PDF"""
    company_name: str = Field(..., description="Full company name")
    quarter: FinancialQuarter = Field(..., description="Financial quarter")
    financial_year: str = Field(..., pattern=r"FY\d{2,4}", description="Financial year (e.g., FY25)")
    
    @validator('company_name')
    def validate_company_name(cls, v):
        if len(v.strip()) < 3:
            raise ValueError('Company name must be at least 3 characters')
        return v.strip()
    
    def __str__(self):
        return f"{self.company_name}-{self.quarter}{self.financial_year}"

class RedFlag(BaseModel):
    """Model for individual red flag"""
    id: int = Field(..., description="Unique identifier for the flag")
    description: str = Field(..., min_length=10, description="Description of the red flag")
    original_quote: Optional[str] = Field(None, description="Original quote from document")
    page_number: Optional[int] = Field(None, ge=1, description="Page number where flag was found")
    speaker: Optional[str] = Field(None, description="Speaker name if from transcript")
    confidence_score: Optional[float] = Field(None, ge=0.0, le=1.0, description="Confidence score")
    
    @validator('description')
    def validate_description(cls, v):
        cleaned = re.sub(r'\s+', ' ', v.strip())
        if len(cleaned) < 10:
            raise ValueError('Flag description must be at least 10 characters')
        return cleaned

class ClassificationResult(BaseModel):
    """Model for flag classification result"""
    flag_id: int = Field(..., description="ID of the classified flag")
    matched_criteria: Optional[CriteriaType] = Field(None, description="Matched criteria type")
    risk_level: RiskLevel = Field(..., description="Assigned risk level")
    reasoning: str = Field(..., min_length=10, description="Reasoning for classification")
    numerical_evidence: Optional[List[str]] = Field(default_factory=list, description="Numerical values found")
    
    @validator('reasoning')
    def validate_reasoning(cls, v):
        return re.sub(r'\s+', ' ', v.strip())

class CategorySummary(BaseModel):
    """Model for categorized summary"""
    category_name: str = Field(..., description="Name of the category")
    flags: List[str] = Field(default_factory=list, description="List of flags in this category")
    summary_points: List[str] = Field(default_factory=list, description="Summary bullet points")
    
    @validator('category_name')
    def validate_category_name(cls, v):
        allowed_categories = [
            "Balance Sheet Issues",
            "P&L (Income Statement) Issues", 
            "Liquidity Issues",
            "Management and Strategy related Issues",
            "Regulatory Issues",
            "Industry and Market Issues",
            "Operational Issues"
        ]
        if v not in allowed_categories:
            # Try to map to closest category
            v_lower = v.lower()
            if any(word in v_lower for word in ['balance', 'sheet', 'asset', 'debt']):
                return "Balance Sheet Issues"
            elif any(word in v_lower for word in ['income', 'revenue', 'profit', 'p&l']):
                return "P&L (Income Statement) Issues"
            elif any(word in v_lower for word in ['liquidity', 'cash', 'flow']):
                return "Liquidity Issues"
            elif any(word in v_lower for word in ['management', 'strategy', 'leadership']):
                return "Management and Strategy related Issues"
            elif any(word in v_lower for word in ['regulatory', 'compliance', 'legal']):
                return "Regulatory Issues"
            elif any(word in v_lower for word in ['market', 'industry', 'competition']):
                return "Industry and Market Issues"
            else:
                return "Operational Issues"  # Default fallback
        return v

class IterationResult(BaseModel):
    """Model for individual iteration results"""
    iteration_number: int = Field(..., ge=1, le=5, description="Iteration number (1-5)")
    stage_name: str = Field(..., description="Name of the processing stage")
    timestamp: datetime = Field(default_factory=datetime.now, description="Processing timestamp")
    processing_time_seconds: Optional[float] = Field(None, ge=0, description="Processing time")
    
    red_flags: Optional[List[RedFlag]] = Field(default_factory=list, description="Extracted red flags")
    categories: Optional[List[CategorySummary]] = Field(default_factory=list, description="Categorized results")
    classifications: Optional[List[ClassificationResult]] = Field(default_factory=list, description="Classification results")
    raw_response: Optional[str] = Field(None, description="Raw LLM response")
    
    @validator('stage_name')
    def validate_stage_name(cls, v):
        allowed_stages = [
            "Initial Analysis",
            "Deduplication",
            "Categorization", 
            "Summary Generation",
            "Enhanced Unique Flags Classification"
        ]
        if v not in allowed_stages:
            return "Initial Analysis"  # Default fallback
        return v

class RiskDistribution(BaseModel):
    """Model for risk distribution statistics"""
    high_risk_count: int = Field(0, ge=0, description="Number of high risk flags")
    low_risk_count: int = Field(0, ge=0, description="Number of low risk flags")
    total_count: int = Field(0, ge=0, description="Total number of flags")
    
    @root_validator
    def validate_totals(cls, values):
        high = values.get('high_risk_count', 0)
        low = values.get('low_risk_count', 0)
        total = values.get('total_count', 0)
        
        if high + low != total:
            values['total_count'] = high + low  # Auto-correct instead of raising error
        return values

class ProcessingConfig(BaseModel):
    """Model for processing configuration"""
    pdf_path: str = Field(..., description="Path to PDF file")
    queries_csv_path: str = Field(..., description="Path to queries CSV/Excel file")
    output_folder: str = Field("results", description="Output folder path")
    
    api_key: str = Field(..., description="Azure OpenAI API key")
    azure_endpoint: str = Field(..., description="Azure OpenAI endpoint")
    api_version: str = Field("2024-02-01", description="Azure OpenAI API version")
    deployment_name: str = Field("gpt-4.1-mini", description="Deployment name")
    
    max_tokens: int = Field(4000, ge=100, le=8000, description="Maximum tokens per API call")
    temperature: float = Field(0.1, ge=0.0, le=2.0, description="Temperature for API calls")
    similarity_threshold: float = Field(0.75, ge=0.0, le=1.0, description="Similarity threshold for deduplication")
    max_flags: int = Field(10, ge=1, le=50, description="Maximum number of flags to extract")

class PipelineResult(BaseModel):
    """Model for complete pipeline results"""
    pdf_name: str = Field(..., description="Name of processed PDF")
    company_info: CompanyInfo = Field(..., description="Extracted company information")
    financial_metrics: FinancialMetrics = Field(..., description="Previous year financial data")
    
    iterations: List[IterationResult] = Field(default_factory=list, description="Results from each iteration")
    risk_distribution: RiskDistribution = Field(..., description="Risk distribution statistics")
    high_risk_flags: List[RedFlag] = Field(default_factory=list, description="High risk flags")
    
    word_document_path: Optional[str] = Field(None, description="Path to generated Word document")
    csv_results_path: Optional[str] = Field(None, description="Path to CSV results file")
    
    total_processing_time: Optional[float] = Field(None, ge=0, description="Total processing time in seconds")
    success: bool = Field(True, description="Whether processing was successful")
    error_message: Optional[str] = Field(None, description="Error message if processing failed")

# ================================
# UTILITY FUNCTIONS
# ================================

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

# ================================
# AZURE OPENAI LLM CLASS
# ================================

class AzureOpenAILLM:
    """Azure OpenAI LLM class with Pydantic integration"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1-mini"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

# ================================
# PDF PROCESSING
# ================================

class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

# ================================
# RULE-BASED DEDUPLICATION
# ================================

class RuleBasedDeduplicator:
    """Rule-based deduplication system for financial red flags"""
    
    def __init__(self):
        if NLTK_AVAILABLE:
            self.stop_words = set(stopwords.words('english'))
            self.stemmer = PorterStemmer()
        else:
            self.stop_words = set(['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'])
            self.stemmer = None
        
        self.financial_synonyms = {
            'revenue': ['sales', 'income', 'top line', 'turnover', 'receipts'],
            'profit': ['earnings', 'net income', 'bottom line', 'profitability'],
            'debt': ['borrowing', 'liabilities', 'leverage', 'loans'],
            'margin': ['profitability', 'markup', 'spread'],
            'decline': ['decrease', 'fall', 'drop', 'reduction', 'deterioration'],
            'increase': ['rise', 'growth', 'expansion', 'surge', 'escalation'],
            'cash': ['liquidity', 'funds', 'capital'],
            'assets': ['holdings', 'resources', 'investments'],
            'expenses': ['costs', 'expenditure', 'outgoings'],
            'management': ['leadership', 'executives', 'administration']
        }
        
        self.synonym_map = {}
        for key, synonyms in self.financial_synonyms.items():
            self.synonym_map[key] = key
            for synonym in synonyms:
                self.synonym_map[synonym] = key
    
    def preprocess_text(self, text: str) -> str:
        """Normalize text for comparison"""
        text = text.lower()
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        words = text.split()
        normalized_words = []
        for word in words:
            if word in self.synonym_map:
                normalized_words.append(self.synonym_map[word])
            else:
                normalized_words.append(word)
        
        return ' '.join(normalized_words)
    
    def extract_key_features(self, text: str) -> Dict[str, any]:
        """Extract key features from flag text for comparison"""
        normalized_text = self.preprocess_text(text)
        
        if NLTK_AVAILABLE:
            words = word_tokenize(normalized_text)
            meaningful_words = [self.stemmer.stem(word) for word in words 
                              if word not in self.stop_words and len(word) > 2]
        else:
            words = normalized_text.split()
            meaningful_words = [word for word in words 
                              if word not in self.stop_words and len(word) > 2]
        
        numbers = re.findall(r'\d+\.?\d*', text)
        percentages = re.findall(r'\d+\.?\d*\s*%', text)
        
        financial_keywords = set()
        for word in meaningful_words:
            if word in ['debt', 'revenue', 'profit', 'margin', 'cash', 'asset', 'expense']:
                financial_keywords.add(word)
        
        return {
            'stemmed_words': set(meaningful_words),
            'numbers': numbers,
            'percentages': percentages,
            'financial_keywords': financial_keywords,
            'word_count': len(meaningful_words),
            'normalized_text': normalized_text
        }
    
    def calculate_similarity_score(self, features1: Dict, features2: Dict) -> float:
        """Calculate similarity between two flag features"""
        scores = []
        
        words1, words2 = features1['stemmed_words'], features2['stemmed_words']
        if words1 or words2:
            word_similarity = len(words1.intersection(words2)) / len(words1.union(words2))
            scores.append(word_similarity * 0.4)
        
        keywords1, keywords2 = features1['financial_keywords'], features2['financial_keywords']
        if keywords1 or keywords2:
            keyword_similarity = len(keywords1.intersection(keywords2)) / max(len(keywords1.union(keywords2)), 1)
            scores.append(keyword_similarity * 0.3)
        
        nums1, nums2 = set(features1['numbers']), set(features2['numbers'])
        if nums1 or nums2:
            num_similarity = len(nums1.intersection(nums2)) / max(len(nums1.union(nums2)), 1)
            scores.append(num_similarity * 0.2)
        
        len1, len2 = features1['word_count'], features2['word_count']
        if len1 > 0 and len2 > 0:
            length_similarity = 1 - abs(len1 - len2) / max(len1, len2)
            scores.append(length_similarity * 0.1)
        
        return sum(scores) if scores else 0.0
    
    def is_duplicate(self, flag1: str, flag2: str, threshold: float = 0.75) -> bool:
        """Determine if two flags are duplicates"""
        features1 = self.extract_key_features(flag1)
        features2 = self.extract_key_features(flag2)
        
        similarity = self.calculate_similarity_score(features1, features2)
        return similarity >= threshold
    
    def is_subset_flag(self, short_flag: str, long_flag: str) -> bool:
        """Check if short_flag is a subset of long_flag"""
        short_normalized = self.preprocess_text(short_flag)
        long_normalized = self.preprocess_text(long_flag)
        
        short_words = set(short_normalized.split())
        long_words = set(long_normalized.split())
        
        if len(short_words) <= 3 and len(short_words) > 0:
            overlap_ratio = len(short_words.intersection(long_words)) / len(short_words)
            return overlap_ratio >= 0.8
        
        return False

class RuleBasedFlagExtractor:
    """Rule-based system for extracting unique financial red flags"""
    
    def __init__(self):
        self.deduplicator = RuleBasedDeduplicator()
        
        self.flag_patterns = {
            'debt_related': [
                r'debt\s+(?:increased?|rose?|grew?|went up)',
                r'(?:higher|rising|growing)\s+debt',
                r'debt.*(?:levels?|amounts?|burden)',
                r'borrowing.*(?:increased?|rose?|grew?)'
            ],
            'revenue_related': [
                r'revenue\s+(?:declined?|fell|decreased?|dropped?)',
                r'(?:sales|income|top line).*(?:declined?|fell|decreased?)',
                r'(?:lower|falling|declining)\s+(?:revenue|sales|income)'
            ],
            'margin_related': [
                r'margin\s+(?:declined?|fell|decreased?|compressed?)',
                r'(?:profitability|margins).*(?:pressure|decline|fall)',
                r'(?:gross|operating|net)\s+margin.*(?:declined?|fell)'
            ],
            'cash_related': [
                r'cash\s+(?:declined?|fell|decreased?|shortage)',
                r'(?:liquidity|cash flow).*(?:issues?|problems?|concerns?)',
                r'(?:working capital|cash position).*(?:declined?|fell)'
            ],
            'asset_related': [
                r'asset\s+(?:declined?|fell|decreased?|impairment)',
                r'(?:write[- ]?off|impairment|provision)',
                r'asset.*(?:quality|deterioration)'
            ]
        }
    
    def extract_flags_from_text(self, text: str) -> List[str]:
        """Extract potential red flags using pattern matching"""
        flags = []
        sentences = re.split(r'[.!?]+', text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20:
                continue
            
            for category, patterns in self.flag_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, sentence, re.IGNORECASE):
                        flag_text = self._extract_flag_context(sentence, pattern)
                        if flag_text and len(flag_text) > 15:
                            flags.append(flag_text)
                        break
        
        return flags
    
    def _extract_flag_context(self, sentence: str, pattern: str) -> str:
        """Extract meaningful context around a matched pattern"""
        sentence = re.sub(r'\s+', ' ', sentence).strip()
        
        if 20 <= len(sentence) <= 200:
            return sentence
        
        if len(sentence) > 200:
            match = re.search(pattern, sentence, re.IGNORECASE)
            if match:
                start = max(0, match.start() - 50)
                end = min(len(sentence), match.end() + 50)
                return sentence[start:end].strip()
        
        return sentence
    
    def deduplicate_flags(self, flags: List[str], similarity_threshold: float = 0.75) -> List[str]:
        """Remove duplicate flags using rule-based approach"""
        if not flags:
            return []
        
        unique_flags = []
        
        for flag in flags:
            is_duplicate = False
            is_subset = False
            
            for i, existing_flag in enumerate(unique_flags):
                if self.deduplicator.is_duplicate(flag, existing_flag, similarity_threshold):
                    is_duplicate = True
                    break
                
                if self.deduplicator.is_subset_flag(flag, existing_flag):
                    is_subset = True
                    break
                
                if self.deduplicator.is_subset_flag(existing_flag, flag):
                    unique_flags[i] = flag  # Replace with longer flag
                    is_duplicate = True
                    break
            
            if not is_duplicate and not is_subset:
                unique_flags.append(flag)
        
        return unique_flags
    
    def rank_flags_by_severity(self, flags: List[str]) -> List[str]:
        """Rank flags by potential severity"""
        severity_keywords = {
            'high': ['significant', 'substantial', 'major', 'severe', 'critical', 'dramatic'],
            'medium': ['moderate', 'noticeable', 'considerable', 'increased', 'decreased'],
            'low': ['slight', 'minor', 'small', 'marginal']
        }
        
        def get_severity_score(flag: str) -> int:
            flag_lower = flag.lower()
            score = 0
            
            for keyword in severity_keywords['high']:
                if keyword in flag_lower:
                    score += 3
            
            for keyword in severity_keywords['medium']:
                if keyword in flag_lower:
                    score += 2
            
            for keyword in severity_keywords['low']:
                if keyword in flag_lower:
                    score += 1
            
            numbers = re.findall(r'\d+', flag)
            if numbers:
                max_num = max(int(num) for num in numbers)
                if max_num > 50:
                    score += 2
                elif max_num > 20:
                    score += 1
            
            if '%' in flag:
                score += 1
            
            return score
        
        ranked_flags = sorted(flags, key=get_severity_score, reverse=True)
        return ranked_flags

# ================================
# PYDANTIC ANALYZER CLASS
# ================================

class PydanticFinancialAnalyzer:
    """Pydantic-based financial document analyzer"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.flag_extractor = RuleBasedFlagExtractor()
        self.validate_config()
    
    def validate_config(self):
        """Validate configuration"""
        pass
    
    def parse_company_info_response(self, llm_response: str) -> CompanyInfo:
        """Parse LLM response to extract company information"""
        try:
            pattern = r'(.+?)-([QR]\d+)(FY\d{2,4})'
            match = re.search(pattern, llm_response)
            
            if match:
                company_name = match.group(1).strip()
                quarter = match.group(2).replace('R', 'Q')
                financial_year = match.group(3)
                
                return CompanyInfo(
                    company_name=company_name,
                    quarter=FinancialQuarter(quarter),
                    financial_year=financial_year
                )
            else:
                lines = llm_response.strip().split('\n')
                for line in lines:
                    if '-Q' in line and 'FY' in line:
                        parts = line.split('-')
                        if len(parts) >= 2:
                            company_name = parts[0].strip()
                            quarter_year = parts[1].strip()
                            
                            quarter_match = re.search(r'(Q\d+)', quarter_year)
                            year_match = re.search(r'(FY\d{2,4})', quarter_year)
                            
                            if quarter_match and year_match:
                                return CompanyInfo(
                                    company_name=company_name,
                                    quarter=FinancialQuarter(quarter_match.group(1)),
                                    financial_year=year_match.group(1)
                                )
                
                return CompanyInfo(
                    company_name="Unknown Company",
                    quarter=FinancialQuarter.Q1,
                    financial_year="FY25"
                )
                
        except Exception as e:
            return CompanyInfo(
                company_name="Unknown Company",
                quarter=FinancialQuarter.Q1,
                financial_year="FY25"
            )
    
    def parse_red_flags_from_response(self, response: str, iteration_number: int) -> List[RedFlag]:
        """Parse red flags from LLM response"""
        flags = []
        lines = response.split('\n')
        current_flag = None
        flag_id = 1
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            numbered_match = re.match(r'^(\d+)\.?\s*(.+)', line)
            bullet_match = re.match(r'^[-*]\s*(.+)', line)
            
            if numbered_match:
                if current_flag:
                    try:
                        flags.append(current_flag)
                    except:
                        pass
                
                description = numbered_match.group(2).strip()
                try:
                    current_flag = RedFlag(
                        id=flag_id,
                        description=description
                    )
                    flag_id += 1
                except:
                    current_flag = None
                    
            elif bullet_match:
                if current_flag:
                    try:
                        flags.append(current_flag)
                    except:
                        pass
                
                description = bullet_match.group(1).strip()
                try:
                    current_flag = RedFlag(
                        id=flag_id,
                        description=description
                    )
                    flag_id += 1
                except:
                    current_flag = None
                    
            elif line.lower().startswith('original quote:'):
                if current_flag:
                    quote = line[len('original quote:'):].strip()
                    current_flag.original_quote = quote
                    
            elif line.lower().startswith('page'):
                if current_flag:
                    page_match = re.search(r'page\s+(\d+)', line.lower())
                    if page_match:
                        current_flag.page_number = int(page_match.group(1))
            
            elif current_flag and len(line) > 20:
                current_flag.description += " " + line
        
        if current_flag:
            try:
                flags.append(current_flag)
            except:
                pass
        
        validated_flags = []
        for flag in flags:
            try:
                validated_flag = RedFlag(
                    id=flag.id,
                    description=flag.description,
                    original_quote=flag.original_quote,
                    page_number=flag.page_number
                )
                validated_flags.append(validated_flag)
            except Exception as e:
                print(f"Warning: Skipping invalid flag {flag.id}: {e}")
        
        return validated_flags
    
    def parse_categorized_response(self, response: str) -> List[CategorySummary]:
        """Parse categorized response into structured categories"""
        categories = []
        current_category = None
        
        lines = response.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('###'):
                if current_category:
                    categories.append(current_category)
                
                category_name = line.replace('###', '').strip()
                try:
                    current_category = CategorySummary(category_name=category_name)
                except ValueError:
                    current_category = CategorySummary(category_name="Operational Issues")
            
            elif line.startswith('-') or line.startswith('*'):
                if current_category:
                    flag_text = line[1:].strip()
                    if len(flag_text) > 5:
                        current_category.flags.append(flag_text)
        
        if current_category:
            categories.append(current_category)
        
        return categories
    
    def parse_summary_response(self, response: str) -> List[CategorySummary]:
        """Parse summary response with bullet points"""
        categories = []
        sections = response.split('###')
        
        for section in sections:
            if not section.strip():
                continue
            
            lines = section.split('\n')
            category_name = ""
            summary_points = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if not category_name and not line.startswith('*') and not line.startswith('-'):
                    category_name = line.strip()
                elif line.startswith('*') or line.startswith('-'):
                    bullet_text = line[1:].strip()
                    if bullet_text:
                        summary_points.append(bullet_text)
            
            if category_name and summary_points:
                try:
                    category = CategorySummary(
                        category_name=category_name,
                        summary_points=summary_points
                    )
                    categories.append(category)
                except ValueError:
                    category = CategorySummary(
                        category_name="Operational Issues",
                        summary_points=summary_points
                    )
                    categories.append(category)
        
        return categories
    
    def classify_flag_with_pydantic(self, flag: RedFlag, criteria_definitions: Dict[str, str], 
                                  financial_metrics: FinancialMetrics) -> ClassificationResult:
        """Classify flag using rule-based approach with Pydantic validation"""
        
        numerical_evidence = re.findall(r'\d+\.?\d*%?', flag.description)
        
        flag_lower = flag.description.lower()
        matched_criteria = None
        risk_level = RiskLevel.LOW
        reasoning = "No specific criteria matched"
        
        # Rule-based classification logic
        if any(word in flag_lower for word in ['debt', 'borrowing', 'leverage']):
            matched_criteria = CriteriaType.DEBT_INCREASE
            if any(word in flag_lower for word in ['increase', 'rose', 'grew', 'higher']):
                # Check numerical thresholds if available
                numbers = [float(re.sub(r'[^\d.]', '', num)) for num in numerical_evidence if re.search(r'\d', num)]
                if numbers and max(numbers) >= 30:  # 30% threshold
                    risk_level = RiskLevel.HIGH
                    reasoning = f"Debt increase detected with {max(numbers)}% threshold met"
                elif any(word in flag_lower for word in ['significant', 'substantial', 'major']):
                    risk_level = RiskLevel.HIGH
                    reasoning = "Debt increase with significant qualitative indicators"
                else:
                    reasoning = "Debt increase detected but below threshold"
        
        elif any(word in flag_lower for word in ['revenue', 'sales', 'income']):
            matched_criteria = CriteriaType.REVENUE_DECLINE
            if any(word in flag_lower for word in ['decline', 'fell', 'decrease', 'drop']):
                numbers = [float(re.sub(r'[^\d.]', '', num)) for num in numerical_evidence if re.search(r'\d', num)]
                if numbers and max(numbers) >= 25:  # 25% threshold
                    risk_level = RiskLevel.HIGH
                    reasoning = f"Revenue decline of {max(numbers)}% exceeds threshold"
                elif any(word in flag_lower for word in ['significant', 'substantial', 'major']):
                    risk_level = RiskLevel.HIGH
                    reasoning = "Revenue decline with significant qualitative indicators"
                else:
                    reasoning = "Revenue decline detected but below threshold"
        
        elif any(word in flag_lower for word in ['margin', 'profitability']):
            matched_criteria = CriteriaType.MARGIN_DECLINE
            if any(word in flag_lower for word in ['decline', 'pressure', 'compression']):
                numbers = [float(re.sub(r'[^\d.]', '', num)) for num in numerical_evidence if re.search(r'\d', num)]
                if numbers and max(numbers) >= 25:  # 25% threshold
                    risk_level = RiskLevel.HIGH
                    reasoning = f"Margin decline of {max(numbers)}% exceeds threshold"
                elif any(word in flag_lower for word in ['significant', 'substantial', 'major']):
                    risk_level = RiskLevel.HIGH
                    reasoning = "Margin pressure with significant qualitative indicators"
                else:
                    reasoning = "Margin pressure detected but below threshold"
        
        elif any(word in flag_lower for word in ['cash', 'liquidity']):
            matched_criteria = CriteriaType.CASH_BALANCE
            if any(word in flag_lower for word in ['decline', 'shortage', 'issues', 'problems']):
                risk_level = RiskLevel.HIGH
                reasoning = "Cash/liquidity issues detected"
        
        elif any(word in flag_lower for word in ['provision', 'write-off', 'writeoff', 'impairment']):
            matched_criteria = CriteriaType.PROVISIONING
            risk_level = RiskLevel.HIGH
            reasoning = "Provisioning or write-offs detected"
        
        elif any(word in flag_lower for word in ['management', 'leadership', 'ceo', 'cfo']):
            matched_criteria = CriteriaType.MANAGEMENT_ISSUES
            if any(word in flag_lower for word in ['change', 'turnover', 'departure', 'resignation']):
                risk_level = RiskLevel.HIGH
                reasoning = "Management turnover detected"
        
        elif any(word in flag_lower for word in ['regulatory', 'compliance', 'legal']):
            matched_criteria = CriteriaType.REGULATORY_COMPLIANCE
            if any(word in flag_lower for word in ['issues', 'concerns', 'violation', 'penalty']):
                risk_level = RiskLevel.HIGH
                reasoning = "Regulatory/compliance issues detected"
        
        elif any(word in flag_lower for word in ['competition', 'competitive', 'market share']):
            matched_criteria = CriteriaType.MARKET_COMPETITION
            if any(word in flag_lower for word in ['pressure', 'decline', 'loss', 'intense']):
                risk_level = RiskLevel.HIGH
                reasoning = "Competitive pressure detected"
        
        elif any(word in flag_lower for word in ['operational', 'supply chain', 'production']):
            matched_criteria = CriteriaType.OPERATIONAL_DISRUPTIONS
            if any(word in flag_lower for word in ['issues', 'problems', 'disruption', 'difficulties']):
                risk_level = RiskLevel.HIGH
                reasoning = "Operational disruptions detected"
        
        return ClassificationResult(
            flag_id=flag.id,
            matched_criteria=matched_criteria,
            risk_level=risk_level,
            reasoning=reasoning,
            numerical_evidence=numerical_evidence
        )

# ================================
# ITERATION FUNCTIONS
# ================================

def iteration_1_initial_analysis(context: str, first_query: str, 
                               llm: AzureOpenAILLM, analyzer: PydanticFinancialAnalyzer) -> IterationResult:
    """Pydantic-enhanced Iteration 1"""
    start_time = time.time()
    
    sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.

COMPLETE DOCUMENT TO ANALYZE:
{context}

Your task is to analyze the ENTIRE document above and identify ALL potential red flags.

CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
"""
    
    first_prompt = f"{sys_prompt}\n\nQuestion: {first_query}\n\nAnswer:"
    first_response = llm._call(first_prompt, max_tokens=4000)
    
    red_flags = analyzer.parse_red_flags_from_response(first_response, 1)
    
    processing_time = time.time() - start_time
    
    return IterationResult(
        iteration_number=1,
        stage_name="Initial Analysis",
        processing_time_seconds=processing_time,
        red_flags=red_flags,
        raw_response=first_response
    )

def iteration_2_deduplication(iteration_1_result: IterationResult, 
                            analyzer: PydanticFinancialAnalyzer) -> IterationResult:
    """Pydantic-enhanced Iteration 2 with rule-based deduplication"""
    start_time = time.time()
    
    flag_descriptions = [flag.description for flag in iteration_1_result.red_flags]
    
    unique_descriptions = analyzer.flag_extractor.deduplicate_flags(flag_descriptions, similarity_threshold=0.75)
    
    unique_flags = []
    for i, description in enumerate(unique_descriptions, 1):
        original_flag = next(
            (flag for flag in iteration_1_result.red_flags if flag.description == description), 
            None
        )
        
        try:
            unique_flag = RedFlag(
                id=i,
                description=description,
                original_quote=original_flag.original_quote if original_flag else None,
                page_number=original_flag.page_number if original_flag else None
            )
            unique_flags.append(unique_flag)
        except:
            pass
    
    processing_time = time.time() - start_time
    
    return IterationResult(
        iteration_number=2,
        stage_name="Deduplication",
        processing_time_seconds=processing_time,
        red_flags=unique_flags,
        raw_response=f"Deduplicated from {len(iteration_1_result.red_flags)} to {len(unique_flags)} flags"
    )

def iteration_3_categorization(context: str, iteration_2_result: IterationResult,
                             llm: AzureOpenAILLM, analyzer: PydanticFinancialAnalyzer) -> IterationResult:
    """Pydantic-enhanced Iteration 3"""
    start_time = time.time()
    
    flags_text = "\n".join([f"{flag.id}. {flag.description}" for flag in iteration_2_result.red_flags])
    
    third_prompt = f"""You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.

- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction.
- Regulatory Issues: Concerns related Compliance with laws, regulations.
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape.
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.

While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.

**Output Format**:
### Balance Sheet Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
    
    third_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{context}

Previous Analysis: {flags_text}

Based on the above analysis and the original context, please answer: {third_prompt}

Answer:"""
    
    third_response = llm._call(third_full_prompt, max_tokens=4000)
    
    categories = analyzer.parse_categorized_response(third_response)
    
    processing_time = time.time() - start_time
    
    return IterationResult(
        iteration_number=3,
        stage_name="Categorization",
        processing_time_seconds=processing_time,
        categories=categories,
        raw_response=third_response
    )

def iteration_4_summary(context: str, iteration_3_result: IterationResult,
                       llm: AzureOpenAILLM, analyzer: PydanticFinancialAnalyzer) -> IterationResult:
    """Pydantic-enhanced Iteration 4"""
    start_time = time.time()
    
    categorized_input = ""
    for category in iteration_3_result.categories:
        categorized_input += f"### {category.category_name}\n"
        for flag in category.flags:
            categorized_input += f"- {flag}\n"
        categorized_input += "\n"
    
    fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:

1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
3. **Focus on factual content**: Base the summary solely on the information associated with each red flag, without introducing external opinions or assumptions
4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
8. **Summary should be factual**: Avoid any subjective interpretations or opinions
9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points

Format the output exactly like this example:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
    
    fourth_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{context}

Previous Analysis: {categorized_input}

Based on the above analysis and the original context, please answer: {fourth_prompt}

Answer:"""
    
    fourth_response = llm._call(fourth_full_prompt, max_tokens=4000)
    
    summary_categories = analyzer.parse_summary_response(fourth_response)
    
    processing_time = time.time() - start_time
    
    return IterationResult(
        iteration_number=4,
        stage_name="Summary Generation", 
        processing_time_seconds=processing_time,
        categories=summary_categories,
        raw_response=fourth_response
    )

def iteration_5_classification(iteration_2_result: IterationResult, 
                             financial_metrics: FinancialMetrics,
                             analyzer: PydanticFinancialAnalyzer) -> IterationResult:
    """Pydantic-enhanced Iteration 5 with rule-based classification"""
    start_time = time.time()
    
    criteria_definitions = {
        "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
        "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBITDA; Low: provisioning or write-offs less than 25% of current quarter's EBITDA",
        "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
        "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
        "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
        "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
        "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
        "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBITDA; Low: one-time expenses or losses less than 25% of current quarter's EBITDA",
        "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
        "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
        "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
        "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
        "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues",
        "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
        "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
    }
    
    classifications = []
    high_risk_flags = []
    
    for flag in iteration_2_result.red_flags:
        try:
            classification = analyzer.classify_flag_with_pydantic(
                flag, criteria_definitions, financial_metrics
            )
            classifications.append(classification)
            
            if classification.risk_level == RiskLevel.HIGH:
                high_risk_flags.append(flag)
                
        except Exception as e:
            logger.error(f"Error classifying flag {flag.id}: {e}")
            classification = ClassificationResult(
                flag_id=flag.id,
                matched_criteria=None,
                risk_level=RiskLevel.LOW,
                reasoning=f"Classification failed: {str(e)}"
            )
            classifications.append(classification)
        
        time.sleep(0.1)  # Small delay to prevent overwhelming
    
    processing_time = time.time() - start_time
    
    return IterationResult(
        iteration_number=5,
        stage_name="Enhanced Unique Flags Classification",
        processing_time_seconds=processing_time,
        red_flags=high_risk_flags,
        classifications=classifications,
        raw_response=f"Classified {len(classifications)} flags: {len(high_risk_flags)} high risk, {len(classifications) - len(high_risk_flags)} low risk"
    )

# ================================
# HELPER FUNCTIONS
# ================================

def extract_company_info_pydantic(pdf_path: str, llm: AzureOpenAILLM, 
                                analyzer: PydanticFinancialAnalyzer) -> CompanyInfo:
    """Extract company information with Pydantic validation"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()[:2000]
        doc.close()
        
        prompt = f"""Extract the company name, quarter, and financial year from this text from an earnings call transcript.

Text: {first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)  
3. Financial Year (FY23/FY24/FY25 etc.)

Format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

Response:"""
        
        response = llm._call(prompt, max_tokens=200)
        return analyzer.parse_company_info_response(response)
        
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return CompanyInfo(
            company_name="Unknown Company",
            quarter=FinancialQuarter.Q1,
            financial_year="FY25"
        )

def parse_financial_metrics_from_string(previous_year_data: str) -> FinancialMetrics:
    """Parse financial metrics from the previous year data string"""
    try:
        metrics = FinancialMetrics()
        
        lines = previous_year_data.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if 'debt' in line.lower() and 'cr' in line.lower():
                debt_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if debt_match:
                    metrics.debt = float(debt_match.group(1))
            
            elif 'ebitda' in line.lower() and 'cr' in line.lower():
                ebitda_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if ebitda_match:
                    metrics.ebitda = float(ebitda_match.group(1))
            
            elif 'asset value' in line.lower() and 'cr' in line.lower():
                asset_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if asset_match:
                    metrics.asset_value = float(asset_match.group(1))
            
            elif 'receivable days' in line.lower():
                rec_match = re.search(r'(\d+)', line)
                if rec_match:
                    metrics.receivable_days = int(rec_match.group(1))
            
            elif 'payable days' in line.lower():
                pay_match = re.search(r'(\d+)', line)
                if pay_match:
                    metrics.payable_days = int(pay_match.group(1))
            
            elif 'revenue' in line.lower() and 'cr' in line.lower():
                rev_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if rev_match:
                    metrics.revenue = float(rev_match.group(1))
            
            elif 'operating margin' in line.lower() and '%' in line:
                margin_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if margin_match:
                    metrics.operating_margin = float(margin_match.group(1))
            
            elif 'cash balance' in line.lower() and 'cr' in line.lower():
                cash_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if cash_match:
                    metrics.cash_balance = float(cash_match.group(1))
            
            elif 'current liabilities' in line.lower() and 'cr' in line.lower():
                liab_match = re.search(r'(\d+(?:\.\d+)?)', line)
                if liab_match:
                    metrics.current_liabilities = float(liab_match.group(1))
        
        return metrics
        
    except Exception as e:
        logger.error(f"Error parsing financial metrics: {e}")
        return FinancialMetrics()

def create_pydantic_pipeline_result(pdf_name: str, 
                                   company_info: CompanyInfo,
                                   financial_metrics: FinancialMetrics,
                                   iterations: List[IterationResult],
                                   high_risk_flags: List[RedFlag]) -> PipelineResult:
    """Create a complete pipeline result with validation"""
    
    high_count = len(high_risk_flags)
    all_flags = []
    for iteration in iterations:
        if iteration.red_flags:
            all_flags.extend(iteration.red_flags)
    
    total_count = len(set(flag.id for flag in all_flags))
    low_count = total_count - high_count
    
    risk_distribution = RiskDistribution(
        high_risk_count=high_count,
        low_risk_count=low_count,
        total_count=total_count
    )
    
    total_time = sum(
        iteration.processing_time_seconds 
        for iteration in iterations 
        if iteration.processing_time_seconds
    )
    
    return PipelineResult(
        pdf_name=pdf_name,
        company_info=company_info,
        financial_metrics=financial_metrics,
        iterations=iterations,
        risk_distribution=risk_distribution,
        high_risk_flags=high_risk_flags,
        total_processing_time=total_time,
        success=True
    )

def create_word_document_pydantic(result: PipelineResult, output_folder: str, 
                                context: str, llm: AzureOpenAILLM) -> str:
    """Create Word document from Pydantic pipeline result"""
    try:
        os.makedirs(output_folder, exist_ok=True)
        doc = Document()
        
        # Title
        title = doc.add_heading(str(result.company_info), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Risk Distribution
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = 'High Risk'
        table.cell(0, 1).text = str(result.risk_distribution.high_risk_count)
        table.cell(1, 0).text = 'Low Risk'
        table.cell(1, 1).text = str(result.risk_distribution.low_risk_count)
        table.cell(2, 0).text = 'Total Flags'
        table.cell(2, 1).text = str(result.risk_distribution.total_count)
        
        for i in range(3):
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph('')
        
        # High Risk Summary
        high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
        high_risk_heading.runs[0].bold = True
        
        if result.high_risk_flags:
            for flag in result.high_risk_flags:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(flag.description)
        else:
            doc.add_paragraph('No high risk flags identified.')
        
        # Summary section
        doc.add_paragraph('_' * 50)
        summary_heading = doc.add_heading('Summary', level=1)
        summary_heading.runs[0].bold = True
        
        summary_iteration = next(
            (iter for iter in result.iterations if iter.iteration_number == 4), 
            None
        )
        
        if summary_iteration and summary_iteration.categories:
            for category in summary_iteration.categories:
                cat_heading = doc.add_heading(category.category_name, level=2)
                cat_heading.runs[0].bold = True
                
                for point in category.summary_points:
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(point)
                
                doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
        
        # Save document
        doc_filename = f"{result.pdf_name}_Pydantic_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
        
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        try:
            doc = Document()
            doc.add_heading(f"{result.pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {result.risk_distribution.high_risk_count}")
            doc.add_paragraph(f"Low Risk Flags: {result.risk_distribution.low_risk_count}")
            doc.add_paragraph(f"Total Flags: {result.risk_distribution.total_count}")
            
            doc_filename = f"{result.pdf_name}_Fallback_Report.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

def save_results_to_csv_pydantic(result: PipelineResult, output_folder: str) -> str:
    """Save Pydantic pipeline results to CSV"""
    try:
        # Prepare data for CSV
        csv_data = []
        
        for iteration in result.iterations:
            csv_data.append({
                "pdf_name": result.pdf_name,
                "iteration": iteration.iteration_number,
                "stage": iteration.stage_name,
                "processing_time": iteration.processing_time_seconds,
                "flags_count": len(iteration.red_flags) if iteration.red_flags else 0,
                "categories_count": len(iteration.categories) if iteration.categories else 0,
                "timestamp": iteration.timestamp.isoformat()
            })
        
        # Create DataFrame and save
        df = pd.DataFrame(csv_data)
        csv_filename = f"{result.pdf_name}_pydantic_results.csv"
        csv_path = os.path.join(output_folder, csv_filename)
        df.to_csv(csv_path, index=False)
        
        # Save detailed flag information
        if result.high_risk_flags:
            flags_data = []
            for flag in result.high_risk_flags:
                flags_data.append({
                    "flag_id": flag.id,
                    "description": flag.description,
                    "original_quote": flag.original_quote,
                    "page_number": flag.page_number,
                    "risk_level": "High"
                })
            
            flags_df = pd.DataFrame(flags_data)
            flags_csv_path = os.path.join(output_folder, f"{result.pdf_name}_high_risk_flags.csv")
            flags_df.to_csv(flags_csv_path, index=False)
        
        # Save classification results
        classification_iteration = next(
            (iter for iter in result.iterations if iter.iteration_number == 5), 
            None
        )
        
        if classification_iteration and classification_iteration.classifications:
            classification_data = []
            for classification in classification_iteration.classifications:
                classification_data.append({
                    "flag_id": classification.flag_id,
                    "matched_criteria": classification.matched_criteria.value if classification.matched_criteria else None,
                    "risk_level": classification.risk_level.value,
                    "reasoning": classification.reasoning,
                    "numerical_evidence": ", ".join(classification.numerical_evidence)
                })
            
            classification_df = pd.DataFrame(classification_data)
            classification_csv_path = os.path.join(output_folder, f"{result.pdf_name}_classifications.csv")
            classification_df.to_csv(classification_csv_path, index=False)
        
        return csv_path
        
    except Exception as e:
        logger.error(f"Error saving CSV results: {e}")
        return None

# ================================
# MAIN PIPELINE FUNCTION
# ================================

def process_pdf_enhanced_pipeline_pydantic(config: ProcessingConfig, 
                                         financial_metrics: FinancialMetrics) -> PipelineResult:
    """Complete Pydantic-based pipeline for PDF processing"""
    start_time = time.time()
    iterations = []
    
    try:
        # Initialize components
        llm = AzureOpenAILLM(
            api_key=config.api_key,
            azure_endpoint=config.azure_endpoint,
            api_version=config.api_version,
            deployment_name=config.deployment_name
        )
        
        analyzer = PydanticFinancialAnalyzer(config)
        
        # Load and process PDF
        docs = mergeDocs(config.pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load first query
        try:
            if config.queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(config.queries_csv_path)
            else:
                queries_df = pd.read_csv(config.queries_csv_path)
            
            first_query = queries_df["prompt"].tolist()[0] if len(queries_df) > 0 else "Analyze this document for potential red flags."
        except Exception:
            first_query = "Analyze this document for potential red flags."
        
        # Extract company info
        company_info = extract_company_info_pydantic(config.pdf_path, llm, analyzer)
        
        # Run all iterations
        print("Running Iteration 1 - Initial Analysis...")
        iteration_1 = iteration_1_initial_analysis(context, first_query, llm, analyzer)
        iterations.append(iteration_1)
        
        print("Running Iteration 2 - Rule-based Deduplication...")
        iteration_2 = iteration_2_deduplication(iteration_1, analyzer)
        iterations.append(iteration_2)
        
        print("Running Iteration 3 - Categorization...")
        iteration_3 = iteration_3_categorization(context, iteration_2, llm, analyzer)
        iterations.append(iteration_3)
        
        print("Running Iteration 4 - Summary Generation...")
        iteration_4 = iteration_4_summary(context, iteration_3, llm, analyzer)
        iterations.append(iteration_4)
        
        print("Running Iteration 5 - Classification...")
        iteration_5 = iteration_5_classification(iteration_2, financial_metrics, analyzer)
        iterations.append(iteration_5)
        
        # Create pipeline result
        result = create_pydantic_pipeline_result(
            pdf_name=Path(config.pdf_path).stem,
            company_info=company_info,
            financial_metrics=financial_metrics,
            iterations=iterations,
            high_risk_flags=iteration_5.red_flags or []
        )
        
        # Generate outputs
        result.word_document_path = create_word_document_pydantic(result, config.output_folder, context, llm)
        result.csv_results_path = save_results_to_csv_pydantic(result, config.output_folder)
        
        result.total_processing_time = time.time() - start_time
        
        # Print results summary
        print(f"\n=== FINAL CLASSIFICATION RESULTS ===")
        print(f"High Risk Flags: {result.risk_distribution.high_risk_count}")
        print(f"Low Risk Flags: {result.risk_distribution.low_risk_count}")
        print(f"Total Flags: {result.risk_distribution.total_count}")
        
        if result.high_risk_flags:
            print(f"\n--- HIGH RISK FLAGS ---")
            for i, flag in enumerate(result.high_risk_flags, 1):
                print(f"  {i}. {flag.description}")
        else:
            print(f"\n--- HIGH RISK FLAGS ---")
            print("  No high risk flags identified")
        
        return result
        
    except Exception as e:
        # Create error result
        error_result = PipelineResult(
            pdf_name=Path(config.pdf_path).stem,
            company_info=CompanyInfo(
                company_name="Error",
                quarter=FinancialQuarter.Q1,
                financial_year="FY25"
            ),
            financial_metrics=financial_metrics,
            iterations=iterations,
            risk_distribution=RiskDistribution(
                high_risk_count=0,
                low_risk_count=0,
                total_count=0
            ),
            high_risk_flags=[],
            success=False,
            error_message=str(e),
            total_processing_time=time.time() - start_time
        )
        
        logger.error(f"Pipeline failed: {e}")
        return error_result

# ================================
# MAIN FUNCTION
# ================================

def main_pydantic():
    """Main function using Pydantic models"""
    
    # Create configuration
    config = ProcessingConfig(
        pdf_path=r"chemplast_pdf",  # This will be updated for each PDF
        queries_csv_path=r"EWS_prompts_v2_2.xlsx",
        output_folder=r"chemplast_results_pydantic",
        api_key="8496bd1d98c",
        azure_endpoint="https://crisil-pp-gpt.openai.azure.com",
        api_version="2025-01-01-preview",
        deployment_name="gpt-4.1-mini",
        max_tokens=4000,
        temperature=0.1,
        similarity_threshold=0.75,
        max_flags=10
    )
    
    # Parse financial metrics
    previous_year_data_str = """
Previous reported Debt	Mar-22	882Cr
Current quarter ebitda	March-23 130Cr
Previous reported asset value	Mar-22	5602Cr
Previous reported receivable days	Mar-22	12days
Previous reported payable days	Mar-22	189days
Previous reported revenue	Dec-22	1189Cr
Previous reported profitability	Dec-22	27Cr
Previous reported operating margin	Dec-22	7%
Previous reported cash balance	Mar-22	1229Cr
Previous reported current liabilities	Mar-22	68Cr
"""
    
    financial_metrics = parse_financial_metrics_from_string(previous_year_data_str)
    
    # Create output folder
    os.makedirs(config.output_folder, exist_ok=True)
    
    # Process all PDFs
    pdf_files = glob.glob(os.path.join(config.pdf_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {config.pdf_path}")
        return
    
    results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        # Update config for current PDF
        current_config = config.copy(update={"pdf_path": pdf_file})
        
        # Process with Pydantic pipeline
        result = process_pdf_enhanced_pipeline_pydantic(current_config, financial_metrics)
        results.append(result)
        
        if result.success:
            print(f"✅ Successfully processed {pdf_file}")
            print(f"📊 Risk Distribution: {result.risk_distribution.high_risk_count} High, {result.risk_distribution.low_risk_count} Low")
            print(f"⏱️  Processing Time: {result.total_processing_time:.2f} seconds")
            if result.word_document_path:
                print(f"📄 Word Document: {result.word_document_path}")
            if result.csv_results_path:
                print(f"📊 CSV Results: {result.csv_results_path}")
        else:
            print(f"❌ Failed to process {pdf_file}: {result.error_message}")
    
    # Save combined results
    if results:
        combined_data = []
        for result in results:
            combined_data.append({
                "pdf_name": result.pdf_name,
                "company": str(result.company_info),
                "high_risk_count": result.risk_distribution.high_risk_count,
                "low_risk_count": result.risk_distribution.low_risk_count,
                "total_count": result.risk_distribution.total_count,
                "processing_time": result.total_processing_time,
                "success": result.success,
                "word_doc_path": result.word_document_path,
                "csv_path": result.csv_results_path,
                "error_message": result.error_message if not result.success else None
            })
        
        combined_df = pd.DataFrame(combined_data)
        combined_csv_path = os.path.join(config.output_folder, "combined_pydantic_results.csv")
        combined_df.to_csv(combined_csv_path, index=False)
        print(f"\n📁 Combined results saved to: {combined_csv_path}")
        
        # Print final summary
        successful_results = [r for r in results if r.success]
        failed_results = [r for r in results if not r.success]
        
        print(f"\n{'='*60}")
        print(f"PROCESSING SUMMARY")
        print(f"{'='*60}")
        print(f"Total PDFs processed: {len(results)}")
        print(f"Successful: {len(successful_results)}")
        print(f"Failed: {len(failed_results)}")
        
        if successful_results:
            total_high_risk = sum(r.risk_distribution.high_risk_count for r in successful_results)
            total_low_risk = sum(r.risk_distribution.low_risk_count for r in successful_results)
            avg_processing_time = sum(r.total_processing_time for r in successful_results) / len(successful_results)
            
            print(f"Total High Risk Flags: {total_high_risk}")
            print(f"Total Low Risk Flags: {total_low_risk}")
            print(f"Average Processing Time: {avg_processing_time:.2f} seconds")
        
        if failed_results:
            print(f"\nFailed PDFs:")
            for result in failed_results:
                print(f"  - {result.pdf_name}: {result.error_message}")

# ================================
# LEGACY COMPATIBILITY FUNCTIONS
# ================================

def process_pdf_enhanced_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Legacy function wrapper for backward compatibility
    """
    try:
        # Create Pydantic config from legacy parameters
        config = ProcessingConfig(
            pdf_path=pdf_path,
            queries_csv_path=queries_csv_path,
            output_folder=output_folder,
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        # Parse financial metrics
        financial_metrics = parse_financial_metrics_from_string(previous_year_data)
        
        # Run Pydantic pipeline
        result = process_pdf_enhanced_pipeline_pydantic(config, financial_metrics)
        
        # Convert to legacy format for compatibility
        if result.success:
            legacy_summary = pd.DataFrame({
                "pdf_name": [result.pdf_name] * len(result.iterations),
                "iteration": [iter.iteration_number for iter in result.iterations],
                "stage": [iter.stage_name for iter in result.iterations],
                "response": [iter.raw_response or f"Processed {len(iter.red_flags or [])} flags" for iter in result.iterations],
                "timestamp": [iter.timestamp.isoformat() for iter in result.iterations]
            })
            return legacy_summary
        else:
            return None
            
    except Exception as e:
        logger.error(f"Legacy wrapper failed: {e}")
        return None

def main():
    """Legacy main function for backward compatibility"""
    return main_pydantic()

# ================================
# ENTRY POINT
# ================================

if __name__ == "__main__":
    main_pydantic()
