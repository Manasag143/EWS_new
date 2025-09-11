"""__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')"""

import json
import os
import re
import time
import docx
import tempfile
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from crisil_langchain import CLLMGateway
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers.merger_retriever import MergerRetriever
from typing import List
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import SpacyTextSplitter, RecursiveCharacterTextSplitter
from langchain_chroma.vectorstores import Chroma
from langchain.vectorstores.base import VectorStoreRetriever
from llms import HostedLLM, HostedLLM_Perplexity
from langchain.chains import RetrievalQA
from dotenv import load_dotenv
import warnings

warnings.filterwarnings('ignore')

load_dotenv()

# pdf_path = r"C:\Users\c-RatingsGe\Downloads\Alok Industries FY 16_v1.pdf"
# LLM_ENDPOINT = os.getenv("llama_url")
# EMBEDDING_ENDPOINT = os.getenv("embedding_url")
# CHROMA_DIR = "chroma_store"


def load_pdf(path: str) -> List[Document]:
    loader = PyPDFLoader(path)
    return loader.load()


def split_by_headers(docs: List[Document], min_header_len=5) -> List[Document]:
    header_regex = re.compile(r"""\n(?=)[A-Z][]A-Za-z\s]{""" + str(min_header_len) + r""",100}""")
    sectioned = []
    for doc in docs:
        text = doc.page_content
        sections = header_regex.split(text)
        for section in sections:
            if section.strip():
                sectioned.append(Document(page_content=section.strip(), metadata=doc.metadata))
    return sectioned


def sentence_split(docs: List[Document]) -> List[Document]:
    return SpacyTextSplitter(chunk_size=2048, chunk_overlap=500).split_documents(docs)


def token_split(docs: List[Document]) -> List[Document]:
    return RecursiveCharacterTextSplitter(chunk_size=2500, chunk_overlap=200).split_documents(docs)


def build_chroma_vectorstore(chunks: List[Document]):
    api_token = "ae3ea1c14b2e496eacd117064cb78cd1"
    embedding = CLLMGateway(api_type="embedding", base_url="https://llmgateway.crisil.local/api/", api_version="v1",
                            provider="tgi", deployment="bge-m3", spec_version=1,
                            max_tokens=2000, api_token=api_token, tls_verify=False).load_client()
    vectordb = Chroma.from_documents(documents=chunks, embedding=embedding)
    return vectordb


def rag_chain(vectordb: Chroma, chunks: List[Document]):
    retriever: VectorStoreRetriever = vectordb.as_retriever(search_type="mmr", search_kwargs={"k": 10})

    bm25_retriever = BM25Retriever.from_documents(chunks)
    bm25_retriever.k = 5

    # hybrid_retriever = MergerRetriever(retrievers=[retriever, bm25_retriever])

    # llm = HostedLLM(endpoint=LLM_ENDPOINT)
    llm = HostedLLM_Perplexity(endpoint="https://as1-lower-llm.crisil.local/perplexity/llama/70b/llm/")

    qa = RetrievalQA.from_chain_type(llm=llm, retriever=retriever, return_source_documents=True)
    return qa


def run_rag(pdf_path: str):

    docs = load_pdf(pdf_path)
    chunks = token_split(docs)
    vectordb = build_chroma_vectorstore(chunks)
    rag_builder = rag_chain(vectordb, chunks)

    return rag_builder


def flag_detector(pdf_path):
    document = docx.Document()
    title = document.add_heading('Audit Report Flag QnA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    question_list = [
        "What is the nature of the audit opinion, are there any qualifications, exceptions, or limitations noted in the audit report?",
        "Has the company defaulted on any loan or borrowing repayments, or failed to pay interest on time?",
        "Have any instances of fraud or theft reported or identified during the audit?",
        "Are there any outstanding fines, penalties, or regulatory actions against the company?",
        "Are there any overdue or unpaid statutory dues, taxes, or other liabilities?",
        "Are there any accumulated losses in subsidiaries?",
        "Are there any assets that are not operational or have been suspended?",
        "Are there any other compliance issues or regulatory breaches noted in the audit report?",
        "Have any assets been revalued downward, and if so, what is the impact on the company's financial position?",
        "Are there any inventory valuation issues?",
        "Are there any concerns regarding the utilization of short-term funds for long-term purposes or to support subsidiaries, associates, or joint ventures?",
        "Are there any material weaknesses or deficiencies in the company's internal controls or accounting processes?",
        "Did the auditor consider any issues or concerns raised by the outgoing auditors? And are there any observations or comments from regulatory bodies or third party auditor?",
        "Are there any balance reconciliations or unaccounted balances in the financial statements?",
        "Are there any title deeds or lease agreements that are not in the name of the company?",
        "Is there any mention of any credit rating downgrades?",
        "Are there any concerns about the appropriateness of the quantum of outstanding trade receivables, particularly with related parties?",
        "Are there any pending litigations or legal proceedings?",
        "Are there any provisions for losses or contingencies that have not been adequately disclosed or accounted for in the financial statements?",
        "Does the company maintain a deposit repayment reserve account?",
        "Are there any related-party transactions, such as discounts given to group companies, that may impact the company's financial position?",
        "Is there anything for which audit trail is not maintained by the company?"
    ]
    cnt = 1
    for k in question_list:
        query = f"""
         You are a financial analyst who expertises in extracting or answer specific questions looking at an annual rpeort. Below is the question for which you have to find the answer :
         {k}
         While answering, adhere to the following guidelines:
         1. Keep your answer short, crisp and clear.
         """

        qa_chain = run_rag(pdf_path)
        result = qa_chain.invoke(query)



        question = document.add_paragraph()
        question.add_run(f'Q{cnt}: {k}').font.size = Pt(12)

        answer = document.add_paragraph()
        answer.add_run(f'A{cnt}: {result["result"]}').font.size = Pt(12)

        document.add_paragraph()
        cnt += 1

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")

    # Save the document to the temporary file
    with open(temp_file.name, "wb") as f:
        document.save(f)

    return temp_file.name
