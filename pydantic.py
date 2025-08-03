import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any, Optional, Union
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
from pydantic import BaseModel, Field, validator
from enum import Enum

warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# ================== PYDANTIC SCHEMAS ==================

class RiskLevel(str, Enum):
    """Enumeration for risk levels"""
    HIGH = "High"
    LOW = "Low"

class CriteriaType(str, Enum):
    """Enumeration for criteria types"""
    DEBT_INCREASE = "debt_increase"
    PROVISIONING = "provisioning"
    ASSET_DECLINE = "asset_decline"
    RECEIVABLE_DAYS = "receivable_days"
    PAYABLE_DAYS = "payable_days"
    DEBT_EBITDA = "debt_ebitda"
    REVENUE_DECLINE = "revenue_decline"
    ONETIME_EXPENSES = "onetime_expenses"
    MARGIN_DECLINE = "margin_decline"
    CASH_BALANCE = "cash_balance"
    SHORT_TERM_DEBT = "short_term_debt"
    MANAGEMENT_ISSUES = "management_issues"
    REGULATORY_COMPLIANCE = "regulatory_compliance"
    MARKET_COMPETITION = "market_competition"
    OPERATIONAL_DISRUPTIONS = "operational_disruptions"

class PDFPage(BaseModel):
    """Schema for PDF page data"""
    page_num: int = Field(..., ge=1, description="Page number (1-indexed)")
    text: str = Field(..., description="Extracted text from the page")
    
    @validator('text')
    def text_not_empty(cls, v):
        if not v.strip():
            raise ValueError('Page text cannot be empty')
        return v

class PDFDocument(BaseModel):
    """Schema for merged PDF document"""
    context: str = Field(..., min_length=1, description="Combined text content")
    page_num: Optional[int] = Field(None, description="Page number if split by pages")

class RedFlag(BaseModel):
    """Schema for individual red flags"""
    flag_id: int = Field(..., ge=1, description="Unique identifier for the flag")
    description: str = Field(..., min_length=5, max_length=500, description="Brief description of the red flag")
    original_quote: str = Field(..., min_length=1, description="Original quote from the document")
    page_reference: Optional[str] = Field(None, description="Page reference where flag was found")
    keywords: List[str] = Field(default_factory=list, description="Associated keywords")
    
    @validator('description')
    def description_valid(cls, v):
        if not v.strip():
            raise ValueError('Description cannot be empty')
        return v.strip()

class FlagClassification(BaseModel):
    """Schema for flag classification results"""
    flag: str = Field(..., min_length=1, description="The red flag text")
    matched_criteria: str = Field(..., description="Matched criteria name or 'None'")
    risk_level: RiskLevel = Field(..., description="Risk level classification")
    reasoning: str = Field(..., min_length=1, description="Reasoning for the classification")
    confidence_score: float = Field(default=0.0, ge=0.0, le=1.0, description="Confidence in classification")

class CompanyInfo(BaseModel):
    """Schema for company information extracted from PDF"""
    company_name: str = Field(..., min_length=1, description="Full company name")
    quarter: str = Field(..., regex=r'^Q[1-4]$', description="Quarter (Q1, Q2, Q3, Q4)")
    financial_year: str = Field(..., regex=r'^FY\d{2,4}$', description="Financial year (e.g., FY25)")
    formatted_name: str = Field(..., description="Formatted company identifier")
    
    @validator('formatted_name')
    def format_name(cls, v, values):
        if 'company_name' in values and 'quarter' in values and 'financial_year' in values:
            return f"{values['company_name']}-{values['quarter']}{values['financial_year']}"
        return v

class RiskCounts(BaseModel):
    """Schema for risk count summary"""
    high: int = Field(default=0, ge=0, description="Number of high risk flags")
    low: int = Field(default=0, ge=0, description="Number of low risk flags")
    total: int = Field(default=0, ge=0, description="Total number of flags")
    
    @validator('total')
    def total_matches_sum(cls, v, values):
        if 'high' in values and 'low' in values:
            expected_total = values['high'] + values['low']
            if v != expected_total:
                raise ValueError(f'Total ({v}) must equal high + low ({expected_total})')
        return v

class CategorySummary(BaseModel):
    """Schema for categorized summary"""
    category_name: str = Field(..., min_length=1, description="Name of the category")
    bullet_points: List[str] = Field(..., min_items=1, description="List of bullet points for this category")
    
    @validator('bullet_points')
    def bullets_not_empty(cls, v):
        return [bullet.strip() for bullet in v if bullet.strip()]

class ProcessingResult(BaseModel):
    """Schema for overall processing results"""
    pdf_name: str = Field(..., min_length=1, description="Name of the processed PDF")
    company_info: CompanyInfo
    risk_counts: RiskCounts
    high_risk_flags: List[str] = Field(default_factory=list, description="List of high risk flags")
    low_risk_flags: List[str] = Field(default_factory=list, description="List of low risk flags")
    classification_results: List[FlagClassification] = Field(default_factory=list, description="Detailed classification results")
    category_summaries: List[CategorySummary] = Field(default_factory=list, description="Categorized summaries")
    word_doc_path: Optional[str] = Field(None, description="Path to generated Word document")
    processing_time: float = Field(..., ge=0, description="Processing time in seconds")

# ================== UTILITY FUNCTIONS ==================

def getFilehash(file_path: str) -> str:
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

# ================== AZURE OPENAI CLASS ==================

class AzureOpenAILLM:
    """Azure OpenAI GPT-4.1-mini LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

# ================== PDF EXTRACTION ==================

class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[PDFPage]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pdf_page = PDFPage(page_num=page_num + 1, text=text)
                pages.append(pdf_page)
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[PDFDocument]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [PDFDocument(context=page.text, page_num=page.page_num) for page in pages]
    else:
        all_text = "\n".join([page.text for page in pages])
        return [PDFDocument(context=all_text)]

# ================== RULE-BASED DEDUPLICATION ==================

class RuleBasedDeduplicator:
    """Rule-based deduplication without GPT calls"""
    
    def __init__(self):
        self.similarity_threshold = 0.7
        self.keyword_overlap_threshold = 0.6
        
    def normalize_text(self, text: str) -> str:
        """Normalize text for comparison"""
        # Remove punctuation, convert to lowercase, remove extra spaces
        normalized = re.sub(r'[^\w\s]', '', text.lower())
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        return normalized
    
    def extract_keywords(self, text: str) -> set:
        """Extract meaningful keywords from text"""
        # Remove common stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
            'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
            'should', 'may', 'might', 'must', 'shall', 'can', 'this', 'that',
            'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they'
        }
        
        words = self.normalize_text(text).split()
        keywords = {word for word in words if len(word) > 2 and word not in stop_words}
        return keywords
    
    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts"""
        keywords1 = self.extract_keywords(text1)
        keywords2 = self.extract_keywords(text2)
        
        if not keywords1 or not keywords2:
            return 0.0
        
        intersection = keywords1.intersection(keywords2)
        union = keywords1.union(keywords2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def remove_duplicates(self, text_items: List[str]) -> List[str]:
        """Remove duplicate text items based on similarity"""
        unique_items = []
        
        for item in text_items:
            if not item.strip():
                continue
                
            is_duplicate = False
            for existing_item in unique_items:
                similarity = self.calculate_similarity(item, existing_item)
                if similarity > self.similarity_threshold:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                unique_items.append(item.strip())
        
        return unique_items
    
    def deduplicate_red_flags(self, analysis_text: str) -> str:
        """Extract and deduplicate red flags from analysis text"""
        # Extract individual red flags using patterns
        flag_patterns = [
            r'\d+\.\s*([^.]+\.)',  # Numbered items
            r'-\s*([^.\n]+\.)',    # Bullet points
            r'\*\s*([^.\n]+\.)',   # Asterisk bullets
        ]
        
        extracted_flags = []
        for pattern in flag_patterns:
            matches = re.findall(pattern, analysis_text, re.MULTILINE)
            extracted_flags.extend(matches)
        
        if not extracted_flags:
            # Fallback: split by sentences and filter
            sentences = analysis_text.split('.')
            extracted_flags = [s.strip() + '.' for s in sentences if len(s.strip()) > 20]
        
        # Remove duplicates
        unique_flags = self.remove_duplicates(extracted_flags)
        
        # Reconstruct the text
        if unique_flags:
            deduplicated_text = "Deduplicated Red Flags:\n\n"
            for i, flag in enumerate(unique_flags, 1):
                deduplicated_text += f"{i}. {flag}\n"
            return deduplicated_text
        else:
            return "No duplicates found in the analysis."

# ================== UNIQUE FLAG EXTRACTOR ==================

class UniqueFlageExtractor:
    """Rule-based unique flag extraction without GPT"""
    
    def __init__(self):
        self.financial_keywords = {
            'debt', 'revenue', 'profit', 'loss', 'margin', 'cash', 'asset', 'liability',
            'ebitda', 'interest', 'expense', 'income', 'balance', 'receivable', 'payable',
            'provision', 'write-off', 'impairment', 'credit', 'liquidity', 'capital'
        }
        
        self.risk_indicators = {
            'decline', 'decrease', 'fall', 'drop', 'reduce', 'lower', 'down', 'weak',
            'negative', 'concern', 'issue', 'problem', 'risk', 'challenge', 'pressure',
            'deteriorate', 'worsen', 'struggle', 'difficulty', 'shortfall', 'deficit'
        }
    
    def extract_financial_flags(self, text: str) -> List[str]:
        """Extract potential financial red flags using rule-based approach"""
        sentences = text.split('.')
        potential_flags = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20:  # Skip very short sentences
                continue
            
            sentence_lower = sentence.lower()
            
            # Check if sentence contains financial keywords and risk indicators
            has_financial_keyword = any(keyword in sentence_lower for keyword in self.financial_keywords)
            has_risk_indicator = any(indicator in sentence_lower for indicator in self.risk_indicators)
            
            if has_financial_keyword and has_risk_indicator:
                potential_flags.append(sentence.strip() + '.')
        
        return potential_flags
    
    def categorize_flags(self, flags: List[str]) -> Dict[str, List[str]]:
        """Categorize flags based on keywords"""
        categories = {
            'debt_related': ['debt', 'borrowing', 'leverage', 'loan'],
            'revenue_related': ['revenue', 'sales', 'income', 'turnover'],
            'profitability_related': ['profit', 'margin', 'ebitda', 'profitability'],
            'liquidity_related': ['cash', 'liquidity', 'working capital'],
            'asset_related': ['asset', 'receivable', 'inventory', 'goodwill'],
            'operational_related': ['operation', 'supply chain', 'production', 'efficiency']
        }
        
        categorized_flags = {category: [] for category in categories}
        uncategorized = []
        
        for flag in flags:
            flag_lower = flag.lower()
            categorized = False
            
            for category, keywords in categories.items():
                if any(keyword in flag_lower for keyword in keywords):
                    categorized_flags[category].append(flag)
                    categorized = True
                    break
            
            if not categorized:
                uncategorized.append(flag)
        
        if uncategorized:
            categorized_flags['other'] = uncategorized
        
        return categorized_flags
    
    def extract_unique_flags_rule_based(self, response_text: str) -> List[str]:
        """Extract unique flags using rule-based approach"""
        # Extract potential flags
        potential_flags = self.extract_financial_flags(response_text)
        
        # Use deduplicator to remove similar flags
        deduplicator = RuleBasedDeduplicator()
        unique_flags = deduplicator.remove_duplicates(potential_flags)
        
        # Limit to reasonable number
        if len(unique_flags) > 12:
            unique_flags = unique_flags[:12]
        
        return unique_flags if unique_flags else ["No specific red flags identified"]

# ================== COMPANY INFO EXTRACTION ==================

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> CompanyInfo:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""
Extract the company name, quarter, and financial year from this text from an earnings call transcript.

Text: {first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)  
3. Financial Year (FY23/FY24/FY25 etc.)

Format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

Response:"""
       
        response = llm._call(prompt, max_tokens=200)
        
        # Parse response to extract components
        company_name = "Unknown Company"
        quarter = "Q1"
        financial_year = "FY25"
        
        # Try to extract from response
        if '-Q' in response and 'FY' in response:
            parts = response.strip().split('-')
            if len(parts) >= 2:
                company_name = parts[0].strip()
                quarter_fy = parts[1].strip()
                
                # Extract quarter and FY
                quarter_match = re.search(r'Q[1-4]', quarter_fy)
                fy_match = re.search(r'FY\d{2,4}', quarter_fy)
                
                if quarter_match:
                    quarter = quarter_match.group()
                if fy_match:
                    financial_year = fy_match.group()
        
        return CompanyInfo(
            company_name=company_name,
            quarter=quarter,
            financial_year=financial_year,
            formatted_name=f"{company_name}-{quarter}{financial_year}"
        )
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return CompanyInfo(
            company_name="Unknown Company",
            quarter="Q1",
            financial_year="FY25",
            formatted_name="Unknown Company-Q1FY25"
        )

# ================== FLAG CLASSIFICATION ==================

def classify_flag_against_criteria_strict(flag: str, criteria_definitions: Dict[str, str], 
                                 previous_year_data: str, llm: AzureOpenAILLM) -> FlagClassification:
    """Strictly classify a single flag against 15 criteria - only High if exact match found"""
    criteria_keywords = {
        "debt_increase": ["debt increase", "debt increased", "debt rising", "debt growth", "higher debt", "debt went up", "debt levels", "borrowing increase"],
        "provisioning": ["provision", "write-off", "write off", "writeoff", "bad debt", "impairment", "credit loss"],
        "asset_decline": ["asset decline", "asset fall", "asset decrease", "asset value down", "asset reduction", "asset impairment"],
        "receivable_days": ["receivable days", "collection period", "DSO", "days sales outstanding", "collection time"],
        "payable_days": ["payable days", "payment period", "DPO", "days payable outstanding", "payment delay"],
        "debt_ebitda": ["debt to ebitda", "debt/ebitda", "debt ebitda ratio", "leverage ratio", "debt multiple"],
        "revenue_decline": ["revenue decline", "revenue fall", "revenue decrease", "sales decline", "top line decline", "income reduction"],
        "onetime_expenses": ["one-time", "onetime", "exceptional", "extraordinary", "non-recurring", "special charges"],
        "margin_decline": ["margin decline", "margin fall", "margin pressure", "margin compression", "profitability decline", "margin squeeze"],
        "cash_balance": ["cash decline", "cash decrease", "cash balance fall", "liquidity issue", "cash shortage", "cash position"],
        "short_term_debt": ["short-term debt", "current liabilities", "working capital", "short term borrowing", "immediate obligations"],
        "management_issues": ["management change", "leadership change", "CEO", "CFO", "resignation", "departure", "management turnover"],
        "regulatory_compliance": ["regulatory", "compliance", "regulation", "regulator", "legal", "penalty", "violation", "sanctions"],
        "market_competition": ["competition", "competitive", "market share", "competitor", "market pressure", "competitive pressure"],
        "operational_disruptions": ["operational", "supply chain", "production", "manufacturing", "disruption", "operational issues"]
    }
    
    criteria_list = "\n".join([f"{i+1}. {name}: {desc}" for i, (name, desc) in enumerate(criteria_definitions.items())])
    
    # Build keyword list for prompt
    keywords_section = "\nKEYWORDS FOR EACH CRITERIA:\n"
    for criteria, keywords in criteria_keywords.items():
        keywords_section += f"  * {criteria}: {', '.join(keywords)}\n"
        
    prompt = f"""
You are a STRICT financial risk classifier. Follow these EXACT rules with NO exceptions:

RED FLAG TO CLASSIFY: "{flag}"

KEYWORDS FOR EACH CRITERIA:
{keywords_section}

CRITERIA DEFINITIONS:
{criteria_list}

PREVIOUS YEAR DATA FOR THRESHOLD CHECKING:
{previous_year_data}

STRICT CLASSIFICATION ALGORITHM:
Step 1: Check for EXACT KEYWORD MATCH using the keyword listed above
Step 2: If NO keyword match found → AUTOMATICALLY classify as "Low"
Step 3: If keyword match found → Check threshold criteria against previous year data
Step 4: Classify as "High" ONLY if BOTH conditions met:
   a) Exact keyword match exists
   b) Threshold criteria is satisfied

DEFAULT RULE: When in doubt → classify as "Low"

OUTPUT FORMAT (follow exactly):
Matched_Criteria: [exact criteria name if keyword found, otherwise "None"]
Risk_Level: [High only if both keyword AND threshold met, otherwise Low]
Reasoning: [Explain keyword search result and threshold check]
Confidence: [0.0 to 1.0]
"""
    
    response = llm._call(prompt, max_tokens=300, temperature=0.0)
    
    # Initialize with safe defaults
    matched_criteria = 'None'
    risk_level = RiskLevel.LOW
    reasoning = 'No exact keyword match found for any criteria'
    confidence = 0.0
    
    # Parse response
    lines = response.strip().split('\n')
    for line in lines:
        if line.startswith('Matched_Criteria:'):
            matched = line.split(':', 1)[1].strip()
            matched_criteria = matched if matched not in ["None", ""] else 'None'
        elif line.startswith('Risk_Level:'):
            risk_level_str = line.split(':', 1)[1].strip()
            if matched_criteria == 'None':
                risk_level = RiskLevel.LOW
            else:
                risk_level = RiskLevel.HIGH if risk_level_str.lower() == 'high' else RiskLevel.LOW
        elif line.startswith('Reasoning:'):
            reasoning = line.split(':', 1)[1].strip()
        elif line.startswith('Confidence:'):
            try:
                confidence = float(line.split(':', 1)[1].strip())
            except:
                confidence = 0.0
    
    # Final safety check
    if matched_criteria == 'None':
        risk_level = RiskLevel.LOW
        reasoning = 'No criteria keyword match - defaulted to Low risk'
        confidence = 0.0
    
    return FlagClassification(
        flag=flag,
        matched_criteria=matched_criteria,
        risk_level=risk_level,
        reasoning=reasoning,
        confidence_score=confidence
    )

# ================== SUMMARY PARSING ==================

def parse_summary_by_categories(fourth_response: str) -> List[CategorySummary]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = []
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            try:
                category_summary = CategorySummary(
                    category_name=category_name,
                    bullet_points=bullets
                )
                categories_summary.append(category_summary)
            except Exception as e:
                logger.warning(f"Invalid category summary format: {e}")
   
    return categories_summary

# ================== WORD DOCUMENT GENERATION ==================

def generate_strict_high_risk_summary(high_risk_flags: List[str], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    # Use deduplicator for high risk flags
    deduplicator = RuleBasedDeduplicator()
    unique_high_risk_flags = deduplicator.remove_duplicates(high_risk_flags)
    
    concise_summaries = []
    
    for flag in unique_high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context}

HIGH RISK FLAG: "{flag}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Do NOT exceed 2 lines under any circumstances
6. Do NOT start with "Summary:" or any prefix
7. Provide ONLY the factual summary content
8. Make it UNIQUE - avoid repeating information from other summaries

OUTPUT FORMAT: [Direct factual summary only, no labels or prefixes]
"""
        
        try:
            response = llm._call(prompt, max_tokens=100, temperature=0.1)
            
            # Clean response - remove any prefixes or labels
            clean_response = response.strip()
            
            # Remove common prefixes that might appear
            prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:", "Flag summary:", "The flag:", "This flag:"]
            for prefix in prefixes_to_remove:
                if clean_response.startswith(prefix):
                    clean_response = clean_response[len(prefix):].strip()
            
            # Split into lines and take first 2
            summary_lines = [line.strip() for line in clean_response.split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                concise_summary = f"{flag}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
            
            concise_summaries.append(concise_summary)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag}': {e}")
            concise_summaries.append(f"{flag}. Review required based on analysis.")
    
    # Final deduplication of summaries
    return deduplicator.remove_duplicates(concise_summaries)

def create_word_document(pdf_name: str, company_info: CompanyInfo, risk_counts: RiskCounts,
                        high_risk_flags: List[str], category_summaries: List[CategorySummary], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> Optional[str]:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info.formatted_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(risk_counts.high)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(risk_counts.low)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(risk_counts.total)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            for summary in concise_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if category_summaries and len(category_summaries) > 0:
            for category_summary in category_summaries:
                cat_heading = doc.add_heading(category_summary.category_name, level=2)
                if len(cat_heading.runs) > 0:
                    cat_heading.runs[0].bold = True
               
                for bullet in category_summary.bullet_points:
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(bullet)
               
                doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.high}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.low}")
            doc.add_paragraph(f"Total Flags: {risk_counts.total}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

# ================== MAIN PROCESSING PIPELINE ==================

def process_pdf_enhanced_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1") -> Optional[ProcessingResult]:
    """
    Process PDF through enhanced 5-iteration pipeline with Pydantic schemas and rule-based deduplication
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
    start_time = time.time()
   
    try:
        # Initialize LLM and load PDF
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        # Load and validate PDF documents
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0].context
        
        # Load first query from CSV/Excel
        try:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            
            if len(queries_df) == 0 or "prompt" not in queries_df.columns:
                first_query = "Analyze this document for potential red flags."
            else:
                first_query = queries_df["prompt"].tolist()[0]
        except Exception as e:
            logger.warning(f"Error loading queries file: {e}. Using default query.")
            first_query = "Analyze this document for potential red flags."
        
        # ITERATION 1: Initial red flag identification
        print("Running 1st iteration - Initial Analysis...")
        sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.
 
COMPLETE DOCUMENT TO ANALYZE:
{context}
 
Your task is to analyze the ENTIRE document above and identify ALL potential red flags.
 
CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
"""
        
        first_prompt = f"{sys_prompt}\n\nQuestion: {first_query}\n\nAnswer:"
        first_response = llm._call(first_prompt, max_tokens=4000)
        
        # ITERATION 2: Rule-based Deduplication (NO GPT)
        print("Running 2nd iteration - Rule-based Deduplication...")
        deduplicator = RuleBasedDeduplicator()
        second_response = deduplicator.deduplicate_red_flags(first_response)
        
        # ITERATION 3: Categorization (still using GPT as it requires domain expertise)
        print("Running 3rd iteration - Categorization...")
        third_prompt = """You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.
 
- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction.
- Regulatory Issues: Concerns related Compliance with laws, regulations.
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape.
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.
 
While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.
 
**Output Format**:
### Balance Sheet Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### P&L (Income Statement) Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
        
        third_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {second_response}
 
Based on the above analysis and the original context, please answer: {third_prompt}
 
Answer:"""
        
        third_response = llm._call(third_full_prompt, max_tokens=4000)
        
        # ITERATION 4: Summary generation
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:
 
1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
3. **Focus on factual content**: Base the summary solely on the information associated with each red flag, without introducing external opinions or assumptions
4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
8. **Summary should be factual**: Avoid any subjective interpretations or opinions
9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points
 
Format the output exactly like this example:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]
 
Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
        
        fourth_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {third_response}
 
Based on the above analysis and the original context, please answer: {fourth_prompt}
 
Answer:"""
        
        fourth_response = llm._call(fourth_full_prompt, max_tokens=4000)
        
        # ITERATION 5: Rule-based Unique Flags Extraction and Classification
        print("Running 5th iteration - Rule-based Unique Flags Classification...")
        
        # Step 1: Extract unique flags using rule-based approach (NO GPT)
        flag_extractor = UniqueFlageExtractor()
        unique_flags = flag_extractor.extract_unique_flags_rule_based(second_response)
        print(f"\nUnique flags extracted: {len(unique_flags)}")
        
        # Define 15 criteria definitions
        criteria_definitions = {
            "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
            "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
            "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
            "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
            "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational or supply chain issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
        }
        
        # Step 2: Classify each unique flag with STRICT criteria matching
        classification_results = []
        high_risk_flags = []
        low_risk_flags = []
        
        if len(unique_flags) > 0 and unique_flags[0] != "No specific red flags identified":
            for i, flag in enumerate(unique_flags, 1):
                try:
                    classification = classify_flag_against_criteria_strict(
                        flag=flag,
                        criteria_definitions=criteria_definitions,
                        previous_year_data=previous_year_data,
                        llm=llm
                    )
                    
                    classification_results.append(classification)
                    
                    # Only add to high risk if explicitly classified as High AND criteria matched
                    if (classification.risk_level == RiskLevel.HIGH and 
                        classification.matched_criteria != 'None'):
                        high_risk_flags.append(flag)
                    else:
                        low_risk_flags.append(flag)
                        
                except Exception as e:
                    logger.error(f"Error classifying flag {i}: {e}")
                    # Always default to low risk if classification fails
                    default_classification = FlagClassification(
                        flag=flag,
                        matched_criteria='None',
                        risk_level=RiskLevel.LOW,
                        reasoning=f'Classification failed: {str(e)}',
                        confidence_score=0.0
                    )
                    classification_results.append(default_classification)
                    low_risk_flags.append(flag)
                  
                time.sleep(0.3)
        
        # Create risk counts using Pydantic schema
        risk_counts = RiskCounts(
            high=len(high_risk_flags),
            low=len(low_risk_flags),
            total=len(unique_flags) if unique_flags and unique_flags[0] != "No specific red flags identified" else 0
        )
        
        print(f"\n=== FINAL CLASSIFICATION RESULTS ===")
        print(f"High Risk Flags: {risk_counts.high}")
        print(f"Low Risk Flags: {risk_counts.low}")
        print(f"Total Flags: {risk_counts.total}")
        
        if high_risk_flags:
            print(f"\n--- HIGH RISK FLAGS ---")
            for i, flag in enumerate(high_risk_flags, 1):
                print(f"  {i}. {flag}")
        else:
            print(f"\n--- HIGH RISK FLAGS ---")
            print("  No high risk flags identified")
        
        if low_risk_flags:
            print(f"\n--- LOW RISK FLAGS ---")
            for i, flag in enumerate(low_risk_flags, 1):
                print(f"  {i}. {flag}")
        else:
            print(f"\n--- LOW RISK FLAGS ---")
            print("  No low risk flags identified")
        
        # Extract company info and create Word document
        print("\nCreating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            category_summaries = parse_summary_by_categories(fourth_response)
           
            # Create Word document with strict high risk summaries
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                high_risk_flags=high_risk_flags,
                category_summaries=category_summaries,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Calculate processing time
        processing_time = time.time() - start_time
        
        # Create ProcessingResult using Pydantic schema
        result = ProcessingResult(
            pdf_name=pdf_name,
            company_info=company_info,
            risk_counts=risk_counts,
            high_risk_flags=high_risk_flags,
            low_risk_flags=low_risk_flags,
            classification_results=classification_results,
            category_summaries=category_summaries,
            word_doc_path=word_doc_path,
            processing_time=processing_time
        )
        
        # Save all results to CSV files
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Save pipeline results
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Rule-based Deduplication", 
                "Categorization",
                "Summary Generation",
                "Rule-based Unique Flags Classification"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Rule-based Classification: {risk_counts.high} High Risk, {risk_counts.low} Low Risk flags from {risk_counts.total} total unique flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_enhanced_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        # Save detailed classification results
        if len(classification_results) > 0:
            classification_data = []
            for classification in classification_results:
                classification_data.append({
                    'flag': classification.flag,
                    'matched_criteria': classification.matched_criteria,
                    'risk_level': classification.risk_level.value,
                    'reasoning': classification.reasoning,
                    'confidence_score': classification.confidence_score
                })
            
            classification_df = pd.DataFrame(classification_data)
            classification_file = os.path.join(output_folder, f"{pdf_name}_enhanced_flag_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        return result
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

def main():
    """Main function to process all PDFs in the specified folder"""
    
    # Configuration
    pdf_folder_path = r"sterlin_q2_pdf" 
    queries_csv_path = r"EWS_prompts_v2_2.xlsx"
    output_folder = r"sterlin_results_neww_sept"
    
    api_key = "8496bd1d1361e498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01"
    deployment_name = "gpt-4.1"
  
    previous_year_data = """
Previous reported Debt Mar-22 446
Current quarter ebidta Sept-22 -341
Previous reported asset value Mar-22 3500
Previous reported receivable days Mar-22 55
Previous reported payable days	Mar-22	
Previous reported revenue June-22 1207
Previous reported profitability June-22 -356
Previous reported operating margin June-22 -28%
Previous reported cash balance Mar-22 504
Previous reported current liabilities Mar-22 435
"""
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    all_results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        result = process_pdf_enhanced_pipeline(
            pdf_path=pdf_file,
            queries_csv_path=queries_csv_path,
            previous_year_data=previous_year_data,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        if result is not None:
            all_results.append(result)
            print(f"✅ Successfully processed {pdf_file} in {result.processing_time:.2f} seconds")
            print(f"   High Risk: {result.risk_counts.high}, Low Risk: {result.risk_counts.low}, Total: {result.risk_counts.total}")
        else:
            print(f"❌ Failed to process {pdf_file}")
    
    # Generate summary report of all processed files
    if all_results:
        summary_data = []
        for result in all_results:
            summary_data.append({
                'pdf_name': result.pdf_name,
                'company': result.company_info.company_name,
                'quarter': result.company_info.quarter,
                'financial_year': result.company_info.financial_year,
                'high_risk_count': result.risk_counts.high,
                'low_risk_count': result.risk_counts.low,
                'total_flags': result.risk_counts.total,
                'processing_time': result.processing_time,
                'word_doc_created': result.word_doc_path is not None
            })
        
        summary_df = pd.DataFrame(summary_data)
        summary_file = os.path.join(output_folder, "processing_summary.csv")
        summary_df.to_csv(summary_file, index=False)
        print(f"\n📊 Summary report saved to: {summary_file}")

if __name__ == "__main__":
    main()
