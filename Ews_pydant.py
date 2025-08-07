import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any, Optional, Union, Set
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
from pydantic import BaseModel, Field, validator
from difflib import SequenceMatcher
from collections import Counter
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except:
    pass

warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# Pydantic Models
class RedFlag(BaseModel):
    """Model for a single red flag"""
    flag_id: int = Field(..., description="Unique identifier for the flag")
    content: str = Field(..., min_length=5, description="The red flag content")
    original_quote: Optional[str] = Field(None, description="Original quote from document")
    page_number: Optional[int] = Field(None, description="Page number reference")
    speaker: Optional[str] = Field(None, description="Speaker attribution")
    
    @validator('content')
    def validate_content(cls, v):
        if not v or len(v.strip()) < 5:
            raise ValueError('Content must be at least 5 characters long')
        return v.strip()

class DeduplicationResult(BaseModel):
    """Model for deduplication results"""
    original_count: int = Field(..., description="Original number of flags")
    deduplicated_count: int = Field(..., description="Number after deduplication")
    unique_flags: List[RedFlag] = Field(..., description="List of unique flags")
    duplicate_groups: List[List[int]] = Field(default=[], description="Groups of duplicate flag IDs")
    similarity_threshold: float = Field(default=0.8, description="Similarity threshold used")

class ClassificationResult(BaseModel):
    """Model for flag classification results"""
    flag: RedFlag
    matched_criteria: str = Field(..., description="Matched criteria name or 'None'")
    risk_level: str = Field(..., description="Risk level: High or Low")
    reasoning: str = Field(..., description="Classification reasoning")
    confidence_score: Optional[float] = Field(None, ge=0.0, le=1.0, description="Classification confidence")
    
    @validator('risk_level')
    def validate_risk_level(cls, v):
        if v not in ['High', 'Low']:
            raise ValueError('Risk level must be either "High" or "Low"')
        return v

class IterationResult(BaseModel):
    """Model for iteration results"""
    iteration_number: int = Field(..., ge=1, le=5, description="Iteration number (1-5)")
    stage_name: str = Field(..., description="Name of the processing stage")
    response_text: str = Field(..., description="Response from this iteration")
    processing_time: float = Field(..., ge=0, description="Processing time in seconds")
    timestamp: str = Field(..., description="Processing timestamp")
    flags_extracted: Optional[List[RedFlag]] = Field(None, description="Flags extracted in this iteration")

class ProcessingResults(BaseModel):
    """Model for complete processing results"""
    pdf_name: str = Field(..., description="Name of the processed PDF")
    company_info: str = Field(..., description="Extracted company information")
    iterations: List[IterationResult] = Field(..., description="Results from all iterations")
    deduplication_result: Optional[DeduplicationResult] = Field(None, description="Deduplication results")
    classification_results: List[ClassificationResult] = Field(default=[], description="Classification results")
    high_risk_count: int = Field(default=0, ge=0, description="Number of high risk flags")
    low_risk_count: int = Field(default=0, ge=0, description="Number of low risk flags")
    total_processing_time: float = Field(..., ge=0, description="Total processing time")

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

class AzureOpenAILLM:
    """Azure OpenAI gpt-4.1-mini-mini LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1-mini"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI gpt-4.1-mini-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

class PythonDeduplicator:
    """Python-based deduplication class using text normalization and keyword extraction"""
    
    def __init__(self, similarity_threshold: float = 0.7):
        self.similarity_threshold = similarity_threshold
        
        # Initialize NLTK components
        try:
            self.stop_words = set(stopwords.words('english'))
        except:
            self.stop_words = {
                'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 'yours',
                'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 'her', 'hers',
                'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves',
                'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those', 'am', 'is', 'are',
                'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does',
                'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until',
                'while', 'of', 'at', 'by', 'for', 'with', 'through', 'during', 'before', 'after',
                'above', 'below', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again',
                'further', 'then', 'once'
            }
        
        try:
            self.stemmer = PorterStemmer()
        except:
            self.stemmer = None
            
        try:
            self.lemmatizer = WordNetLemmatizer()
        except:
            self.lemmatizer = None
    
    def normalize_text(self, text: str) -> str:
        """Normalize text by cleaning and standardizing"""
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters but keep spaces, numbers, and letters
        text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common business terms that don't add meaning for deduplication
        business_noise = {
            'quarter', 'fiscal', 'year', 'company', 'business', 'management', 'financial',
            'report', 'statement', 'analysis', 'discussion', 'results', 'performance'
        }
        
        words = text.split()
        words = [word for word in words if word not in business_noise]
        
        return ' '.join(words).strip()
    
    def extract_keywords(self, text: str) -> Set[str]:
        """Extract meaningful keywords from text using NLTK"""
        normalized_text = self.normalize_text(text)
        
        if not normalized_text:
            return set()
        
        # Tokenize
        try:
            tokens = word_tokenize(normalized_text)
        except:
            tokens = normalized_text.split()
        
        # Remove stop words and short words
        keywords = []
        for token in tokens:
            if (len(token) >= 3 and 
                token not in self.stop_words and 
                not token.isdigit()):
                keywords.append(token)
        
        # Apply stemming/lemmatization if available
        processed_keywords = set()
        for keyword in keywords:
            if self.lemmatizer:
                try:
                    processed_keyword = self.lemmatizer.lemmatize(keyword)
                except:
                    processed_keyword = keyword
            elif self.stemmer:
                try:
                    processed_keyword = self.stemmer.stem(keyword)
                except:
                    processed_keyword = keyword
            else:
                processed_keyword = keyword
            
            processed_keywords.add(processed_keyword)
        
        return processed_keywords
    
    def calculate_keyword_similarity(self, keywords1: Set[str], keywords2: Set[str]) -> float:
        """Calculate similarity between two sets of keywords using Jaccard similarity"""
        if not keywords1 and not keywords2:
            return 1.0
        
        if not keywords1 or not keywords2:
            return 0.0
        
        # Jaccard similarity: |intersection| / |union|
        intersection = keywords1.intersection(keywords2)
        union = keywords1.union(keywords2)
        
        if len(union) == 0:
            return 0.0
        
        jaccard_similarity = len(intersection) / len(union)
        
        # Also consider overlap percentage from each set's perspective
        overlap1 = len(intersection) / len(keywords1) if keywords1 else 0
        overlap2 = len(intersection) / len(keywords2) if keywords2 else 0
        
        # Take the maximum overlap percentage
        max_overlap = max(overlap1, overlap2)
        
        # Weighted combination: 60% Jaccard + 40% max overlap
        combined_similarity = (jaccard_similarity * 0.6) + (max_overlap * 0.4)
        
        return combined_similarity
    
    def calculate_sequence_similarity(self, text1: str, text2: str) -> float:
        """Calculate sequence similarity using difflib"""
        norm1 = self.normalize_text(text1)
        norm2 = self.normalize_text(text2)
        
        if not norm1 and not norm2:
            return 1.0
        
        if not norm1 or not norm2:
            return 0.0
        
        return SequenceMatcher(None, norm1, norm2).ratio()
    
    def extract_financial_terms(self, text: str) -> Set[str]:
        """Extract financial-specific terms that are important for similarity"""
        financial_patterns = [
            r'\b(debt|revenue|profit|loss|margin|cash|asset|liability)\b',
            r'\b(ebitda|earnings|sales|income|expense)\b',
            r'\b(increase|decrease|decline|fall|rise|growth)\b',
            r'\b(provision|write-off|impairment)\b',
            r'\b(ratio|percentage|basis\s+points?)\b',
            r'\b\d+(?:\.\d+)?%\b',  # Percentages
            r'\b\d+(?:,\d{3})*(?:\.\d+)?\s*(?:cr|crore|million|billion|k|m|b)\b'  # Numbers with units
        ]
        
        financial_terms = set()
        text_lower = text.lower()
        
        for pattern in financial_patterns:
            matches = re.findall(pattern, text_lower)
            if isinstance(matches, list):
                for match in matches:
                    if isinstance(match, str):
                        financial_terms.add(match.strip())
                    elif isinstance(match, tuple):
                        for item in match:
                            if item:
                                financial_terms.add(item.strip())
        
        return financial_terms
    
    def is_similar(self, flag1: RedFlag, flag2: RedFlag) -> bool:
        """Determine if two flags are similar using keyword-based approach"""
        
        # Extract keywords from both flags
        keywords1 = self.extract_keywords(flag1.content)
        keywords2 = self.extract_keywords(flag2.content)
        
        # Extract financial terms
        financial1 = self.extract_financial_terms(flag1.content)
        financial2 = self.extract_financial_terms(flag2.content)
        
        # Combine regular keywords with financial terms
        all_keywords1 = keywords1.union(financial1)
        all_keywords2 = keywords2.union(financial2)
        
        # Calculate keyword similarity
        keyword_sim = self.calculate_keyword_similarity(all_keywords1, all_keywords2)
        
        # Calculate sequence similarity as a backup
        sequence_sim = self.calculate_sequence_similarity(flag1.content, flag2.content)
        
        # If either similarity is very high, consider them similar
        if keyword_sim >= self.similarity_threshold or sequence_sim >= 0.8:
            return True
        
        # Additional check for short texts (less than 10 words)
        words1 = len(flag1.content.split())
        words2 = len(flag2.content.split())
        
        if words1 <= 10 or words2 <= 10:
            # For short texts, use a higher threshold for sequence similarity
            return sequence_sim >= 0.7
        
        # Weighted combination for longer texts
        combined_similarity = (keyword_sim * 0.7) + (sequence_sim * 0.3)
        
        return combined_similarity >= self.similarity_threshold
    
    def find_duplicates(self, flags: List[RedFlag]) -> List[List[int]]:
        """Find groups of duplicate flags using keyword-based similarity"""
        if len(flags) <= 1:
            return []
        
        duplicate_groups = []
        processed_indices = set()
        
        for i in range(len(flags)):
            if i in processed_indices:
                continue
                
            current_group = [i]
            
            for j in range(i + 1, len(flags)):
                if j in processed_indices:
                    continue
                    
                if self.is_similar(flags[i], flags[j]):
                    current_group.append(j)
                    processed_indices.add(j)
            
            if len(current_group) > 1:
                duplicate_groups.append(current_group)
                processed_indices.update(current_group)
        
        return duplicate_groups
    
    def deduplicate_flags(self, flags: List[RedFlag]) -> DeduplicationResult:
        """Deduplicate flags and return results"""
        original_count = len(flags)
        
        if original_count <= 1:
            return DeduplicationResult(
                original_count=original_count,
                deduplicated_count=original_count,
                unique_flags=flags,
                duplicate_groups=[],
                similarity_threshold=self.similarity_threshold
            )
        
        # Find duplicate groups
        duplicate_groups = self.find_duplicates(flags)
        
        # Create list of unique flags (keep first from each duplicate group)
        unique_indices = set(range(len(flags)))
        
        for group in duplicate_groups:
            # Keep the first flag from each group (usually the most complete), remove others
            for idx in group[1:]:
                unique_indices.discard(idx)
        
        unique_flags = [flags[i] for i in sorted(unique_indices)]
        
        # Log deduplication details
        if duplicate_groups:
            print(f"  Deduplication found {len(duplicate_groups)} duplicate groups:")
            for i, group in enumerate(duplicate_groups, 1):
                print(f"    Group {i}: {len(group)} similar flags (keeping index {group[0]})")
        else:
            print("  No duplicates found")
        
        return DeduplicationResult(
            original_count=original_count,
            deduplicated_count=len(unique_flags),
            unique_flags=unique_flags,
            duplicate_groups=duplicate_groups,
            similarity_threshold=self.similarity_threshold
        )

def parse_flags_from_response(response_text: str) -> List[RedFlag]:
    """Parse red flags from LLM response text"""
    flags = []
    lines = response_text.split('\n')
    current_flag = None
    flag_counter = 1
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Look for numbered flags (1., 2., etc.)
        flag_match = re.match(r'^(\d+)\.\s*(.+)', line)
        if flag_match:
            if current_flag:
                # Save previous flag
                try:
                    flags.append(current_flag)
                except Exception as e:
                    logger.warning(f"Invalid flag data: {e}")
            
            # Start new flag
            flag_content = flag_match.group(2).strip()
            current_flag = RedFlag(
                flag_id=flag_counter,
                content=flag_content
            )
            flag_counter += 1
            continue
        
        # Look for "Original Quote:" lines
        if line.startswith('Original Quote:') and current_flag:
            quote_match = re.search(r'Original Quote:\s*"([^"]+)"\s*(?:\(([^)]+)\))?', line)
            if quote_match:
                current_flag.original_quote = quote_match.group(1)
                page_info = quote_match.group(2)
                if page_info:
                    # Extract page number and speaker
                    page_match = re.search(r'Page\s+(\d+)', page_info)
                    if page_match:
                        current_flag.page_number = int(page_match.group(1))
            continue
        
        # If we have a current flag and this line looks like content, append it
        if current_flag and not line.startswith(('Original Quote:', 'Page', 'Speaker:')):
            current_flag.content += " " + line
    
    # Don't forget the last flag
    if current_flag:
        try:
            flags.append(current_flag)
        except Exception as e:
            logger.warning(f"Invalid flag data: {e}")
    
    return flags

def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""<role>
You are an expert document analyst specializing in extracting key corporate information from financial documents and earnings transcripts.
</role>

<system_prompt>
You excel at quickly identifying and extracting specific corporate identifiers from financial documents with high accuracy and consistency.
</system_prompt>

<instruction>
Extract the company name, quarter, and financial year from the provided text.

EXTRACTION REQUIREMENTS:
1. Company Name: Full legal company name including suffixes (Ltd/Limited/Inc/Corp etc.)
2. Quarter: Identify the quarter (Q1/Q2/Q3/Q4)  
3. Financial Year: Extract financial year (FY23/FY24/FY25 etc.)

OUTPUT FORMAT:
Provide ONLY the result in this exact format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If any component cannot be clearly identified, use reasonable defaults based on context.
</instruction>

<context>
DOCUMENT TEXT TO ANALYZE:
{first_page_text}
</context>

Extract company information:"""
       
        response = llm._call(prompt, max_tokens=200)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def classify_flag_against_criteria_strict(flag: RedFlag, criteria_definitions: Dict[str, str], 
                                 previous_year_data: str, llm: AzureOpenAILLM) -> ClassificationResult:
    """Strictly classify a single flag against 15 criteria with enhanced accuracy"""
    
    # Enhanced keyword mapping with more comprehensive coverage
    criteria_keywords = {
        "debt_increase": [
            "debt increase", "debt increased", "debt rising", "debt growth", "higher debt", "debt went up", 
            "debt levels increased", "borrowing increase", "borrowings increased", "borrowing levels",
            "leverage increased", "leverage higher", "total debt", "gross debt", "net debt increase",
            "debt position", "debt burden", "indebtedness", "borrowed funds", "credit facilities",
            "term loans", "working capital loans", "bank borrowings", "financial liabilities increased"
        ],
        "provisioning": [
            "provision", "provisions", "provisioning", "write-off", "write off", "writeoff", "written off",
            "bad debt", "doubtful debt", "impairment", "credit loss", "expected credit loss",
            "loan loss provision", "bad debt provision", "impairment provision", "credit impairment",
            "provision for doubtful debts", "allowance for credit losses", "provision expense",
            "asset write-off", "inventory write-off", "receivables written off"
        ],
        "asset_decline": [
            "asset decline", "asset fall", "asset decrease", "asset value down", "asset reduction",
            "asset impairment", "asset deterioration", "asset value fell", "assets decreased",
            "fixed assets", "current assets", "total assets", "net assets", "asset base",
            "property values", "equipment value", "inventory value", "investment value"
        ],
        "receivable_days": [
            "receivable days", "collection period", "DSO", "days sales outstanding", "collection time",
            "receivables collection", "collection efficiency", "debtor days", "trade receivables days",
            "slow collections", "delayed collections", "collection challenges", "extended credit terms",
            "receivables aging", "overdue receivables", "collection period increased"
        ],
        "payable_days": [
            "payable days", "payment period", "DPO", "days payable outstanding", "payment delay",
            "creditor days", "trade payables days", "supplier payment terms", "payment cycles",
            "extended payment terms", "delayed payments", "payment scheduling", "payables management"
        ],
        "debt_ebitda": [
            "debt to ebitda", "debt/ebitda", "debt ebitda ratio", "leverage ratio", "debt multiple",
            "debt to earnings ratio", "leverage multiple", "net debt to ebitda", "gross debt to ebitda",
            "high leverage", "leverage levels", "debt coverage", "debt service coverage",
            "financial leverage", "debt capacity", "borrowing capacity"
        ],
        "revenue_decline": [
            "revenue decline", "revenue fall", "revenue decrease", "sales decline", "top line decline",
            "income reduction", "turnover decline", "revenue drop", "sales fall", "sales decrease",
            "lower sales", "reduced revenue", "revenue pressure", "sales pressure", "revenue challenges",
            "revenue contraction", "sales contraction", "business decline", "revenue performance"
        ],
        "onetime_expenses": [
            "one-time", "onetime", "one time", "exceptional", "extraordinary", "non-recurring",
            "special charges", "unusual items", "exceptional items", "one-off", "oneoff",
            "restructuring costs", "impairment charges", "settlement costs", "litigation costs",
            "acquisition costs", "integration costs", "exit costs", "disposal costs"
        ],
        "margin_decline": [
            "margin decline", "margin fall", "margin pressure", "margin compression", "profitability decline",
            "margin squeeze", "gross margin", "operating margin", "profit margins", "margin erosion",
            "profitability pressure", "profit decline", "earnings pressure", "margin contraction",
            "cost pressures", "pricing pressure", "margin deterioration", "profit margin fell"
        ],
        "cash_balance": [
            "cash decline", "cash decrease", "cash balance fall", "liquidity issue", "cash shortage",
            "cash position", "cash flow problems", "cash constraints", "liquidity constraints",
            "cash management", "cash availability", "cash reserves", "liquid assets", "cash resources",
            "working capital", "free cash flow", "operating cash flow", "cash generation"
        ],
        "short_term_debt": [
            "short-term debt", "current liabilities", "short term borrowing", "immediate obligations",
            "current debt", "near-term debt", "short-term loans", "current portion",
            "working capital deficit", "current ratio", "quick ratio", "liquidity ratio",
            "short-term financing", "bridge financing", "temporary financing"
        ],
        "management_issues": [
            "management change", "leadership change", "CEO", "CFO", "resignation", "departure",
            "management turnover", "executive changes", "board changes", "leadership transition",
            "management performance", "leadership issues", "governance issues", "management quality",
            "execution issues", "strategic direction", "management effectiveness"
        ],
        "regulatory_compliance": [
            "regulatory", "regulation", "regulator", "compliance", "legal", "penalty", "violation",
            "sanctions", "regulatory action", "compliance issues", "regulatory risk",
            "license", "permit", "authorization", "regulatory approval", "government action",
            "regulatory investigation", "enforcement action", "regulatory scrutiny"
        ],
        "market_competition": [
            "competition", "competitive", "market share", "competitor", "market pressure",
            "competitive pressure", "competitive landscape", "market dynamics", "competitive position",
            "market challenges", "industry challenges", "competitive threats", "market disruption",
            "pricing competition", "competitive intensity", "market saturation"
        ],
        "operational_disruptions": [
            "operational", "supply chain", "production", "manufacturing", "disruption",
            "operational issues", "operational challenges", "supply chain issues", "production issues",
            "operational efficiency", "process issues", "system issues", "infrastructure issues",
            "operational performance", "capacity utilization", "operational difficulties"
        ]
    }
    
    criteria_list = "\n".join([f"{i+1}. {name}: {desc}" for i, (name, desc) in enumerate(criteria_definitions.items())])
    
    # Build comprehensive keyword list for prompt
    keywords_section = "\nCOMPREHENSIVE KEYWORDS FOR EACH CRITERIA:\n"
    for criteria, keywords in criteria_keywords.items():
        keywords_section += f"  * {criteria}: {', '.join(keywords[:10])}...\n"
        
    prompt = f"""<role>
You are a STRICT financial risk classifier and expert quantitative analyst with 20+ years of experience in financial risk assessment and criteria-based classification systems.
</role>

<system_prompt>
You excel at precise classification using comprehensive keyword matching and rigorous threshold-based risk assessment. You have zero tolerance for ambiguous classifications and always err on the side of accuracy over false positives.
</system_prompt>

<instruction>
Classify the given red flag against the 15 criteria using ENHANCED STRICT rules for maximum accuracy in identifying HIGH RISK flags.

OUTPUT FORMAT (follow exactly):
Matched_Criteria: [exact criteria name if keyword found, otherwise "None"]
Risk_Level: [High if keyword match + threshold/severity criteria met, otherwise Low]
Reasoning: [Detailed explanation of keyword analysis, threshold check, and classification logic]
Confidence_Score: [0.0 to 1.0 representing classification confidence]
</instruction>

<context>
RED FLAG TO CLASSIFY: "{flag.content}"

{keywords_section}

FULL CRITERIA DEFINITIONS:
{criteria_list}

PREVIOUS YEAR DATA FOR THRESHOLD CHECKING:
{previous_year_data}
</context>

Classify the red flag:"""
    
    response = llm._call(prompt, max_tokens=400, temperature=0.0)
    
    # Initialize with safe defaults
    matched_criteria = 'None'
    risk_level = 'Low'
    reasoning = 'No keyword match found for any criteria'
    confidence_score = 0.5
    
    # Parse response
    lines = response.strip().split('\n')
    for line in lines:
        if line.startswith('Matched_Criteria:'):
            matched = line.split(':', 1)[1].strip()
            matched_criteria = matched if matched not in ["None", ""] else 'None'
        elif line.startswith('Risk_Level:'):
            risk_level_raw = line.split(':', 1)[1].strip()
            risk_level = risk_level_raw if risk_level_raw in ['High', 'Low'] else 'Low'
        elif line.startswith('Reasoning:'):
            reasoning = line.split(':', 1)[1].strip()
        elif line.startswith('Confidence_Score:'):
            try:
                confidence_score = float(line.split(':', 1)[1].strip())
                confidence_score = max(0.0, min(1.0, confidence_score))  # Clamp to [0,1]
            except:
                confidence_score = 0.5
    
    # Enhanced post-processing for better High risk detection
    flag_lower = flag.content.lower()
    
    # Check for any keyword matches that might have been missed
    for criteria_name, keywords in criteria_keywords.items():
        for keyword in keywords:
            if keyword.lower() in flag_lower:
                if matched_criteria == 'None':
                    matched_criteria = criteria_name
                    # If we found a keyword match, be more lenient with High classification
                    if any(severity_word in flag_lower for severity_word in 
                           ['significant', 'major', 'substantial', 'critical', 'severe', 
                            'increased', 'higher', 'declined', 'fell', 'decreased', 'dropped']):
                        risk_level = 'High'
                        reasoning = f"Keyword match found for {criteria_name}: '{keyword}' with severity indicators"
                        confidence_score = 0.8
                break
        if matched_criteria != 'None':
            break
    
    # Final safety check
    if matched_criteria == 'None':
        risk_level = 'Low'
        reasoning = 'No criteria keyword match - defaulted to Low risk'
        confidence_score = 0.9
    
    return ClassificationResult(
        flag=flag,
        matched_criteria=matched_criteria,
        risk_level=risk_level,
        reasoning=reasoning,
        confidence_score=confidence_score
    )

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(high_risk_flags: List[RedFlag], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    # Use Python deduplicator for final deduplication
    deduplicator = PythonDeduplicator(similarity_threshold=0.9)
    dedup_result = deduplicator.deduplicate_flags(high_risk_flags)
    unique_high_risk_flags = dedup_result.unique_flags
    
    concise_summaries = []
    
    for flag in unique_high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context}

HIGH RISK FLAG: "{flag.content}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Do NOT exceed 2 lines under any circumstances

OUTPUT FORMAT: [Direct factual summary only, no labels or prefixes]
"""
        
        try:
            response = llm._call(prompt, max_tokens=100, temperature=0.1)
            
            # Clean response
            clean_response = response.strip()
            
            # Remove common prefixes
            prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:"]
            for prefix in prefixes_to_remove:
                if clean_response.startswith(prefix):
                    clean_response = clean_response[len(prefix):].strip()
            
            # Split into lines and take first 2
            summary_lines = [line.strip() for line in clean_response.split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                concise_summary = f"{flag.content}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
            
            concise_summaries.append(concise_summary)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag.content}': {e}")
            concise_summaries.append(f"{flag.content}. Review required based on analysis.")
    
    return concise_summaries

def create_word_document(pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                        high_risk_flags: List[RedFlag], summary_by_categories: Dict[str, List[str]], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> str:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        high_count = risk_counts.get('High', 0)
        low_count = risk_counts.get('Low', 0)
        total_count = high_count + low_count
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(high_count)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(low_count)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(total_count)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            for summary in concise_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.get('High', 0)}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.get('Low', 0)}")
            doc.add_paragraph(f"Total Flags: {risk_counts.get('Total', 0)}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

def process_pdf_enhanced_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1-mini") -> Optional[ProcessingResults]:
    """
    Process PDF through enhanced 5-iteration pipeline with Python deduplication and Pydantic validation
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
    start_time = time.time()
    
    try:
        # Initialize LLM and load PDF
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load first query from CSV/Excel
        try:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            
            if len(queries_df) == 0 or "prompt" not in queries_df.columns:
                first_query = "Analyze this document for potential red flags."
            else:
                first_query = queries_df["prompt"].tolist()[0]
        except Exception as e:
            logger.warning(f"Error loading queries file: {e}. Using default query.")
            first_query = "Analyze this document for potential red flags."
        
        iterations = []
        
        # ITERATION 1: Initial red flag identification with structured prompt
        print("Running 1st iteration - Initial Analysis...")
        iter1_start = time.time()
        
        first_prompt = f"""<role>
You are an expert financial analyst with 15+ years of experience specializing in identifying red flags from earnings call transcripts and financial documents.
</role>

<system_prompt>
You excel at comprehensive document analysis, identifying subtle financial risks, and providing detailed evidence-based assessments with precise documentation.
</system_prompt>

<instruction>
Analyze the ENTIRE document and identify ALL potential red flags comprehensively.

ANALYSIS REQUIREMENTS:
- Review every section of the document thoroughly
- Identify financial, operational, strategic, and management risks
- Focus on quantitative concerns with specific data points
- Document exact quotes with speaker attribution
- Number each red flag sequentially (1, 2, 3, etc.)
- Include page references where available

OUTPUT FORMAT:
For each red flag:
1. The potential red flag you observed - [brief description]
Original Quote: "[exact quote with speaker name]" (Page X)

CRITICAL: Ensure comprehensive analysis of the entire document.
</instruction>

<context>
COMPLETE DOCUMENT TO ANALYZE:
{context}

SPECIFIC QUESTION: {first_query}
</context>

Provide comprehensive red flag analysis:"""
        
        first_response = llm._call(first_prompt, max_tokens=4000)
        iter1_time = time.time() - iter1_start
        
        # Parse flags from first iteration
        flags_iter1 = parse_flags_from_response(first_response)
        
        iterations.append(IterationResult(
            iteration_number=1,
            stage_name="Initial Analysis",
            response_text=first_response,
            processing_time=iter1_time,
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S"),
            flags_extracted=flags_iter1
        ))
        
        # ITERATION 2: Python-based Deduplication (NO GPT)
        print("Running 2nd iteration - Python Deduplication...")
        iter2_start = time.time()
        
        # Initialize Python deduplicator with keyword-based approach
        deduplicator = PythonDeduplicator(similarity_threshold=0.7)
        
        # Perform deduplication
        deduplication_result = deduplicator.deduplicate_flags(flags_iter1)
        
        # Create response text for consistency
        second_response = f"""Python Keyword-Based Deduplication Results:
Original flags count: {deduplication_result.original_count}
After deduplication: {deduplication_result.deduplicated_count}
Similarity threshold: {deduplication_result.similarity_threshold}
Duplicate groups found: {len(deduplication_result.duplicate_groups)}

Deduplication Method: Keyword extraction + Set operations (Union/Intersection)
- Text normalization (lowercase, remove special chars)
- NLTK tokenization and stop word removal  
- Stemming/Lemmatization for keyword standardization
- Financial term extraction with regex patterns
- Jaccard similarity on keyword sets
- Sequence similarity as backup for short texts

Deduplicated Flags:
"""
        
        for i, flag in enumerate(deduplication_result.unique_flags, 1):
            second_response += f"{i}. {flag.content}\n"
            if flag.original_quote:
                second_response += f"Original Quote: \"{flag.original_quote}\""
                if flag.page_number:
                    second_response += f" (Page {flag.page_number})"
                second_response += "\n"
            second_response += "\n"
        
        if deduplication_result.duplicate_groups:
            second_response += "Duplicate Groups Removed:\n"
            for i, group in enumerate(deduplication_result.duplicate_groups, 1):
                second_response += f"Group {i}: Kept flag {group[0]}, removed flags {group[1:]}\n"
        
        iter2_time = time.time() - iter2_start
        
        iterations.append(IterationResult(
            iteration_number=2,
            stage_name="Python Deduplication",
            response_text=second_response,
            processing_time=iter2_time,
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S"),
            flags_extracted=deduplication_result.unique_flags
        ))
        
        # ITERATION 3: Categorization with structured prompt
        print("Running 3rd iteration - Categorization...")
        iter3_start = time.time()
        
        third_prompt = f"""<role>
You are a senior financial analyst expert in financial risk categorization with deep knowledge of balance sheet analysis, P&L assessment, and corporate risk frameworks.
</role>

<system_prompt>
You excel at organizing financial risks into standardized categories, ensuring comprehensive coverage of all financial risk areas, and maintaining accuracy in risk classification.
</system_prompt>

<instruction>
Categorize the identified red flags into the following 7 standardized categories based on their financial nature and business impact.

MANDATORY CATEGORIES:
1. Balance Sheet Issues: Assets, liabilities, equity, debt, and overall financial position concerns
2. P&L (Income Statement) Issues: Revenue, expenses, profits, and financial performance concerns  
3. Liquidity Issues: Short-term obligations, cash flow, debt repayment, working capital concerns
4. Management and Strategy Issues: Leadership, governance, decision-making, strategy, and vision concerns
5. Regulatory Issues: Compliance, laws, regulations, and regulatory body concerns
6. Industry and Market Issues: Market position, industry trends, competitive landscape concerns
7. Operational Issues: Internal processes, systems, infrastructure, and operational efficiency concerns

OUTPUT FORMAT:
### Balance Sheet Issues
- [Red flag 1 with original quote and page reference]
- [Red flag 2 with original quote and page reference]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote and page reference]

Continue this format for all applicable categories.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}
</context>

Provide categorized analysis:"""
        
        third_response = llm._call(third_prompt, max_tokens=4000)
        iter3_time = time.time() - iter3_start
        
        iterations.append(IterationResult(
            iteration_number=3,
            stage_name="Categorization",
            response_text=third_response,
            processing_time=iter3_time,
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S")
        ))
        
        # ITERATION 4: Summary generation with structured prompt
        print("Running 4th iteration - Summary Generation...")
        iter4_start = time.time()
        
        fourth_prompt = f"""<role>
You are an expert financial summarization specialist with expertise in creating concise, factual, and comprehensive summaries that preserve critical quantitative data and key insights.
</role>

<system_prompt>
You excel at distilling complex financial analysis into clear, actionable summaries while maintaining objectivity, preserving all quantitative details, and ensuring no critical information is lost.
</system_prompt>

<instruction>
Create a comprehensive summary of each category of red flags in bullet point format following these strict guidelines.

SUMMARY REQUIREMENTS:
1. Retain ALL quantitative information (numbers, percentages, ratios, dates)
2. Maintain completely neutral, factual tone - no opinions or interpretations
3. Include every red flag from each category - no omissions
4. Base content solely on the provided categorized analysis
5. Preserve specific data points and statistics wherever mentioned
6. Each bullet point should capture key details of individual red flags
7. Balance thoroughness with conciseness
8. Use professional financial terminology
9. Ensure category-specific content alignment

OUTPUT FORMAT:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories that contain red flags.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        fourth_response = llm._call(fourth_prompt, max_tokens=4000)
        iter4_time = time.time() - iter4_start
        
        iterations.append(IterationResult(
            iteration_number=4,
            stage_name="Summary Generation",
            response_text=fourth_response,
            processing_time=iter4_time,
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S")
        ))
        
        # ITERATION 5: Classification with enhanced criteria matching
        print("Running 5th iteration - Enhanced Classification...")
        iter5_start = time.time()
        
        # Define 15 criteria definitions
        criteria_definitions = {
            "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
            "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
            "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
            "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
            "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
        }
        
        # Classify each unique flag
        classification_results = []
        high_risk_flags = []
        low_risk_flags = []
        
        unique_flags = deduplication_result.unique_flags
        
        if len(unique_flags) > 0:
            for i, flag in enumerate(unique_flags, 1):
                try:
                    classification = classify_flag_against_criteria_strict(
                        flag=flag,
                        criteria_definitions=criteria_definitions,
                        previous_year_data=previous_year_data,
                        llm=llm
                    )
                    
                    classification_results.append(classification)
                    
                    # Add to appropriate risk category
                    if (classification.risk_level.lower() == 'high' and 
                        classification.matched_criteria != 'None'):
                        high_risk_flags.append(flag)
                    else:
                        low_risk_flags.append(flag)
                        
                except Exception as e:
                    logger.error(f"Error classifying flag {i}: {e}")
                    # Always default to low risk if classification fails
                    classification_results.append(ClassificationResult(
                        flag=flag,
                        matched_criteria='None',
                        risk_level='Low',
                        reasoning=f'Classification failed: {str(e)}',
                        confidence_score=0.0
                    ))
                    low_risk_flags.append(flag)
                  
                time.sleep(0.3)
        
        iter5_time = time.time() - iter5_start
        
        # Create fifth iteration response text
        fifth_response = f"""Enhanced Classification Results:
Total flags classified: {len(unique_flags)}
High risk flags: {len(high_risk_flags)}
Low risk flags: {len(low_risk_flags)}

Classification Details:
"""
        
        for result in classification_results:
            fifth_response += f"Flag: {result.flag.content}\n"
            fifth_response += f"Criteria: {result.matched_criteria}\n"
            fifth_response += f"Risk Level: {result.risk_level}\n"
            fifth_response += f"Confidence: {result.confidence_score:.2f}\n"
            fifth_response += f"Reasoning: {result.reasoning}\n\n"
        
        iterations.append(IterationResult(
            iteration_number=5,
            stage_name="Enhanced Classification",
            response_text=fifth_response,
            processing_time=iter5_time,
            timestamp=time.strftime("%Y-%m-%d %H:%M:%S")
        ))
        
        risk_counts = {
            'High': len(high_risk_flags),
            'Low': len(low_risk_flags),
            'Total': len(unique_flags)
        }
        
        print(f"\n=== FINAL CLASSIFICATION RESULTS ===")
        print(f"High Risk Flags: {risk_counts['High']}")
        print(f"Low Risk Flags: {risk_counts['Low']}")
        print(f"Total Flags: {risk_counts['Total']}")
        
        if high_risk_flags:
            print(f"\n--- HIGH RISK FLAGS ---")
            for i, flag in enumerate(high_risk_flags, 1):
                print(f"  {i}. {flag.content}")
        
        # Extract company info and create Word document
        print("\nCreating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
           
            # Create Word document
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                high_risk_flags=high_risk_flags,
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Create ProcessingResults with Pydantic validation
        total_time = time.time() - start_time
        
        processing_results = ProcessingResults(
            pdf_name=pdf_name,
            company_info=company_info,
            iterations=iterations,
            deduplication_result=deduplication_result,
            classification_results=classification_results,
            high_risk_count=len(high_risk_flags),
            low_risk_count=len(low_risk_flags),
            total_processing_time=total_time
        )
        
        # Save results to files
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Save pipeline results as JSON (Pydantic model)
        json_file = os.path.join(output_folder, f"{pdf_name}_processing_results.json")
        with open(json_file, 'w') as f:
            f.write(processing_results.json(indent=2))
        
        # Save pipeline results as CSV (for backward compatibility)
        results_summary = pd.DataFrame([{
            "pdf_name": result.stage_name,
            "iteration": result.iteration_number,
            "stage": result.stage_name,
            "response": result.response_text,
            "processing_time": result.processing_time,
            "timestamp": result.timestamp
        } for result in iterations])
       
        results_file = os.path.join(output_folder, f"{pdf_name}_enhanced_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        # Save detailed classification results
        if len(classification_results) > 0:
            classification_data = []
            for result in classification_results:
                classification_data.append({
                    'flag_id': result.flag.flag_id,
                    'flag_content': result.flag.content,
                    'original_quote': result.flag.original_quote,
                    'page_number': result.flag.page_number,
                    'matched_criteria': result.matched_criteria,
                    'risk_level': result.risk_level,
                    'reasoning': result.reasoning,
                    'confidence_score': result.confidence_score
                })
            
            classification_df = pd.DataFrame(classification_data)
            classification_file = os.path.join(output_folder, f"{pdf_name}_enhanced_flag_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        return processing_results
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

def main():
    """Main function to process all PDFs in the specified folder"""
    
    # Configuration
    pdf_folder_path = r"chemplast_pdf" 
    queries_csv_path = r"EWS_prompts_v2_2.xlsx"
    output_folder = r"chemplast_results_04"

    api_key = "8496bd1da4e498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01-preview"
    deployment_name = "gpt-4.1-mini"
  
    previous_year_data = """
Previous reported Debt	Mar-22	882Cr
Current quarter ebidta	March-23 130Cr
Previous reported asset value	Mar-22	5602Cr
Previous reported receivable days	Mar-22	12days
Previous reported payable days	Mar-22	189days
Previous reported revenue	Dec-22	1189Cr
Previous reported profitability	Dec-22	27Cr
Previous reported operating margin	Dec-22	7%
Previous reported cash balance	Mar-22	1229Cr
Previous reported current liabilities	Mar-22	68Cr
"""
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    all_results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        result = process_pdf_enhanced_pipeline(
            pdf_path=pdf_file,
            queries_csv_path=queries_csv_path,
            previous_year_data=previous_year_data,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
            all_results.append(result)
        else:
            print(f"❌ Failed to process {pdf_file}")
    
    # Save consolidated results
    if all_results:
        consolidated_file = os.path.join(output_folder, "consolidated_results.json")
        consolidated_data = {
            "processing_summary": {
                "total_pdfs": len(pdf_files),
                "successful_pdfs": len(all_results),
                "failed_pdfs": len(pdf_files) - len(all_results),
                "total_high_risk_flags": sum(r.high_risk_count for r in all_results),
                "total_low_risk_flags": sum(r.low_risk_count for r in all_results),
                "average_processing_time": sum(r.total_processing_time for r in all_results) / len(all_results)
            },
            "individual_results": [result.dict() for result in all_results]
        }
        
        with open(consolidated_file, 'w') as f:
            json.dump(consolidated_data, f, indent=2)
        
        print(f"\n📊 Consolidated results saved to: {consolidated_file}")

if __name__ == "__main__":
    main()
