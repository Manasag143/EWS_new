import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
 
def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class AzureOpenAILLM:
    """Azure OpenAI GPT-4.1-mini LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1-mini"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            
            if hasattr(response, 'usage'):
                logger.info(f"Tokens used - Prompt: {response.usage.prompt_tokens}, "
                          f"Completion: {response.usage.completion_tokens}, "
                          f"Total: {response.usage.total_tokens}")
            
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        start_time = time.time()
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            logger.info(f"PDF text extraction took {time.time() - start_time:.2f} seconds")
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""
Extract the company name, quarter, and financial year from this text from an earnings call transcript.

Text: {first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)  
3. Financial Year (FY23/FY24/FY25 etc.)

Format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

Response:"""
       
        response = llm._call(prompt, max_tokens=200)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def extract_unique_flags_with_strict_deduplication(response_text: str, llm: AzureOpenAILLM) -> List[str]:
    """Enhanced extraction with STRICT deduplication to prevent duplicates"""
    
    prompt = f"""
You are an expert financial analyst tasked with extracting TRULY UNIQUE red flags with ZERO duplicates.

RED FLAGS ANALYSIS TO PROCESS:
{response_text}

STRICT DEDUPLICATION RULES:
1. If multiple flags refer to the SAME underlying financial issue, merge them into ONE comprehensive flag
2. Remove any flag that is a subset or variation of a broader flag
3. Combine similar concepts: "Debt increased" + "Higher debt levels" + "Rising debt burden" → "Debt levels increased significantly"
4. Remove generic flags that don't add specific value
5. Maximum 10-12 truly unique, specific flags only
6. Each flag must represent a DISTINCT financial concern
7. Prioritize flags with specific numbers/percentages over generic statements

EXAMPLES OF WHAT TO MERGE:
- "Revenue declined" + "Sales decreased" + "Top line fell" → "Revenue/sales declined"
- "Margin pressure" + "Profitability issues" + "Reduced margins" → "Margin compression and profitability pressure"
- "Cash flow problems" + "Liquidity concerns" + "Working capital issues" → "Cash flow and liquidity challenges"

OUTPUT REQUIREMENTS:
- Return ONLY a clean Python list format
- Each flag should be a concise, specific statement (5-15 words)
- No duplicates, no similar flags, no overlapping concerns
- Focus on the most critical and distinct issues only

Format: ["unique flag 1", "unique flag 2", "unique flag 3", ...]

EXTRACT UNIQUE FLAGS:
"""
    
    try:
        response = llm._call(prompt, max_tokens=800, temperature=0.0)  # Temperature 0 for consistency
        
        # Try to parse as Python list
        try:
            unique_flags = ast.literal_eval(response.strip())
            if isinstance(unique_flags, list) and len(unique_flags) <= 15:  # Reasonable limit
                return [flag.strip() for flag in unique_flags if flag.strip()]
        except:
            pass
        
        # Fallback parsing if ast.literal_eval fails
        lines = response.strip().split('\n')
        unique_flags = []
        
        for line in lines:
            line = line.strip()
            # Look for quoted strings
            if (line.startswith('"') and line.endswith('"')) or (line.startswith("'") and line.endswith("'")):
                flag = line[1:-1].strip()
                if flag and len(flag) > 5:  # Avoid very short, meaningless flags
                    unique_flags.append(flag)
            # Look for list items
            elif line.startswith('- ') or line.startswith('* '):
                flag = line[2:].strip()
                if flag and len(flag) > 5:
                    unique_flags.append(flag)
        
        # Remove duplicates and limit count
        seen = set()
        final_flags = []
        for flag in unique_flags:
            flag_lower = flag.lower()
            if flag_lower not in seen and len(final_flags) < 12:
                seen.add(flag_lower)
                final_flags.append(flag)
        
        return final_flags if final_flags else ["No specific red flags identified"]
        
    except Exception as e:
        logger.error(f"Error in strict deduplication: {e}")
        return ["Error in flag extraction"]

def classify_flag_against_criteria_strict(flag: str, criteria_definitions: Dict[str, str], 
                                 previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, str]:
    """Strictly classify a single flag against 15 criteria - only High if exact match found"""
    
    # Define strict keyword patterns for each criteria
    criteria_keywords = {
        "debt_increase": ["debt increase", "debt increased", "debt rising", "debt growth", "higher debt", "debt went up", "debt levels", "borrowing increase"],
        "provisioning": ["provision", "write-off", "write off", "writeoff", "bad debt", "impairment", "credit loss"],
        "asset_decline": ["asset decline", "asset fall", "asset decrease", "asset value down", "asset reduction", "asset impairment"],
        "receivable_days": ["receivable days", "collection period", "DSO", "days sales outstanding", "collection time"],
        "payable_days": ["payable days", "payment period", "DPO", "days payable outstanding", "payment delay"],
        "debt_ebitda": ["debt to ebitda", "debt/ebitda", "debt ebitda ratio", "leverage ratio", "debt multiple"],
        "revenue_decline": ["revenue decline", "revenue fall", "revenue decrease", "sales decline", "top line decline", "income reduction"],
        "onetime_expenses": ["one-time", "onetime", "exceptional", "extraordinary", "non-recurring", "special charges"],
        "margin_decline": ["margin decline", "margin fall", "margin pressure", "margin compression", "profitability decline", "margin squeeze"],
        "cash_balance": ["cash decline", "cash decrease", "cash balance fall", "liquidity issue", "cash shortage", "cash position"],
        "short_term_debt": ["short-term debt", "current liabilities", "working capital", "short term borrowing", "immediate obligations"],
        "management_issues": ["management change", "leadership change", "CEO", "CFO", "resignation", "departure", "management turnover"],
        "regulatory_compliance": ["regulatory", "compliance", "regulation", "regulator", "legal", "penalty", "violation", "sanctions"],
        "market_competition": ["competition", "competitive", "market share", "competitor", "market pressure", "competitive pressure"],
        "operational_disruptions": ["operational", "supply chain", "production", "manufacturing", "disruption", "operational issues"]
    }
    
    criteria_list = "\n".join([f"{i+1}. {name}: {desc}" for i, (name, desc) in enumerate(criteria_definitions.items())])
    
    prompt = f"""
You are a STRICT financial risk classifier. Follow these EXACT rules with NO exceptions:

RED FLAG TO CLASSIFY: "{flag}"

CRITERIA DEFINITIONS:
{criteria_list}

PREVIOUS YEAR DATA FOR THRESHOLD CHECKING:
{previous_year_data}

STRICT CLASSIFICATION ALGORITHM:
Step 1: Check for EXACT KEYWORD MATCH
- Search for specific keywords in the red flag text
- Keywords for each criteria:
  * debt_increase: "debt increase", "debt increased", "higher debt", "debt rising"
  * provisioning: "provision", "write-off", "impairment", "bad debt"
  * revenue_decline: "revenue decline", "sales decline", "top line decline"
  * margin_decline: "margin decline", "margin pressure", "profitability decline"
  * management_issues: "management change", "CEO", "CFO", "resignation"
  * regulatory_compliance: "regulatory", "compliance", "penalty", "violation"
  * operational_disruptions: "operational", "supply chain", "production disruption"
  * (and so on for all 15 criteria)

Step 2: If NO keyword match found → AUTOMATICALLY classify as "Low"

Step 3: If keyword match found → Check threshold criteria against previous year data

Step 4: Classify as "High" ONLY if BOTH conditions met:
   a) Exact keyword match exists
   b) Threshold criteria is satisfied

DEFAULT RULE: When in doubt → classify as "Low"

OUTPUT FORMAT (follow exactly):
Matched_Criteria: [exact criteria name if keyword found, otherwise "None"]
Risk_Level: [High only if both keyword AND threshold met, otherwise Low]
Reasoning: [Explain keyword search result and threshold check]
"""
    
    response = llm._call(prompt, max_tokens=300, temperature=0.0)
    
    # Initialize with safe defaults
    result = {
        'matched_criteria': 'None',
        'risk_level': 'Low',
        'reasoning': 'No exact keyword match found for any criteria'
    }
    
    # Parse response
    lines = response.strip().split('\n')
    for line in lines:
        if line.startswith('Matched_Criteria:'):
            matched = line.split(':', 1)[1].strip()
            result['matched_criteria'] = matched if matched not in ["None", ""] else 'None'
        elif line.startswith('Risk_Level:'):
            risk_level = line.split(':', 1)[1].strip()
            # Double validation: force Low if no criteria matched
            if result['matched_criteria'] == 'None':
                result['risk_level'] = 'Low'
            else:
                result['risk_level'] = risk_level if risk_level in ['High', 'Low'] else 'Low'
        elif line.startswith('Reasoning:'):
            result['reasoning'] = line.split(':', 1)[1].strip()
    
    # Final safety check
    if result['matched_criteria'] == 'None':
        result['risk_level'] = 'Low'
        result['reasoning'] = 'No criteria keyword match - defaulted to Low risk'
    
    return result

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(high_risk_flags: List[str], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    concise_summaries = []
    
    for flag in high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context[:3000]}  # Limit context to avoid token limits

HIGH RISK FLAG: "{flag}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Do NOT exceed 2 lines under any circumstances

OUTPUT FORMAT: [1-2 line summary only, no other text]

Summary:"""
        
        try:
            response = llm._call(prompt, max_tokens=100, temperature=0.1)
            
            # Clean and validate response
            summary_lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                # Take first 2 lines only
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                # Fallback
                concise_summary = f"{flag}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
                
            concise_summaries.append(concise_summary)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag}': {e}")
            concise_summaries.append(f"{flag}. Review required based on analysis.")
    
    return concise_summaries

def create_word_document(pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                        high_risk_flags: List[str], summary_by_categories: Dict[str, List[str]], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> str:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        high_count = risk_counts.get('High', 0)
        low_count = risk_counts.get('Low', 0)
        total_count = high_count + low_count
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(high_count)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(low_count)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(total_count)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            print("Generating concise summaries for high risk flags...")
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            for summary in concise_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(str(summary))
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.get('High', 0)}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.get('Low', 0)}")
            doc.add_paragraph(f"Total Flags: {risk_counts.get('Total', 0)}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            print(f"Error creating fallback document: {e2}")
            return None

def process_pdf_enhanced_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Process PDF through enhanced 5-iteration pipeline with strict deduplication and classification
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM and load PDF
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load first query from CSV/Excel
        try:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            
            if len(queries_df) == 0 or "prompt" not in queries_df.columns:
                print("Warning: No prompts found in queries file. Using default query.")
                first_query = "Analyze this document for potential red flags."
            else:
                first_query = queries_df["prompt"].tolist()[0]
        except Exception as e:
            print(f"Error loading queries file: {e}. Using default query.")
            first_query = "Analyze this document for potential red flags."
        
        # ITERATION 1: Initial red flag identification
        print("Running 1st iteration - Initial Analysis...")
        sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.
 
COMPLETE DOCUMENT TO ANALYZE:
{context}
 
Your task is to analyze the ENTIRE document above and identify ALL potential red flags.
 
CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
"""
        
        first_prompt = f"{sys_prompt}\n\nQuestion: {first_query}\n\nAnswer:"
        first_response = llm._call(first_prompt, max_tokens=4000)
        
        # ITERATION 2: Deduplication 
        print("Running 2nd iteration - Deduplication...")
        second_prompt = "Remove the duplicates from the above context. Also if the Original Quote and Keyword identifies is same remove them. Do not lose data if duplicates are not found."
        
        second_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {first_response}
 
Based on the above analysis and the original context, please answer: {second_prompt}
 
Answer:"""
        
        second_response = llm._call(second_full_prompt, max_tokens=4000)
        
        # ITERATION 3: Categorization
        print("Running 3rd iteration - Categorization...")
        third_prompt = """You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.
 
- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction.
- Regulatory Issues: Concerns related Compliance with laws, regulations.
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape.
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.
 
While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.
 
**Output Format**:
### Balance Sheet Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### P&L (Income Statement) Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
        
        third_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {second_response}
 
Based on the above analysis and the original context, please answer: {third_prompt}
 
Answer:"""
        
        third_response = llm._call(third_full_prompt, max_tokens=4000)
        
        # ITERATION 4: Summary generation
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:
 
1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
3. **Focus on factual content**: Base the summary solely on the information associated with each red flag, without introducing external opinions or assumptions
4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
8. **Summary should be factual**: Avoid any subjective interpretations or opinions
9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points
 
Format the output exactly like this example:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]
 
Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
        
        fourth_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {third_response}
 
Based on the above analysis and the original context, please answer: {fourth_prompt}
 
Answer:"""
        
        fourth_response = llm._call(fourth_full_prompt, max_tokens=4000)
        
        # ITERATION 5: Extract unique flags with ENHANCED DEDUPLICATION and classify
        print("Running 5th iteration - Enhanced Unique Flags Classification...")
        
        # Step 1: Extract unique flags with STRICT deduplication
        print("  Extracting unique flags with strict deduplication...")
        try:
            unique_flags = extract_unique_flags_with_strict_deduplication(second_response, llm)
            print(f"  Extracted {len(unique_flags)} unique flags:")
            for i, flag in enumerate(unique_flags, 1):
                print(f"    {i}. {flag}")
        except Exception as e:
            print(f"  Error extracting flags: {e}")
            unique_flags = ["Error in flag extraction"]
        
        # Define 15 criteria definitions
        criteria_definitions = {
            "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
            "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
            "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
            "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
            "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
        }
        
        # Step 2: Classify each unique flag with STRICT criteria matching
        print("  Classifying flags against 15 criteria with strict matching...")
        classification_results = []
        high_risk_flags = []
        low_risk_flags = []
        
        if len(unique_flags) > 0 and unique_flags[0] != "Error in flag extraction":
            for i, flag in enumerate(unique_flags, 1):
                print(f"    Classifying flag {i}: {flag[:50]}...")
                
                try:
                    classification = classify_flag_against_criteria_strict(
                        flag=flag,
                        criteria_definitions=criteria_definitions,
                        previous_year_data=previous_year_data,
                        llm=llm
                    )
                    
                    classification_results.append({
                        'flag': flag,
                        'matched_criteria': classification['matched_criteria'],
                        'risk_level': classification['risk_level'],
                        'reasoning': classification['reasoning']
                    })
                    
                    # Only add to high risk if explicitly classified as High AND criteria matched
                    if (classification['risk_level'].lower() == 'high' and 
                        classification['matched_criteria'] != 'None'):
                        high_risk_flags.append(flag)
                        print(f"      → HIGH RISK (Criteria: {classification['matched_criteria']})")
                    else:
                        low_risk_flags.append(flag)
                        print(f"      → Low Risk (Reason: {classification['reasoning'][:50]}...)")
                        
                except Exception as e:
                    print(f"    Error classifying flag {i}: {e}")
                    # Always default to low risk if classification fails
                    classification_results.append({
                        'flag': flag,
                        'matched_criteria': 'None',
                        'risk_level': 'Low',
                        'reasoning': f'Classification failed: {str(e)}'
                    })
                    low_risk_flags.append(flag)
                    
                time.sleep(0.3)  # Small delay to avoid rate limits
        else:
            print("  No valid unique flags found to classify")
        
        # Calculate final counts
        risk_counts = {
            'High': len(high_risk_flags),
            'Low': len(low_risk_flags),
            'Total': len(unique_flags) if unique_flags and unique_flags[0] != "Error in flag extraction" else 0
        }
        
        print(f"\n=== FINAL CLASSIFICATION RESULTS ===")
        print(f"High Risk Flags: {risk_counts['High']}")
        print(f"Low Risk Flags: {risk_counts['Low']}")
        print(f"Total Flags: {risk_counts['Total']}")
        
        if high_risk_flags:
            print(f"\nHigh Risk Flags List:")
            for i, flag in enumerate(high_risk_flags, 1):
                print(f"  {i}. {flag}")
        
        # Extract company info and create Word document
        print("\nCreating Word document with enhanced analysis...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
           
            # Create Word document with strict high risk summaries
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                high_risk_flags=high_risk_flags,
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            print(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Save all results to CSV files
        print("Saving results to CSV files...")
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Save pipeline results
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Deduplication", 
                "Categorization",
                "Summary Generation",
                "Enhanced Unique Flags Classification"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Enhanced Classification: {risk_counts['High']} High Risk, {risk_counts['Low']} Low Risk flags from {risk_counts['Total']} total unique flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_enhanced_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        print(f"Pipeline results saved: {results_file}")
        
        # Save detailed classification results
        if len(classification_results) > 0:
            classification_df = pd.DataFrame(classification_results)
            classification_file = os.path.join(output_folder, f"{pdf_name}_enhanced_flag_classification.csv")
            classification_df.to_csv(classification_file, index=False)
            print(f"Classification results saved: {classification_file}")
        else:
            print("No classification results to save")

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        return results_summary
       
    except Exception as e:
        print(f"Error processing {pdf_name}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """Main function to process all PDFs in the specified folder"""
    
    # Configuration
    pdf_folder_path = r"ola_pdf" 
    queries_csv_path = r"EWS_prompts_v2.xlsx"
    output_folder = r"ola_results_enhanced"
    
    # Azure OpenAI Configuration
    api_key = "8496bd1da41e498c"
    azure_endpoint = "https://crisil-pp-gp"  # Complete your endpoint
    api_version = "2025-01-01-preview"
    deployment_name = "gpt-4.1-mini"
  
    # Previous year data for threshold comparisons
    previous_year_data = """
Previous reported Debt 5,684 Cr Mar-24
Current quarter ebidta (525)Cr Mar-25
Previous reported asset value 7,735Cr Mar-24
Previous reported receivable days 12 days Mar-24
Previous reported payable days 112 days Mar-24
Previous reported revenue 1,045 Cr Dec-24
Previous reported profitability (460) Cr Dec-24
Previous reported operating margin -44.00% Dec-24
Previous reported cash balance 1,663 Cr Mar-24
Previous reported current liabilities 1,071 Cr Mar-24
"""
 
    print("=== ENHANCED PDF PROCESSING PIPELINE ===")
    print(f"PDF Folder: {pdf_folder_path}")
    print(f"Output Folder: {output_folder}")
    print(f"Queries File: {queries_csv_path}")
    print("="*50)
    
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    successful_processed = 0
    failed_processed = 0
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        result = process_pdf_enhanced_pipeline(
            pdf_path=pdf_file,
            queries_csv_path=queries_csv_path,
            previous_year_data=previous_year_data,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
            successful_processed += 1
        else:
            print(f"❌ Failed to process {pdf_file}")
            failed_processed += 1
        
        print("-" * 60)
    
    # Final summary
    print(f"\n{'='*60}")
    print(f"PROCESSING SUMMARY")
    print(f"{'='*60}")
    print(f"Total PDFs: {len(pdf_files)}")
    print(f"Successfully Processed: {successful_processed}")
    print(f"Failed: {failed_processed}")
    print(f"Output Directory: {output_folder}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
