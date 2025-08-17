import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)
 
def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class AzureOpenAILLM:
    """Azure OpenAI gpt-4.1 LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI gpt-4.1"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""<role>
You are an expert document analyst specializing in extracting key corporate information from financial documents and earnings transcripts.
</role>

<system_prompt>
You excel at quickly identifying and extracting specific corporate identifiers from financial documents with high accuracy and consistency.
</system_prompt>

<instruction>
Extract the company name, quarter, and financial year from the provided text.

EXTRACTION REQUIREMENTS:
1. Company Name: Full legal company name including suffixes (Ltd/Limited/Inc/Corp etc.)
2. Quarter: Identify the quarter (Q1/Q2/Q3/Q4)  
3. Financial Year: Extract financial year (FY23/FY24/FY25 etc.)

OUTPUT FORMAT:
Provide ONLY the result in this exact format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If any component cannot be clearly identified, use reasonable defaults based on context.
</instruction>

<context>
DOCUMENT TEXT TO ANALYZE:
{first_page_text}
</context>

Extract company information:"""
       
        response = llm._call(prompt, max_tokens=200)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def enhanced_deduplication_iteration(first_response: str, context: str, llm: AzureOpenAILLM) -> str:
    """
    Enhanced 2nd iteration that does intelligent deduplication while preserving 
    all original context, quotes, and supporting text
    """
    
    second_prompt = f"""<role>
You are an expert financial analyst specializing in document analysis and content organization with 15+ years of experience in identifying and consolidating duplicate financial concerns.
</role>

<system_prompt>
You excel at identifying duplicate red flags while preserving all original supporting evidence, quotes, and contextual information that led to each unique concern identification.
</system_prompt>

<instruction>
Analyze the red flag analysis below and remove ONLY duplicate red flags while preserving ALL original context and supporting evidence.

DEDUPLICATION RULES:
1. **Identify Duplicate Flags**: Find red flags that represent the SAME underlying financial concern
2. **Merge Duplicates**: When duplicates are found, combine them into ONE comprehensive red flag entry
3. **Preserve ALL Evidence**: Keep ALL original quotes, speaker attributions, page references, and supporting text
4. **Maintain Structure**: Keep the numbered format and detailed explanations
5. **No Data Loss**: Do NOT remove any original quotes, context, or supporting evidence
6. **Consolidate Quotes**: If multiple quotes support the same concern, include ALL of them in the merged entry

MERGING APPROACH:
- Combine similar concerns into one comprehensive red flag
- Include ALL original quotes from duplicate entries
- Preserve ALL speaker attributions and page references  
- Maintain ALL supporting context and explanations
- Create comprehensive descriptions that capture the full scope

EXAMPLES OF WHAT TO MERGE:
- "Revenue declined 20%" + "Sales performance weak" + "Top line pressure" → Merge into comprehensive revenue concern with ALL quotes
- "Debt increased significantly" + "Higher borrowing levels" + "Leverage concerns" → Merge into comprehensive debt concern with ALL quotes
- "Cash flow issues" + "Liquidity problems" + "Working capital constraints" → Merge into comprehensive liquidity concern with ALL quotes

OUTPUT FORMAT:
Maintain the same detailed format as the input:
1. [Comprehensive merged red flag description] - [detailed explanation]
Original Quotes: 
- "[exact quote 1 with speaker name]" (Page X)
- "[exact quote 2 with speaker name]" (Page Y)  
- "[exact quote 3 with speaker name]" (Page Z)
[Any additional supporting context and analysis]

CRITICAL REQUIREMENTS:
- Do NOT lose any original quotes or supporting text
- Do NOT change the meaning or context of any red flag
- Do NOT remove red flags that are genuinely different concerns
- Only merge flags that represent the SAME underlying issue
- Preserve exact quotes with proper attribution
- Maintain comprehensive explanations for each unique concern
</instruction>

<context>
ORIGINAL DOCUMENT CONTEXT:
{context}

RED FLAG ANALYSIS TO DEDUPLICATE:
{first_response}
</context>

Provide deduplicated analysis with preserved context:"""
    
    try:
        enhanced_second_response = llm._call(second_prompt, max_tokens=4000, temperature=0.1)
        return enhanced_second_response
    except Exception as e:
        logger.error(f"Error in enhanced deduplication: {e}")
        return first_response  # Return original if deduplication fails

def classify_bucket_risks(deduplicated_analysis: str, bucket_name: str, bucket_config: Dict, 
                         previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, Any]:
    """
    Classify red flags for a specific risk bucket using 6-bucket system
    """
    
    risk_indicators = bucket_config["risk_indicators"]
    related_data = bucket_config["related_data"]
    
    # Create risk indicators text
    indicators_text = "\n".join([f"* `{key}`: {value}" for key, value in risk_indicators.items()])
    
    # Create related data text
    data_text = "\n".join([f"* {item}" for item in related_data])
    
    classification_prompt = f"""<role>
You are a senior financial risk analyst with 15+ years of experience specializing in {bucket_name.lower()} assessment and quantitative risk evaluation.
</role>

<system_prompt>
You excel at identifying and classifying financial risks within specific risk categories, applying quantitative thresholds, and providing evidence-based risk assessments with precise reasoning.
</system_prompt>

<instruction>
Analyze the deduplicated red flag analysis and identify ALL red flags that belong to **{bucket_name}** category.
Classify each identified flag as High Risk or Low Risk based on the specific criteria provided.

**{bucket_name}**
{indicators_text}

**Related Previous Data:**
{data_text}

CLASSIFICATION PROCESS:
1. **Identify Relevant Flags**: Find all red flags from the analysis that relate to {bucket_name.lower()}
2. **Apply Quantitative Criteria**: Use the specific thresholds provided for each risk indicator
3. **Evidence-Based Assessment**: Base classification on actual data and quotes from the document
4. **Risk Level Assignment**: Classify as High or Low based on meeting the specified thresholds
5. **Detailed Reasoning**: Provide specific reasoning with supporting quotes and calculations

CLASSIFICATION RULES:
- Only classify flags that clearly belong to this risk category
- Use ONLY the risk indicators and thresholds provided above
- Include exact quotes and calculations that support the classification
- If data is insufficient for quantitative assessment, classify as Low Risk
- Provide specific reasoning for each classification decision

OUTPUT FORMAT:
**BUCKET: {bucket_name}**

**Identified Flags in this Category:**
1. [Flag description with original quotes]
   - **Risk Indicator:** [Which specific indicator from the list above]
   - **Risk Level:** [High/Low]
   - **Supporting Data:** [Specific numbers/percentages from document]
   - **Calculation/Reasoning:** [Show calculation or reasoning for classification]
   - **Original Quotes:** [Exact quotes that support this classification]

2. [Next flag if any...]

**Summary:**
- Total Flags in {bucket_name}: [number]
- High Risk: [number]
- Low Risk: [number]

**IMPORTANT:** If NO flags are found for this category, respond with: "No flags identified for {bucket_name} category."
</instruction>

<context>
PREVIOUS YEAR DATA:
{previous_year_data}

DEDUPLICATED RED FLAG ANALYSIS TO CLASSIFY:
{deduplicated_analysis}
</context>

Classify {bucket_name} risks:"""
    
    try:
        classification_response = llm._call(classification_prompt, max_tokens=1500, temperature=0.0)
        
        # Parse the response to extract structured data
        result = {
            "bucket_name": bucket_name,
            "classification_response": classification_response,
            "flags_identified": [],
            "high_risk_count": 0,
            "low_risk_count": 0,
            "total_count": 0
        }
        
        # Simple parsing to extract counts (you can make this more sophisticated)
        lines = classification_response.lower().split('\n')
        for line in lines:
            if 'high risk:' in line:
                try:
                    result["high_risk_count"] = int(line.split(':')[1].strip())
                except:
                    pass
            elif 'low risk:' in line:
                try:
                    result["low_risk_count"] = int(line.split(':')[1].strip())
                except:
                    pass
            elif 'total flags' in line and bucket_name.lower() in line:
                try:
                    result["total_count"] = int(line.split(':')[1].strip())
                except:
                    pass
        
        return result
        
    except Exception as e:
        logger.error(f"Error classifying {bucket_name}: {e}")
        return {
            "bucket_name": bucket_name,
            "classification_response": f"Error in classification: {str(e)}",
            "flags_identified": [],
            "high_risk_count": 0,
            "low_risk_count": 0,
            "total_count": 0
        }

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(high_risk_flags: List[str], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    concise_summaries = []
    
    for flag in high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context}

HIGH RISK FLAG: "{flag}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Ensure subsequent statements are cautious and do not downplay the risk. Avoid neutral/positive statements that contradict the warning.
6. Do NOT start with "Summary:" or any prefix
7. Provide ONLY the factual summary content
8. Make it UNIQUE - avoid repeating information from other summaries
9. If applicable Specify whether the flag is for : A specific business unit/division, Consolidated financials, Standalone financials, geographical region.

OUTPUT FORMAT: [Direct factual summary only, no labels or prefixes]
"""
        
        try:
            response = llm._call(prompt, max_tokens=100, temperature=0.1)
            
            # Clean response - remove any prefixes or labels
            clean_response = response.strip()
            
            # Remove common prefixes that might appear
            prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:", "Flag summary:", "The flag:", "This flag:"]
            for prefix in prefixes_to_remove:
                if clean_response.startswith(prefix):
                    clean_response = clean_response[len(prefix):].strip()
            
            # Split into lines and take first 2
            summary_lines = [line.strip() for line in clean_response.split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                concise_summary = f"{flag}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
            
            concise_summaries.append(concise_summary)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag}': {e}")
            concise_summaries.append(f"{flag}. Review required based on analysis.")
    
    return concise_summaries

def create_word_document(pdf_name: str, company_info: str, bucket_results: List[Dict],
                        summary_by_categories: Dict[str, List[str]], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> str:
    """Create a formatted Word document with 6-bucket classification results"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Calculate total risk counts from all buckets
        total_high_risk = sum(bucket["high_risk_count"] for bucket in bucket_results)
        total_low_risk = sum(bucket["low_risk_count"] for bucket in bucket_results)
        total_flags = sum(bucket["total_count"] for bucket in bucket_results)
        
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(total_high_risk)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(total_low_risk)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(total_flags)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        high_risk_flags = []
        for bucket in bucket_results:
            if bucket["high_risk_count"] > 0:
                # Extract high risk flags from bucket response (simplified - you may want to improve this parsing)
                high_risk_flags.append(f"High risk identified in {bucket['bucket_name']}")
        
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            for summary in concise_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {total_high_risk}")
            doc.add_paragraph(f"Low Risk Flags: {total_low_risk}")
            doc.add_paragraph(f"Total Flags: {total_flags}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

def process_pdf_6bucket_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1"):
    """
    Process PDF through 4-iteration pipeline + 6-bucket classification system
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM and load PDF
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load first query from CSV/Excel
        try:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            
            if len(queries_df) == 0 or "prompt" not in queries_df.columns:
                first_query = "Analyze this document for potential red flags."
            else:
                first_query = queries_df["prompt"].tolist()[0]
        except Exception as e:
            logger.warning(f"Error loading queries file: {e}. Using default query.")
            first_query = "Analyze this document for potential red flags."
        
        # ITERATION 1: Initial red flag identification
        print("Running 1st iteration - Initial Analysis...")
        first_prompt = f"""<role>
You are an expert financial analyst with 15+ years of experience specializing in identifying red flags from earnings call transcripts and financial documents.
</role>

<system_prompt>
You excel at comprehensive document analysis, identifying subtle financial risks, and providing detailed evidence-based assessments with precise documentation.
</system_prompt>

<instruction>
Analyze the ENTIRE document and identify ALL potential red flags comprehensively.

ANALYSIS REQUIREMENTS:
- Review every section of the document thoroughly
- Identify financial, operational, strategic, and management risks
- Focus on quantitative concerns with specific data points
- Document exact quotes with speaker attribution
- Number each red flag sequentially (1, 2, 3, etc.)
- Include page references where available

OUTPUT FORMAT:
For each red flag:
1. The potential red flag you observed - [brief description]
Original Quote: "[exact quote with speaker name]" (Page X)

CRITICAL: Ensure comprehensive analysis of the entire document.
</instruction>

<context>
COMPLETE DOCUMENT TO ANALYZE:
{context}

SPECIFIC QUESTION: {first_query}
</context>

Provide comprehensive red flag analysis:"""
        
        first_response = llm._call(first_prompt, max_tokens=4000)
        
        # ITERATION 2: Enhanced Deduplication
        print("Running 2nd iteration - Enhanced Deduplication...")
        second_response = enhanced_deduplication_iteration(first_response, context, llm)
        
        # ITERATION 3: Categorization
        print("Running 3rd iteration - Categorization...")
        third_prompt = f"""<role>
You are a senior financial analyst expert in financial risk categorization with deep knowledge of balance sheet analysis, P&L assessment, and corporate risk frameworks.
</role>

<system_prompt>
You excel at organizing financial risks into standardized categories, ensuring comprehensive coverage of all financial risk areas, and maintaining accuracy in risk classification.
</system_prompt>

<instruction>
Categorize the identified red flags into the following 7 standardized categories based on their financial nature and business impact.

MANDATORY CATEGORIES:
1. Balance Sheet Issues: Assets, liabilities, equity, debt, and overall financial position concerns
2. P&L (Income Statement) Issues: Revenue, expenses, profits, and financial performance concerns  
3. Liquidity Issues: Short-term obligations, cash flow, debt repayment, working capital concerns
4. Management and Strategy Issues: Leadership, governance, decision-making, strategy, and vision concerns
5. Regulatory Issues: Compliance, laws, regulations, and regulatory body concerns
6. Industry and Market Issues: Market position, industry trends, competitive landscape concerns
7. Operational Issues: Internal processes, systems, infrastructure, and operational efficiency concerns

CATEGORIZATION RULES:
- Assign each red flag to the MOST relevant category only
- Do not create new categories - use only the 7 listed above
- Preserve all Original Quotes exactly as provided
- If a red flag could fit multiple categories, choose the primary/most relevant one
- Do not leave any red flag unclassified
- Do not repeat categories in the output
- Maintain sequential organization within each category

OUTPUT FORMAT:
### Balance Sheet Issues
- [Red flag 1 with original quotes and page references]
- [Red flag 2 with original quotes and page references]

### P&L (Income Statement) Issues
- [Red flag 1 with original quotes and page references]

Continue this format for all applicable categories.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}
</context>

Provide categorized analysis:"""
        
        third_response = llm._call(third_prompt, max_tokens=4000)
        
        # ITERATION 4: Summary generation
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = f"""<role>
You are an expert financial summarization specialist with expertise in creating concise, factual, and comprehensive summaries that preserve critical quantitative data and key insights.
</role>

<system_prompt>
You excel at distilling complex financial analysis into clear, actionable summaries while maintaining objectivity, preserving all quantitative details, and ensuring no critical information is lost.
</system_prompt>

<instruction>
Create a comprehensive summary of each category of red flags in bullet point format following these strict guidelines.

SUMMARY REQUIREMENTS:
1. Retain ALL quantitative information (numbers, percentages, ratios, dates)
2. Maintain completely neutral, factual tone - no opinions or interpretations
3. Include every red flag from each category - no omissions
4. Base content solely on the provided categorized analysis
5. Preserve specific data points and statistics wherever mentioned
6. Each bullet point should capture key details of individual red flags
7. Balance thoroughness with conciseness
8. Use professional financial terminology
9. Ensure category-specific content alignment

OUTPUT FORMAT:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories that contain red flags.

CRITICAL: Each bullet point represents a concise summary of individual red flags with preserved quantitative details.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        fourth_response = llm._call(fourth_prompt, max_tokens=4000)
        
        # 6-BUCKET CLASSIFICATION SYSTEM (6 separate LLM calls)
        print("Running 6-Bucket Classification System...")
        
        # Define the 6 buckets with their configurations
        bucket_configs = {
            "Debt and Leverage Risk": {
                "risk_indicators": {
                    "debt_increase": "High: Debt increased >30% vs previous balance sheet; Low: <30%",
                    "debt_ebitda": "High: Debt/EBITDA >3x; Low: Debt/EBITDA <3x",
                    "short_term_borrowings": "High: Short-term borrowings increased >30% vs previous balance sheet; Low: <30%"
                },
                "related_data": [
                    "Debt as per Previous reported balance sheet number (Mar-23): 80,329 Cr",
                    "Current quarter EBITDA (March-24): 11,511 Cr",
                    "Short term borrowings as per previous balance sheet (Mar-23): 36,407 Cr"
                ]
            },
            "Asset Quality and Impairment Risk": {
                "risk_indicators": {
                    "asset_decline": "High: Asset value falls >30% vs previous balance sheet; Low: <30%",
                    "impairment": "High: Impairment >25% of previous net worth; Low: <25%",
                    "provisioning": "High: Provisioning/write-offs >25% of current quarter EBITDA; Low: <25%"
                },
                "related_data": [
                    "Asset value as per previous balance sheet (Mar-23): 189,455 Cr",
                    "Previous reported net worth from balance sheet (Mar-23): 47,896 Cr",
                    "Current quarter EBITDA (March-24): 11,511 Cr"
                ]
            },
            "Working Capital and Liquidity Risk": {
                "risk_indicators": {
                    "receivable_days": "High: Receivable days increased >30% vs previous balance sheet; Low: <30%",
                    "payable_days": "High: Payable days increased >30% vs previous balance sheet; Low: <30%",
                    "receivables": "High: Receivables increased >30% vs previous balance sheet; Low: <30%",
                    "payables": "High: Payables increased >30% vs previous balance sheet; Low: <30%",
                    "cash_balance": "High: Cash balance falling >25% vs previous balance sheet; Low: <25%"
                },
                "related_data": [
                    "Receivable days as per previous balance sheet (Mar-23): 10 days",
                    "Payable days as per previous balance sheet (Mar-23): 91 days",
                    "Receivables as per previous balance sheet (Mar-23): 6,414 Cr",
                    "Payables as per previous balance sheet (Mar-23): 11,043 Cr",
                    "Cash balance as per previous balance sheet (Mar-23): 9,254 Cr"
                ]
            },
            "Revenue and Profitability Risk": {
                "risk_indicators": {
                    "revenue_decline": "High: Revenue falls >25% vs previous quarter; Low: <25%",
                    "profit_before_tax_decline": "High: PBT falls >25% vs previous quarter; Low: <25%",
                    "profit_after_tax_decline": "High: PAT falls >25% vs previous quarter; Low: <25%",
                    "EBIDTA_decline": "High: EBITDA falls >25% vs previous quarter; Low: <25%"
                },
                "related_data": [
                    "Revenue as per previous quarter (Dec-23): 35,541 Cr",
                    "Profit before tax as per previous quarter (Dec-23): 4,105 Cr",
                    "Profit after tax as per previous quarter (Dec-23): 2,868 Cr",
                    "EBITDA as per previous quarter (Dec-23): 8,531 Cr"
                ]
            },
            "Margin and Operational Efficiency Risk": {
                "risk_indicators": {
                    "margin_decline": "High: Operating margin falling >25% vs previous quarter; Low: <25%",
                    "gross_margin": "High: Gross margin falling >100bps; Low: <100bps",
                    "one-time_expenses": "High: One-time expenses >25% of current quarter EBITDA; Low: <25%"
                },
                "related_data": [
                    "Operating margin as per previous quarter (Dec-23): 25%",
                    "Current quarter EBITDA (March-24): 11,511 Cr"
                ]
            },
            "Strategic and External Risk": {
                "risk_indicators": {
                    "management_issues": "High: Management/strategy concerns found; Low: No clear concerns",
                    "regulatory_compliance": "High: Regulatory issues/warnings found; Low: No clear concerns",
                    "market_competition": "High: Competitive intensity/market share decline; Low: Stable/increasing share",
                    "operational_disruptions": "High: Operational/supply chain concerns; Low: No clear concerns",
                    "others": "High: Other significant business impact issues; Low: No other concerns"
                },
                "related_data": [
                    "Current quarter EBITDA (March-24): 11,511 Cr (for significance comparison)",
                    "Previous reported net worth (Mar-23): 47,896 Cr (for context on business scale)"
                ]
            }
        }
        
        # Perform 6 separate LLM calls for each bucket
        bucket_results = []
        total_high_risk = 0
        total_low_risk = 0
        total_flags = 0
        
        for i, (bucket_name, bucket_config) in enumerate(bucket_configs.items(), 1):
            print(f"  Bucket {i}/6: {bucket_name}...")
            
            bucket_result = classify_bucket_risks(
                deduplicated_analysis=second_response,
                bucket_name=bucket_name,
                bucket_config=bucket_config,
                previous_year_data=previous_year_data,
                llm=llm
            )
            
            bucket_results.append(bucket_result)
            total_high_risk += bucket_result["high_risk_count"]
            total_low_risk += bucket_result["low_risk_count"]
            total_flags += bucket_result["total_count"]
            
            print(f"    {bucket_name}: {bucket_result['high_risk_count']} High, {bucket_result['low_risk_count']} Low, {bucket_result['total_count']} Total")
            
            time.sleep(0.5)  # Small delay between API calls
        
        print(f"\n=== 6-BUCKET CLASSIFICATION RESULTS ===")
        print(f"Total High Risk Flags: {total_high_risk}")
        print(f"Total Low Risk Flags: {total_low_risk}")
        print(f"Total Flags: {total_flags}")
        
        print(f"\n--- BUCKET BREAKDOWN ---")
        for bucket in bucket_results:
            print(f"  {bucket['bucket_name']}: {bucket['high_risk_count']} High, {bucket['low_risk_count']} Low")
        
        # Extract company info and create Word document
        print("\nCreating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
           
            # Create Word document with 6-bucket results
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                bucket_results=bucket_results,
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Save all results to CSV files
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Save pipeline results
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 4,
            "iteration": [1, 2, 3, 4],
            "stage": [
                "Initial Analysis",
                "Enhanced Deduplication", 
                "Categorization",
                "Summary Generation"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response
            ],
            "timestamp": [timestamp] * 4
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        # Save 6-bucket classification results
        bucket_classification_data = []
        for bucket in bucket_results:
            bucket_classification_data.append({
                'pdf_name': pdf_name,
                'bucket_name': bucket['bucket_name'],
                'high_risk_count': bucket['high_risk_count'],
                'low_risk_count': bucket['low_risk_count'],
                'total_count': bucket['total_count'],
                'classification_response': bucket['classification_response'],
                'timestamp': timestamp
            })
        
        bucket_df = pd.DataFrame(bucket_classification_data)
        bucket_file = os.path.join(output_folder, f"{pdf_name}_6bucket_classification.csv")
        bucket_df.to_csv(bucket_file, index=False)
        
        # Save summary statistics
        summary_stats = pd.DataFrame({
            'pdf_name': [pdf_name],
            'total_high_risk': [total_high_risk],
            'total_low_risk': [total_low_risk],
            'total_flags': [total_flags],
            'debt_leverage_high': [bucket_results[0]['high_risk_count']],
            'asset_quality_high': [bucket_results[1]['high_risk_count']],
            'working_capital_high': [bucket_results[2]['high_risk_count']],
            'revenue_profitability_high': [bucket_results[3]['high_risk_count']],
            'margin_efficiency_high': [bucket_results[4]['high_risk_count']],
            'strategic_external_high': [bucket_results[5]['high_risk_count']],
            'word_document_path': [word_doc_path],
            'timestamp': [timestamp]
        })
        
        summary_file = os.path.join(output_folder, f"{pdf_name}_summary_stats.csv")
        summary_stats.to_csv(summary_file, index=False)

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        print(f"Files created:")
        print(f"  - Pipeline results: {results_file}")
        print(f"  - 6-bucket classification: {bucket_file}")
        print(f"  - Summary statistics: {summary_file}")
        if word_doc_path:
            print(f"  - Word document: {word_doc_path}")
        
        return {
            'pdf_name': pdf_name,
            'pipeline_results': results_summary,
            'bucket_results': bucket_results,
            'summary_stats': summary_stats,
            'word_document_path': word_doc_path,
            'total_high_risk': total_high_risk,
            'total_low_risk': total_low_risk,
            'total_flags': total_flags
        }
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

def main():
    """Main function to process all PDFs in the specified folder with 6-bucket classification"""
    
    # Configuration
    pdf_folder_path = r"kalyan_pdf" 
    queries_csv_path = r"EWS_prompts_v2_2.xlsx"
    output_folder = r"kalyan_results_6bucket"

    api_key = "8496498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01-preview"
    deployment_name = "gpt-4.1"
  
    # Enhanced structured previous year data for 6-bucket system
    previous_year_data = """
Debt as per Previous reported balance sheet number (Mar-23): 80,329 Cr
Current quarter EBITDA (March-24): 11,511 Cr
Short term borrowings as per previous balance sheet (Mar-23): 36,407 Cr
Asset value as per previous balance sheet (Mar-23): 189,455 Cr
Previous reported net worth from balance sheet (Mar-23): 47,896 Cr
Receivable days as per previous balance sheet (Mar-23): 10 days
Payable days as per previous balance sheet (Mar-23): 91 days
Receivables as per previous balance sheet (Mar-23): 6,414 Cr
Payables as per previous balance sheet (Mar-23): 11,043 Cr
Cash balance as per previous balance sheet (Mar-23): 9,254 Cr
Revenue as per previous quarter (Dec-23): 35,541 Cr
Profit before tax as per previous quarter (Dec-23): 4,105 Cr
Profit after tax as per previous quarter (Dec-23): 2,868 Cr
EBITDA as per previous quarter (Dec-23): 8,531 Cr
Operating margin as per previous quarter (Dec-23): 25%
"""
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    all_results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        result = process_pdf_6bucket_pipeline(
            pdf_path=pdf_file,
            queries_csv_path=queries_csv_path,
            previous_year_data=previous_year_data,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            all_results.append(result)
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
        else:
            print(f"❌ Failed to process {pdf_file}")
    
    # Create consolidated summary report
    if all_results:
        print(f"\n{'='*60}")
        print("CREATING CONSOLIDATED SUMMARY REPORT")
        print(f"{'='*60}")
        
        consolidated_data = []
        for result in all_results:
            consolidated_data.append({
                'pdf_name': result['pdf_name'],
                'total_high_risk': result['total_high_risk'],
                'total_low_risk': result['total_low_risk'],
                'total_flags': result['total_flags'],
                'processing_status': 'Success',
                'word_document': result['word_document_path'] is not None
            })
        
        consolidated_df = pd.DataFrame(consolidated_data)
        consolidated_file = os.path.join(output_folder, "consolidated_summary.csv")
        consolidated_df.to_csv(consolidated_file, index=False)
        
        print(f"\n=== BATCH PROCESSING COMPLETE ===")
        print(f"Total PDFs processed: {len(all_results)}")
        print(f"Consolidated summary: {consolidated_file}")
        print(f"Average flags per document: {consolidated_df['total_flags'].mean():.1f}")
        print(f"Average high risk per document: {consolidated_df['total_high_risk'].mean():.1f}")
        
if __name__ == "__main__":
    main()
                    "
