import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

# ==============================================================================
# AZURE OPENAI LLM CLASS
# ==============================================================================

class AzureOpenAILLM:
    """Azure OpenAI gpt-4.1 LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = None, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI gpt-4.1"""
        try:
            kwargs = {
                "model": self.deployment_name,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "top_p": 0.95,
                "frequency_penalty": 0,
                "presence_penalty": 0
            }
            
            if max_tokens:
                kwargs["max_tokens"] = max_tokens
                
            response = self.client.chat.completions.create(**kwargs)
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

# ==============================================================================
# PDF PROCESSING FUNCTIONS
# ==============================================================================


keywords_part1 = """
Attrition: Refers to the increasing or high loss of employees, customers, or revenue due to various reasons such as resignation, retirement, or competition, which can negatively impact a company's financial performance. 
Adverse: Describes an unfavorable or negative situation, event, or trend, such as adverse market conditions or regulatory changes. 
Cautious outlook: Indicates a company's conservative or pessimistic view of its future financial performance, often due to uncertainty or potential risks. 
Challenging environment: Refers to a difficult or competitive market situation.
Competition intensifying: Describes an increase in competition in a market or industry, which can lead to decreased market share, revenue, or profitability for a company. 
Corporate governance: Refers to the system of rules, practices, and processes by which a company is directed and controlled, including issues related to board composition, executive compensation, and audit committee independence.
Cost inflation: Describes an increase in costs, such as labor, materials, or overheads. 
Customer confidence: Refers to the level of trust and faith that customers have in a company's products or services, which can impact sales and revenue. 
Debt repayment challenges: Describes difficulties a company faces in repaying its debt, which can lead to default, restructuring, or other negative consequences. 
Decline: Describes a decrease in a company's financial performance, such as revenue, profitability, or market share. 
Delay: Refers to a postponement or deferral of a project, investment, or other business initiative, which can impact a company's financial performance.
Group company exposure: Describes a company's financial exposure to its subsidiaries, affiliates, or joint ventures, which can impact its consolidated financial performance.
Guidance revision: Refers to a change in a company's financial guidance
Impairment charges: Refers to non-cash charges taken by a company to reflect the decline in value of its assets, such as goodwill, property, or equipment. 
Increase provisions: Describes an increase in a company's provisions for bad debts, warranties, or other contingent liabilities.
Increasing working capital: Describes an increase in a company's working capital requirements, such as accounts receivable, inventory, or accounts payable.
Inventory levels gone up: Refers to an increase in a company's inventory levels, which can indicate slower sales, overproduction, or supply chain disruptions.
Liquidity concerns: Describes a company's difficulties in meeting its short-term financial obligations, such as paying debts or meeting working capital requirements.
Margin pressure: Describes a decline in a company's profit or EBIDTA margins.
New management: Refers to the appointment of new executives or managers to a company's leadership team, which can impact its strategy, culture, and financial performance.
One-off expenses: Refers to non-recurring expenses or charges taken by a company, such as restructuring costs, impairment charges, or litigation expenses.
One-time write-offs: Refers to non-recurring write-offs or charges taken by a company, such as asset impairments, inventory write-offs, or accounts receivable write-offs.
Operational issues: Describes challenges or problems a company faces in its operations.
Regulatory uncertainty: Describes uncertainty or ambiguity related to regulatory requirements, laws, or policies.
Related party transaction: Refers to a transaction between a company and its related parties, such as subsidiaries, affiliates, or joint ventures, which can impact its financial performance and transparency.
Restructuring efforts: Refers to a company's plans or actions to reorganize its operations, finances, or management structure to improve its performance, efficiency, or competitiveness.
Scale down: Describes a company's decision to reduce its operations, investments, or workforce to conserve resources, cut costs, or adapt to changing market conditions.
Service issue: Refers to problems or difficulties a company faces in delivering its services.
Shortage: Describes a situation where a company faces a lack of supply, resources, or personnel.
Stress: Refers to a company's financial difficulties or challenges, such as debt, cash flow problems etc.
Supply chain disruptions: Refers to interruptions or problems in a company's supply chain, which can impact its ability to produce, deliver, or distribute its products or services.
Warranty cost: Refers to the expenses or provisions a company makes for warranties or guarantees provided to its customers.
Misappropriation of funds: Describes the unauthorized or improper use of a company's funds, assets, or resources.
"""
keywords_part2 = """
Increase in borrowing cost: Refers to a rise in the cost of borrowing for a company. 
One time reversal: Describes a non-recurring or one-time adjustment to a company's financial statements.
Bloated balance sheet: Refers to a company's balance sheet that is overly leveraged, inefficient, or burdened with debt.
Reversal: a credit or refund to the customer, which reduces the original sale and is recorded as a reduction in revenue.
Debtors increasing or going up: Refers to an increase in a company's accounts receivable or debtors.
Receivables increase: Describes an increase in a company's accounts receivable.
Challenges in collections: Refers to difficulties a company faces in collecting its accounts receivable or debtors, which can impact its cash flow, liquidity, or financial performance. 
Slow down on disbursement: A reduction in the rate at which loans or funds are disbursed.
Write-offs: The process of removing a debt or asset from a company's balance sheet or Profit and loss statement.
Increase of provisioning: An increase in the amount of money set aside by a financial institution to cover potential losses on loans or assets.
Delinquency increase: A rise in the number of borrowers who are late or behind on their loan payments, often indicating a deterioration in credit quality.
GNPA increasing: An increase in Gross Non-Performing Assets (GNPA), which refers to the total value of loans that are overdue or in default.
Slippages: The reclassification of loans from a performing to a non-performing category
High credit deposit ratio: A situation where a bank's credit growth exceeds its deposit growth.
CAR decreasing: A decline in the Capital Adequacy Ratio (CAR), which measures a bank's capital as a percentage of its risk-weighted assets.
Provision coverage falling: A decline in the provision coverage ratio, indicating that the provisions made for potential losses are decreasing relative to the growth in non-performing assets.
Low Profitability: A state where a business, project, or investment generates revenue, but the net income or return on investment (ROI) is significantly lower than expected, industry average, or benchmark. 
Falling Net Interest Margin (NIM): A decrease in the difference between the interest income earned by a financial institution and the interest expense paid on deposits and other borrowings due to changes in interest or deposit rate, reduced profitability etc.
Negative Capital Employed: Statements that indicate a company's liabilities exceed its assets, or its return on capital employed is negative.
Capacity Utilisation falling: Refers to the extent to which a company's production facilities or resources are being used, with low utilisation indicating underproduction or declining demand.
Destocking: The process of reducing inventory levels, often due to decreased demand or overstocking, which can indicate a decline in sales or shift in market trends.
Pricing Pressure: Downward pressure on a company's prices due to competition or market conditions.
Renegotiation: The process of revising or re-evaluating existing contracts or agreements, which can indicate disputes or changes in market conditions.
Credit rating action/Rating downgrade/Watch negative: A change in a company's credit rating, indicating a higher risk of default or negative outlook.
Weakening/softening of demand: A decline in customer demand or slowdown in sales growth, indicating a decline in market share or shift in market trends.
Long recovery time: A prolonged period required for a company to recover from a downturn or disruption, indicating significant challenges or reduced competitiveness.
Capex plan mentioned but no roadmap/clarity of funding: A capital expenditure plan without a clear plan for funding or implementation, indicating a lack of financial resources or unclear priorities.
Loss: A financial loss incurred by a company, indicating poor financial management or reduced competitiveness.
Anti-dumping: Measures taken to prevent the importation of goods at below-normal prices, which can indicate trade tensions or protectionism.
Demerger: The separation of a company into independent entities, often to improve focus or reduce complexity, but can also indicate a lack of synergy or decline in profitability.
"""



class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""<role>
You are an expert document analyst specializing in extracting key corporate information from financial documents and earnings transcripts.
</role>

<system_prompt>
You excel at quickly identifying and extracting specific corporate identifiers from financial documents with high accuracy and consistency.
</system_prompt>

<instruction>
Extract the company name, quarter, and financial year from the provided text.

EXTRACTION REQUIREMENTS:
1. Company Name: Full legal company name including suffixes (Ltd/Limited/Inc/Corp etc.)
2. Quarter: Identify the quarter (Q1/Q2/Q3/Q4)  
3. Financial Year: Extract financial year (FY23/FY24/FY25 etc.)

OUTPUT FORMAT:
Provide ONLY the result in this exact format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If any component cannot be clearly identified, use reasonable defaults based on context.
</instruction>

<context>
DOCUMENT TEXT TO ANALYZE:
{first_page_text}
</context>

Extract company information:"""
       
        response = llm._call(prompt)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

# ==============================================================================
# NEW EXCEL READING FUNCTIONS
# ==============================================================================

def read_prompts_and_keywords_from_excel(queries_csv_path: str):
    """Read main prompt and split keywords from existing Excel structure"""
    
    try:
        if queries_csv_path.endswith('.xlsx'):
            queries_df = pd.read_excel(queries_csv_path)
        else:
            queries_df = pd.read_csv(queries_csv_path)
        
        if len(queries_df) < 3 or "prompt" not in queries_df.columns:
            # Fallback if structure is not as expected
            return (
                "Analyze this document for potential red flags.",
                "<reference>No keywords part 1 available</reference>",
                "<reference>No keywords part 2 available</reference>"
            )
        
        # Extract the three rows
        main_prompt = queries_df["prompt"].tolist()[0]        # Row 1: Main prompt
        keywords_part_1 = queries_df["prompt"].tolist()[1]    # Row 2: Keywords 1-32
        keywords_part_2 = queries_df["prompt"].tolist()[2]    # Row 3: Keywords 33-63
        
        return main_prompt, keywords_part_1, keywords_part_2
        
    except Exception as e:
        logger.warning(f"Error loading queries file: {e}. Using defaults.")
        return (
            "Analyze this document for potential red flags.",
            "<reference>No keywords part 1 available</reference>", 
            "<reference>No keywords part 2 available</reference>"
        )

# ==============================================================================
# CRITERIA BUCKETS AND CLASSIFICATION FUNCTIONS
# ==============================================================================

def create_criteria_buckets():
    """Organize 23 criteria into 6 buckets for better LLM classification"""
    
    # Bucket 1: Core Debt & Leverage (4 criteria)
    bucket_1 = {
        "debt_increase": "High: Debt is increased more than 30% compared to previous reported balance sheet number; Low: Debt increased less than 30% compared to previous reported balance sheet number",
        "debt_ebitda": "High: Debt/EBITDA > 3x i.e. Debt to EBITDA ratio is above (greater than) three times; Low: Debt/EBITDA < 3x i.e. Debt to EBITDA ratio is less than three times",
        "short_term_borrowings": "High: Short-term borrowings or current liabilities increase by more than 30% compared to previous reported balance sheet number; Low: Short-term borrowings or current liabilities increase is less than 30% compared to previous reported balance sheet number",
        "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number"
    }
    
    # Bucket 2: Profitability & Performance (4 criteria)
    bucket_2 = {
        "revenue_decline": "High: revenue falls by more than 25% compared to previous reported quarter number; Low: revenue falls by less than 25% compared to previous reported quarter number",
        "profit_before_tax_decline": "High: profitability or profit before tax (PBT) falls by more than 25% compared to previous reported quarter number; Low: profitability or profit before tax (PBT) falls by less than 25% compared to previous reported quarter number",
        "profit_after_tax_decline": "High: Profit after tax (PAT) falls by more than 25% compared to previous reported quarter number; Low: Profit after tax (PAT) falls by less than 25% compared to previous reported quarter number",
        "EBIDTA_decline": "High: EBITDA falls by more than 25% compared to previous reported quarter number; Low: EBITDA falls by less than 25% compared to previous reported quarter number"
    }
    
    # Bucket 3: Margins & Operational Efficiency (4 criteria)
    bucket_3 = {
        "margin_decline": "High: operating margin falling more than 25% compared to previous reported quarter number; Low: Operating margin falling less than 25% compared to previous reported quarter number",
        "gross_margin": "High: gross margin falling more than 100bps (basis points) ; Low: gross margin falling less than 100bps (basis points)",
        "one-time_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBITDA; Low: one-time expenses or losses less than 25% of current quarter's EBITDA",
        "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBITDA; Low: provisioning or write-offs less than 25% of current quarter's EBITDA"
    }
    
    # Bucket 4: Working Capital & Asset Management (4 criteria)
    bucket_4 = {
        "receivable_days": "High: receivable days OR debtor days are increased more than 30% compared to previous reported balance sheet number; Low: receivable days or debtor's days are increased but less than 30% compared to previous reported balance sheet number",
        "payable_days": "High: payable days or creditors days increase by more than 30% compared to previous reported balance sheet number; Low: payable days or creditors days increase is less than 30% compared to previous reported balance sheet number",
        "receivables": "High: receivables or debtors are increased more than 30% compared to previous reported balance sheet number; Low: receivables or debtors are increase is less than 30% compared to previous reported balance sheet number",
        "payables": "High: payables or creditors increase by greater than 30% compared to previous reported balance sheet number; Low: payables or creditors is less than 30% compared to previous reported balance sheet number"
    }
    
    # Bucket 5: Asset Quality & Governance (4 criteria) - ENHANCED FOR QUALITATIVE
    bucket_5 = {
        "asset_decline": "High: Asset value falls by more than 30% compared to the previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
        "impairment": "High: Impairment or devaluation more than 25% of previous reported net worth from balance sheet; Low: Impairment or devaluation less than 25% of previous reported net worth from balance sheet",
        "management_issues": "High: If CEO/CFO resignation, fraud/misconduct investigations, major governance failures, significant strategy failures, or loss of institutional confidence; Low: Routine management changes, normal strategy discussions, minor organizational adjustments",
        "regulatory_compliance": "High: If regulatory penalties, violations, warnings from regulators, license threats, or material compliance failures; Low: Routine compliance discussions, minor filing delays, normal regulatory interactions"
    }
    
    # Bucket 6: Market & Operational Risks (3 criteria) - ENHANCED FOR QUALITATIVE
    bucket_6 = {
        "market_competition": "High: If significant market share loss mentioned with numbers, major competitor entry causing material impact, or pricing wars affecting margins significantly; Low: General competitive pressure, normal market dynamics, routine competitive discussions",
        "operational_disruptions": "High: If major facility shutdowns, critical supply chain breaks with material impact, widespread strikes, or significant operational failures mentioned; Low: Minor operational adjustments, routine efficiency discussions, normal operational challenges",
        "others": "High: Other material issues with quantified impact > 25% of current quarter EBITDA or clear material business impact; Low: Other minor issues or concerns without material impact"
    }
    
    return [bucket_1, bucket_2, bucket_3, bucket_4, bucket_5, bucket_6]

def create_previous_data_buckets(previous_year_data: str):
    """Organize previous year data into 6 buckets matching the criteria buckets"""
    
    # Parse the previous year data to extract relevant metrics for each bucket
    lines = previous_year_data.strip().split('\n')
    data_dict = {}
    
    for line in lines:
        if line.strip():
            parts = line.split('\t')
            if len(parts) >= 3:
                metric = parts[0].strip()
                value = '\t'.join(parts[1:]).strip()
                data_dict[metric.lower()] = f"{metric}\t{value}"
    
    # Bucket 1: Core Debt & Leverage
    bucket_1_data = ""
    for key in ['debt as per previous reported balance sheet number', 'current quarter ebitda', 'cash balance as per previous reported balance sheet number', 'short term borrowings as per the previous reported balance sheet number', 'ebitda as per previous reported quarter number']:
        if key in data_dict:
            bucket_1_data += data_dict[key] + "\n"
    
    # Bucket 2: Profitability & Performance  
    bucket_2_data = ""
    for key in ['revenue as per previous reported quarter number', 'profit before tax as per previous reported quarter number', 'profit after tax as per previous reported quarter number', 'ebitda as per previous reported quarter number', 'current quarter ebitda']:
        if key in data_dict:
            bucket_2_data += data_dict[key] + "\n"
    
    # Bucket 3: Margins & Operational Efficiency
    bucket_3_data = ""
    for key in ['operating margin as per previous quarter number', 'current quarter ebitda', 'ebitda as per previous reported quarter number', 'revenue as per previous reported quarter number']:
        if key in data_dict:
            bucket_3_data += data_dict[key] + "\n"
    
    # Bucket 4: Working Capital & Asset Management
    bucket_4_data = ""
    for key in ['receivable days as per previous reported balance sheet number', 'payable days as per previous reported balance sheet number', 'receivables as per previous reported balance sheet number', 'payables as per previous reported balance sheet number']:
        if key in data_dict:
            bucket_4_data += data_dict[key] + "\n"
    
    # Bucket 5: Asset Quality & Governance
    bucket_5_data = ""
    for key in ['asset value as per previous reported balance sheet number', 'previous reported net worth from balance sheet', 'current quarter ebitda']:
        if key in data_dict:
            bucket_5_data += data_dict[key] + "\n"
    
    # Bucket 6: Market & Operational Risks
    bucket_6_data = ""
    for key in ['current quarter ebitda', 'revenue as per previous reported quarter number', 'ebitda as per previous reported quarter number']:
        if key in data_dict:
            bucket_6_data += data_dict[key] + "\n"
    
    return [bucket_1_data, bucket_2_data, bucket_3_data, bucket_4_data, bucket_5_data, bucket_6_data]

def classify_all_flags_with_buckets(all_flags_with_context: List[str], previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, List[Dict[str, str]]]:
    """
    Efficient classification using 6 total LLM calls for all flags combined - one call per bucket
    """
    
    criteria_buckets = create_criteria_buckets()
    data_buckets = create_previous_data_buckets(previous_year_data)
    
    bucket_names = [
        "Core Debt & Leverage",
        "Profitability & Performance", 
        "Margins & Operational Efficiency",
        "Working Capital & Asset Management",
        "Asset Quality & Governance",
        "Market & Operational Risks"
    ]
    
    # Prepare all flags text for analysis
    all_flags_text = ""
    for i, flag in enumerate(all_flags_with_context, 1):
        all_flags_text += f"\n--- FLAG {i} ---\n{flag}\n"
    
    bucket_results = {}
    
    for i, (criteria_bucket, data_bucket, bucket_name) in enumerate(zip(criteria_buckets, data_buckets, bucket_names)):
        criteria_list = "\n".join([f"{name}: {desc}" for name, desc in criteria_bucket.items()])
        
        # Enhanced prompt for bulk analysis
        if i >= 4:  # Buckets 5 and 6 contain qualitative criteria
            prompt = f"""Analyze ALL red flags below against the {bucket_name} criteria with ENHANCED QUALITATIVE ASSESSMENT.

ALL RED FLAGS TO ANALYZE:
{all_flags_text}

RELEVANT CRITERIA FOR THIS BUCKET:
{criteria_list}

RELEVANT PREVIOUS YEAR DATA FOR COMPARISON:
{data_bucket}

SPECIAL INSTRUCTIONS FOR QUALITATIVE ASSESSMENT:
- For management_issues: Look for SPECIFIC evidence of senior leadership changes, investigations, or governance failures
- For regulatory_compliance: Look for SPECIFIC penalties, violations, or regulatory warnings mentioned
- For market_competition: Look for QUANTIFIED market share losses or material competitive impacts
- For operational_disruptions: Look for SPECIFIC facility issues, supply chain breaks, or operational failures

ANALYSIS TASK:
1. Review each flag against the criteria for this bucket
2. Extract specific evidence from original quotes where available
3. Apply STRICT HIGH RISK thresholds for qualitative criteria
4. Provide analysis for each flag that matches any criteria in this bucket

RESPONSE FORMAT - For each matching flag:
FLAG_[number]: [flag description]
Matched_Criteria: [exact criteria name]
Risk_Level: [High or Low]
Confidence: [1-10]
Quote_Evidence: [specific quotes that support the classification]
Reasoning: [detailed explanation with evidence]

---

Only include flags that match criteria in this bucket. If no flags match, respond with "No flags match {bucket_name} criteria."

IMPORTANT: Be CONSERVATIVE with qualitative High Risk classifications."""
        else:
            prompt = f"""Analyze ALL red flags below against the {bucket_name} criteria.

ALL RED FLAGS TO ANALYZE:
{all_flags_text}

RELEVANT CRITERIA FOR THIS BUCKET:
{criteria_list}

RELEVANT PREVIOUS YEAR DATA FOR COMPARISON:
{data_bucket}

ANALYSIS INSTRUCTIONS:
1. Extract EXACT numbers and percentages from original quotes in each flag
2. Calculate percentage changes using previous year data benchmarks
3. Compare against criteria thresholds (>25%, >30%, >3x, etc.)
4. Analyze each flag independently against this bucket's criteria

ANALYSIS TASK:
1. Review each flag for numerical data that matches this bucket's criteria
2. Extract specific values from original quotes
3. Calculate percentage changes where applicable
4. Determine High/Low risk based on thresholds

RESPONSE FORMAT - For each matching flag:
FLAG_[number]: [flag description]
Matched_Criteria: [exact criteria name]
Risk_Level: [High or Low]
Confidence: [1-10]
Extracted_Data: [specific numbers from quotes]
Calculation: [percentage change calculations]
Reasoning: [explanation with calculations]

---

Only include flags that match criteria in this bucket. If no flags match, respond with "No flags match {bucket_name} criteria."

IMPORTANT: Only consider the criteria listed above for this bucket."""

        try:
            print(f"Analyzing all flags against {bucket_name} bucket...")
            response = llm._call(prompt, temperature=0.0)
            bucket_results[bucket_name] = response
            
        except Exception as e:
            logger.error(f"Error analyzing {bucket_name}: {e}")
            bucket_results[bucket_name] = f"Error in {bucket_name}: {str(e)}"
        
        time.sleep(0.5)  # Rate limiting between bucket calls
    
    return bucket_results

def parse_bucket_results_to_classifications(bucket_results: Dict[str, str], all_flags_with_context: List[str]) -> List[Dict[str, str]]:
    """
    Parse the bucket analysis results into individual flag classifications
    """
    flag_classifications = []
    
    # Initialize all flags as Low risk
    for i, flag_with_context in enumerate(all_flags_with_context, 1):
        flag_description = flag_with_context.split('\n')[0]
        flag_description = re.sub(r'^\d+\.\s+', '', flag_description).strip()
        
        flag_classifications.append({
            'flag': flag_description,
            'flag_with_context': flag_with_context,
            'matched_criteria': 'None',
            'risk_level': 'Low',
            'reasoning': 'No matching criteria found across all buckets',
            'bucket': 'None'
        })
    
    # Parse bucket results to update classifications
    for bucket_name, bucket_response in bucket_results.items():
        if isinstance(bucket_response, str) and "No flags match" not in bucket_response and "Error" not in bucket_response:
            lines = bucket_response.strip().split('\n')
            current_flag_index = None
            current_flag_data = {}
            
            for line in lines:
                line = line.strip()
                if line.startswith('FLAG_'):
                    # Save previous flag data if exists
                    if current_flag_index is not None and current_flag_data:
                        try:
                            idx = current_flag_index - 1  # Convert to 0-based index
                            if 0 <= idx < len(flag_classifications):
                                # Only update if this is higher risk or higher confidence
                                current_risk = current_flag_data.get('risk_level', 'Low')
                                current_confidence = current_flag_data.get('confidence', 0)
                                existing_risk = flag_classifications[idx]['risk_level']
                                
                                if (current_risk == 'High' and existing_risk == 'Low') or \
                                   (current_risk == existing_risk and current_confidence > 5):
                                    flag_classifications[idx].update({
                                        'matched_criteria': current_flag_data.get('matched_criteria', 'None'),
                                        'risk_level': current_risk,
                                        'reasoning': current_flag_data.get('reasoning', 'Bucket analysis'),
                                        'bucket': bucket_name
                                    })
                        except:
                            pass
                    
                    # Start new flag
                    try:
                        current_flag_index = int(line.split('_')[1].split(':')[0])
                        current_flag_data = {}
                    except:
                        current_flag_index = None
                        
                elif current_flag_index and ':' in line:
                    key_value = line.split(':', 1)
                    if len(key_value) == 2:
                        key = key_value[0].strip().lower().replace(' ', '_')
                        value = key_value[1].strip()
                        
                        if key == 'matched_criteria':
                            current_flag_data['matched_criteria'] = value
                        elif key == 'risk_level':
                            current_flag_data['risk_level'] = value
                        elif key == 'confidence':
                            try:
                                current_flag_data['confidence'] = int(''.join(filter(str.isdigit, value)))
                            except:
                                current_flag_data['confidence'] = 5
                        elif key in ['reasoning', 'quote_evidence', 'extracted_data', 'calculation']:
                            if 'reasoning' not in current_flag_data:
                                current_flag_data['reasoning'] = value
                            else:
                                current_flag_data['reasoning'] += f" {value}"
            
            # Don't forget the last flag
            if current_flag_index is not None and current_flag_data:
                try:
                    idx = current_flag_index - 1
                    if 0 <= idx < len(flag_classifications):
                        current_risk = current_flag_data.get('risk_level', 'Low')
                        current_confidence = current_flag_data.get('confidence', 0)
                        existing_risk = flag_classifications[idx]['risk_level']
                        
                        if (current_risk == 'High' and existing_risk == 'Low') or \
                           (current_risk == existing_risk and current_confidence > 5):
                            flag_classifications[idx].update({
                                'matched_criteria': current_flag_data.get('matched_criteria', 'None'),
                                'risk_level': current_risk,
                                'reasoning': current_flag_data.get('reasoning', 'Bucket analysis'),
                                'bucket': bucket_name
                            })
                except:
                    pass
    
    return flag_classifications

def extract_flags_with_complete_context(second_response: str) -> List[str]:
    """
    Enhanced flag extraction that preserves complete context including original quotes and page references
    """
    flags_with_context = []
    lines = second_response.split('\n')
    current_flag = ""
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Check if this is the start of a new flag
        if re.match(r'^\d+\.\s+', line):
            # Save previous flag if it exists
            if current_flag.strip():
                flags_with_context.append(current_flag.strip())
            
            # Start new flag
            current_flag = line
            
            # Look ahead to capture original quotes and page references
            j = i + 1
            while j < len(lines) and not re.match(r'^\d+\.\s+', lines[j].strip()):
                next_line = lines[j].strip()
                if next_line:  # Only add non-empty lines
                    current_flag += "\n" + next_line
                j += 1
        
    # Don't forget the last flag
    if current_flag.strip():
        flags_with_context.append(current_flag.strip())
    
    # Clean and validate flags
    cleaned_flags = []
    for flag in flags_with_context:
        # Remove any prefixes but keep the complete context
        flag = re.sub(r'^The potential red flag you observed - ', '', flag)
        flag = flag.strip()
        
        if flag and len(flag) > 10:  # Minimum length check
            cleaned_flags.append(flag)
    
    return cleaned_flags

# ==============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ==============================================================================

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(high_risk_flags: List[str], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    # Deduplicate the high_risk_flags
    unique_high_risk_flags = []
    seen_flag_keywords = []
    
    for flag in high_risk_flags:
        normalized_flag = re.sub(r'[^\w\s]', '', flag.lower()).strip()
        flag_words = set(normalized_flag.split())
        
        # Check for keyword overlap with existing flags
        is_duplicate_flag = False
        for existing_keywords in seen_flag_keywords:
            overlap = len(flag_words.intersection(existing_keywords)) / max(len(flag_words), len(existing_keywords))
            if overlap > 0.7:  # High threshold for flag deduplication
                is_duplicate_flag = True
                break
        
        if not is_duplicate_flag:
            unique_high_risk_flags.append(flag)
            seen_flag_keywords.append(flag_words)
    
    concise_summaries = []
    seen_summary_keywords = []
    
    for flag in unique_high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context}

HIGH RISK FLAG: "{flag}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Ensure subsequent statements are cautious and do not downplay the risk. Avoid neutral/positive statements that contradict the warning.
6. Do NOT start with "Summary:" or any prefix
7. Provide ONLY the factual summary content
8. Make it UNIQUE - avoid repeating information from other summaries
9. If applicable Specify whether the flag is for : A specific business unit/division, Consolidated financials, Standalone financials, geographical region.

OUTPUT FORMAT: [Direct factual summary only, no labels or prefixes]

"""
        
        try:
            response = llm._call(prompt, temperature=0.1)
            
            # Clean response - remove any prefixes or labels
            clean_response = response.strip()
            
            # Remove common prefixes that might appear
            prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:", "Flag summary:", "The flag:", "This flag:"]
            for prefix in prefixes_to_remove:
                if clean_response.startswith(prefix):
                    clean_response = clean_response[len(prefix):].strip()
            
            # Split into lines and take first 2
            summary_lines = [line.strip() for line in clean_response.split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                concise_summary = f"{flag}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
            
            # Check for duplicate content in summaries
            normalized_summary = re.sub(r'[^\w\s]', '', concise_summary.lower()).strip()
            summary_words = set(normalized_summary.split())
            
            is_duplicate_summary = False
            for existing_keywords in seen_summary_keywords:
                overlap = len(summary_words.intersection(existing_keywords)) / max(len(summary_words), len(existing_keywords))
                if overlap > 0.8:  # Very high threshold for summary deduplication
                    is_duplicate_summary = True
                    break
            
            if not is_duplicate_summary:
                concise_summaries.append(concise_summary)
                seen_summary_keywords.append(summary_words)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag}': {e}")
            if len(concise_summaries) < len(unique_high_risk_flags):
                concise_summaries.append(f"{flag}. Review required based on analysis.")
    
    return concise_summaries

def create_word_document(pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                        high_risk_flags: List[str], summary_by_categories: Dict[str, List[str]], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> str:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        high_count = risk_counts.get('High', 0)
        low_count = risk_counts.get('Low', 0)
        total_count = high_count + low_count
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(high_count)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(low_count)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(total_count)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            # Final deduplication check at Word document level
            final_unique_summaries = []
            seen_content = set()
            
            for summary in concise_summaries:
                if not summary or not summary.strip():
                    continue
                    
                # Create multiple normalized versions for comparison
                normalized1 = re.sub(r'[^\w\s]', '', summary.lower()).strip()
                normalized2 = re.sub(r'\b(the|a|an|and|or|but|in|on|at|to|for|of|with|by)\b', '', normalized1)
                
                # Check if this content is substantially different
                is_unique = True
                for seen in seen_content:
                    # Calculate similarity
                    words1 = set(normalized2.split())
                    words2 = set(seen.split())
                    if len(words1) == 0 or len(words2) == 0:
                        continue
                    similarity = len(words1.intersection(words2)) / len(words1.union(words2))
                    if similarity > 0.6:  # If more than 60% similar, consider duplicate
                        is_unique = False
                        break
                
                if is_unique:
                    final_unique_summaries.append(summary)
                    seen_content.add(normalized2)
            
            for summary in final_unique_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.get('High', 0)}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.get('Low', 0)}")
            doc.add_paragraph(f"Total Flags: {risk_counts.get('Total', 0)}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None


system_prompt_step_1 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern directly, and refrain from highlighting any positive or neutral flags.
7. Extract all original quotes and contexts explaining the identified red flag. No quotes need to be missed that can help explain the red flag. 
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. Ensure the original quotes are comprehensive and all inclusive for the red flags identified. 
5. No explanation needed.
</review>

####
{keywords_part1}
####

"""

system_prompt_step_1_2 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags. 
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. No explanation needed.
</review>

####
{keywords_part2}
####

"""

# ==============================================================================
# MAIN PROCESSING PIPELINE WITH SPLIT FIRST ITERATION
# ==============================================================================

def process_pdf_enhanced_pipeline_with_split_iteration(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4o"):
    """
    Enhanced processing pipeline with split first iteration using existing Excel structure
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM and load PDF
        llm_client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint, 
            api_version=api_version,
        )
        
        llm = AzureOpenAILLM(
            api_key=api_key,
            azure_endpoint=azure_endpoint, 
            api_version=api_version,
            deployment_name=deployment_name
        )

        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Make a chat completions call
        response1 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_1},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )
        print(response1.choices[0].message.content)
        print("******************************************")
        # Make a chat completions call
        response2 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_1_2},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )

        print(response2.choices[0].message.content)
        print("******************************************")
        first_response = response1.choices[0].message.content + "\n\n" + response2.choices[0].message.content
        
        print(f"Split first iteration completed. Combined response length: {len(first_response)} characters")
        
        # ITERATION 2: Enhanced Deduplication (UNCHANGED)
        print("Running 2nd iteration - Enhanced Deduplication...")
        second_prompt = f"""<role>
You are an expert financial analyst specializing in identifying and eliminating duplicate red flags while maintaining comprehensive analysis integrity.
</role>

<system_prompt>
You excel at recognizing when multiple red flags describe the same underlying financial issue, even when worded differently, and consolidating them into single, comprehensive red flags while preserving all supporting evidence.
</system_prompt>

<instruction>
Analyze the red flags and remove duplicates that describe the same underlying financial concern. Consolidate similar issues into single, comprehensive red flags.

DEDUPLICATION RULES:
1. MERGE red flags that refer to the same financial metric, issue, or concern
2. COMBINE red flags about the same business area/division/segment  
3. CONSOLIDATE similar operational or strategic concerns
4. ELIMINATE redundant mentions of the same data point or statistic
5. KEEP the most comprehensive version with the best supporting evidence
6. PRESERVE all original quotes, speaker attributions, and page references from merged items
7. MAINTAIN sequential numbering (1, 2, 3, etc.) after deduplication
8. DO NOT lose any substantive financial concerns - only remove true duplicates

OUTPUT FORMAT:
1. [Comprehensive red flag description covering all related issues]
Original Quotes: "[Combined relevant quotes with speaker names]" (Page X, Y, Z)

2. [Next unique red flag]
Original Quote: "[quote with speaker name]" (Page X)

CRITICAL INSTRUCTIONS:
- START YOUR RESPONSE IMMEDIATELY WITH "1." - NO INTRODUCTION OR EXPLANATION
- DO NOT include any introductory text, summaries, or conclusions
- PRESERVE ALL original quotes and page references
- Be aggressive in removing duplicates while preserving all important context and evidence
</instruction>

<context>
ORIGINAL DOCUMENT CONTEXT:
{context}

FIRST ITERATION ANALYSIS TO DEDUPLICATE:
{first_response}
</context>

Provide deduplicated analysis with merged duplicates and preserved evidence:"""
        
        second_response = llm._call(second_prompt)
        
        # ITERATION 3: Categorization (UNCHANGED)
        print("Running 3rd iteration - Categorization...")
        third_prompt = f"""<role>
You are a senior financial analyst expert in financial risk categorization with deep knowledge of balance sheet analysis, P&L assessment, and corporate risk frameworks.
</role>

<system_prompt>
You excel at organizing financial risks into standardized categories, ensuring comprehensive coverage of all financial risk areas, and maintaining accuracy in risk classification.
</system_prompt>

<instruction>
Categorize the identified red flags into the following 7 standardized categories based on their financial nature and business impact.

MANDATORY CATEGORIES:
1. Balance Sheet Issues: Assets, liabilities, equity, debt, and overall financial position concerns
2. P&L (Income Statement) Issues: Revenue, expenses, profits, and financial performance concerns  
3. Liquidity Issues: Short-term obligations, cash flow, debt repayment, working capital concerns
4. Management and Strategy Issues: Leadership, governance, decision-making, strategy, and vision concerns
5. Regulatory Issues: Compliance, laws, regulations, and regulatory body concerns
6. Industry and Market Issues: Market position, industry trends, competitive landscape concerns
7. Operational Issues: Internal processes, systems, infrastructure, and operational efficiency concerns

CATEGORIZATION RULES:
- Assign each red flag to the MOST relevant category only
- Do not create new categories - use only the 7 listed above
- Preserve all Original Quotes exactly as provided
- Maintain sequential organization within each category

OUTPUT FORMAT:
### Balance Sheet Issues
- [Red flag 1 with original quote and page reference]
- [Red flag 2 with original quote and page reference]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote and page reference]

Continue this format for all applicable categories.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}
</context>

Provide categorized analysis:"""
        
        third_response = llm._call(third_prompt)
        
        # ITERATION 4: Summary Generation (UNCHANGED)
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = f"""<role>
You are an expert financial summarization specialist with expertise in creating concise, factual, and comprehensive summaries that preserve critical quantitative data and key insights.
</role>

<system_prompt>
You excel at distilling complex financial analysis into clear, actionable summaries while maintaining objectivity, preserving all quantitative details, and ensuring no critical information is lost.
</system_prompt>

<instruction>
Create a comprehensive summary of each category of red flags in bullet point format following these strict guidelines.

SUMMARY REQUIREMENTS:
1. Retain ALL quantitative information (numbers, percentages, ratios, dates)
2. Maintain completely neutral, factual tone - no opinions or interpretations
3. Include every red flag from each category - no omissions
4. Base content solely on the provided categorized analysis
5. Preserve specific data points and statistics wherever mentioned
6. Each bullet point should capture key details of individual red flags
7. Balance thoroughness with conciseness
8. Use professional financial terminology
9. Ensure category-specific content alignment

OUTPUT FORMAT:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories that contain red flags.

CRITICAL: Each bullet point represents a concise summary of individual red flags with preserved quantitative details.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        fourth_response = llm._call(fourth_prompt)
        
        # ITERATION 5: Efficient Bucket-Based Classification (UNCHANGED)
        print("Running 5th iteration - Efficient Bucket-Based Classification...")
        
        try:
            flags_with_context = extract_flags_with_complete_context(second_response)
            print(f"\nFlags with context extracted: {len(flags_with_context)}")
            
            if flags_with_context:
                print(f"Example flag with context:\n{flags_with_context[0][:200]}...")
            
        except Exception as e:
            logger.error(f"Error parsing flags with context: {e}")
            flags_with_context = ["Error in flag parsing"]

        classification_results = []
        high_risk_flags = []
        low_risk_flags = []

        if len(flags_with_context) > 0 and flags_with_context[0] != "Error in flag parsing":
            try:
                print(f"Analyzing all {len(flags_with_context)} flags using 6 bucket calls.")
                
                bucket_results = classify_all_flags_with_buckets(flags_with_context, previous_year_data, llm)
                classification_results = parse_bucket_results_to_classifications(bucket_results, flags_with_context)
                
                for result in classification_results:
                    if (result['risk_level'].lower() == 'high' and 
                        result['matched_criteria'] != 'None'):
                        high_risk_flags.append(result['flag'])
                    else:
                        low_risk_flags.append(result['flag'])
                        
            except Exception as e:
                logger.error(f"Error in efficient bucket classification: {e}")
                for flag_with_context in flags_with_context:
                    flag_description = flag_with_context.split('\n')[0]
                    flag_description = re.sub(r'^\d+\.\s+', '', flag_description).strip()
                    
                    classification_results.append({
                        'flag': flag_description,
                        'flag_with_context': flag_with_context,
                        'matched_criteria': 'None',
                        'risk_level': 'Low',
                        'reasoning': f'Classification failed: {str(e)}',
                        'bucket': 'Error'
                    })
                    low_risk_flags.append(flag_description)

        risk_counts = {
            'High': len(high_risk_flags),
            'Low': len(low_risk_flags),
            'Total': len(flags_with_context) if flags_with_context and flags_with_context[0] != "Error in flag parsing" else 0
        }
        
        print(f"\n=== SPLIT ITERATION CLASSIFICATION RESULTS (2+6 LLM calls total) ===")
        print(f"High Risk Flags: {risk_counts['High']}")
        print(f"Low Risk Flags: {risk_counts['Low']}")
        print(f"Total Flags: {risk_counts['Total']}")
        
        if high_risk_flags:
            print(f"\n--- HIGH RISK FLAGS (classified using original quotes) ---")
            for i, flag in enumerate(high_risk_flags, 1):
                print(f"  {i}. {flag}")
        else:
            print(f"\n--- HIGH RISK FLAGS ---")
            print("  No high risk flags identified using efficient bucket analysis")
        
        # Word Document Creation (UNCHANGED)
        print("\nCreating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
           
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                high_risk_flags=high_risk_flags,
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Save all results to CSV files (MODIFIED)
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Split Initial Analysis (2 calls)",
                "Enhanced Deduplication",
                "Categorization",
                "Summary Generation", 
                "Enhanced Context-Based Classification"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Enhanced Context-Based Classification: {risk_counts['High']} High Risk, {risk_counts['Low']} Low Risk flags from {risk_counts['Total']} total flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_split_iteration_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        if len(classification_results) > 0:
            classification_df = pd.DataFrame(classification_results)
            classification_file = os.path.join(output_folder, f"{pdf_name}_split_iteration_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        print(f"\n=== SPLIT ITERATION PROCESSING COMPLETE FOR {pdf_name} ===")
        return results_summary
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

# ==============================================================================
# MAIN FUNCTION
# ==============================================================================

def main(): 
    API_CONFIG = {
        "api_key": "8498c",
        "azure_endpoint": "https://crisil-pp-gpt.openai.azure.com",
        "api_version": "2025-01-01-preview",
        "deployment_name": "gpt-4.1"
    }
    
    PATHS_CONFIG = {
        "pdf_folder_path": r"vedanta_pdf",
        "queries_csv_path": r"EWS_prompts_v2_2.xlsx", 
        "output_folder": r"vedanta_results_split_iteration"
    }
    
    PREVIOUS_YEAR_DATA = """
Debt as per Previous reported balance sheet number	Mar-23	80329Cr
Current quarter ebidta	March-24	11511Cr
Asset value as per previous reported balance sheet number	Mar-23	189455Cr
Receivable days as per previous reported balance sheet number	Mar-23	10days
Payable days as per Previous reported balance sheet number	Mar-23	91days
Revenue as per previous reported quarter number	Dec-23	35541Cr
profit before tax as per previous reported quarter number	Dec-23	4105Cr
profit after tax as per previous reported quarter number	Dec-23	2868Cr
EBIDTA as per previous reported quarter number	Dec-23	8531Cr
Operating margin as per previous quarter number	Dec-23	25%	
Cash balance as per previous reported balance sheet number	Mar-23	9254Cr
Short term borrowings as per the previous reported balance sheet number	Mar-23	36407Cr
previous reported net worth from balance sheet	Mar-23	47896Cr
Receivables as per previous reported balance sheet number	Mar-23	6414Cr
Payables as per Previous reported balance sheet number	Mar-23	11043Cr
"""
   
    os.makedirs(PATHS_CONFIG["output_folder"], exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(PATHS_CONFIG["pdf_folder_path"], "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {PATHS_CONFIG['pdf_folder_path']}")
        return    

    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        # Use the new split iteration function
        result = process_pdf_enhanced_pipeline_with_split_iteration(
            pdf_path=pdf_file,
            queries_csv_path=PATHS_CONFIG["queries_csv_path"],
            previous_year_data=PREVIOUS_YEAR_DATA,
            output_folder=PATHS_CONFIG["output_folder"],
            api_key=API_CONFIG["api_key"],
            azure_endpoint=API_CONFIG["azure_endpoint"],
            api_version=API_CONFIG["api_version"],
            deployment_name=API_CONFIG["deployment_name"]
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
        else:
            print(f"❌ Failed to process {pdf_file}")

if __name__ == "__main__":
    main()
