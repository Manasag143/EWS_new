import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx

warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

class AzureOpenAILLM:
    """Azure OpenAI GPT-4.1-mini LLM class"""
    
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1-mini"):
        """
        Initialize Azure OpenAI client
        
        Args:
            api_key: Your Azure OpenAI API key
            azure_endpoint: Your Azure OpenAI endpoint URL
            api_version: API version (e.g., "2024-02-01")
            deployment_name: Your deployment name (default: "gpt-4.1-mini")
        """
        self.deployment_name = deployment_name
        
        # Create httpx client with SSL verification disabled (if needed)
        httpx_client = httpx.Client(verify=False)
        
        # Initialize Azure OpenAI client
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
    
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            # Extract the response text
            response_text = response.choices[0].message.content
            
            # Log token usage
            if hasattr(response, 'usage'):
                logger.info(f"Tokens used - Prompt: {response.usage.prompt_tokens}, "
                          f"Completion: {response.usage.completion_tokens}, "
                          f"Total: {response.usage.total_tokens}")
            
            return response_text.strip() if response_text else ""
            
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

class PDFExtractor:
    """Class for extracting text from PDF files"""
    
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        start_time = time.time()
        try:
            doc = fitz.open(pdf_path)
            pages = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
            
            # Explicitly close the document to free memory
            doc.close()
            logger.info(f"PDF text extraction took {time.time() - start_time:.2f} seconds")
            return pages
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
    
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        # Merge all pages into single context
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

class LlamaQueryPipeline:
    """Main pipeline class for querying PDF content with Azure OpenAI GPT-4.1-mini"""
    
    def __init__(self, pdf_path: str, queries_csv_path: str = None, 
                 api_key: str = None, azure_endpoint: str = None, 
                 api_version: str = None, deployment_name: str = "gpt-4.1-mini",
                 previous_results_path: str = None):
        """
        Initialize the pipeline with Azure OpenAI
        
        Args:
            pdf_path: Path to PDF file
            queries_csv_path: Path to queries CSV/Excel file
            api_key: Azure OpenAI API key
            azure_endpoint: Azure OpenAI endpoint
            api_version: API version
            deployment_name: Deployment name
            previous_results_path: Path to previous results
        """
        # Initialize Azure OpenAI LLM
        self.llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        self.docs = mergeDocs(pdf_path, split_pages=False)
        
        # Load queries from Excel or CSV (only if provided)
        if queries_csv_path:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            self.queries = queries_df["prompt"].tolist()
        else:
            self.queries = []
            
        self.pdf_path = pdf_path
        self.pdf_name = Path(pdf_path).stem  # Get filename without extension
        
        # Load previous results if provided
        self.previous_results = None
        if previous_results_path and os.path.exists(previous_results_path):
            self.previous_results = pd.read_csv(previous_results_path)
    
    def query_llama_with_chaining(self, new_queries_csv_path: str, iteration_number: int = 2) -> pd.DataFrame:
        """Query the Azure OpenAI API using previous results for chaining"""
        if self.previous_results is None:
            raise ValueError("No previous results loaded. Please provide previous_results_path in __init__")
        
        # Load new queries
        if new_queries_csv_path.endswith('.xlsx'):
            new_queries_df = pd.read_excel(new_queries_csv_path)
        else:
            new_queries_df = pd.read_csv(new_queries_csv_path)
        
        new_queries = new_queries_df["prompt"].tolist()
        
        sys_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{self.docs[0]["context"]}

"""
        
        results = []
        
        # Process each new query with corresponding previous response
        for i, new_query in enumerate(new_queries):
            start = time.perf_counter()
            
            try:
                # Get the corresponding previous response (if available)
                if i < len(self.previous_results):
                    # Check if it's 3rd iteration (chained results) or 2nd iteration (original results)
                    if 'chained_response' in self.previous_results.columns:
                        previous_response = self.previous_results.iloc[i]['chained_response']
                    else:
                        previous_response = self.previous_results.iloc[i]['response']
                    
                    # Create chained prompt: new query + previous response
                    chained_prompt = f"""Previous Analysis: {previous_response}

Based on the above analysis and the original context, please answer: {new_query}

Answer:"""
                else:
                    # If no previous response available, use regular prompt
                    chained_prompt = f"""Question: {new_query}
Answer:"""
                
                full_prompt = f"{sys_prompt}\n{chained_prompt}"
                
                # Get response from Azure OpenAI
                response_text = self.llm._call(full_prompt, max_tokens=4000)
                end = time.perf_counter()
                
                # Calculate token approximations (more accurate with OpenAI response)
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
                
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
                
            except Exception as e:
                end = time.perf_counter()
                
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
            
            results.append(usage)
        
        return pd.DataFrame(results)

    def query_llama(self, maintain_conversation: bool = True, enable_chaining: bool = False) -> pd.DataFrame:
        """Query the Azure OpenAI API for a list of queries using the provided context"""
        sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.

COMPLETE DOCUMENT TO ANALYZE:
{self.docs[0]["context"]}

Your task is to analyze the ENTIRE document above and identify ALL potential red flags.

CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
- Do not miss any sections or concerning statements

EXAMPLE FORMAT:
1. The potential red flag you observed - Debt reduction is lower than expected  
Original Quote:  
"Vikrant Kashyap: Have you -- are you able to reduce any debt in quarter one in Middle East and India?  
Ramesh Kalyanaraman: So India, we are not reduced, but the cash balance has been increased to around INR75 crores, but we have not reduced any debt in Q1 in India and Middle East because Middle East we have not converted any showroom in Q1." (Page 9)  

2. The potential red flag you observed - Margin pressure/Competition intensifying/Cost inflation  
Original Quote:  
"Pulkit Singhal: Thank you for the opportunity and congrats on the good set of performance. Just the first question is really on the margins, which seems to have been elusive. I mean looking at the company for the last two years, we seem to be executing quite well in terms of store expansions and revenue growth and clearly delivering higher than expected there. But margin expansion just has been completely elusive.  
And I find it surprising your comment that while growing at 30% growth and 12% SSSG, I mean, which is quite healthy, you're still talking about high competitive intensity kind of quite contradictory that with such high growth rates, we have to invest so high. So can you talk about this a bit more? I mean we don't expect with lower revenue growth rates that you would not have to invest in the business. And it's only during a higher revenue growth that you expect margin expansion.  
Ramesh Kalyanaraman: Yes. So you're right. We are -- meaning somewhere we have missed out on the operating leverage for advertisements that's why I told you that even Q1, it was a miss. And regarding competition, I will tell you where in new markets, where we assume that we will not spend too much because the brand is already aware and the location is the only thing which has to be communicated.  
When you see the local players, regional players or the micro market players there becoming extremely active because of our showroom launch then we will have to increase the noise level there. Otherwise, we will lose our market share. And existing local players they increase their activity around our launch time. So that is where we also put more money so that we don't end up losing the market share or we don't end up taking out lesser from the competition." (Page 12-13)  

3. The potential red flag you observed - Debt high/Increase in borrowing cost  
Original Quote:  
"Vikrant Kashyap: My question is how are you going to address this? Because if you continue to grow at a higher level, but bottom line is not expanding related to the top line, it will going to impact your overall performance? So, what are the steps you are taking to improve the bottom line in the businesses  
Sanjay Raghuraman: Finance costs will be taken care because we told you when we convert stores, that money is going to reduce our debt. Okay. And again, FOCO, when you do FOCO showrooms, the margins will come down. And surely, that will have an impact on the gross margin. Okay. And interest, if you look at actually, the interest rates have been going up last year. So next year, that will be the base, right? So then again, we will not have this kind of issue is what we feel. So interest rates have been going up over the past year, one year, in that region.  
And we are also beginning to repay loans now because of conversion. So all put together, interest part will be taken care but other area where FOCO showrooms will surely reduce our margin. We cannot have the own store margin. So that should be the way we should look at it." (Page 9)

Continue this exact format for ALL red flags identified throughout the document.

"""
        
        prompt_template = """Question: {query}

Analyze the complete document and provide ALL red flags in the exact numbered format specified above. Be thorough and comprehensive - cover the entire document.

Answer:"""
        
        conversation_history = ""
        results = []
        previous_output = ""  # For prompt chaining
        
        for i, query in enumerate(self.queries, 1):
            start = time.perf_counter()
            
            try:
                if enable_chaining and i > 1 and previous_output:
                    chained_query = f"{query}\n\nPrevious context: {previous_output}"
                else:
                    chained_query = query
                
                if maintain_conversation and conversation_history:
                    full_prompt = f"{sys_prompt}\n{conversation_history}\n{prompt_template.format(query=chained_query)}"
                else:
                    full_prompt = f"{sys_prompt}\n{prompt_template.format(query=chained_query)}"
                
                # Use higher max_tokens for comprehensive analysis
                response_text = self.llm._call(full_prompt, max_tokens=4000)
                end = time.perf_counter()
                
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
                
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": chained_query if enable_chaining else query,
                    "response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
                
                # Update conversation history for next iteration
                if maintain_conversation:
                    conversation_history += f"\nQuestion: {chained_query}\nAnswer: {response_text}\n"
                
                # Store output for next chaining iteration
                if enable_chaining:
                    previous_output = response_text
                
            except Exception as e:
                end = time.perf_counter()
                
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": query,
                    "response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
                
                if enable_chaining:
                    previous_output = ""
            
            results.append(usage)
        
        return pd.DataFrame(results)
    
    def save_results(self, results_df: pd.DataFrame, output_path: str = None):
        """Save results to CSV file"""
        if output_path is None:
            # Generate filename based on PDF name and timestamp
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_path = f"azure_openai_query_results_{self.pdf_name}_{timestamp}.csv"
        
        results_df.to_csv(output_path, index=False)
        return output_path

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
        
        # Limit text to first 2000 characters to avoid token limits
        first_page_text = first_page_text[:2000]
        
        prompt = f"""
You are a financial document analyst. Extract the company name, quarter, and financial year from the following text which is from the first page of an earnings call transcript or financial document.

Text from first page:
{first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)
3. Financial Year (FY23/FY24/FY25 etc.)

Format your response as: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If you cannot find clear information, make the best estimate based on available data.

Response:"""
        
        response = llm._call(prompt, max_tokens=200)
        # Clean the response to get just the formatted string
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
        
        # Fallback - return first non-empty line
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
        
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def debug_risk_data_structure(risk_data):
    """Debug function to inspect the structure of risk_data"""
    print("=== DEBUGGING RISK_DATA STRUCTURE ===")
    print(f"Type of risk_data: {type(risk_data)}")
    print(f"Risk_data content: {risk_data}")
    
    if isinstance(risk_data, dict):
        print(f"Keys in risk_data: {list(risk_data.keys())}")
        
        # Check 'High' key
        if 'High' in risk_data:
            print(f"Type of risk_data['High']: {type(risk_data['High'])}")
            print(f"Content of risk_data['High']: {risk_data['High']}")
            if isinstance(risk_data['High'], list):
                print(f"Length of High list: {len(risk_data['High'])}")
                for i, item in enumerate(risk_data['High']):
                    print(f"  High[{i}]: {type(item)} - {item}")
        
        # Check 'Low' key
        if 'Low' in risk_data:
            print(f"Type of risk_data['Low']: {type(risk_data['Low'])}")
            print(f"Content of risk_data['Low']: {risk_data['Low']}")
        
        # Check 'counts' key
        if 'counts' in risk_data:
            print(f"Type of risk_data['counts']: {type(risk_data['counts'])}")
            print(f"Content of risk_data['counts']: {risk_data['counts']}")
    
    print("=== END DEBUG ===")

def parse_risk_classification_response(response_text: str) -> Dict:
    """
    Parse the 5th iteration response to extract risk flag counts and HIGH RISK summaries
    This replaces LLM calls for counting and summary extraction
    """
    
    # Convert to lowercase for case-insensitive matching
    original_resp = response_text
    response_lower = response_text.lower()
    
    # Initialize counters
    count_high = 0
    count_low = 0
    
    # Only store HIGH RISK summaries (without category)
    high_summaries = []
    low_summaries = []
    
    # Split response into lines
    output_lines = response_lower.splitlines()
    original_lines = original_resp.splitlines()
    
    # Track current red flag summary
    current_summary = ""
    
    for idx, line in enumerate(output_lines):
        line = line.strip()
        original_line = original_lines[idx].strip() if idx < len(original_lines) else ""
        
        # Detect red flag summary (lines starting with * or •)
        if line.startswith("*") or line.startswith("•"):
            # Extract summary (remove bullet point and get text until risk classification)
            summary_match = re.match(r'^[*•]\s*(.+?)(?:\s*-\s*high:|$)', original_line, re.IGNORECASE)
            if summary_match:
                current_summary = summary_match.group(1).strip()
            continue
        
        # Check for risk level classifications and count them
        if "high:" in line and "yes" in line:
            count_high += 1
            # Only store summary for HIGH RISK (without category)
            if current_summary:
                high_summaries.append(current_summary)
                
        elif "low:" in line and "yes" in line:
            count_low += 1
            if current_summary:
                low_summaries.append(current_summary)

    high_summaries_deduped = list(dict.fromkeys(high_summaries))
    low_summaries_deduped = list(dict.fromkeys(low_summaries))

    print("summary:------", high_summaries_deduped)
    
    # Safe LLM initialization with error handling
    try:
        llm = AzureOpenAILLM(
            api_key="8496bd1da40242a29de18fe1361e498c",
            azure_endpoint="https://crisil-pp-gpt.openai.azure.com",
            api_version="2025-01-01-preview",
            deployment_name="gpt-4.1-mini"
        )
        
        prompt1 = f"""
        Remove duplicate from the following list and keep only the unique values. No elements should talk about the same theme/topic.     
        {high_summaries_deduped}
        Output Format:
        [updated_list]
        """
        
        prompt3 = f"""
        Remove duplicate from the following list and keep only the unique values. No elements should talk about the same theme/topic.     
        {low_summaries_deduped}
        Output Format:
        [updated_list]
        """
        
        response1 = llm._call(prompt1, max_tokens=4000)
        response3 = llm._call(prompt3, max_tokens=4000)
        
        print("high.....", response1)
        print("High...Deduped.", high_summaries_deduped)
        print("/nlow.....", response3)
        print("Low...Deduped.", low_summaries_deduped)
        
        # Safe parsing of LLM responses
        try:
            final_high = ast.literal_eval(response1)
        except:
            print("Warning: Could not parse high risk response, using original list")
            final_high = high_summaries_deduped
            
        try:
            final_low = ast.literal_eval(response3)
        except:
            print("Warning: Could not parse low risk response, using original list")
            final_low = low_summaries_deduped
        
    except Exception as e:
        print(f"Error in LLM deduplication: {e}")
        final_high = high_summaries_deduped
        final_low = low_summaries_deduped
    
    return {
        'High': final_high,
        'Low': final_low,
        'counts': {
            'High': len(final_high),
            'Low': len(final_low)
        }
    }

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    
    # Split by ### headers
    sections = fourth_response.split('###')
    
    for section in sections:
        if not section.strip():
            continue
            
        lines = section.split('\n')
        category_name = ""
        bullets = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                # This is likely the category name
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                # This is a bullet point
                bullet_text = line[1:].strip()  # Remove bullet symbol
                if bullet_text:
                    bullets.append(bullet_text)
        
        if category_name and bullets:
            categories_summary[category_name] = bullets
    
    return categories_summary

def create_word_document(pdf_name: str, company_info: str, risk_data: Dict[str, Any],
                        summary_by_categories: Dict[str, List[str]], output_folder: str) -> str:
    """Create a formatted Word document with the analysis results"""
    
    print(f"Creating document for: {pdf_name}")
    print(f"Company info: {company_info}")
    print(f"Risk data keys: {list(risk_data.keys()) if isinstance(risk_data, dict) else 'Not a dict'}")
    
    # Debug the risk data structure
    debug_risk_data_structure(risk_data)
    
    # Create new document
    doc = Document()
    
    # Set document title
    title = doc.add_heading(company_info, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Flag Distribution section
    flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
    flag_dist_heading.runs[0].bold = True
    
    # Create table for flag distribution with proper structure
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Safely extract counts with error handling
    try:
        high_count = risk_data.get('counts', {}).get('High', 0)
        low_count = risk_data.get('counts', {}).get('Low', 0)
        total_count = high_count + low_count
    except Exception as e:
        print(f"Error extracting risk counts: {e}")
        high_count = low_count = total_count = 0
    
    # Populate table cells with error handling
    try:
        # Row 0: High Risk
        table.cell(0, 0).text = 'High Risk'
        table.cell(0, 1).text = str(high_count)
        
        # Row 1: Empty row (or you can add Medium Risk if needed)
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''
        
        # Row 2: Low Risk  
        table.cell(2, 0).text = 'Low Risk'
        table.cell(2, 1).text = str(low_count)
        
        # Row 3: Total Flags
        table.cell(3, 0).text = 'Total Flags'
        table.cell(3, 1).text = str(total_count)
        
        # Make table headers bold with error handling
        for i in [0, 2, 3]:  # Only bold the rows with actual data
            try:
                if table.cell(i, 0).paragraphs:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
            except (IndexError, AttributeError):
                pass  # Skip if cell doesn't have the expected structure
                
    except IndexError as e:
        print(f"Error populating table cells: {e}")
        # Create a simpler table structure
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'High Risk'
        table.cell(0, 1).text = str(high_count)
        table.cell(1, 0).text = 'Low Risk'
        table.cell(1, 1).text = str(low_count)
        table.cell(2, 0).text = 'Total Flags'
        table.cell(2, 1).text = str(total_count)
    
    # Add space
    doc.add_paragraph('')
    
    # Debug: Print what we're about to add to the document
    print(f"DEBUG WORD: About to add {high_count} high risk flags to document")
    print(f"DEBUG WORD: High risk flags: {risk_data.get('High', [])}")
    
    # Add High Risk Flags section only with error handling
    high_risk_flags = risk_data.get('High', [])
    if high_risk_flags and len(high_risk_flags) > 0:
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
        
        for flag in high_risk_flags:
            try:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(str(flag))  # Convert to string to handle any data type
                print(f"DEBUG WORD: Added flag to document: {flag}")
            except Exception as e:
                print(f"Error adding flag to document: {e}")
                continue
    else:
        # If no high risk flags, add a note
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
        no_flags_para = doc.add_paragraph('No high risk flags identified.')
        print("DEBUG WORD: No high risk flags found - added 'No high risk flags identified' message")
    
    # Add horizontal line
    doc.add_paragraph('_' * 50)
    
    # Add Summary section
    summary_heading = doc.add_heading('Summary', level=1)
    summary_heading.runs[0].bold = True
    
    # Add categorized summary with error handling
    if summary_by_categories:
        for category, bullets in summary_by_categories.items():
            if bullets and len(bullets) > 0:  # Only add if there are bullets
                try:
                    # Add category as subheading
                    cat_heading = doc.add_heading(str(category), level=2)
                    cat_heading.runs[0].bold = True
                    
                    # Add bullet points for this category
                    for bullet in bullets:
                        try:
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'
                            p.add_run(str(bullet))  # Convert to string
                        except Exception as e:
                            print(f"Error adding bullet point: {e}")
                            continue
                    
                    # Add space between categories
                    doc.add_paragraph('')
                    
                except Exception as e:
                    print(f"Error processing category {category}: {e}")
                    continue
    else:
        # Add a note if no summary categories found
        doc.add_paragraph('No summary categories found.')
    
    # Save document with error handling
    try:
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
        print(f"Document saved successfully: {doc_path}")
        return doc_path
    except Exception as e:
        print(f"Error saving document: {e}")
        # Try with a simpler filename
        try:
            simple_filename = f"Report_{int(time.time())}.docx"
            simple_path = os.path.join(output_folder, simple_filename)
            doc.save(simple_path)
            print(f"Document saved with alternative name: {simple_path}")
            return simple_path
        except Exception as e2:
            print(f"Failed to save document with alternative name: {e2}")
            raise e2

def evaluate_individual_criteria(llm: AzureOpenAILLM, context: str, fourth_response: str, previous_year_data: str, criteria_name: str, criteria_description: str) -> str:
    """
    Evaluate a single criterion with specific targeted prompts to avoid over-counting
    """
    
    # Define specific prompts for each criterion to be more targeted
    specific_prompts = {
        "debt_increase": f"""
Analyze ONLY debt-related red flags from the summary. Look specifically for:
- Total debt increases
- Long-term debt changes
- Borrowing increases
- Debt restructuring issues

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention debt increases or borrowing increases. Ignore margin issues, revenue declines, or operational problems.

Output Format:
### Debt Increase Risk Classification
* [Only debt increase related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no debt increase red flags found, state: "No debt increase red flags present."
""",
        
        "provisioning": f"""
Analyze ONLY provisioning and write-off related red flags from the summary. Look specifically for:
- Provisioning for bad debts
- Write-offs mentioned
- Impairment charges
- Credit loss provisions

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention provisioning, write-offs, or impairment. Ignore other financial issues.

Output Format:
### Provisioning Risk Classification
* [Only provisioning/write-off related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no provisioning red flags found, state: "No provisioning red flags present."
""",

        "asset_decline": f"""
Analyze ONLY asset value decline related red flags from the summary. Look specifically for:
- Asset value decreases
- Asset impairment
- Fixed asset write-downs
- Investment value declines

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention asset value declines or asset impairments. Ignore debt, revenue, or operational issues.

Output Format:
### Asset Decline Risk Classification
* [Only asset decline related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no asset decline red flags found, state: "No asset decline red flags present."
""",

        "receivable_days": f"""
Analyze ONLY receivable days or accounts receivable related red flags from the summary. Look specifically for:
- Receivable days increase
- Collection period extension
- Accounts receivable growth
- Customer payment delays

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention receivable days, collection periods, or accounts receivable issues.

Output Format:
### Receivable Days Risk Classification
* [Only receivable days related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no receivable days red flags found, state: "No receivable days red flags present."
""",

        "payable_days": f"""
Analyze ONLY payable days or accounts payable related red flags from the summary. Look specifically for:
- Payable days increase
- Payment period extension
- Accounts payable growth
- Supplier payment delays

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention payable days, payment periods, or accounts payable issues.

Output Format:
### Payable Days Risk Classification
* [Only payable days related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no payable days red flags found, state: "No payable days red flags present."
""",

        "debt_ebitda": f"""
Analyze ONLY debt-to-EBITDA ratio related red flags from the summary. Look specifically for:
- Debt/EBITDA ratio mentions
- High leverage concerns
- Coverage ratio issues
- Debt serviceability problems

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention debt-to-EBITDA ratios or leverage ratios.

Output Format:
### Debt EBITDA Risk Classification
* [Only debt/EBITDA ratio related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no debt/EBITDA red flags found, state: "No debt/EBITDA red flags present."
""",

        "revenue_decline": f"""
Analyze ONLY revenue decline related red flags from the summary. Look specifically for:
- Revenue decreases
- Sales decline
- Top-line reduction
- Income statement revenue issues

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention revenue declines or sales decreases. Ignore margin compression or cost issues.

Output Format:
### Revenue Decline Risk Classification
* [Only revenue decline related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no revenue decline red flags found, state: "No revenue decline red flags present."
""",

        "onetime_expenses": f"""
Analyze ONLY one-time expenses or exceptional items related red flags from the summary. Look specifically for:
- One-time charges
- Exceptional expenses
- Non-recurring costs
- Special items or extraordinary expenses

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention one-time, exceptional, or non-recurring expenses.

Output Format:
### One-time Expenses Risk Classification
* [Only one-time expenses related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no one-time expenses red flags found, state: "No one-time expenses red flags present."
""",

        "margin_decline": f"""
Analyze ONLY margin decline related red flags from the summary. Look specifically for:
- Gross margin decline
- Operating margin compression
- EBITDA margin reduction
- Profitability margin issues

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention margin declines or margin compression. Ignore revenue or cost issues separately.

Output Format:
### Margin Decline Risk Classification
* [Only margin decline related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no margin decline red flags found, state: "No margin decline red flags present."
""",

        "cash_balance": f"""
Analyze ONLY cash balance decline related red flags from the summary. Look specifically for:
- Cash balance decreases
- Cash flow negative trends
- Liquidity reduction
- Cash position deterioration

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention cash balance declines or cash position issues.

Output Format:
### Cash Balance Risk Classification
* [Only cash balance decline related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no cash balance red flags found, state: "No cash balance red flags present."
""",

        "short_term_debt": f"""
Analyze ONLY short-term debt or current liabilities related red flags from the summary. Look specifically for:
- Short-term debt increases
- Current liabilities growth
- Working capital issues
- Short-term borrowing increases

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention short-term debt or current liabilities increases.

Output Format:
### Short-term Debt Risk Classification
* [Only short-term debt related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no short-term debt red flags found, state: "No short-term debt red flags present."
""",

        "management_issues": f"""
Analyze ONLY management and leadership related red flags from the summary. Look specifically for:
- Management turnover
- Key personnel departures
- Leadership changes
- Governance issues
- Strategic execution problems

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention management changes, leadership issues, or governance problems.

Output Format:
### Management Issues Risk Classification
* [Only management/leadership related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no management issues red flags found, state: "No management issues red flags present."
""",

        "regulatory_compliance": f"""
Analyze ONLY regulatory compliance related red flags from the summary. Look specifically for:
- Regulatory violations
- Compliance issues
- Legal problems
- Regulatory warnings or penalties

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention regulatory or compliance issues.

Output Format:
### Regulatory Compliance Risk Classification
* [Only regulatory compliance related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no regulatory compliance red flags found, state: "No regulatory compliance red flags present."
""",

        "market_competition": f"""
Analyze ONLY market competition related red flags from the summary. Look specifically for:
- Competitive pressure
- Market share loss
- New competitors entering
- Industry competition intensifying

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention competitive pressures or market share issues.

Output Format:
### Market Competition Risk Classification
* [Only market competition related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no market competition red flags found, state: "No market competition red flags present."
""",

        "operational_disruptions": f"""
Analyze ONLY operational disruption related red flags from the summary. Look specifically for:
- Supply chain disruptions
- Production issues
- IT system failures
- Operational inefficiencies
- Infrastructure problems

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

From the summary, identify ONLY red flags that specifically mention operational disruptions or system failures.

Output Format:
### Operational Disruptions Risk Classification
* [Only operational disruption related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no operational disruption red flags found, state: "No operational disruption red flags present."
"""
    }
    
    dedup_instruction = (
        "REMOVE ALL DUPLICATES: Report each distinct high risk flag only once (do not include repetitive or paraphrased statements). "
        "If multiple flags describe the same underlying issue, only include the most representative one."
    )
    
    # Get the specific prompt for this criterion
    specific_prompt = specific_prompts.get(criteria_name, f"""
Analyze red flags related to {criteria_name} from the summary.

Criterion: {criteria_description}

Previous Year Data: {previous_year_data}

{dedup_instruction}

Output Format:
### {criteria_name.title()} Risk Classification
* [Only {criteria_name} related red flag if found]
  - High: yes/no
  - Low: yes/no
  - Not_Applicable: yes/no

If no {criteria_name} red flags found, state: "No {criteria_name} red flags present."
""")

    full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{context}

Summary from previous iteration:
{fourth_response}

Based on the above context and analysis, please answer: {specific_prompt}

Answer:"""
    
    return llm._call(full_prompt, max_tokens=1000)

def process_single_pdf_five_iterations(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                                     output_folder: str = "results", 
                                     api_key: str = None, azure_endpoint: str = None, 
                                     api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Process a single PDF through the 5-iteration pipeline with individual criteria evaluation using Azure OpenAI
    """
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Get PDF name without extension
    pdf_name = Path(pdf_path).stem
    
    print(f"\nProcessing: {pdf_name}")
    print("=" * 50)
    
    try:
        # ITERATION 1: Initial red flag identification
        print("Running 1st iteration - Initial Analysis...")
        pipeline_1st = LlamaQueryPipeline(
            pdf_path=pdf_path,
            queries_csv_path=queries_csv_path,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        # Run 1st iteration
        first_results_df = pipeline_1st.query_llama(maintain_conversation=True, enable_chaining=False)
        
        # Get first response for chaining
        first_response = first_results_df.iloc[0]['response']
        
        # ITERATION 2: Deduplication and cleanup
        print("Running 2nd iteration - Deduplication...")
        second_prompt = """Remove the duplicates from the above context. Also if the Original Quote and Keyword identifies is same remove them.Do not lose data if duplicates are not found."""
        
        second_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{pipeline_1st.docs[0]["context"]}

Previous Analysis: {first_response}

Based on the above analysis and the original context, please answer: {second_prompt}

Answer:"""
        
        second_response = pipeline_1st.llm._call(second_full_prompt, max_tokens=4000)
        
        # ITERATION 3: Categorization of red flags
        print("Running 3rd iteration - Categorization...")
        third_prompt = """You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.

- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction.
- Regulatory Issues: Concerns related Compliance with laws, regulations.
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape.
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.

While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.

**Output Format**:
### Balance Sheet Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### Liquidity Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### Management and Strategy related Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### Regulatory Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### Industry and Market Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

### Operational Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]

Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
        
        third_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{pipeline_1st.docs[0]["context"]}

Previous Analysis: {second_response}

Based on the above analysis and the original context, please answer: {third_prompt}

Answer:"""
        
        third_response = pipeline_1st.llm._call(third_full_prompt, max_tokens=4000)
        
        # ITERATION 4: Detailed summary generation
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:

1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
3. **Focus on factual content**: Base the summary solely on the information associated with each red flag, without introducing external opinions or assumptions
4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
8. **Summary should be factual**: Avoid any subjective interpretations or opinions
9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points

Format the output exactly like this example:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
* [Summary of red flag 3 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
* [Summary of red flag 3 with specific data points and factual information]

### Liquidity Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### Management and Strategy related Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### Regulatory Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### Industry and Market Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### Operational Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
        
        fourth_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{pipeline_1st.docs[0]["context"]}

Previous Analysis: {third_response}

Based on the above analysis and the original context, please answer: {fourth_prompt}

Answer:"""
        
        fourth_response = pipeline_1st.llm._call(fourth_full_prompt, max_tokens=4000)
        
        # ITERATION 5: Individual Risk Classification for 15 Criteria
        print("Running 5th iteration - Individual Criteria Risk Classification...")
        
        criteria_definitions = {
            "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
            "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
            "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
            "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
            "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues ",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues "
        }

        # Process each criterion individually
        fifth_results = []
        all_criteria_responses = ""
        
        for criteria_name, criteria_description in criteria_definitions.items():
            print(f"Evaluating criterion: {criteria_name}")
            
            # Get individual criterion response
            criterion_response = evaluate_individual_criteria(
                llm=pipeline_1st.llm,
                context=pipeline_1st.docs[0]["context"],
                fourth_response=fourth_response,
                previous_year_data=previous_year_data,
                criteria_name=criteria_name,
                criteria_description=criteria_description
            )
            
            # Store individual results
            fifth_results.append({
                "criteria": criteria_name,
                "description": criteria_description,
                "classification_response": criterion_response
            })
            
            # Combine all responses for parsing
            all_criteria_responses += f"\n{criterion_response}\n"
            
            # Small delay to avoid overwhelming the API
            time.sleep(1)
        
        # Extract company information from first page
        print("Extracting company information...")
        company_info = extract_company_info_from_pdf(pdf_path, pipeline_1st.llm)
        risk_data = parse_risk_classification_response(all_criteria_responses)
        
        summary_by_categories = parse_summary_by_categories(fourth_response)
        
        # Create Word document
        print("Creating Word document...")
        word_doc_path = create_word_document(
            pdf_name=pdf_name,
            company_info=company_info,
            risk_data=risk_data,
            summary_by_categories=summary_by_categories,
            output_folder=output_folder
        )
        
        # Save all results together
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        all_results = pd.DataFrame({
            "pdf_name": [pdf_name, pdf_name, pdf_name, pdf_name, pdf_name],
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Deduplication",
                "Categorization",
                "Summary Generation",
                "Individual Criteria Risk Classification"
            ],
            "prompt": [
                first_results_df.iloc[0]['query'],  # Original query from 1st iteration
                second_prompt,
                third_prompt,
                fourth_prompt,
                "Individual Risk Classification for 15 criteria"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                all_criteria_responses
            ],
            "timestamp": [timestamp, timestamp, timestamp, timestamp, timestamp]
        })
        
        # Save complete results
        complete_output_file = os.path.join(output_folder, f"{pdf_name}_complete_5iteration_pipeline_results.csv")
        all_results.to_csv(complete_output_file, index=False)
        
        # Save individual risk classification results with detailed breakdown
        detailed_risk_df = pd.DataFrame(fifth_results)
        detailed_risk_output_file = os.path.join(output_folder, f"{pdf_name}_detailed_risk_classification.csv")
        detailed_risk_df.to_csv(detailed_risk_output_file, index=False)

        print(f"Complete 5-iteration pipeline finished for {pdf_name}!")
        print(f"CSV Results saved to: {complete_output_file}")
        print(f"Detailed risk classification saved to: {detailed_risk_output_file}")
        print(f"Word document saved to: {word_doc_path}")

        return all_results
        
    except Exception as e:
        print(f"Error processing {pdf_name}: {str(e)}")
        # Save error log
        error_df = pd.DataFrame({
            "pdf_name": [pdf_name],
            "error": [str(e)],
            "timestamp": [time.strftime("%Y%m%d_%H%M%S")]
        })
        error_file = os.path.join(output_folder, f"{pdf_name}_error_log.csv")
        error_df.to_csv(error_file, index=False)
        return None

def run_multiple_pdfs_five_iterations_pipeline(pdf_folder_path: str, queries_csv_path: str, previous_year_data: str, 
                                              output_folder: str = "results",
                                              api_key: str = None, azure_endpoint: str = None, 
                                              api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Process multiple PDFs from a folder through the 5-iteration pipeline with individual criteria evaluation using Azure OpenAI
    
    Args:
        pdf_folder_path: Path to folder containing PDF files
        queries_csv_path: Path to CSV/Excel file containing queries
        previous_year_data: String containing previous year financial data
        output_folder: Path to output folder for results
        api_key: Azure OpenAI API key
        azure_endpoint: Azure OpenAI endpoint
        api_version: API version
        deployment_name: Deployment name
    """
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Get all PDF files from the folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return
    
    print(f"Found {len(pdf_files)} PDF files to process:")
    for pdf_file in pdf_files:
        print(f"  - {os.path.basename(pdf_file)}")
    
    print(f"\nStarting batch processing with 5 iterations and 15 individual criteria using Azure OpenAI GPT-4.1-mini...")
    print(f"
