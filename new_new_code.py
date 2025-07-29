import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
 
def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class AzureOpenAILLM:
    """Azure OpenAI GPT-4.1-mini LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1-mini"):
        """
        Initialize Azure OpenAI client
        
        Args:
            api_key: Your Azure OpenAI API key
            azure_endpoint: Your Azure OpenAI endpoint URL
            api_version: API version (e.g., "2024-02-01")
            deployment_name: Your deployment name (default: "gpt-4.1-mini")
        """
        self.deployment_name = deployment_name
        
        # Create httpx client with SSL verification disabled (if needed)
        httpx_client = httpx.Client(verify=False)
        
        # Initialize Azure OpenAI client
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            # Extract the response text
            response_text = response.choices[0].message.content
            
            # Log token usage
            if hasattr(response, 'usage'):
                logger.info(f"Tokens used - Prompt: {response.usage.prompt_tokens}, "
                          f"Completion: {response.usage.completion_tokens}, "
                          f"Total: {response.usage.total_tokens}")
            
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        start_time = time.time()
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            # Explicitly close the document to free memory
            doc.close()
            logger.info(f"PDF text extraction took {time.time() - start_time:.2f} seconds")
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        # Merge all pages into single context
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]
 
class LlamaQueryPipeline:
    """Main pipeline class for querying PDF content with Azure OpenAI GPT-4.1-mini"""
   
    def __init__(self, pdf_path: str, queries_csv_path: str = None, 
                 api_key: str = None, azure_endpoint: str = None, 
                 api_version: str = None, deployment_name: str = "gpt-4.1-mini",
                 previous_results_path: str = None):
        """
        Initialize the pipeline with Azure OpenAI
        
        Args:
            pdf_path: Path to PDF file
            queries_csv_path: Path to queries CSV/Excel file
            api_key: Azure OpenAI API key
            azure_endpoint: Azure OpenAI endpoint
            api_version: API version
            deployment_name: Deployment name
            previous_results_path: Path to previous results
        """
        # Initialize Azure OpenAI LLM
        self.llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        self.docs = mergeDocs(pdf_path, split_pages=False)
       
        # Load queries from Excel or CSV (only if provided)
        if queries_csv_path:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            self.queries = queries_df["prompt"].tolist()
        else:
            self.queries = []
           
        self.pdf_path = pdf_path
        self.pdf_name = Path(pdf_path).stem  # Get filename without extension
       
        # Load previous results if provided
        self.previous_results = None
        if previous_results_path and os.path.exists(previous_results_path):
            self.previous_results = pd.read_csv(previous_results_path)
   
    def query_llama_with_chaining(self, new_queries_csv_path: str, iteration_number: int = 2) -> pd.DataFrame:
        """Query the Azure OpenAI API using previous results for chaining"""
        if self.previous_results is None:
            raise ValueError("No previous results loaded. Please provide previous_results_path in __init__")
       
        # Load new queries
        if new_queries_csv_path.endswith('.xlsx'):
            new_queries_df = pd.read_excel(new_queries_csv_path)
        else:
            new_queries_df = pd.read_csv(new_queries_csv_path)
       
        new_queries = new_queries_df["prompt"].tolist()
       
        sys_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{self.docs[0]["context"]}
 
"""
       
        results = []
       
        # Process each new query with corresponding previous response
        for i, new_query in enumerate(new_queries):
            start = time.perf_counter()
           
            try:
                # Get the corresponding previous response (if available)
                if i < len(self.previous_results):
                    # Check if it's 3rd iteration (chained results) or 2nd iteration (original results)
                    if 'chained_response' in self.previous_results.columns:
                        previous_response = self.previous_results.iloc[i]['chained_response']
                    else:
                        previous_response = self.previous_results.iloc[i]['response']
                   
                    # Create chained prompt: new query + previous response
                    chained_prompt = f"""Previous Analysis: {previous_response}
 
Based on the above analysis and the original context, please answer: {new_query}
 
Answer:"""
                else:
                    # If no previous response available, use regular prompt
                    chained_prompt = f"""Question: {new_query}
Answer:"""
               
                full_prompt = f"{sys_prompt}\n{chained_prompt}"
               
                # Get response from Azure OpenAI
                response_text = self.llm._call(full_prompt, max_tokens=4000)
                end = time.perf_counter()
               
                # Calculate token approximations (more accurate with OpenAI response)
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
               
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
               
            except Exception as e:
                end = time.perf_counter()
               
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
           
            results.append(usage)
       
        return pd.DataFrame(results)
 
    def query_llama(self, maintain_conversation: bool = True, enable_chaining: bool = False) -> pd.DataFrame:
        """Query the Azure OpenAI API for a list of queries using the provided context"""
        sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.
 
COMPLETE DOCUMENT TO ANALYZE:
{self.docs[0]["context"]}
 
Your task is to analyze the ENTIRE document above and identify ALL potential red flags.
 
CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
- Do not miss any sections or concerning statements
 
EXAMPLE FORMAT:
1. The potential red flag you observed - Debt reduction is lower than expected  
Original Quote:  
"Vikrant Kashyap: Have you -- are you able to reduce any debt in quarter one in Middle East and India?  
Ramesh Kalyanaraman: So India, we are not reduced, but the cash balance has been increased to around INR75 crores, but we have not reduced any debt in Q1 in India and Middle East because Middle East we have not converted any showroom in Q1." (Page 9)  
 
2. The potential red flag you observed - Margin pressure/Competition intensifying/Cost inflation  
Original Quote:  
"Pulkit Singhal: Thank you for the opportunity and congrats on the good set of performance. Just the first question is really on the margins, which seems to have been elusive. I mean looking at the company for the last two years, we seem to be executing quite well in terms of store expansions and revenue growth and clearly delivering higher than expected there. But margin expansion just has been completely elusive.  
And I find it surprising your comment that while growing at 30% growth and 12% SSSG, I mean, which is quite healthy, you're still talking about high competitive intensity kind of quite contradictory that with such high growth rates, we have to invest so high. So can you talk about this a bit more? I mean we don't expect with lower revenue growth rates that you would not have to invest in the business. And it's only during a higher revenue growth that you expect margin expansion.  
Ramesh Kalyanaraman: Yes. So you're right. We are -- meaning somewhere we have missed out on the operating leverage for advertisements that's why I told you that even Q1, it was a miss. And regarding competition, I will tell you where in new markets, where we assume that we will not spend too much because the brand is already aware and the location is the only thing which has to be communicated.  
When you see the local players, regional players or the micro market players there becoming extremely active because of our showroom launch then we will have to increase the noise level there. Otherwise, we will lose our market share. And existing local players they increase their activity around our launch time. So that is where we also put more money so that we don't end up losing the market share or we don't end up taking out lesser from the competition." (Page 12-13)  
 
3. The potential red flag you observed - Debt high/Increase in borrowing cost  
Original Quote:  
"Vikrant Kashyap: My question is how are you going to address this? Because if you continue to grow at a higher level, but bottom line is not expanding related to the top line, it will going to impact your overall performance? So, what are the steps you are taking to improve the bottom line in the businesses  
Sanjay Raghuraman: Finance costs will be taken care because we told you when we convert stores, that money is going to reduce our debt. Okay. And again, FOCO, when you do FOCO showrooms, the margins will come down. And surely, that will have an impact on the gross margin. Okay. And interest, if you look at actually, the interest rates have been going up last year. So next year, that will be the base, right? So then again, we will not have this kind of issue is what we feel. So interest rates have been going up over the past year, one year, in that region.  
And we are also beginning to repay loans now because of conversion. So all put together, interest part will be taken care but other area where FOCO showrooms will surely reduce our margin. We cannot have the own store margin. So that should be the way we should look at it." (Page 9)
 
Continue this exact format for ALL red flags identified throughout the document.
 
"""
       
        prompt_template = """Question: {query}
 
Analyze the complete document and provide ALL red flags in the exact numbered format specified above. Be thorough and comprehensive - cover the entire document.
 
Answer:"""
       
        conversation_history = ""
        results = []
        previous_output = ""  # For prompt chaining
       
        for i, query in enumerate(self.queries, 1):
            start = time.perf_counter()
           
            try:
                if enable_chaining and i > 1 and previous_output:
                    chained_query = f"{query}\n\nPrevious context: {previous_output}"
                else:
                    chained_query = query
               
                if maintain_conversation and conversation_history:
                    full_prompt = f"{sys_prompt}\n{conversation_history}\n{prompt_template.format(query=chained_query)}"
                else:
                    full_prompt = f"{sys_prompt}\n{prompt_template.format(query=chained_query)}"
               
                # Use higher max_tokens for comprehensive analysis
                response_text = self.llm._call(full_prompt, max_tokens=4000)
                end = time.perf_counter()
               
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
               
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": chained_query if enable_chaining else query,
                    "response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
               
                # Update conversation history for next iteration
                if maintain_conversation:
                    conversation_history += f"\nQuestion: {chained_query}\nAnswer: {response_text}\n"
               
                # Store output for next chaining iteration
                if enable_chaining:
                    previous_output = response_text
               
            except Exception as e:
                end = time.perf_counter()
               
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": query,
                    "response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
               
                if enable_chaining:
                    previous_output = ""
           
            results.append(usage)
       
        return pd.DataFrame(results)
   
    def save_results(self, results_df: pd.DataFrame, output_path: str = None):
        """Save results to CSV file"""
        if output_path is None:
            # Generate filename based on PDF name and timestamp
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_path = f"azure_openai_query_results_{self.pdf_name}_{timestamp}.csv"
       
        results_df.to_csv(output_path, index=False)
        return output_path
 
def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        # Limit text to first 2000 characters to avoid token limits
        first_page_text = first_page_text[:2000]
       
        prompt = f"""
You are a financial document analyst. Extract the company name, quarter, and financial year from the following text which is from the first page of an earnings call transcript or financial document.
 
Text from first page:
{first_page_text}
 
Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)
3. Financial Year (FY23/FY24/FY25 etc.)
 
Format your response as: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25
 
If you cannot find clear information, make the best estimate based on available data.
 
Response:"""
       
        response = llm._call(prompt, max_tokens=200)
        # Clean the response to get just the formatted string
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        # Fallback - return first non-empty line
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def extract_unique_flags_from_second_iteration(second_response: str, llm: AzureOpenAILLM) -> List[str]:
    """
    NEW METHOD: Extract unique flags from 2nd iteration response
    """
    prompt = f"""
Analyze the following red flags analysis and extract a clean list of unique red flags.

Red Flags Analysis:
{second_response}

Your task:
1. Identify each unique red flag mentioned in the analysis
2. Extract only the brief description of each red flag (not the full quotes)
3. Remove any duplicates or similar flags
4. Present each flag as a concise, clear statement

Output format - return ONLY a Python list format:
["flag description 1", "flag description 2", "flag description 3", ...]

Example output:
["Debt reduction lower than expected", "Margin pressure due to competition", "High borrowing costs", "Revenue decline", "Cash flow issues"]

Return only the list, no other text or explanation.
"""
    
    response = llm._call(prompt, max_tokens=1000)
    
    try:
        # Try to parse as Python list
        unique_flags = ast.literal_eval(response.strip())
        if isinstance(unique_flags, list):
            return unique_flags
    except:
        # If parsing fails, try to extract manually
        lines = response.strip().split('\n')
        unique_flags = []
        for line in lines:
            line = line.strip()
            if line.startswith('"') and line.endswith('"'):
                unique_flags.append(line[1:-1])
            elif line.startswith("'") and line.endswith("'"):
                unique_flags.append(line[1:-1])
        
        if unique_flags:
            return unique_flags
    
    # Fallback: return empty list
    return []

def classify_flag_against_criteria(flag: str, criteria_definitions: Dict[str, str], previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, str]:
    """
    NEW METHOD: Classify a single flag against all 15 criteria
    """
    criteria_list = "\n".join([f"{i+1}. {name}: {desc}" for i, (name, desc) in enumerate(criteria_definitions.items())])
    
    prompt = f"""
You are a financial risk analyst. Your task is to classify the following red flag against the given criteria.

RED FLAG TO CLASSIFY: "{flag}"

CRITERIA DEFINITIONS:
{criteria_list}

PREVIOUS YEAR DATA:
{previous_year_data}

INSTRUCTIONS:
1. Check if this red flag matches ANY of the 15 criteria above
2. If it matches a criterion, classify it as either "High" or "Low" risk based on the criterion definition
3. If it doesn't match any specific criterion, classify it as "Low" risk by default
4. Return the result in the specified format

OUTPUT FORMAT:
Matched_Criteria: [criteria_name if matched, or "None" if no match]
Risk_Level: [High/Low]
Reasoning: [Brief explanation of why this classification was chosen]

Example outputs:
Matched_Criteria: debt_increase
Risk_Level: High
Reasoning: Flag indicates debt increase of 40% which exceeds the 30% threshold for high risk

OR

Matched_Criteria: None
Risk_Level: Low
Reasoning: Flag does not match any specific criteria, defaulting to low risk
"""
    
    response = llm._call(prompt, max_tokens=500)
    
    # Parse the response
    result = {
        'matched_criteria': 'None',
        'risk_level': 'Low',
        'reasoning': 'No specific reasoning provided'
    }
    
    lines = response.strip().split('\n')
    for line in lines:
        if line.startswith('Matched_Criteria:'):
            result['matched_criteria'] = line.split(':', 1)[1].strip()
        elif line.startswith('Risk_Level:'):
            result['risk_level'] = line.split(':', 1)[1].strip()
        elif line.startswith('Reasoning:'):
            result['reasoning'] = line.split(':', 1)[1].strip()
    
    return result

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
   
    # Split by ### headers
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                # This is likely the category name
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                # This is a bullet point
                bullet_text = line[1:].strip()  # Remove bullet symbol
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary
 
def create_word_document(pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                        high_risk_flags: List[str], summary_by_categories: Dict[str, List[str]], 
                        output_folder: str) -> str:
    """Create a formatted Word document with the analysis results"""
   
    # Create new document
    doc = Document()
   
    # Set document title
    title = doc.add_heading(company_info, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    # Add Flag Distribution section
    flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
    flag_dist_heading.runs[0].bold = True
   
    # Create table for flag distribution
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
   
    # Add table headers and data using counts
    high_count = risk_counts['High']
    low_count = risk_counts['Low']
    total_count = high_count + low_count
   
    table.cell(0, 0).text = 'High Risk'
    table.cell(0, 1).text = str(high_count)
    table.cell(2, 0).text = 'Low Risk'
    table.cell(2, 1).text = str(low_count)
    table.cell(3, 0).text = 'Total Flags'
    table.cell(3, 1).text = str(total_count)
   
    # Make table headers bold
    for i in range(4):
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
   
    # Add space
    doc.add_paragraph('')
   
    # Add High Risk Flags section
    if high_risk_flags:
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
       
        for flag in high_risk_flags:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(flag)
    else:
        # If no high risk flags, add a note
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
        no_flags_para = doc.add_paragraph('No high risk flags identified.')
   
    # Add horizontal line
    doc.add_paragraph('_' * 50)
   
    # Add Summary section (4th iteration results)
    summary_heading = doc.add_heading('Summary', level=1)
    summary_heading.runs[0].bold = True
   
    # Add categorized summary
    for category, bullets in summary_by_categories.items():
        if bullets:  # Only add if there are bullets
            # Add category as subheading
            cat_heading = doc.add_heading(category, level=2)
            cat_heading.runs[0].bold = True
           
            # Add bullet points for this category
            for bullet in bullets:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(bullet)
           
            # Add space between categories
            doc.add_paragraph('')
   
    # Save document
    doc_filename = f"{pdf_name}_Report.docx"
    doc_path = os.path.join(output_folder, doc_filename)
    doc.save(doc_path)
   
    return doc_path

def process_single_pdf_five_iterations(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                                     output_folder: str = "results", 
                                     api_key: str = None, azure_endpoint: str = None, 
                                     api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Process a single PDF through the NEW 5-iteration pipeline using Azure OpenAI
    """
   
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
   
    # Get PDF name without extension
    pdf_name = Path(pdf_path).stem
   
    print(f"\nProcessing: {pdf_name}")
    print("=" * 50)
   
    try:
        # ITERATION 1: Initial red flag identification (NO CHANGES)
        print("Running 1st iteration - Initial Analysis...")
        pipeline_1st = LlamaQueryPipeline(
            pdf_path=pdf_path,
            queries_csv_path=queries_csv_path,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
       
        # Run 1st iteration
        first_results_df = pipeline_1st.query_llama(maintain_conversation=True, enable_chaining=False)
       
        # Get first response for chaining
        first_response = first_results_df.iloc[0]['response']
       
        # ITERATION 2: Deduplication and cleanup (NO CHANGES)
        print("Running 2nd iteration - Deduplication...")
        second_prompt = """Remove the duplicates from the above context. Also if the Original Quote and Keyword identifies is same remove them.Do not lose data if duplicates are not found."""
       
        second_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{pipeline_1st.docs[0]["context"]}
 
Previous Analysis: {first_response}
 
Based on the above analysis and the original context, please answer: {second_prompt}
 
Answer:"""
       
        second_response = pipeline_1st.llm._call(second_full_prompt, max_tokens=4000)
       
        # ITERATION 3: Categorization of red flags (NO CHANGES)
        print("Running 3rd iteration - Categorization...")
        third_prompt = """You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.
 
- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction.
- Regulatory Issues: Concerns related Compliance with laws, regulations.
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape.
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.
 
While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.
 
**Output Format**:
### Balance Sheet Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### P&L (Income Statement) Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### Liquidity Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### Management and Strategy related Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### Regulatory Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### Industry and Market Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
### Operational Issues
- [Red flag 1 with original quote]
- [Red flag 2 with original quote]
 
Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
       
        third_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{pipeline_1st.docs[0]["context"]}
 
Previous Analysis: {second_response}
 
Based on the above analysis and the original context, please answer: {third_prompt}
 
Answer:"""
       
        third_response = pipeline_1st.llm._call(third_full_prompt, max_tokens=4000)
       
        # ITERATION 4: Detailed summary generation (NO CHANGES)
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:
 
1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
3. **Focus on factual content**: Base the summary solely on the information associated with each red flag, without introducing external opinions or assumptions
4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
8. **Summary should be factual**: Avoid any subjective interpretations or opinions
9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points
 
Format the output exactly like this example:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
* [Summary of red flag 3 with specific data points and factual information]
 
### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
* [Summary of red flag 3 with specific data points and factual information]
 
### Liquidity Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### Management and Strategy related Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### Regulatory Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### Industry and Market Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
### Operational Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]
 
Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
       
        fourth_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{pipeline_1st.docs[0]["context"]}
 
Previous Analysis: {third_response}
 
Based on the above analysis and the original context, please answer: {fourth_prompt}
 
Answer:"""
       
        fourth_response = pipeline_1st.llm._call(fourth_full_prompt, max_tokens=4000)
       
        # NEW ITERATION 5: Extract unique flags and classify against 15 criteria
        print("Running 5th iteration - NEW METHOD: Unique Flags Classification...")
        
        # Step 1: Extract unique flags from 2nd iteration
        print("  Step 1: Extracting unique flags from 2nd iteration...")
        unique_flags = extract_unique_flags_from_second_iteration(second_response, pipeline_1st.llm)
        print(f"  Found {len(unique_flags)} unique flags")
        
        # Step 2: Define the 15 criteria
        criteria_definitions = {
            "debt_increase": "High: Debt increase by >=30% compared to previous reported balance sheet number; Low: Debt increase is less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by >=30% compared to previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days increase by >=30% compared to previous reported balance sheet number; Low: receivable days increase is less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days increase by >=30% compared to previous reported balance sheet number; Low: payable days increase is less than 30% compared to previous reported balance sheet number",
            "debt_ebitda": "High: Debt/EBITDA >= 3x; Low: Debt/EBITDA < 3x",
            "revenue_decline": "High: revenue or profitability falls by >=25% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
            "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_debt": "High: Short-term debt or current liabilities increase by >=30% compared to previous reported balance sheet number; Low: Short-term debt or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "management_issues": "High: Any management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Low: No management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates",
            "regulatory_compliance": "High: if found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: if there is a no clear concern for the company basis the discussion on the regulatory issues",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
        }
        
        # Step 3: Classify each unique flag against the 15 criteria
        print("  Step 2: Classifying each flag against 15 criteria...")
        classification_results = []
        high_risk_flags = []
        low_risk_flags = []
        
        for i, flag in enumerate(unique_flags, 1):
            print(f"    Classifying flag {i}/{len(unique_flags)}: {flag[:50]}...")
            
            classification = classify_flag_against_criteria(
                flag=flag,
                criteria_definitions=criteria_definitions,
                previous_year_data=previous_year_data,
                llm=pipeline_1st.llm
            )
            
            classification_results.append({
                'flag': flag,
                'matched_criteria': classification['matched_criteria'],
                'risk_level': classification['risk_level'],
                'reasoning': classification['reasoning']
            })
            
            # Categorize flags by risk level
            if classification['risk_level'].lower() == 'high':
                high_risk_flags.append(flag)
            else:
                low_risk_flags.append(flag)
            
            # Small delay to avoid overwhelming the API
            time.sleep(0.5)
        
        # Step 4: Calculate counts
        risk_counts = {
            'High': len(high_risk_flags),
            'Low': len(low_risk_flags),
            'Total': len(unique_flags)
        }
        
        print(f"  Classification complete: {risk_counts['High']} High, {risk_counts['Low']} Low, {risk_counts['Total']} Total")
        
        # Extract company information from first page
        print("Extracting company information...")
        company_info = extract_company_info_from_pdf(pdf_path, pipeline_1st.llm)
        
        # Parse 4th iteration summary for Word document
        summary_by_categories = parse_summary_by_categories(fourth_response)
       
        # Create Word document with new structure
        print("Creating Word document...")
        word_doc_path = create_word_document(
            pdf_name=pdf_name,
            company_info=company_info,
            risk_counts=risk_counts,
            high_risk_flags=high_risk_flags,
            summary_by_categories=summary_by_categories,
            output_folder=output_folder
        )
       
        # Save all results together
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        all_results = pd.DataFrame({
            "pdf_name": [pdf_name, pdf_name, pdf_name, pdf_name, pdf_name],
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Deduplication",
                "Categorization",
                "Summary Generation",
                "NEW: Unique Flags Classification"
            ],
            "prompt": [
                first_results_df.iloc[0]['query'],  # Original query from 1st iteration
                second_prompt,
                third_prompt,
                fourth_prompt,
                f"NEW METHOD: Extracted {len(unique_flags)} unique flags and classified against 15 criteria"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Classification Results: {risk_counts['High']} High Risk, {risk_counts['Low']} Low Risk flags"
            ],
            "timestamp": [timestamp, timestamp, timestamp, timestamp, timestamp]
        })
       
        # Save complete results
        complete_output_file = os.path.join(output_folder, f"{pdf_name}_complete_5iteration_pipeline_results.csv")
        all_results.to_csv(complete_output_file, index=False)
       
        # Save detailed classification results
        detailed_classification_df = pd.DataFrame(classification_results)
        detailed_classification_output_file = os.path.join(output_folder, f"{pdf_name}_detailed_flag_classification.csv")
        detailed_classification_df.to_csv(detailed_classification_output_file, index=False)

        print(f"NEW 5-iteration pipeline finished for {pdf_name}!")
        print(f"CSV Results saved to: {complete_output_file}")
        print(f"Detailed classification saved to: {detailed_classification_output_file}")
        print(f"Word document saved to: {word_doc_path}")
        print(f"Risk Distribution: {risk_counts['High']} High, {risk_counts['Low']} Low, {risk_counts['Total']} Total")

        return all_results
       
    except Exception as e:
        print(f"Error processing {pdf_name}: {str(e)}")
        # Save error log
        error_df = pd.DataFrame({
            "pdf_name": [pdf_name],
            "error": [str(e)],
            "timestamp": [time.strftime("%Y%m%d_%H%M%S")]
        })
        error_file = os.path.join(output_folder, f"{pdf_name}_error_log.csv")
        error_df.to_csv(error_file, index=False)
        return None
 
def run_multiple_pdfs_five_iterations_pipeline(pdf_folder_path: str, queries_csv_path: str, previous_year_data: str, 
                                              output_folder: str = "results",
                                              api_key: str = None, azure_endpoint: str = None, 
                                              api_version: str = None, deployment_name: str = "gpt-4.1-mini"):
    """
    Process multiple PDFs from a folder through the NEW 5-iteration pipeline using Azure OpenAI
   
    Args:
        pdf_folder_path: Path to folder containing PDF files
        queries_csv_path: Path to CSV/Excel file containing queries
        previous_year_data: String containing previous year financial data
        output_folder: Path to output folder for results
        api_key: Azure OpenAI API key
        azure_endpoint: Azure OpenAI endpoint
        api_version: API version
        deployment_name: Deployment name
    """
   
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
   
    # Get all PDF files from the folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
   
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return
   
    print(f"Found {len(pdf_files)} PDF files to process:")
    for pdf_file in pdf_files:
        print(f"  - {os.path.basename(pdf_file)}")
   
    print(f"\nStarting NEW batch processing with 5 iterations using unique flags classification...")
    print(f"Output folder: {output_folder}")
    print("=" * 60)
   
    # Process each PDF
    successful_processing = []
    failed_processing = []
   
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {os.path.basename(pdf_file)}")
       
        try:
            result = process_single_pdf_five_iterations(
                pdf_path=pdf_file,
                queries_csv_path=queries_csv_path,
                previous_year_data=previous_year_data,
                output_folder=output_folder,
                api_key=api_key,
                azure_endpoint=azure_endpoint,
                api_version=api_version,
                deployment_name=deployment_name
            )
           
            if result is not None:
                successful_processing.append(os.path.basename(pdf_file))
            else:
                failed_processing.append(os.path.basename(pdf_file))
               
        except Exception as e:
            print(f"Failed to process {os.path.basename(pdf_file)}: {str(e)}")
            failed_processing.append(os.path.basename(pdf_file))
   
    print("\n" + "=" * 60)
    print("BATCH PROCESSING COMPLETE")
    print("=" * 60)
    
    if successful_processing:
        print(f"\nSuccessfully processed ({len(successful_processing)}):")
        for file in successful_processing:
            print(f"  ✓ {file}")
   
    if failed_processing:
        print(f"\nFailed to process ({len(failed_processing)}):")
        for file in failed_processing:
            print(f"  ✗ {file}")

def main_batch_processing_five_iterations():
    """
    Main function to run the NEW batch processing with Azure OpenAI GPT-4.1-mini
    Configure your Azure OpenAI credentials here
    """
    
    # Configuration
    pdf_folder_path = r"ola_pdf"
    queries_csv_path = r"EWS_prompts_v2.xlsx"      
    output_folder = r"ola_results_NEW_unique_flags_classification"
    
    api_key = "8496bd1da40242a29de18fe1361e498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01-preview"
    deployment_name = "gpt-4.1-mini"
  
    previous_year_data = """
Previous reported Debt 5,684 Cr Mar-24
Current quarter ebidta (525)Cr Mar-25
Previous reported asset value 7,735Cr Mar-24
Previous reported receivable days 12 days Mar-24
Previous reported payable days 112 days Mar-24
Previous reported revenue 1,045 Cr Dec-24
Previous reported profitability (460) Cr Dec-24
Previous reported operating margin -44.00% Dec-24
Previous reported cash balance 1,663 Cr Mar-24
Previous reported current liabilities 1,071 Cr Mar-24

"""
 
    print("=" * 60)
    print("NEW 5-ITERATION PIPELINE WITH UNIQUE FLAGS CLASSIFICATION")
    print("=" * 60)
    print("Changes implemented:")
    print("1. Iterations 1-4: No changes (same as before)")
    print("2. NEW Iteration 5:")
    print("   - Extract unique flags from 2nd iteration")
    print("   - Classify each flag against 15 criteria")
    print("   - Auto-assign 'Low' risk if no criteria match")
    print("3. Word document: Count + High flags + 4th iteration summary")
    print("=" * 60)
    
    run_multiple_pdfs_five_iterations_pipeline(
        pdf_folder_path=pdf_folder_path,
        queries_csv_path=queries_csv_path,
        previous_year_data=previous_year_data,
        output_folder=output_folder,
        api_key=api_key,
        azure_endpoint=azure_endpoint,
        api_version=api_version,
        deployment_name=deployment_name
    )
 
if __name__ == "__main__":
    main_batch_processing_five_iterations()
