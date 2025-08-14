import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)
 
def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class AzureOpenAILLM:
    """Azure OpenAI gpt-4.1 LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI gpt-4.1"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""<role>
You are an expert document analyst specializing in extracting key corporate information from financial documents and earnings transcripts.
</role>

<system_prompt>
You excel at quickly identifying and extracting specific corporate identifiers from financial documents with high accuracy and consistency.
</system_prompt>

<instruction>
Extract the company name, quarter, and financial year from the provided text.

EXTRACTION REQUIREMENTS:
1. Company Name: Full legal company name including suffixes (Ltd/Limited/Inc/Corp etc.)
2. Quarter: Identify the quarter (Q1/Q2/Q3/Q4)  
3. Financial Year: Extract financial year (FY23/FY24/FY25 etc.)

OUTPUT FORMAT:
Provide ONLY the result in this exact format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If any component cannot be clearly identified, use reasonable defaults based on context.
</instruction>

<context>
DOCUMENT TEXT TO ANALYZE:
{first_page_text}
</context>

Extract company information:"""
       
        response = llm._call(prompt, max_tokens=200)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def extract_unique_flags_with_strict_deduplication(response_text: str, llm: AzureOpenAILLM) -> List[str]:
    """Enhanced extraction with STRICT deduplication to prevent duplicates"""
    
    prompt = f"""<role>
You are an expert financial analyst specializing in red flag extraction and deduplication with 15+ years of experience in financial risk assessment.
</role>

<system_prompt>
You excel at identifying unique financial concerns, eliminating redundancy, and extracting the most critical risk factors that require management attention and investor awareness.
</system_prompt>

<instruction>
Extract UNIQUE financial red flags from the analysis text with ZERO duplicates or overlapping concerns.

EXTRACTION RULES:
1. Extract all distinct financial red flags mentioned in the text
2. Each flag must represent a COMPLETELY separate financial concern
3. Merge similar flags into one comprehensive statement
4. Remove generic statements without specific value
5. Prioritize flags with quantitative data over vague statements
6. Focus on actionable, specific concerns
7. Use clear, concise business language
8. Maximum 15 most critical flags

DEDUPLICATION ALGORITHM:
- If 2+ flags address the same underlying issue → MERGE into one comprehensive flag
- If one flag is a subset of another → REMOVE the subset
- If flags share >60% similar keywords → CONSOLIDATE
- Example merges:
  • "Revenue declined 20%" + "Sales performance weak" + "Top line pressure" → "Revenue declined 20% with continued sales performance pressure"
  • "Debt increased significantly" + "Higher borrowing levels" + "Leverage concerns" → "Debt levels increased significantly raising leverage concerns"
  • "Cash flow issues" + "Liquidity problems" + "Working capital constraints" → "Cash flow and liquidity challenges with working capital constraints"

OUTPUT FORMAT:
Return ONLY a clean Python list with no additional text or explanations.
Format: ["flag 1", "flag 2", "flag 3", ...]

QUALITY CRITERIA:
- Each flag = distinct financial risk
- No redundancy or overlap between flags
- Specific and actionable statements
- Include numbers/percentages when available
- Professional financial terminology
</instruction>

<context>
FINANCIAL ANALYSIS TO PROCESS:
{response_text}
</context>

Extract unique flags:"""
    
    try:
        response = llm._call(prompt, max_tokens=600, temperature=0.0)
        
        # Try to parse as Python list
        try:
            unique_flags = ast.literal_eval(response.strip())
            if isinstance(unique_flags, list) and len(unique_flags) <= 12:
                flags_list = [flag.strip() for flag in unique_flags if flag.strip()]
            else:
                flags_list = unique_flags[:10] if len(unique_flags) > 10 else unique_flags
        except:
            # Fallback parsing if ast.literal_eval fails
            lines = response.strip().split('\n')
            flags_list = []
            
            for line in lines:
                line = line.strip()
                # Look for quoted strings
                if (line.startswith('"') and line.endswith('"')) or (line.startswith("'") and line.endswith("'")):
                    flag = line[1:-1].strip()
                    if flag and len(flag) > 5:
                        flags_list.append(flag)
                # Look for list items
                elif line.startswith('- ') or line.startswith('* '):
                    flag = line[2:].strip()
                    if flag and len(flag) > 5:
                        flags_list.append(flag)
        
        # Apply additional aggressive deduplication
        final_flags = []
        seen_keywords = []
        
        for flag in flags_list:
            if not flag or len(flag) <= 5:
                continue
                
            # Create normalized version for comparison
            normalized = re.sub(r'[^\w\s]', '', flag.lower()).strip()
            words = set(normalized.split())
            
            # Check for keyword overlap with existing flags
            is_duplicate = False
            for existing_keywords in seen_keywords:
                # If more than 60% of words overlap, consider it duplicate
                overlap = len(words.intersection(existing_keywords)) / max(len(words), len(existing_keywords))
                if overlap > 0.6:
                    is_duplicate = True
                    break
            
            if not is_duplicate and len(final_flags) < 10:
                final_flags.append(flag)
                seen_keywords.append(words)
        
        return final_flags if final_flags else ["No specific red flags identified"]
        
    except Exception as e:
        logger.error(f"Error in strict deduplication: {e}")
        return ["Error in flag extraction"]

def parse_previous_year_data(previous_year_data: str) -> Dict[str, float]:
    """Parse the previous year data string into a dictionary for calculations"""
    data_dict = {}
    lines = previous_year_data.strip().split('\n')
    
    for line in lines:
        if '\t' in line:
            parts = line.split('\t')
            if len(parts) >= 3:
                key = parts[0].strip().lower()
                value_str = parts[2].strip()
                
                # Extract numeric value (remove 'Cr', '%', etc.)
                value_str = re.sub(r'[^\d.-]', '', value_str)
                try:
                    value = float(value_str)
                    data_dict[key] = value
                except ValueError:
                    continue
    
    return data_dict

def extract_numbers_from_flag(flag: str) -> List[float]:
    """Extract numeric values from flag text"""
    numbers = re.findall(r'\d+(?:\.\d+)?', flag)
    return [float(num) for num in numbers]

def get_criteria_keywords():
    """Get keyword mappings for criteria matching"""
    return {
        "debt_increase": ["debt increase", "debt increased", "higher debt", "borrowing increase", "debt rise", "leverage increase"],
        "provisioning": ["provision", "write-off", "bad debt", "impairment", "writeoff", "provisions"],
        "asset_decline": ["asset decline", "asset fall", "asset decrease", "asset value decline", "asset reduction"],
        "receivable_days": ["receivable days", "collection period", "DSO", "debtor days", "collection days"],
        "payable_days": ["payable days", "payment period", "DPO", "creditor days", "payment days"],
        "debt_ebitda": ["debt to ebitda", "leverage ratio", "debt multiple", "debt ebitda", "leverage"],
        "revenue_decline": ["revenue decline", "sales decline", "revenue fall", "top line decline", "revenue drop"],
        "profit_before_tax_decline": ["profit before tax", "PBT decline", "pre-tax profit", "profit fall"],
        "profit_after_tax_decline": ["profit after tax", "PAT decline", "net profit", "bottom line"],
        "ebidta_decline": ["ebitda decline", "ebitda fall", "ebitda drop", "operating profit"],
        "margin_decline": ["margin decline", "margin pressure", "profitability decline", "margin compression"],
        "cash_balance": ["cash decline", "liquidity issue", "cash shortage", "cash position", "cash flow"],
        "short_term_borrowings": ["short-term debt", "current liabilities", "short term borrowing"],
        "receivables": ["receivables increase", "debtors increase", "trade receivables"],
        "payables": ["payables increase", "creditors increase", "trade payables"],
        "one-time_expenses": ["one-time", "exceptional", "non-recurring", "extraordinary"],
        "impairment": ["impairment", "writedown", "asset impairment", "goodwill impairment"],
        "gross_margin": ["gross margin", "gross profit margin", "gross profitability"],
        "management_issues": ["management change", "CEO", "CFO", "resignation", "manpower", "leadership"],
        "regulatory_compliance": ["regulatory", "compliance", "penalty", "violation", "regulator"],
        "market_competition": ["competition", "market share", "competitor", "competitive pressure"],
        "operational_disruptions": ["operational", "supply chain", "production issues", "operations"]
    }

def perform_detailed_calculation_analysis(flag: str, criteria_name: str, previous_data: Dict[str, float]) -> Dict[str, Any]:
    """Perform detailed calculations with comprehensive print statements"""
    
    print(f"\n{'-'*80}")
    print(f"DETAILED CALCULATION FOR FLAG: {flag}")
    print(f"MATCHED CRITERIA: {criteria_name}")
    print(f"{'-'*80}")
    
    flag_numbers = extract_numbers_from_flag(flag)
    print(f"Numbers extracted from flag: {flag_numbers}")
    
    result = {
        'calculation_performed': False,
        'threshold_met': False,
        'percentage_change': 0.0,
        'threshold': 0.0,
        'risk_level': 'Low',
        'calculation_type': '',
        'previous_value': 0.0,
        'current_value': 0.0,
        'calculation_details': ''
    }
    
    # Define calculation mappings with detailed info
    calculation_mappings = {
        'debt_increase': {
            'prev_key': 'debt as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Debt Increase Analysis'
        },
        'revenue_decline': {
            'prev_key': 'revenue as per previous reported quarter number', 
            'threshold': 25,
            'comparison': 'decline',
            'description': 'Revenue Decline Analysis'
        },
        'profit_before_tax_decline': {
            'prev_key': 'profit before tax as per previous reported quarter number',
            'threshold': 25,
            'comparison': 'decline',
            'description': 'Profit Before Tax Decline Analysis'
        },
        'profit_after_tax_decline': {
            'prev_key': 'profit after tax as per previous reported quarter number',
            'threshold': 25,
            'comparison': 'decline',
            'description': 'Profit After Tax Decline Analysis'
        },
        'ebidta_decline': {
            'prev_key': 'ebidta as per previous reported quarter number',
            'threshold': 25,
            'comparison': 'decline',
            'description': 'EBITDA Decline Analysis'
        },
        'asset_decline': {
            'prev_key': 'asset value as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'decline',
            'description': 'Asset Value Decline Analysis'
        },
        'receivable_days': {
            'prev_key': 'receivable days as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Receivable Days Increase Analysis'
        },
        'payable_days': {
            'prev_key': 'payable days as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Payable Days Increase Analysis'
        },
        'margin_decline': {
            'prev_key': 'operating margin as per previous quarter number',
            'threshold': 25,
            'comparison': 'decline',
            'description': 'Operating Margin Decline Analysis'
        },
        'cash_balance': {
            'prev_key': 'cash balance as per previous reported balance sheet number',
            'threshold': 25,
            'comparison': 'decline',
            'description': 'Cash Balance Decline Analysis'
        },
        'short_term_borrowings': {
            'prev_key': 'short term borrowings as per the previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Short-term Borrowings Increase Analysis'
        },
        'receivables': {
            'prev_key': 'receivables as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Receivables Increase Analysis'
        },
        'payables': {
            'prev_key': 'payables as per previous reported balance sheet number',
            'threshold': 30,
            'comparison': 'increase',
            'description': 'Payables Increase Analysis'
        }
    }
    
    if criteria_name in calculation_mappings:
        mapping = calculation_mappings[criteria_name]
        prev_key = mapping['prev_key']
        threshold = mapping['threshold']
        comparison_type = mapping['comparison']
        description = mapping['description']
        
        print(f"Analysis Type: {description}")
        print(f"Previous data key: '{prev_key}'")
        print(f"Threshold for High Risk: {threshold}%")
        print(f"Comparison Type: {comparison_type}")
        
        if prev_key in previous_data:
            previous_value = previous_data[prev_key]
            print(f"Previous period value: {previous_value}")
            
            if flag_numbers:
                print(f"Attempting calculation with flag numbers: {flag_numbers}")
                
                for current_value in flag_numbers:
                    print(f"Testing current value: {current_value}")
                    
                    if previous_value != 0:
                        if comparison_type == 'decline':
                            percentage_change = ((previous_value - current_value) / previous_value) * 100
                            direction = "decline"
                        else:  # increase
                            percentage_change = ((current_value - previous_value) / previous_value) * 100
                            direction = "increase"
                        
                        print(f"Calculation: (({previous_value} - {current_value}) / {previous_value}) × 100 = {percentage_change:.2f}%")
                        print(f"Direction: {direction}")
                        
                        result['calculation_performed'] = True
                        result['percentage_change'] = percentage_change
                        result['threshold'] = threshold
                        result['calculation_type'] = comparison_type
                        result['previous_value'] = previous_value
                        result['current_value'] = current_value
                        result['calculation_details'] = f"{direction.title()}: {percentage_change:.1f}% (threshold: {threshold}%)"
                        
                        if percentage_change >= threshold:
                            result['risk_level'] = 'High'
                            result['threshold_met'] = True
                            print(f"✅ HIGH RISK: {percentage_change:.2f}% {direction} ≥ {threshold}% threshold")
                        else:
                            result['risk_level'] = 'Low'
                            result['threshold_met'] = False
                            print(f"⚠️  LOW RISK: {percentage_change:.2f}% {direction} < {threshold}% threshold")
                        break
                    else:
                        print(f"⚠️  Cannot calculate - previous value is zero")
            else:
                print(f"⚠️  No numbers found in flag text for calculation")
        else:
            print(f"⚠️  Previous data key '{prev_key}' not found")
            print(f"Available keys: {list(previous_data.keys())}")
    
    # Special case for debt_ebitda ratio
    elif criteria_name == 'debt_ebitda':
        print(f"Analysis Type: Debt to EBITDA Ratio Analysis")
        debt_key = 'debt as per previous reported balance sheet number'
        ebitda_key = 'current quarter ebidta'
        
        print(f"Looking for debt key: '{debt_key}'")
        print(f"Looking for EBITDA key: '{ebitda_key}'")
        
        if debt_key in previous_data and ebitda_key in previous_data:
            debt = previous_data[debt_key]
            ebitda = previous_data[ebitda_key]
            
            print(f"Debt value: {debt}")
            print(f"EBITDA value: {ebitda}")
            
            if ebitda != 0:
                debt_ebitda_ratio = debt / ebitda
                print(f"Calculation: {debt} / {ebitda} = {debt_ebitda_ratio:.2f}x")
                print(f"Threshold for High Risk: 3.0x")
                
                result['calculation_performed'] = True
                result['percentage_change'] = debt_ebitda_ratio
                result['threshold'] = 3.0
                result['calculation_type'] = 'ratio'
                result['previous_value'] = ebitda
                result['current_value'] = debt
                result['calculation_details'] = f"Debt/EBITDA: {debt_ebitda_ratio:.1f}x (threshold: 3.0x)"
                
                if debt_ebitda_ratio > 3.0:
                    result['risk_level'] = 'High'
                    result['threshold_met'] = True
                    print(f"✅ HIGH RISK: {debt_ebitda_ratio:.2f}x > 3.0x threshold")
                else:
                    result['risk_level'] = 'Low'
                    result['threshold_met'] = False
                    print(f"⚠️  LOW RISK: {debt_ebitda_ratio:.2f}x ≤ 3.0x threshold")
            else:
                print(f"⚠️  Cannot calculate ratio - EBITDA is zero")
        else:
            print(f"⚠️  Required data not found for debt/EBITDA calculation")
    else:
        print(f"⚠️  No calculation mapping defined for criteria: {criteria_name}")
    
    print(f"FINAL RESULT: {result['risk_level']} Risk")
    if result['calculation_performed']:
        print(f"CALCULATION SUMMARY: {result['calculation_details']}")
    print(f"{'-'*80}")
    
    return result

def classify_flag_against_criteria_strict(flag: str, criteria_definitions: Dict[str, str], 
                                         previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, Any]:
    """Enhanced classification with strict criteria matching and detailed calculations"""
    
    # Parse previous year data
    previous_data = parse_previous_year_data(previous_year_data)
    
    # Get criteria keywords for matching
    criteria_keywords = get_criteria_keywords()
    
    print(f"\n📊 CLASSIFYING FLAG: {flag}")
    print(f"Parsed previous year data keys: {list(previous_data.keys())}")
    
    # First, try keyword-based matching
    matched_criteria = 'None'
    flag_lower = flag.lower()
    
    # Check for keyword matches
    for criteria_name, keywords in criteria_keywords.items():
        for keyword in keywords:
            if keyword.lower() in flag_lower:
                matched_criteria = criteria_name
                print(f"✅ Keyword match found: '{keyword}' → {criteria_name}")
                break
        if matched_criteria != 'None':
            break
    
    # If no keyword match, use LLM for classification
    if matched_criteria == 'None':
        print(f"⚠️  No keyword match found, using LLM classification...")
        criteria_list = "\n".join([f"{name}: {desc}" for name, desc in criteria_definitions.items()])
        
        prompt = f"""Look at this red flag and match it to ONE criteria from the list below.

RED FLAG: "{flag}"

CRITERIA LIST:
{criteria_list}

Give answer in this format:
Matched_Criteria: [criteria name or "None"]"""
        
        try:
            response = llm._call(prompt, max_tokens=200, temperature=0.0)
            lines = response.strip().split('\n')
            for line in lines:
                if 'Matched_Criteria:' in line:
                    matched_criteria = line.split(':', 1)[1].strip()
                    print(f"🤖 LLM matched criteria: {matched_criteria}")
                    break
        except Exception as e:
            print(f"❌ LLM classification failed: {e}")
            matched_criteria = 'None'
    
    # Perform calculations if criteria is matched
    if matched_criteria != 'None' and matched_criteria in criteria_definitions:
        calculation_result = perform_detailed_calculation_analysis(flag, matched_criteria, previous_data)
        
        return {
            'matched_criteria': matched_criteria,
            'risk_level': calculation_result['risk_level'],
            'reasoning': calculation_result['calculation_details'] if calculation_result['calculation_performed'] else "No calculation data available",
            'calculation_performed': str(calculation_result['calculation_performed']),
            'threshold_met': str(calculation_result['threshold_met']),
            'percentage_change': calculation_result['percentage_change'],
            'threshold': calculation_result['threshold'],
            'previous_value': calculation_result['previous_value'],
            'current_value': calculation_result['current_value']
        }
    else:
        print(f"⚠️  No matching criteria found or criteria not in definitions")
        return {
            'matched_criteria': matched_criteria,
            'risk_level': 'Low',
            'reasoning': "No matching criteria found",
            'calculation_performed': 'False',
            'threshold_met': 'False',
            'percentage_change': 0.0,
            'threshold': 0.0,
            'previous_value': 0.0,
            'current_value': 0.0
        }

def print_classification_results(classification_results: List[Dict], unique_flags: List[str]):
    """Print clean, professional classification results"""
    
    print(f"\n{'='*120}")
    print(f"                                    RISK CLASSIFICATION RESULTS")
    print(f"{'='*120}")
    
    high_count = 0
    low_count = 0
    
    for i, (flag, result) in enumerate(zip(unique_flags, classification_results), 1):
        risk_level = result['risk_level']
        criteria = result['matched_criteria']
        
        if risk_level == 'High':
            high_count += 1
        else:
            low_count += 1
        
        # Format the output line
        flag_short = flag[:50] + "..." if len(flag) > 50 else flag
        
        if result['calculation_performed'] == 'True':
            if result['matched_criteria'] == 'debt_ebitda':
                calc_display = f"{result['percentage_change']:.1f}x (threshold: {result['threshold']}x)"
            else:
                calc_display = f"{result['percentage_change']:.1f}% (threshold: {result['threshold']}%)"
        else:
            calc_display = "No calculation"
        
        print(f"{i:2d}. {risk_level:4s} | {criteria:22s} | {calc_display:25s} | {flag_short}")
    
    print(f"{'='*120}")
    print(f"SUMMARY: {high_count} High Risk  |  {low_count} Low Risk  |  {len(unique_flags)} Total Flags")
    print(f"{'='*120}")

def process_flags_with_detailed_calculations(unique_flags, criteria_definitions, previous_year_data, llm):
    """Process flags with detailed calculation output"""
    
    classification_results = []
    high_risk_flags = []
    low_risk_flags = []
    
    if len(unique_flags) > 0 and unique_flags[0] != "Error in flag extraction":
        print(f"\n🔍 STARTING DETAILED CLASSIFICATION OF {len(unique_flags)} FLAGS")
        print(f"{'='*100}")
        
        for i, flag in enumerate(unique_flags, 1):
            print(f"\n🚩 FLAG {i}/{len(unique_flags)}")
            try:
                classification = classify_flag_against_criteria_strict(
                    flag=flag,
                    criteria_definitions=criteria_definitions,
                    previous_year_data=previous_year_data, 
                    llm=llm
                )
                
                classification_results.append(classification)
                
                # Add to appropriate risk category
                if (classification['risk_level'].lower() == 'high' and 
                    classification['matched_criteria'] != 'None'):
                    high_risk_flags.append(flag)
                else:
                    low_risk_flags.append(flag)
                    
            except Exception as e:
                logger.error(f"Error classifying flag {i}: {e}")
                classification_results.append({
                    'flag': flag,
                    'matched_criteria': 'None',
                    'risk_level': 'Low',
                    'reasoning': f'Classification failed: {str(e)}',
                    'calculation_performed': 'False',
                    'threshold_met': 'False',
                    'percentage_change': 0.0,
                    'threshold': 0.0,
                    'previous_value': 0.0,
                    'current_value': 0.0
                })
                low_risk_flags.append(flag)
            
            time.sleep(0.3)
    
    # Print clean results
    print_classification_results(classification_results, unique_flags)
    
    risk_counts = {
        'High': len(high_risk_flags),
        'Low': len(low_risk_flags),
        'Total': len(unique_flags) if unique_flags and unique_flags[0] != "Error in flag extraction" else 0
    }
    
    return classification_results, high_risk_flags, low_risk_flags, risk_counts

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(high_risk_flags: List[str], context: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using original PDF context"""
    if not high_risk_flags:
        return []
    
    # First, deduplicate the high_risk_flags themselves
    unique_high_risk_flags = []
    seen_flag_keywords = []
    
    for flag in high_risk_flags:
        normalized_flag = re.sub(r'[^\w\s]', '', flag.lower()).strip()
        flag_words = set(normalized_flag.split())
        
        # Check for keyword overlap with existing flags
        is_duplicate_flag = False
        for existing_keywords in seen_flag_keywords:
            overlap = len(flag_words.intersection(existing_keywords)) / max(len(flag_words), len(existing_keywords))
            if overlap > 0.7:  # High threshold for flag deduplication
                is_duplicate_flag = True
                break
        
        if not is_duplicate_flag:
            unique_high_risk_flags.append(flag)
            seen_flag_keywords.append(flag_words)
    
    concise_summaries = []
    seen_summary_keywords = []
    
    for flag in unique_high_risk_flags:
        prompt = f"""
Based on the original PDF context, create a VERY concise 1-2 line summary for this high risk flag.

ORIGINAL PDF CONTEXT:
{context}

HIGH RISK FLAG: "{flag}"

STRICT REQUIREMENTS:
1. EXACTLY 1-2 lines (maximum 2 sentences)
2. Use ONLY specific information from the PDF context
3. Include exact numbers/percentages if mentioned
4. Be factual and direct - no speculation
5. Ensure subsequent statements are cautious and do not downplay the risk. Avoid neutral/positive statements that contradict the warning.
6. Do NOT start with "Summary:" or any prefix
7. Provide ONLY the factual summary content
8. Make it UNIQUE - avoid repeating information from other summaries
9. If applicable Specify whether the flag is for : A specific business unit/division, Consolidated financials, Standalone financials, geographical region.

OUTPUT FORMAT: [Direct factual summary only, no labels or prefixes]

"""
        
        try:
            response = llm._call(prompt, max_tokens=100, temperature=0.1)
            
            # Clean response - remove any prefixes or labels
            clean_response = response.strip()
            
            # Remove common prefixes that might appear
            prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:", "Flag summary:", "The flag:", "This flag:"]
            for prefix in prefixes_to_remove:
                if clean_response.startswith(prefix):
                    clean_response = clean_response[len(prefix):].strip()
            
            # Split into lines and take first 2
            summary_lines = [line.strip() for line in clean_response.split('\n') if line.strip()]
            
            if len(summary_lines) > 2:
                concise_summary = '. '.join(summary_lines[:2])
            elif len(summary_lines) == 0:
                concise_summary = f"{flag}. Requires management attention."
            else:
                concise_summary = '. '.join(summary_lines)
            
            # Ensure proper ending
            if not concise_summary.endswith('.'):
                concise_summary += '.'
            
            # Check for duplicate content in summaries
            normalized_summary = re.sub(r'[^\w\s]', '', concise_summary.lower()).strip()
            summary_words = set(normalized_summary.split())
            
            is_duplicate_summary = False
            for existing_keywords in seen_summary_keywords:
                overlap = len(summary_words.intersection(existing_keywords)) / max(len(summary_words), len(existing_keywords))
                if overlap > 0.8:  # Very high threshold for summary deduplication
                    is_duplicate_summary = True
                    break
            
            if not is_duplicate_summary:
                concise_summaries.append(concise_summary)
                seen_summary_keywords.append(summary_words)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{flag}': {e}")
            if len(concise_summaries) < len(unique_high_risk_flags):  # Only add fallback if we haven't added this one yet
                concise_summaries.append(f"{flag}. Review required based on analysis.")
    
    return concise_summaries

def create_word_document(pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                        high_risk_flags: List[str], summary_by_categories: Dict[str, List[str]], 
                        output_folder: str, context: str, llm: AzureOpenAILLM) -> str:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
       
        high_count = risk_counts.get('High', 0)
        low_count = risk_counts.get('Low', 0)
        total_count = high_count + low_count
       
        # Safely set table cells
        if len(table.rows) >= 3 and len(table.columns) >= 2:
            table.cell(0, 0).text = 'High Risk'
            table.cell(0, 1).text = str(high_count)
            table.cell(1, 0).text = 'Low Risk'
            table.cell(1, 1).text = str(low_count)
            table.cell(2, 0).text = 'Total Flags'
            table.cell(2, 1).text = str(total_count)
           
            # Make headers bold
            for i in range(3):
                if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                    table.cell(i, 0).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
       
        # High Risk Flags section with concise summaries
        if high_risk_flags and len(high_risk_flags) > 0:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Generate concise summaries for high risk flags
            concise_summaries = generate_strict_high_risk_summary(high_risk_flags, context, llm)
            
            # Final deduplication check at Word document level - more aggressive
            final_unique_summaries = []
            seen_content = set()
            
            for summary in concise_summaries:
                if not summary or not summary.strip():
                    continue
                    
                # Create multiple normalized versions for comparison
                normalized1 = re.sub(r'[^\w\s]', '', summary.lower()).strip()
                normalized2 = re.sub(r'\b(the|a|an|and|or|but|in|on|at|to|for|of|with|by)\b', '', normalized1)
                
                # Check if this content is substantially different
                is_unique = True
                for seen in seen_content:
                    # Calculate similarity
                    words1 = set(normalized2.split())
                    words2 = set(seen.split())
                    if len(words1) == 0 or len(words2) == 0:
                        continue
                    similarity = len(words1.intersection(words2)) / len(words1.union(words2))
                    if similarity > 0.6:  # If more than 60% similar, consider duplicate
                        is_unique = False
                        break
                
                if is_unique:
                    final_unique_summaries.append(summary)
                    seen_content.add(normalized2)
            
            for summary in final_unique_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        # Create minimal document as fallback
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.get('High', 0)}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.get('Low', 0)}")
            doc.add_paragraph(f"Total Flags: {risk_counts.get('Total', 0)}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

def process_pdf_enhanced_pipeline(pdf_path: str, queries_csv_path: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1"):
    """
    Process PDF through enhanced 5-iteration pipeline with structured prompts and enhanced classification
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM and load PDF
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load first query from CSV/Excel
        try:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            
            if len(queries_df) == 0 or "prompt" not in queries_df.columns:
                first_query = "Analyze this document for potential red flags."
            else:
                first_query = queries_df["prompt"].tolist()[0]
        except Exception as e:
            logger.warning(f"Error loading queries file: {e}. Using default query.")
            first_query = "Analyze this document for potential red flags."
        
        # ITERATION 1: Initial red flag identification with structured prompt (UNCHANGED)
        print("Running 1st iteration - Initial Analysis...")
        first_prompt = f"""<role>
You are an expert financial analyst with 15+ years of experience specializing in identifying red flags from earnings call transcripts and financial documents.
</role>

<system_prompt>
You excel at comprehensive document analysis, identifying subtle financial risks, and providing detailed evidence-based assessments with precise documentation.
</system_prompt>

<instruction>
Analyze the ENTIRE document and identify ALL potential red flags comprehensively.

ANALYSIS REQUIREMENTS:
- Review every section of the document thoroughly
- Identify financial, operational, strategic, and management risks
- Focus on quantitative concerns with specific data points
- Document exact quotes with speaker attribution
- Number each red flag sequentially (1, 2, 3, etc.)
- Include page references where available

OUTPUT FORMAT:
For each red flag:
1. The potential red flag you observed - [brief description]
Original Quote: "[exact quote with speaker name]" (Page X)

CRITICAL: Ensure comprehensive analysis of the entire document.
</instruction>

<context>
COMPLETE DOCUMENT TO ANALYZE:
{context}

SPECIFIC QUESTION: {first_query}
</context>

Provide comprehensive red flag analysis:"""
        
        first_response = llm._call(first_prompt, max_tokens=4000)
        
        # ITERATION 2: Deduplication (UNCHANGED)
        print("Running 2nd iteration - Deduplication...")
        second_prompt = "Remove the duplicates from the above context. Also if the Original Quote and Keyword identifies is same remove them. Do not lose data if duplicates are not found."
        
        second_full_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{context}
 
Previous Analysis: {first_response}
 
Based on the above analysis and the original context, please answer: {second_prompt}
 
Answer:"""
        
        second_response = llm._call(second_full_prompt, max_tokens=4000)
        
        # ITERATION 3: Categorization (UNCHANGED)
        print("Running 3rd iteration - Categorization...")
        third_prompt = f"""<role>
You are a senior financial analyst expert in financial risk categorization with deep knowledge of balance sheet analysis, P&L assessment, and corporate risk frameworks.
</role>

<system_prompt>
You excel at organizing financial risks into standardized categories, ensuring comprehensive coverage of all financial risk areas, and maintaining accuracy in risk classification.
</system_prompt>

<instruction>
Categorize the identified red flags into the following 7 standardized categories based on their financial nature and business impact.

MANDATORY CATEGORIES:
1. Balance Sheet Issues: Assets, liabilities, equity, debt, and overall financial position concerns
2. P&L (Income Statement) Issues: Revenue, expenses, profits, and financial performance concerns  
3. Liquidity Issues: Short-term obligations, cash flow, debt repayment, working capital concerns
4. Management and Strategy Issues: Leadership, governance, decision-making, strategy, and vision concerns
5. Regulatory Issues: Compliance, laws, regulations, and regulatory body concerns
6. Industry and Market Issues: Market position, industry trends, competitive landscape concerns
7. Operational Issues: Internal processes, systems, infrastructure, and operational efficiency concerns

CATEGORIZATION RULES:
- Assign each red flag to the MOST relevant category only
- Do not create new categories - use only the 7 listed above
- Preserve all Original Quotes exactly as provided
- If a red flag could fit multiple categories, choose the primary/most relevant one
- Do not leave any red flag unclassified
- Do not repeat categories in the output
- Maintain sequential organization within each category

OUTPUT FORMAT:
### Balance Sheet Issues
- [Red flag 1 with original quote and page reference]
- [Red flag 2 with original quote and page reference]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote and page reference]

Continue this format for all applicable categories.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}
</context>

Provide categorized analysis:"""
        
        third_response = llm._call(third_prompt, max_tokens=4000)
        
        # ITERATION 4: Summary generation (UNCHANGED)
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = f"""<role>
You are an expert financial summarization specialist with expertise in creating concise, factual, and comprehensive summaries that preserve critical quantitative data and key insights.
</role>

<system_prompt>
You excel at distilling complex financial analysis into clear, actionable summaries while maintaining objectivity, preserving all quantitative details, and ensuring no critical information is lost.
</system_prompt>

<instruction>
Create a comprehensive summary of each category of red flags in bullet point format following these strict guidelines.

SUMMARY REQUIREMENTS:
1. Retain ALL quantitative information (numbers, percentages, ratios, dates)
2. Maintain completely neutral, factual tone - no opinions or interpretations
3. Include every red flag from each category - no omissions
4. Base content solely on the provided categorized analysis
5. Preserve specific data points and statistics wherever mentioned
6. Each bullet point should capture key details of individual red flags
7. Balance thoroughness with conciseness
8. Use professional financial terminology
9. Ensure category-specific content alignment

OUTPUT FORMAT:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories that contain red flags.

CRITICAL: Each bullet point represents a concise summary of individual red flags with preserved quantitative details.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        fourth_response = llm._call(fourth_prompt, max_tokens=4000)
        
        # ITERATION 5: Enhanced flag extraction and classification with detailed calculations
        print("Running 5th iteration - Enhanced Classification with Detailed Calculations...")
        
        # Parse previous year data first and display it
        parsed_previous_data = parse_previous_year_data(previous_year_data)
        print(f"\n📊 PARSED PREVIOUS YEAR DATA:")
        print(f"{'='*60}")
        for key, value in parsed_previous_data.items():
            print(f"{key}: {value}")
        print(f"{'='*60}")
        
        # Step 1: Extract unique flags with STRICT deduplication
        try:
            unique_flags = extract_unique_flags_with_strict_deduplication(second_response, llm)
            print(f"\n🚩 EXTRACTED {len(unique_flags)} UNIQUE FLAGS")
        except Exception as e:
            logger.error(f"Error extracting flags: {e}")
            unique_flags = ["Error in flag extraction"]
        
        # Define criteria definitions
        criteria_definitions = {
            "debt_increase": "High: Debt is increased more than 30% compared to previous reported balance sheet number; Low: Debt increased less than 30% compared to previous reported balance sheet number",
            "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 25% of current quarter's EBIDTA",
            "asset_decline": "High: Asset value falls by more than 30% compared to the previous reported balance sheet number; Low: Asset value falls by less than 30% compared to previous reported balance sheet number",
            "receivable_days": "High: receivable days OR debtor days are increased more than 30% compared to previous reported balance sheet number; Low: receivable days or debtor's days are increased but less than 30% compared to previous reported balance sheet number",
            "payable_days": "High: payable days or creditors days increase by more than 30% compared to previous reported balance sheet number; Low: payable days or creditors days increase is less than 30% compared to previous reported balance sheet number",
            "revenue_decline": "High: revenue falls by more than 25% compared to previous reported quarter number; Low: revenue falls by less than 25% compared to previous reported quarter number",
            "profit_before_tax_decline": "High: profitability or profit before tax (PBT) falls by more than 25% compared to previous reported quarter number; Low: profitability or profit before tax (PBT) falls by less than 25% compared to previous reported quarter number",
            "profit_after_tax_decline": "High: Profit after tax (PAT) falls by more than 25% compared to previous reported quarter number; Low: Profit after tax (PAT) falls by less than 25% compared to previous reported quarter number",
            "EBIDTA_decline": "High: EBIDTA falls by more than 25% compared to previous reported quarter number; Low: EBIDTA falls by less than 25% compared to previous reported quarter number",
            "margin_decline": "High: operating margin falling more than 25% compared to previous reported quarter number; Low: Operating margin falling less than 25% compared to previous reported quarter number",
            "cash_balance": "High: cash balance falling more than 25% compared to previous reported balance sheet number; Low: cash balance falling less than 25% compared to previous reported balance sheet number",
            "short_term_borrowings": "High: Short-term borrowings or current liabilities increase by more than 30% compared to previous reported balance sheet number; Low: Short-term borrowings or current liabilities increase is less than 30% compared to previous reported balance sheet number",
            "impairment": "High: Impairment or devaluation more than 25% of previous reported net worth from balance sheet; Low: Impairment or devaluation less than 25% of previous reported net worth from balance sheet.",
            "receivables": "High: receivables or debtors are increased more than 30% compared to previous reported balance sheet number; Low: receivables or debtors are increase is less than 30% compared to previous reported balance sheet number",
            "payables": "High: payables or creditors increase by greater than 30% compared to previous reported balance sheet number; Low: payables or creditors is less than 30% compared to previous reported balance sheet number",
            "one-time_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 25% of current quarter's EBIDTA",
            "debt_ebitda": "High: Debt/EBIDTA > 3x i.e. Debt to EBITDA ratio is above (greater than) three times; Low: Debt/EBITDA < 3x i.e. Debt to EBITDA ratio is less than three times",
            "gross_margin": "High: gross margin falling more than 100bps (basis points) ; Low: gross margin falling less than 100bps (basis points)",
            "management_issues": "High: If found any management or strategy related issues or concerns or a conclusion of any discussion related to management and strategy; Low: If there is a no clear concern for the company basis the discussion on the management or strategy related issues",
            "regulatory_compliance": "High: If found any regulatory issues as a concern or a conclusion of any discussion related to regulatory issues or warning(s) from the regulators; Low: If there is a no clear concern for the company basis the discussion on the regulatory issues",
            "market_competition": "High: Any competitive intensity or new entrants, any decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share",
            "operational_disruptions": "High: if found any operational or supply chain issues as a concern or a conclusion of any discussion related to operational issues; Low: if there is no clear concern for the company basis the discussion on the operational or supply chain issues"
        }
        
        # Step 2: Classify each unique flag with DETAILED CALCULATIONS
        classification_results, high_risk_flags, low_risk_flags, risk_counts = process_flags_with_detailed_calculations(
            unique_flags, criteria_definitions, previous_year_data, llm
        )
        
        # Extract company info and create Word document
        print("\nCreating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
           
            # Create Word document with strict high risk summaries
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                high_risk_flags=high_risk_flags,
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                context=context,
                llm=llm
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Save all results to CSV files
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        # Save pipeline results
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Deduplication", 
                "Categorization",
                "Summary Generation",
                "Enhanced Classification with Detailed Calculations"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Enhanced Classification: {risk_counts['High']} High Risk, {risk_counts['Low']} Low Risk flags from {risk_counts['Total']} total unique flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_enhanced_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        # Save detailed classification results
        if len(classification_results) > 0:
            classification_df = pd.DataFrame(classification_results)
            classification_file = os.path.join(output_folder, f"{pdf_name}_enhanced_flag_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        return results_summary
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

def main():
    """Main function to process all PDFs in the specified folder"""
    
    # Configuration
    pdf_folder_path = r"vedanta_pdf" 
    queries_csv_path = r"EWS_prompts_v2_2.xlsx"
    output_folder = r"vedanta_results_try2.2.0"

    api_key = "8496498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01-preview"
    deployment_name = "gpt-4.1"
  
    previous_year_data = """
Debt as per Previous reported balance sheet number	Mar-23	80329Cr
Current quarter ebidta	March-24	11511Cr
Asset value as per previous reported balance sheet number	Mar-23	189455Cr
Receivable days as per previous reported balance sheet number	Mar-23	10days
Payable days as per Previous reported balance sheet number	Mar-23	91days
Revenue as per previous reported quarter number	Dec-23	35541Cr
profit before tax as per previous reported quarter number	Dec-23	4105Cr
profit after tax as per previous reported quarter number	Dec-23	2868Cr
EBIDTA as per previous reported quarter number	Dec-23	8531Cr
Operating margin as per previous quarter number	Dec-23	25%
Cash balance as per previous reported balance sheet number	Mar-23	9254Cr
Short term borrowings as per the previous reported balance sheet number	Mar-23	36407Cr
previous reported net worth from balance sheet	Mar-23	47896Cr
Receivables as per previous reported balance sheet number	Mar-23	6414Cr
Payables as per Previous reported balance sheet number	Mar-23	11043Cr

"""
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*60}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        result = process_pdf_enhanced_pipeline(
            pdf_path=pdf_file,
            queries_csv_path=queries_csv_path,
            previous_year_data=previous_year_data,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
        else:
            print(f"❌ Failed to process {pdf_file}")

if __name__ == "__main__":
    main()
