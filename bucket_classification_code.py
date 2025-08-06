import ast
import os
import time
import pandas as pd
import fitz  
import warnings
import hashlib
import logging
import json
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
from openai import AzureOpenAI
import httpx
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# CONSTANT PROMPT TEMPLATE
CONSTANT_PROMPT_TEMPLATE = """<role>Your role is that of an experienced financial analyst tasked with analyzing Earnings call transcripts for a company and identifying potential causes for concern: red flags.</role>
<instructions>1. Search for all keywords mentioned in <reference> within the document. 
2. Ensure that no keyword mentioned in <reference> is overlooked. 
3. While searching for the keywords, consider their definitions as provided in the <reference> section, and focus on those that indicate a potential cause for concern, i.e., a red flag. 
4. A keyword may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.
5. Only identify a keyword if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags.</instructions>
For each identified negative red flag, strictly adhere to the following output format:
<output format>1. The potential red flag you observed - the actual keyword
2. Original Quote: Provide the entire original quote or text that led to the identification of the red flag, along with the page number where the statement was found.</output format>

<reference>
{bucket_reference}
</reference>

Document to Analyze:
{context}

Analysis:"""

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class AzureOpenAILLM:
    """Azure OpenAI GPT-4.1-mini LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI GPT-4.1-mini"""
        try:
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""
Extract the company name, quarter, and financial year from this text from an earnings call transcript.

Text: {first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)  
3. Financial Year (FY23/FY24/FY25 etc.)

Format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

Response:"""
       
        response = llm._call(prompt, max_tokens=200)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def load_bucket_references(excel_path: str) -> Dict[str, Dict[str, Any]]:
    """Load bucket references from Excel file"""
    try:
        df = pd.read_excel(excel_path)
        bucket_data = {}
        for _, row in df.iterrows():
            bucket_data[row['bucket_name']] = {
                'reference': row['bucket_reference'],
                'keyword_count': row['keyword_count']
            }
        return bucket_data
    except Exception as e:
        logger.error(f"Error loading bucket references: {e}")
        return {}

def create_bucket_prompt(bucket_reference: str, context: str) -> str:
    """Create complete prompt by inserting bucket reference and context into constant template"""
    return CONSTANT_PROMPT_TEMPLATE.format(
        bucket_reference=bucket_reference,
        context=context
    )

def process_all_buckets(context: str, bucket_data: Dict[str, Dict[str, Any]], llm: AzureOpenAILLM) -> Dict[str, Any]:
    """Process document with all 7 bucket prompts"""
    
    results = {}
    
    # Process each bucket
    for bucket_name, data in bucket_data.items():
        print(f"Processing {bucket_name} ({data['keyword_count']} keywords)...")
        
        # Create complete prompt
        complete_prompt = create_bucket_prompt(data['reference'], context)
        
        # Call LLM
        response = llm._call(complete_prompt, max_tokens=4000)
        results[bucket_name] = {
            'response': response,
            'keyword_count': data['keyword_count'],
            'timestamp': time.strftime("%Y%m%d_%H%M%S")
        }
        
        time.sleep(0.5)  # Rate limiting
    
    return results

def parse_bucket_responses(bucket_results: Dict[str, Any]) -> Dict[str, List[str]]:
    """Parse bucket responses to extract red flags"""
    parsed_results = {}
    
    for bucket_name, data in bucket_results.items():
        response = data['response']
        red_flags = []
        
        # Split response into lines and look for red flags
        lines = response.split('\n')
        current_flag = ""
        
        for line in lines:
            line = line.strip()
            if line.startswith('1.') or line.startswith('2.') or re.match(r'^\d+\.', line):
                if current_flag:
                    red_flags.append(current_flag.strip())
                current_flag = line
            elif line.startswith('Original Quote:') or line.startswith('2. Original Quote:'):
                current_flag += f" | {line}"
            elif current_flag and line:
                current_flag += f" {line}"
        
        # Add the last flag if exists
        if current_flag:
            red_flags.append(current_flag.strip())
        
        parsed_results[bucket_name] = red_flags
    
    return parsed_results

def create_summary_by_buckets(parsed_results: Dict[str, List[str]], context: str, llm: AzureOpenAILLM) -> str:
    """Create a summary of all red flags organized by buckets"""
    
    summary_prompt = f"""Based on the red flags identified across different categories, provide a comprehensive summary organized by category.

RED FLAGS BY CATEGORY:
"""
    
    for bucket_name, flags in parsed_results.items():
        if flags:
            summary_prompt += f"\n### {bucket_name}:\n"
            for i, flag in enumerate(flags, 1):
                summary_prompt += f"{i}. {flag}\n"
    
    summary_prompt += f"""

ORIGINAL DOCUMENT CONTEXT:
{context}

TASK: Create a comprehensive summary following these guidelines:
1. Organize by the categories shown above
2. Provide factual, concise bullet points for each category
3. Include specific numbers, percentages, and data points where available
4. Maintain neutral, objective tone
5. Focus only on the red flags that were actually identified
6. If a category has no red flags, you can skip it

FORMAT:
### Category Name
* Summary point 1 with specific details
* Summary point 2 with specific details

Generate the summary:"""
    
    summary_response = llm._call(summary_prompt, max_tokens=3000)
    return summary_response

def create_comprehensive_word_document(pdf_name: str, company_info: str, 
                                     bucket_results: Dict[str, Any],
                                     parsed_results: Dict[str, List[str]],
                                     summary: str,
                                     output_folder: str) -> str:
    """Create a comprehensive Word document with all bucket analysis"""
   
    try:
        doc = Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Executive Summary section
        exec_summary_heading = doc.add_heading('Executive Summary', level=1)
        exec_summary_heading.runs[0].bold = True
        
        # Count total flags across all buckets
        total_flags = sum(len(flags) for flags in parsed_results.values())
        
        # Add summary statistics
        summary_stats = doc.add_paragraph()
        summary_stats.add_run(f"Total Red Flags Identified: {total_flags}")
        summary_stats.add_run("\nAnalysis conducted across 7 risk categories using automated keyword detection.")
        
        # Flag Distribution by Category section
        flag_dist_heading = doc.add_heading('Flag Distribution by Category:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=len(bucket_results) + 2, cols=3)
        table.style = 'Table Grid'
       
        # Header row
        table.cell(0, 0).text = 'Category'
        table.cell(0, 1).text = 'Red Flags Found'
        table.cell(0, 2).text = 'Keywords Analyzed'
        
        for i in range(3):
            table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Data rows
        row_idx = 1
        total_keywords = 0
        for bucket_name, data in bucket_results.items():
            flags_count = len(parsed_results.get(bucket_name, []))
            keyword_count = data['keyword_count']
            total_keywords += keyword_count
            
            table.cell(row_idx, 0).text = bucket_name
            table.cell(row_idx, 1).text = str(flags_count)
            table.cell(row_idx, 2).text = str(keyword_count)
            row_idx += 1
        
        # Total row
        table.cell(row_idx, 0).text = 'TOTAL'
        table.cell(row_idx, 1).text = str(total_flags)
        table.cell(row_idx, 2).text = str(total_keywords)
        
        for i in range(3):
            table.cell(row_idx, i).paragraphs[0].runs[0].bold = True
       
        doc.add_paragraph('')
        
        # Detailed Analysis by Category
        detailed_heading = doc.add_heading('Detailed Analysis by Category', level=1)
        detailed_heading.runs[0].bold = True
        
        for bucket_name, flags in parsed_results.items():
            if flags:  # Only add sections for buckets with red flags
                bucket_heading = doc.add_heading(f'{bucket_name} ({len(flags)} flags)', level=2)
                bucket_heading.runs[0].bold = True
                
                for i, flag in enumerate(flags, 1):
                    flag_paragraph = doc.add_paragraph()
                    flag_paragraph.style = 'List Number'
                    flag_paragraph.add_run(flag)
                
                doc.add_paragraph('')
        
        # Horizontal line
        doc.add_paragraph('_' * 60)
       
        # Summary section
        summary_heading = doc.add_heading('Risk Assessment Summary', level=1)
        summary_heading.runs[0].bold = True
        
        # Parse and add the summary
        summary_lines = summary.split('\n')
        current_section = None
        
        for line in summary_lines:
            line = line.strip()
            if line.startswith('###'):
                # New section header
                section_name = line.replace('###', '').strip()
                current_section = doc.add_heading(section_name, level=2)
                current_section.runs[0].bold = True
            elif line.startswith('*') or line.startswith('-'):
                # Bullet point
                bullet_text = line[1:].strip()
                if bullet_text:
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(bullet_text)
            elif line and not line.startswith('#'):
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Processing Information
        doc.add_paragraph('_' * 60)
        processing_heading = doc.add_heading('Processing Information', level=2)
        processing_heading.runs[0].bold = True
        
        processing_info = doc.add_paragraph()
        processing_info.add_run(f"Analysis Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        processing_info.add_run(f"Document: {pdf_name}\n")
        processing_info.add_run(f"Total Categories Analyzed: {len(bucket_results)}\n")
        processing_info.add_run(f"Total Keywords Monitored: {total_keywords}\n")
        processing_info.add_run("Analysis Method: Automated keyword detection with AI-powered context analysis")
       
        # Save document
        doc_filename = f"{pdf_name}_Comprehensive_Analysis.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        doc.save(doc_path)
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating comprehensive Word document: {e}")
        return None

def save_detailed_results(pdf_name: str, bucket_results: Dict[str, Any], 
                         parsed_results: Dict[str, List[str]], 
                         output_folder: str) -> str:
    """Save detailed results to CSV files"""
    
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    
    # Save raw bucket responses
    bucket_responses = []
    for bucket_name, data in bucket_results.items():
        bucket_responses.append({
            'pdf_name': pdf_name,
            'bucket_name': bucket_name,
            'keyword_count': data['keyword_count'],
            'raw_response': data['response'],
            'timestamp': data['timestamp']
        })
    
    responses_df = pd.DataFrame(bucket_responses)
    responses_file = os.path.join(output_folder, f"{pdf_name}_bucket_responses.csv")
    responses_df.to_csv(responses_file, index=False)
    
    # Save parsed red flags
    red_flags_data = []
    for bucket_name, flags in parsed_results.items():
        for i, flag in enumerate(flags, 1):
            red_flags_data.append({
                'pdf_name': pdf_name,
                'bucket_name': bucket_name,
                'flag_number': i,
                'red_flag': flag,
                'timestamp': timestamp
            })
    
    if red_flags_data:
        flags_df = pd.DataFrame(red_flags_data)
        flags_file = os.path.join(output_folder, f"{pdf_name}_red_flags.csv")
        flags_df.to_csv(flags_file, index=False)
    
    # Save summary statistics
    summary_stats = []
    for bucket_name, data in bucket_results.items():
        flags_count = len(parsed_results.get(bucket_name, []))
        summary_stats.append({
            'pdf_name': pdf_name,
            'bucket_name': bucket_name,
            'keywords_analyzed': data['keyword_count'],
            'red_flags_found': flags_count,
            'flag_rate': f"{(flags_count/data['keyword_count']*100):.1f}%" if data['keyword_count'] > 0 else "0%",
            'timestamp': timestamp
        })
    
    stats_df = pd.DataFrame(summary_stats)
    stats_file = os.path.join(output_folder, f"{pdf_name}_analysis_statistics.csv")
    stats_df.to_csv(stats_file, index=False)
    
    return responses_file

def process_pdf_bucket_pipeline(pdf_path: str, excel_path: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1"):
    """
    Complete pipeline to process PDF through 7-bucket analysis
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM
        llm = AzureOpenAILLM(
            api_key=api_key or os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=azure_endpoint or os.getenv("AZURE_OPENAI_ENDPOINT"), 
            api_version=api_version or os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            deployment_name=deployment_name
        )
        
        # Extract PDF content
        print("Extracting PDF content...")
        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Load bucket references from Excel
        print(f"Loading bucket references from {excel_path}...")
        bucket_data = load_bucket_references(excel_path)
        
        if not bucket_data:
            print("Failed to load bucket data from Excel file")
            return None
        
        print(f"Loaded {len(bucket_data)} buckets from Excel")
        
        # Process all buckets
        print("Processing all buckets...")
        bucket_results = process_all_buckets(context, bucket_data, llm)
        
        # Parse responses to extract red flags
        print("Parsing bucket responses...")
        parsed_results = parse_bucket_responses(bucket_results)
        
        # Count total flags
        total_flags = sum(len(flags) for flags in parsed_results.values())
        print(f"Total red flags identified: {total_flags}")
        
        # Print bucket summary
        print("\n=== BUCKET ANALYSIS SUMMARY ===")
        for bucket_name, flags in parsed_results.items():
            keyword_count = bucket_data[bucket_name]['keyword_count']
            print(f"{bucket_name}: {len(flags)} flags found (from {keyword_count} keywords)")
        
        # Create summary
        print("Generating comprehensive summary...")
        summary = create_summary_by_buckets(parsed_results, context, llm)
        
        # Extract company info
        print("Extracting company information...")
        company_info = extract_company_info_from_pdf(pdf_path, llm)
        
        # Create Word document
        print("Creating comprehensive Word document...")
        word_doc_path = create_comprehensive_word_document(
            pdf_name=pdf_name,
            company_info=company_info,
            bucket_results=bucket_results,
            parsed_results=parsed_results,
            summary=summary,
            output_folder=output_folder
        )
        
        if word_doc_path:
            print(f"Word document created: {word_doc_path}")
        else:
            print("Failed to create Word document")
        
        # Save detailed results to CSV
        print("Saving detailed results...")
        csv_path = save_detailed_results(pdf_name, bucket_results, parsed_results, output_folder)
        print(f"CSV files saved in: {output_folder}")

        print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
        print(f"Total Red Flags: {total_flags}")
        print(f"Word Document: {word_doc_path}")
        print(f"CSV Files: {output_folder}")
        
        return {
            'pdf_name': pdf_name,
            'total_flags': total_flags,
            'bucket_results': bucket_results,
            'parsed_results': parsed_results,
            'word_doc_path': word_doc_path,
            'csv_path': csv_path
        }
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

def main():
    """Main function to process all PDFs in the specified folder"""
    
    # Configuration
    pdf_folder_path = r"sterlin_q2_pdf" 
    excel_path = r"EWS_prompt_v3.xlsx"  # Your new Excel file
    output_folder = r"sterlin_results_bucket_analysis"
    
    api_key = "8496bd1d1361e498c"
    azure_endpoint = "https://crisil-pp-gpt.openai.azure.com"
    api_version = "2025-01-01"
    deployment_name = "gpt-4.1"
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all PDFs in folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return    
    
    print(f"Found {len(pdf_files)} PDF files to process")
    print(f"Using Excel file: {excel_path}")
    print(f"Output folder: {output_folder}")
    
    all_results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n{'='*80}")
        print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
        print(f"{'='*80}")
        
        start_time = time.time()
        
        result = process_pdf_bucket_pipeline(
            pdf_path=pdf_file,
            excel_path=excel_path,
            output_folder=output_folder,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            deployment_name=deployment_name
        )
        
        processing_time = time.time() - start_time
        
        if result is not None:
            print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
            all_results.append(result)
        else:
            print(f"❌ Failed to process {pdf_file}")
    
    # Create overall summary
    if all_results:
        print(f"\n{'='*80}")
        print("OVERALL PROCESSING SUMMARY")
        print(f"{'='*80}")
        print(f"Total PDFs processed: {len(all_results)}")
        print(f"Total red flags across all documents: {sum(r['total_flags'] for r in all_results)}")
        print(f"Results saved in: {output_folder}")
        
        # Save overall summary
        overall_summary = pd.DataFrame([
            {
                'pdf_name': r['pdf_name'],
                'total_red_flags': r['total_flags'],
                'processing_status': 'Success'
            } for r in all_results
        ])
        
        overall_summary_path = os.path.join(output_folder, "overall_processing_summary.csv")
        overall_summary.to_csv(overall_summary_path, index=False)
        print(f"Overall summary saved: {overall_summary_path}")

if __name__ == "__main__":
    main()
