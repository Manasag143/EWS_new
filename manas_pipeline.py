"""
pipeline.py - Core pipeline classes for financial analysis
"""

import logging
import time
import pandas as pd
import re
from typing import Dict, List, Tuple
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import (
    AzureOpenAILLM, PDFProcessor, CriteriaBucketManager, 
    PromptGenerator, KeywordDefinitions, FileUtils
)

logger = logging.getLogger(__name__)

# ==============================================================================
# FLAG ANALYSIS CLASS
# ==============================================================================

class FlagAnalyzer:
    """Class for analyzing and classifying red flags"""
    
    def __init__(self, llm: AzureOpenAILLM):
        self.llm = llm
        self.criteria_manager = CriteriaBucketManager()

    def extract_flags_with_complete_context(self, response: str) -> List[str]:
        """Extract flags with complete context including original quotes and page references"""
        flags_with_context = []
        lines = response.split('\n')
        current_flag = ""
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Check if this is the start of a new flag
            if re.match(r'^\d+\.\s+', line):
                # Save previous flag if it exists
                if current_flag.strip():
                    flags_with_context.append(current_flag.strip())
                
                # Start new flag
                current_flag = line
                
                # Look ahead to capture original quotes and page references
                j = i + 1
                while j < len(lines) and not re.match(r'^\d+\.\s+', lines[j].strip()):
                    next_line = lines[j].strip()
                    if next_line:  # Only add non-empty lines
                        current_flag += "\n" + next_line
                    j += 1
            
        # Don't forget the last flag
        if current_flag.strip():
            flags_with_context.append(current_flag.strip())
        
        # Clean and validate flags
        cleaned_flags = []
        for flag in flags_with_context:
            # Remove any prefixes but keep the complete context
            flag = re.sub(r'^The potential red flag you observed - ', '', flag)
            flag = flag.strip()
            
            if flag and len(flag) > 10:  # Minimum length check
                cleaned_flags.append(flag)
        
        return cleaned_flags

    def classify_all_flags_with_enhanced_buckets(self, all_flags_with_context: List[str], 
                                               previous_year_data: str) -> Dict[str, str]:
        """Enhanced classification using 8 total LLM calls for all flags combined"""
        
        criteria_buckets = self.criteria_manager.create_criteria_buckets()
        data_buckets = self.criteria_manager.create_previous_data_buckets(previous_year_data)
        
        bucket_names = [
            "Core Debt & Leverage (Quantitative)",
            "Profitability & Performance (Quantitative)", 
            "Margins & Operational Efficiency (Quantitative)",
            "Working Capital & Asset Management (Quantitative)",
            "Asset Quality & Impairments (Quantitative)",
            "Other Quantitative Risks (Quantitative)",
            "Management & Regulatory Issues (Qualitative)",
            "Qualitative Risk Indicators (Qualitative)"
        ]
        
        # Prepare all flags text for analysis with clear numbering
        all_flags_text = ""
        for i, flag in enumerate(all_flags_with_context, 1):
            all_flags_text += f"\n--- FLAG_{i} ---\n{flag}\n"
        
        bucket_results = {}
        
        for i, (criteria_bucket, data_bucket, bucket_name) in enumerate(zip(criteria_buckets, data_buckets, bucket_names)):
            criteria_list = "\n\n".join(criteria_bucket)
            
            # Different prompts for quantitative vs qualitative buckets
            if "Quantitative" in bucket_name:
                prompt = f"""You are an experienced financial analyst. Your goal is to classify given red flags gathered from earnings call transcript document into High/Low Risk based on QUANTITATIVE thresholds.
 
Red Flags to be analyzed:-
{all_flags_text}
 
High/Low Risk identification criteria (QUANTITATIVE - focus on numbers and percentages):-
{criteria_list}
 
Financial Metrics of the company needed for analysis:-
{data_bucket}
 
<instructions>
1. Review each flag against the above given QUANTITATIVE criteria and the financial metrics.
2. Classify ONLY the red flags that match the criteria in this bucket.
3. For each matching flag, determine if it's High or Low risk based on the EXACT numerical thresholds in the criteria.
4. Use the exact flag numbering format: FLAG_1, FLAG_2, etc.
5. Focus on specific numbers, percentages, ratios mentioned in the flags.
6. If no flags match the criteria in this bucket, respond with "No flags match the criteria in this bucket."
</instructions>
 
Output format - For each matching flag:
Flag_Number: FLAG_X (where X is the flag number)
Matched_Criteria: [exact criteria name from the criteria list]
Risk_Level: [High or Low]
Reasoning: [brief explanation with specific numbers/evidence from the flag and financial metrics]

<review>
1. Only analyze flags that specifically match the QUANTITATIVE criteria in this bucket.
2. Use exact flag numbering: FLAG_1, FLAG_2, FLAG_3, etc.
3. Ensure risk level determination follows the exact numerical thresholds in the criteria.
4. If a flag doesn't match any criteria in this bucket, don't include it in the output.
</review>
"""
            else:  # Qualitative bucket
                prompt = f"""You are an experienced financial analyst. Your goal is to classify given red flags gathered from earnings call transcript document into High/Low Risk based on QUALITATIVE indicators.
 
Red Flags to be analyzed:-
{all_flags_text}
 
High/Low Risk identification criteria (QUALITATIVE - focus on concerns, issues, and strategic matters):-
{criteria_list}
 
<instructions>
1. Review each flag against the above given QUALITATIVE criteria.
2. Classify ONLY the red flags that match the criteria in this bucket.
3. For each matching flag, determine if it's High or Low risk based on the presence/absence of concerns mentioned in the criteria.
4. Use the exact flag numbering format: FLAG_1, FLAG_2, etc.
5. Focus on management issues, regulatory concerns, operational problems, and strategic uncertainties.
6. Refer to the sample examples provided in criteria_list to help identify high risk flags accurately.
7. If no flags match the criteria in this bucket, respond with "No flags match the criteria in this bucket."
</instructions>
 
Output format - For each matching flag:
Flag_Number: FLAG_X (where X is the flag number)
Matched_Criteria: [exact criteria name from the criteria list]
Risk_Level: [High or Low]
Reasoning: [brief explanation with evidence from the flag about the qualitative concern]

<review>
1. Only analyze flags that specifically match the QUALITATIVE criteria in this bucket.
2. Use exact flag numbering: FLAG_1, FLAG_2, FLAG_3, etc.
3. Ensure risk level determination follows the qualitative indicators in the criteria.
4. If a flag doesn't match any criteria in this bucket, don't include it in the output.
</review>
"""

            try:
                print(f"Analyzing all flags against {bucket_name}...")
                response = self.llm.call(prompt, temperature=0.0)
                bucket_results[bucket_name] = response
                
            except Exception as e:
                logger.error(f"Error analyzing {bucket_name}: {e}")
                bucket_results[bucket_name] = f"Error in {bucket_name}: {str(e)}"
        
        return bucket_results

    def parse_bucket_results_to_classifications(self, bucket_results: Dict[str, str], 
                                              all_flags_with_context: List[str]) -> List[Dict[str, str]]:
        """Parse bucket results with explicit flag numbering"""
        flag_classifications = []
        
        # Initialize all flags as Low risk with proper flag descriptions
        for i, flag_with_context in enumerate(all_flags_with_context, 1):
            # Extract the first line as flag description, clean it up
            flag_lines = flag_with_context.strip().split('\n')
            flag_description = flag_lines[0] if flag_lines else flag_with_context
            
            # Remove numbering prefix if it exists (e.g., "1. " or "2. ")
            flag_description = re.sub(r'^\d+\.\s*', '', flag_description).strip()
            
            # Remove common prefixes
            flag_description = re.sub(r'^(The potential red flag you observed - |Red flag: |Flag: )', '', flag_description, flags=re.IGNORECASE).strip()
            
            flag_classifications.append({
                'flag': flag_description,
                'flag_with_context': flag_with_context,
                'matched_criteria': 'None',
                'risk_level': 'Low',
                'reasoning': 'No matching criteria found across all buckets',
                'bucket': 'None'
            })
        
        # Parse bucket results
        for bucket_name, bucket_response in bucket_results.items():
            if isinstance(bucket_response, str) and "No flags match" not in bucket_response and "Error" not in bucket_response:
                
                # Split response into individual flag entries
                sections = re.split(r'\n\s*(?=Flag_Number:\s*FLAG_\d+)', bucket_response.strip())
                
                for section in sections:
                    if not section.strip():
                        continue
                    
                    # Initialize variables
                    flag_number = None
                    matched_criteria = None
                    risk_level = None
                    reasoning = None
                    
                    # Parse each line in the section
                    lines = section.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('Flag_Number:'):
                            flag_number_text = line.replace('Flag_Number:', '').strip()
                            # Extract number from FLAG_X format
                            flag_match = re.search(r'FLAG_(\d+)', flag_number_text)
                            if flag_match:
                                flag_number = int(flag_match.group(1))
                        elif line.startswith('Matched_Criteria:'):
                            matched_criteria = line.replace('Matched_Criteria:', '').strip()
                            # Clean up criteria name
                            matched_criteria = re.sub(r'^\[|\]$', '', matched_criteria).strip()
                        elif line.startswith('Risk_Level:'):
                            risk_level_text = line.replace('Risk_Level:', '').strip()
                            # Extract High or Low
                            if 'High' in risk_level_text:
                                risk_level = 'High'
                            elif 'Low' in risk_level_text:
                                risk_level = 'Low'
                        elif line.startswith('Reasoning:'):
                            reasoning = line.replace('Reasoning:', '').strip()
                            # Clean up reasoning
                            reasoning = re.sub(r'^\[|\]$', '', reasoning).strip()
                    
                    # Update classification if we have all required fields
                    if (flag_number is not None and matched_criteria and 
                        risk_level and reasoning and 
                        1 <= flag_number <= len(flag_classifications)):
                        
                        flag_index = flag_number - 1
                        current_classification = flag_classifications[flag_index]
                        
                        # Update if this is a High risk classification, or if current is still default Low
                        if (risk_level == 'High' or 
                            (current_classification['matched_criteria'] == 'None' and risk_level == 'Low')):
                            
                            flag_classifications[flag_index].update({
                                'matched_criteria': matched_criteria,
                                'risk_level': risk_level,
                                'reasoning': reasoning,
                                'bucket': bucket_name
                            })
                            
                            print(f"Updated FLAG_{flag_number}: {risk_level} risk in {bucket_name}")
        
        return flag_classifications

    def parse_summary_by_categories(self, fourth_response: str) -> Dict[str, List[str]]:
        """Parse the 4th iteration summary response by categories"""
        categories_summary = {}
        sections = fourth_response.split('###')
       
        for section in sections:
            if not section.strip():
                continue
               
            lines = section.split('\n')
            category_name = ""
            bullets = []
           
            for line in lines:
                line = line.strip()
                if line and not line.startswith('*') and not line.startswith('-'):
                    category_name = line.strip()
                elif line.startswith('*') or line.startswith('-'):
                    bullet_text = line[1:].strip()
                    if bullet_text:
                        bullets.append(bullet_text)
           
            if category_name and bullets:
                categories_summary[category_name] = bullets
       
        return categories_summary

# ==============================================================================
# SUMMARY GENERATOR CLASS
# ==============================================================================

class SummaryGenerator:
    """Class for generating high-risk flag summaries"""
    
    def __init__(self, llm: AzureOpenAILLM):
        self.llm = llm

    def generate_strict_high_risk_summary(self, classification_results: List[Dict[str, str]], 
                                        previous_year_data: str) -> List[str]:
        """Generate VERY concise 1-2 line summaries for high risk flags"""
        
        # Filter only high risk flags
        high_risk_classifications = [result for result in classification_results if result['risk_level'] == 'High']
        
        if not high_risk_classifications:
            return []
        
        # Create data buckets for bucket-specific financial metrics
        data_buckets = CriteriaBucketManager.create_previous_data_buckets(previous_year_data)
        bucket_name_to_index = {
            "Core Debt & Leverage": 0,
            "Profitability & Performance": 1, 
            "Margins & Operational Efficiency": 2,
            "Working Capital & Asset Management": 3,
            "Asset Quality & Governance": 4,
            "Market & Operational Risks": 5
        }
        
        # Deduplicate the high_risk_classifications
        unique_high_risk_classifications = self._deduplicate_classifications(high_risk_classifications)
        
        concise_summaries = []
        seen_summary_keywords = []
        
        for classification in unique_high_risk_classifications:
            summary = self._generate_single_summary(classification, data_buckets, bucket_name_to_index, previous_year_data)
            
            if summary:
                # Check for duplicate content in summaries
                normalized_summary = re.sub(r'[^\w\s]', '', summary.lower()).strip()
                summary_words = set(normalized_summary.split())
                
                is_duplicate_summary = False
                for existing_keywords in seen_summary_keywords:
                    overlap = len(summary_words.intersection(existing_keywords)) / max(len(summary_words), len(existing_keywords))
                    if overlap > 0.8:  # Very high threshold for summary deduplication
                        is_duplicate_summary = True
                        break
                
                if not is_duplicate_summary:
                    concise_summaries.append(summary)
                    seen_summary_keywords.append(summary_words)
        
        return concise_summaries

    def _deduplicate_classifications(self, high_risk_classifications: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Remove duplicate classifications based on flag content similarity"""
        unique_high_risk_classifications = []
        seen_flag_keywords = []
        
        for classification in high_risk_classifications:
            flag_text = classification.get('flag', '')
            normalized_flag = re.sub(r'[^\w\s]', '', flag_text.lower()).strip()
            flag_words = set(normalized_flag.split())
            
            # Check for keyword overlap with existing flags
            is_duplicate_flag = False
            for existing_keywords in seen_flag_keywords:
                overlap = len(flag_words.intersection(existing_keywords)) / max(len(flag_words), len(existing_keywords))
                if overlap > 0.7:  # High threshold for flag deduplication
                    is_duplicate_flag = True
                    break
            
            if not is_duplicate_flag:
                unique_high_risk_classifications.append(classification)
                seen_flag_keywords.append(flag_words)
        
        return unique_high_risk_classifications

    def _generate_single_summary(self, classification: Dict[str, str], data_buckets: List[str], 
                                bucket_name_to_index: Dict[str, int], previous_year_data: str) -> str:
        """Generate summary for a single classification"""
        flag_with_context = classification.get('flag_with_context', classification.get('flag', ''))
        matched_criteria = classification.get('matched_criteria', 'Unknown criteria')
        bucket_name = classification.get('bucket', 'Unknown bucket')
        reasoning = classification.get('reasoning', 'No reasoning provided')
        
        # Get bucket-specific financial data
        bucket_index = bucket_name_to_index.get(bucket_name, 0)
        bucket_specific_data = data_buckets[bucket_index] if bucket_index < len(data_buckets) else ""
        
        # If no bucket-specific data, fall back to full previous year data
        financial_data = bucket_specific_data if bucket_specific_data.strip() else previous_year_data

        prompt = f"""<role>
You are an experienced financial analyst working in ratings company. Your goal is to review the high risk red flag identified for accuracy and generate summary of high-risk financial red flag identified from earnings call transcript.
Input document is delimited by ####. It will consist of the red flag identified and corresponding details from the earnings call transcript.
The Criteria which led to high risk identification is delimited by %%%%.
Financial Metrics of the company from previous quarter is delimited by &&&&.
</role>
 
<instructions>
1. Analyze the financials, red flag identified and the contexts, the criteria which led to high risk identification.
2. Ensure the accuracy of the identification of the red flag to be high risk.
3. Create a very concise 1-2 line summary for each high-risk flag.
4. Include exact numbers, percentages, ratios, and dates whenever mentioned which led to identification of high risk flag.
5. Be factual and direct - no speculation or interpretation.
6. Ensure subsequent statements are cautious and do not downplay the risk.
7. Avoid neutral/positive statements that contradict the warning.
8. If applicable, specify whether the flag is for: specific business unit/division, consolidated financials, standalone financials, or geographical region. Maintain professional financial terminology.
</instructions>
 
<context>
Input document:-
####
{flag_with_context}
####
 
Criteria for Risk identification:
%%%%
Criteria Name: {matched_criteria}
Risk Classification Reasoning: {reasoning}
Bucket Category: {bucket_name}
%%%%
 
Financial Metrics of the company from previous quarter:-
&&&&
{financial_data}
&&&&
 
</context>
 
<output_format>
high_risk_flag: yes if it is actually high risk after review, no otherwise.
high_risk_flag_summary: [if high risk, provide factual summary]
</output_format>
 
<review>
1. Ensure summary is exactly 1-2 lines and preserves all quantitative information
2. Confirm that all summaries are based solely on information from the input document context
3. Check that each summary maintains a cautious tone without downplaying risks
4. Ensure proper business unit/division specification where applicable
5. Verify that the summary uses professional financial terminology
6. Check that no speculative or interpretive language is used
7. Ensure all relevant exact numbers, percentages and dates from the context are preserved
8. Verify that the output follows the output format specified above
</review>"""
        
        try:
            response = self.llm.call(prompt, temperature=0.1)
            
            # Parse the response to extract high_risk_flag and summary
            lines = response.strip().split('\n')
            high_risk_flag = None
            high_risk_summary = None
            
            for line in lines:
                line = line.strip()
                if line.lower().startswith('high_risk_flag:'):
                    high_risk_value = line.split(':', 1)[1].strip().lower()
                    high_risk_flag = 'yes' in high_risk_value
                elif line.lower().startswith('high_risk_flag_summary:'):
                    high_risk_summary = line.split(':', 1)[1].strip()
                    # Clean up summary
                    high_risk_summary = re.sub(r'^\[|\]$', '', high_risk_summary).strip()
            
            # Only include if confirmed as high risk and has summary
            if high_risk_flag and high_risk_summary:
                return self._clean_summary(high_risk_summary)
            
        except Exception as e:
            logger.error(f"Error generating summary for flag '{classification.get('flag', 'Unknown')}': {e}")
            # Fallback summary
            return f"High risk identified: {matched_criteria}. Review required based on analysis."
        
        return None

    def _clean_summary(self, summary: str) -> str:
        """Clean and format summary text"""
        # Remove common prefixes that might appear
        prefixes_to_remove = ["Summary:", "The summary:", "Based on", "According to", "Analysis:", "Flag summary:", "The flag:", "This flag:"]
        for prefix in prefixes_to_remove:
            if summary.startswith(prefix):
                summary = summary[len(prefix):].strip()
        
        # Split into lines and take first 2
        summary_lines = [line.strip() for line in summary.split('\n') if line.strip()]
        
        if len(summary_lines) > 2:
            concise_summary = '. '.join(summary_lines[:2])
        elif len(summary_lines) == 0:
            concise_summary = "High risk identified. Requires management attention."
        else:
            concise_summary = '. '.join(summary_lines)
        
        # Ensure proper ending
        if not concise_summary.endswith('.'):
            concise_summary += '.'
        
        return concise_summary

# ==============================================================================
# DOCUMENT GENERATOR CLASS
# ==============================================================================

class DocumentGenerator:
    """Class for generating Word documents and reports"""
    
    def __init__(self, llm: AzureOpenAILLM):
        self.llm = llm
        self.summary_generator = SummaryGenerator(llm)

    def create_word_document(self, pdf_name: str, company_info: str, risk_counts: Dict[str, int],
                            classification_results: List[Dict[str, str]], summary_by_categories: Dict[str, List[str]], 
                            output_folder: str, previous_year_data: str) -> str:
        """Create a formatted Word document with concise high risk summaries"""
       
        try:
            doc = Document()
           
            # Document title
            title = doc.add_heading(company_info, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
           
            # Flag Distribution section
            flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
            flag_dist_heading.runs[0].bold = True
           
            # Create flag distribution table
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
           
            high_count = risk_counts.get('High', 0)
            low_count = risk_counts.get('Low', 0)
            total_count = high_count + low_count
           
            # Safely set table cells
            if len(table.rows) >= 3 and len(table.columns) >= 2:
                table.cell(0, 0).text = 'High Risk'
                table.cell(0, 1).text = str(high_count)
                table.cell(1, 0).text = 'Low Risk'
                table.cell(1, 1).text = str(low_count)
                table.cell(2, 0).text = 'Total Flags'
                table.cell(2, 1).text = str(total_count)
               
                # Make headers bold
                for i in range(3):
                    if len(table.cell(i, 0).paragraphs) > 0 and len(table.cell(i, 0).paragraphs[0].runs) > 0:
                        table.cell(i, 0).paragraphs[0].runs[0].bold = True
           
            doc.add_paragraph('')
           
            # High Risk Flags section with concise summaries
            high_risk_classifications = [result for result in classification_results if result['risk_level'] == 'High']
            if high_risk_classifications and len(high_risk_classifications) > 0:
                high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
                if len(high_risk_heading.runs) > 0:
                    high_risk_heading.runs[0].bold = True
               
                # Generate concise summaries for high risk flags using classification results
                concise_summaries = self.summary_generator.generate_strict_high_risk_summary(classification_results, previous_year_data)
                
                # Final deduplication check at Word document level
                final_unique_summaries = self._deduplicate_summaries(concise_summaries)
                
                for summary in final_unique_summaries:
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(summary)
            else:
                high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
                if len(high_risk_heading.runs) > 0:
                    high_risk_heading.runs[0].bold = True
                doc.add_paragraph('No high risk flags identified.')
           
            # Horizontal line
            doc.add_paragraph('_' * 50)
           
            # Summary section (4th iteration results)
            summary_heading = doc.add_heading('Summary', level=1)
            if len(summary_heading.runs) > 0:
                summary_heading.runs[0].bold = True
           
            # Add categorized summary
            if summary_by_categories and len(summary_by_categories) > 0:
                for category, bullets in summary_by_categories.items():
                    if bullets and len(bullets) > 0:
                        cat_heading = doc.add_heading(str(category), level=2)
                        if len(cat_heading.runs) > 0:
                            cat_heading.runs[0].bold = True
                       
                        for bullet in bullets:
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'
                            p.add_run(str(bullet))
                       
                        doc.add_paragraph('')
            else:
                doc.add_paragraph('No categorized summary available.')
           
            # Save document
            doc_filename = f"{pdf_name}_Report.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
           
            return doc_path
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return self._create_fallback_document(pdf_name, risk_counts, output_folder)

    def _deduplicate_summaries(self, summaries: List[str]) -> List[str]:
        """Final deduplication check for summaries"""
        final_unique_summaries = []
        seen_content = set()
        
        for summary in summaries:
            if not summary or not summary.strip():
                continue
                
            # Create multiple normalized versions for comparison
            normalized1 = re.sub(r'[^\w\s]', '', summary.lower()).strip()
            normalized2 = re.sub(r'\b(the|a|an|and|or|but|in|on|at|to|for|of|with|by)\b', '', normalized1)
            
            # Check if this content is substantially different
            is_unique = True
            for seen in seen_content:
                # Calculate similarity
                words1 = set(normalized2.split())
                words2 = set(seen.split())
                if len(words1) == 0 or len(words2) == 0:
                    continue
                similarity = len(words1.intersection(words2)) / len(words1.union(words2))
                if similarity > 0.6:  # If more than 60% similar, consider duplicate
                    is_unique = False
                    break
            
            if is_unique:
                final_unique_summaries.append(summary)
                seen_content.add(normalized2)
        
        return final_unique_summaries

    def _create_fallback_document(self, pdf_name: str, risk_counts: Dict[str, int], output_folder: str) -> str:
        """Create minimal document as fallback"""
        try:
            doc = Document()
            doc.add_heading(f"{pdf_name} - Analysis Report", 0)
            doc.add_paragraph(f"High Risk Flags: {risk_counts.get('High', 0)}")
            doc.add_paragraph(f"Low Risk Flags: {risk_counts.get('Low', 0)}")
            doc.add_paragraph(f"Total Flags: {risk_counts.get('Total', 0)}")
            
            doc_filename = f"{pdf_name}_Report_Fallback.docx"
            doc_path = os.path.join(output_folder, doc_filename)
            doc.save(doc_path)
            return doc_path
        except Exception as e2:
            logger.error(f"Error creating fallback document: {e2}")
            return None

# ==============================================================================
# RED FLAG IDENTIFICATION PIPELINE CLASS
# ==============================================================================

class RedFlagIdentificationPipeline:
    """Main pipeline class for red flag identification across multiple iterations"""
    
    def __init__(self, llm: AzureOpenAILLM):
        self.llm = llm
        self.keywords = KeywordDefinitions()
        self.prompt_generator = PromptGenerator()

    def run_red_flag_identification(self, context: str) -> str:
        """Run the 4-step red flag identification process"""
        print("Running red flag identification across 4 keyword groups...")
        
        # Step 1: Keywords Part 1
        system_prompt_1 = self.prompt_generator.get_red_flag_identification_prompt(
            self.keywords.get_keywords_part1()
        )
        response1 = self.llm.call_with_system_prompt(
            system_prompt_1, 
            f"%%%%{context}%%%%"
        )
        print("Step 1 completed")
        print("******************************************")
        
        # Step 2: Keywords Part 2
        system_prompt_2 = self.prompt_generator.get_red_flag_identification_prompt(
            self.keywords.get_keywords_part2()
        )
        response2 = self.llm.call_with_system_prompt(
            system_prompt_2, 
            f"%%%%{context}%%%%"
        )
        print("Step 2 completed")
        print("******************************************")
        
        # Step 3: Keywords Part 3
        system_prompt_3 = self.prompt_generator.get_red_flag_identification_prompt(
            self.keywords.get_keywords_part3()
        )
        response3 = self.llm.call_with_system_prompt(
            system_prompt_3, 
            f"%%%%{context}%%%%"
        )
        print("Step 3 completed")
        print("******************************************")
        
        # Step 4: Keywords Part 4
        system_prompt_4 = self.prompt_generator.get_red_flag_identification_prompt(
            self.keywords.get_keywords_part4()
        )
        response4 = self.llm.call_with_system_prompt(
            system_prompt_4, 
            f"%%%%{context}%%%%"
        )
        print("Step 4 completed")
        print("******************************************")
        
        # Combine all responses
        combined_response = (response1 + "\n" + response2 + "\n" + response3 + "\n" + response4)
        return combined_response

    def run_deduplication(self, first_response: str) -> str:
        """Run enhanced deduplication"""
        print("Running 2nd iteration - Enhanced Deduplication...")
        
        system_prompt = self.prompt_generator.get_deduplication_prompt()
        user_content = f"""<context>
first iteration analysis to deduplicate:
{first_response}
</context>

Provide deduplicated analysis with merged duplicates and preserved evidence:"""

        response = self.llm.call_with_system_prompt(
            system_prompt, 
            f"####{user_content}####"
        )
        return response

    def run_categorization(self, context: str, second_response: str) -> str:
        """Run categorization"""
        print("Running 3rd iteration - Categorization...")
        
        system_prompt = self.prompt_generator.get_categorization_prompt()
        user_content = f"""<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}
</context>

Provide categorized analysis:"""
        
        response = self.llm.call(user_content)
        return response

    def run_summary_generation(self, context: str, third_response: str) -> str:
        """Run summary generation"""
        print("Running 4th iteration - Summary Generation...")
        
        system_prompt = self.prompt_generator.get_summary_generation_prompt()
        user_content = f"""<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        response = self.llm.call(user_content)
        return response

# ==============================================================================
# MAIN PROCESSING PIPELINE CLASS
# ==============================================================================

class FinancialAnalysisPipeline:
    """Main pipeline orchestrator class"""
    
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.llm = AzureOpenAILLM(api_key, azure_endpoint, api_version, deployment_name)
        self.pdf_processor = PDFProcessor()
        self.flag_analyzer = FlagAnalyzer(self.llm)
        self.document_generator = DocumentGenerator(self.llm)
        self.red_flag_pipeline = RedFlagIdentificationPipeline(self.llm)

    def process_single_pdf(self, pdf_path: str, previous_year_data: str, output_folder: str) -> pd.DataFrame:
        """Process a single PDF through the complete pipeline"""
        
        FileUtils.ensure_directory(output_folder)
        pdf_name = Path(pdf_path).stem
       
        try:
            # Load PDF content
            docs = self.pdf_processor.merge_docs(pdf_path, split_pages=False)
            context = docs[0]["context"]
            
            # ITERATION 1: Red Flag Identification (4 steps)
            first_response = self.red_flag_pipeline.run_red_flag_identification(context)
            
            # ITERATION 2: Enhanced Deduplication
            second_response = self.red_flag_pipeline.run_deduplication(first_response)
                    
            # ITERATION 3: Categorization
            third_response = self.red_flag_pipeline.run_categorization(context, second_response)
            
            # ITERATION 4: Summary Generation
            fourth_response = self.red_flag_pipeline.run_summary_generation(context, third_response)
            
            # ITERATION 5: Efficient Bucket-Based Classification
            print("Running 5th iteration - Efficient Bucket-Based Classification...")
            
            classification_results, risk_counts = self._run_classification(second_response, previous_year_data)
            
            # Generate Word Document
            self._generate_word_document(pdf_path, pdf_name, classification_results, risk_counts, 
                                       fourth_response, output_folder, previous_year_data)
            
            # Save results to CSV
            results_summary = self._save_results_to_csv(pdf_name, first_response, second_response, 
                                                      third_response, fourth_response, classification_results, 
                                                      risk_counts, output_folder)
            
            print(f"\n=== PROCESSING COMPLETE FOR {pdf_name} ===")
            return results_summary
           
        except Exception as e:
            logger.error(f"Error processing {pdf_name}: {str(e)}")
            return None

    def _run_classification(self, second_response: str, previous_year_data: str) -> Tuple[List[Dict[str, str]], Dict[str, int]]:
        """Run the classification step"""
        try:
            flags_with_context = self.flag_analyzer.extract_flags_with_complete_context(second_response)
            print(f"\nFlags with context extracted: {len(flags_with_context)}")
            
            if flags_with_context:
                print(f"Example flag with context:\n{flags_with_context[0][:200]}...")
            
        except Exception as e:
            logger.error(f"Error parsing flags with context: {e}")
            flags_with_context = ["Error in flag parsing"]

        classification_results = []
        high_risk_flags = []
        low_risk_flags = []

        if len(flags_with_context) > 0 and flags_with_context[0] != "Error in flag parsing":
            try:
                print(f"Analyzing all {len(flags_with_context)} flags using 8 bucket calls.")
                
                bucket_results = self.flag_analyzer.classify_all_flags_with_enhanced_buckets(
                    flags_with_context, previous_year_data
                )
                classification_results = self.flag_analyzer.parse_bucket_results_to_classifications(
                    bucket_results, flags_with_context
                )

                for result in classification_results:
                    if (result['risk_level'].lower() == 'high' and 
                        result['matched_criteria'] != 'None'):
                        high_risk_flags.append(result['flag'])
                    else:
                        low_risk_flags.append(result['flag'])
                        
            except Exception as e:
                logger.error(f"Error in efficient bucket classification: {e}")
                # Create fallback classifications
                for flag_with_context in flags_with_context:
                    flag_description = flag_with_context.split('\n')[0]
                    flag_description = re.sub(r'^\d+\.\s+', '', flag_description).strip()
                    
                    classification_results.append({
                        'flag': flag_description,
                        'flag_with_context': flag_with_context,
                        'matched_criteria': 'None',
                        'risk_level': 'Low',
                        'reasoning': f'Classification failed: {str(e)}',
                        'bucket': 'Error'
                    })
                    low_risk_flags.append(flag_description)

        risk_counts = {
            'High': len(high_risk_flags),
            'Low': len(low_risk_flags),
            'Total': len(flags_with_context) if flags_with_context and flags_with_context[0] != "Error in flag parsing" else 0
        }
        
        print(f"\n=== CLASSIFICATION RESULTS ===")
        print(f"High Risk Flags: {risk_counts['High']}")
        print(f"Low Risk Flags: {risk_counts['Low']}")
        print(f"Total Flags: {risk_counts['Total']}")
        
        if high_risk_flags:
            print(f"\n--- HIGH RISK FLAGS ---")
            for i, flag in enumerate(high_risk_flags, 1):
                print(f"  {i}. {flag}")
        else:
            print(f"\n--- HIGH RISK FLAGS ---")
            print("  No high risk flags identified")
        
        return classification_results, risk_counts

    def _generate_word_document(self, pdf_path: str, pdf_name: str, classification_results: List[Dict[str, str]], 
                               risk_counts: Dict[str, int], fourth_response: str, output_folder: str, 
                               previous_year_data: str) -> None:
        """Generate Word document"""
        print("\nCreating Word document...")
        try:
            company_info = self.pdf_processor.extract_company_info_from_pdf(pdf_path, self.llm)
            summary_by_categories = self.flag_analyzer.parse_summary_by_categories(fourth_response)
        
            word_doc_path = self.document_generator.create_word_document(
                pdf_name=pdf_name,
                company_info=company_info,
                risk_counts=risk_counts,
                classification_results=classification_results,  
                summary_by_categories=summary_by_categories,
                output_folder=output_folder,
                previous_year_data=previous_year_data
            )
            
            if word_doc_path:
                print(f"Word document created: {word_doc_path}")
            else:
                print("Failed to create Word document")
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")

    def _save_results_to_csv(self, pdf_name: str, first_response: str, second_response: str, 
                           third_response: str, fourth_response: str, classification_results: List[Dict[str, str]], 
                           risk_counts: Dict[str, int], output_folder: str) -> pd.DataFrame:
        """Save all results to CSV files"""
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Red Flags",
                "Enhanced Deduplication",
                "Categorization",
                "Summary Generation", 
                "Enhanced Context-Based Classification"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Enhanced Context-Based Classification: {risk_counts['High']} High Risk, {risk_counts['Low']} Low Risk flags from {risk_counts['Total']} total flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        if len(classification_results) > 0:
            classification_df = pd.DataFrame(classification_results)
            classification_file = os.path.join(output_folder, f"{pdf_name}_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        return results_summary

    def process_multiple_pdfs(self, pdf_folder_path: str, previous_year_data: str, output_folder: str) -> None:
        """Process multiple PDFs in a folder"""
        FileUtils.ensure_directory(output_folder)
        
        # Get all PDF files
        import glob
        pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
        if not pdf_files:
            print(f"No PDF files found in {pdf_folder_path}")
            return    

        for i, pdf_file in enumerate(pdf_files, 1):
            print(f"\n{'='*60}")
            print(f"PROCESSING PDF {i}/{len(pdf_files)}: {os.path.basename(pdf_file)}")
            print(f"{'='*60}")
            
            start_time = time.time()
            
            result = self.process_single_pdf(
                pdf_path=pdf_file,
                previous_year_data=previous_year_data,
                output_folder=output_folder
            )
            
            processing_time = time.time() - start_time
            
            if result is not None:
                print(f"✅ Successfully processed {pdf_file} in {processing_time:.2f} seconds")
            else:
                print(f"❌ Failed to process {pdf_file}")
