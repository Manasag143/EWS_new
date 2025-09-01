import os
import time
import pandas as pd
import fitz  
import tempfile
import warnings
import hashlib
from logger_config import logger
from io import BytesIO
from typing import Dict, List, Any
from pathlib import Path
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from openai import AzureOpenAI
import httpx
from constants import *
 
warnings.filterwarnings('ignore')

def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()

# ==============================================================================
# AZURE OPENAI LLM CLASS
# ==============================================================================

class AzureOpenAILLM:
    """Azure OpenAI gpt-4.1 LLM class"""
   
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str = "gpt-4.1"):
        self.deployment_name = deployment_name
        httpx_client = httpx.Client(verify=False)
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
            http_client=httpx_client
        )
   
    def _call(self, prompt: str, max_tokens: int = None, temperature: float = 0.1) -> str:
        """Make API call to Azure OpenAI gpt-4.1"""
        try:
            kwargs = {
                "model": self.deployment_name,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "top_p": 0.95,
                "frequency_penalty": 0,
                "presence_penalty": 0
            }
            
            if max_tokens:
                kwargs["max_tokens"] = max_tokens
                
            response = self.client.chat.completions.create(**kwargs)
            
            response_text = response.choices[0].message.content
            return response_text.strip() if response_text else ""
           
        except Exception as e:
            logger.error(f"Azure OpenAI API call failed: {str(e)}")
            return f"Azure OpenAI Call Failed: {str(e)}"

# ==============================================================================
# PDF PROCESSING FUNCTIONS
# ==============================================================================


keywords_part1 = """
Attrition: Refers to the increasing or high loss of employees, customers, or revenue due to various reasons such as resignation, retirement, or competition, which can negatively impact a company's financial performance. 
Adverse: Describes an unfavorable or negative situation, event, or trend, such as adverse market conditions or regulatory changes. 
Cautious outlook: Indicates a company's conservative or pessimistic view of its future financial performance, often due to uncertainty or potential risks. 
Challenging environment: Refers to a difficult or competitive market situation.
Competition intensifying: Describes an increase in competition in a market or industry, which can lead to decreased market share, revenue, or profitability for a company. 
Corporate governance: Refers to the system of rules, practices, and processes by which a company is directed and controlled, including issues related to board composition, executive compensation, and audit committee independence.
Cost inflation: Describes an increase in costs, such as labor, materials, or overheads. 
Customer confidence: Refers to the level of trust and faith that customers have in a company's products or services, which can impact sales and revenue. 
Debt repayment challenges: Describes difficulties a company faces in repaying its debt, which can lead to default, restructuring, or other negative consequences. 
Decline: Describes a decrease in a company's financial performance, such as revenue, profitability, or market share. 
Delay: Refers to a postponement or deferral of a project, investment, or other business initiative, which can impact a company's financial performance.
Group company exposure: Describes a company's financial exposure to its subsidiaries, affiliates, or joint ventures, which can impact its consolidated financial performance.
Guidance revision: Refers to a change in a company's financial guidance
Impairment charges: Refers to non-cash charges taken by a company to reflect the decline in value of its assets, such as goodwill, property, or equipment. 
Increase provisions: Describes an increase in a company's provisions for bad debts, warranties, or other contingent liabilities.
Increasing working capital: Describes an increase in a company's working capital requirements, such as accounts receivable, inventory, or accounts payable.
"""
keywords_part2="""
Inventory levels gone up: Refers to an increase in a company's inventory levels, which can indicate slower sales, overproduction, or supply chain disruptions.
Liquidity concerns: Describes a company's difficulties in meeting its short-term financial obligations, such as paying debts or meeting working capital requirements.
Margin pressure: Describes a decline in a company's profit or EBIDTA margins.
New management: Refers to the appointment of new executives or managers to a company's leadership team, which can impact its strategy, culture, and financial performance.
One-off expenses: Refers to non-recurring expenses or charges taken by a company, such as restructuring costs, impairment charges, or litigation expenses.
One-time write-offs: Refers to non-recurring write-offs or charges taken by a company, such as asset impairments, inventory write-offs, or accounts receivable write-offs.
Operational issues: Describes challenges or problems a company faces in its operations.
Regulatory uncertainty: Describes uncertainty or ambiguity related to regulatory requirements, laws, or policies.
Related party transaction: Refers to a transaction between a company and its related parties, such as subsidiaries, affiliates, or joint ventures, which can impact its financial performance and transparency.
Restructuring efforts: Refers to a company's plans or actions to reorganize its operations, finances, or management structure to improve its performance, efficiency, or competitiveness.
Scale down: Describes a company's decision to reduce its operations, investments, or workforce to conserve resources, cut costs, or adapt to changing market conditions.
Service issue: Refers to problems or difficulties a company faces in delivering its services.
Shortage: Describes a situation where a company faces a lack of supply, resources, or personnel.
Stress: Refers to a company's financial difficulties or challenges, such as debt, cash flow problems etc.
Supply chain disruptions: Refers to interruptions or problems in a company's supply chain, which can impact its ability to produce, deliver, or distribute its products or services.
Warranty cost: Refers to the expenses or provisions a company makes for warranties or guarantees provided to its customers.
Misappropriation of funds: Describes the unauthorized or improper use of a company's funds, assets, or resources.
"""
keywords_part3 = """
Increase in borrowing cost: Refers to a rise in the cost of borrowing for a company. 
One time reversal: Describes a non-recurring or one-time adjustment to a company's financial statements.
Bloated balance sheet: Refers to a company's balance sheet that is overly leveraged, inefficient, or burdened with debt.
Reversal: a credit or refund to the customer, which reduces the original sale and is recorded as a reduction in revenue.
Debtors increasing or going up: Refers to an increase in a company's accounts receivable or debtors.
Receivables increase: Describes an increase in a company's accounts receivable.
Challenges in collections: Refers to difficulties a company faces in collecting its accounts receivable or debtors, which can impact its cash flow, liquidity, or financial performance. 
Slow down on disbursement: A reduction in the rate at which loans or funds are disbursed.
Write-offs: The process of removing a debt or asset from a company's balance sheet or Profit and loss statement.
Increase of provisioning: An increase in the amount of money set aside by a financial institution to cover potential losses on loans or assets.
Delinquency increase: A rise in the number of borrowers who are late or behind on their loan payments, often indicating a deterioration in credit quality.
GNPA increasing: An increase in Gross Non-Performing Assets (GNPA), which refers to the total value of loans that are overdue or in default.
Slippages: The reclassification of loans from a performing to a non-performing category
High credit deposit ratio: A situation where a bank's credit growth exceeds its deposit growth.
CAR decreasing: A decline in the Capital Adequacy Ratio (CAR), which measures a bank's capital as a percentage of its risk-weighted assets.
"""
keywords_part4 = """
Provision coverage falling: A decline in the provision coverage ratio, indicating that the provisions made for potential losses are decreasing relative to the growth in non-performing assets.
Low Profitability: A state where a business, project, or investment generates revenue, but the net income or return on investment (ROI) is significantly lower than expected, industry average, or benchmark. 
Falling Net Interest Margin (NIM): A decrease in the difference between the interest income earned by a financial institution and the interest expense paid on deposits and other borrowings due to changes in interest or deposit rate, reduced profitability etc.
Negative Capital Employed: Statements that indicate a company's liabilities exceed its assets, or its return on capital employed is negative.
Capacity Utilisation falling: Refers to the extent to which a company's production facilities or resources are being used, with low utilisation indicating underproduction or declining demand.
Destocking: The process of reducing inventory levels, often due to decreased demand or overstocking, which can indicate a decline in sales or shift in market trends.
Pricing Pressure: Downward pressure on a company's prices due to competition or market conditions.
Renegotiation: The process of revising or re-evaluating existing contracts or agreements, which can indicate disputes or changes in market conditions.
Credit rating action/Rating downgrade/Watch negative: A change in a company's credit rating, indicating a higher risk of default or negative outlook.
Weakening/softening of demand: A decline in customer demand or slowdown in sales growth, indicating a decline in market share or shift in market trends.
Long recovery time: A prolonged period required for a company to recover from a downturn or disruption, indicating significant challenges or reduced competitiveness.
Capex plan mentioned but no roadmap/clarity of funding: A capital expenditure plan without a clear plan for funding or implementation, indicating a lack of financial resources or unclear priorities.
Loss: A financial loss incurred by a company, indicating poor financial management or reduced competitiveness.
Anti-dumping: Measures taken to prevent the importation of goods at below-normal prices, which can indicate trade tensions or protectionism.
Demerger: The separation of a company into independent entities, often to improve focus or reduce complexity, but can also indicate a lack of synergy or decline in profitability.
"""

class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            doc.close()
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]

def extract_company_info_from_pdf(pdf_path: str, llm: AzureOpenAILLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
       
        first_page_text = first_page_text[:2000]
       
        prompt = f"""<role>
You are an expert document analyst specializing in extracting key corporate information from financial documents and earnings transcripts.
</role>

<system_prompt>
You excel at quickly identifying and extracting specific corporate identifiers from financial documents with high accuracy and consistency.
</system_prompt>

<instruction>
Extract the company name, quarter, and financial year from the provided text.

EXTRACTION REQUIREMENTS:
1. Company Name: Full legal company name including suffixes (Ltd/Limited/Inc/Corp etc.)
2. Quarter: Identify the quarter (Q1/Q2/Q3/Q4)  
3. Financial Year: Extract financial year (FY23/FY24/FY25 etc.)

OUTPUT FORMAT:
Provide ONLY the result in this exact format: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If any component cannot be clearly identified, use reasonable defaults based on context.
</instruction>

<context>
DOCUMENT TEXT TO ANALYZE:
{first_page_text}
</context>

Extract company information:"""
       
        response = llm._call(prompt)
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
       
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
       
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

# ==============================================================================
# CRITERIA BUCKETS AND CLASSIFICATION FUNCTIONS
# ==============================================================================

def create_criteria_buckets():
    # QUANTITATIVE BUCKETS (6 buckets)
    # Bucket 1: Core Debt & Leverage (Quantitative)
    bucket_1_quant = [    
        """debt_increase: 
            High: Debt is increased for the company or any business line of the company more than 30% compared to previous reported balance sheet number; 
            Low: Debt increased less than 30% compared to previous reported balance sheet number""",
        """debt_ebitda: 
            High: Debt/EBITDA > 3x i.e. Debt to EBITDA ratio is above (greater than) three times for the company or any business line of the company; 
            Low: Debt/EBITDA < 3x i.e. Debt to EBITDA ratio is less than three times""",
        """short_term_borrowings: 
            High: Short-term borrowings or current liabilities increase by more than 30% compared to previous reported balance sheet number for the company or any business line of the company; 
            Low: Short-term borrowings or current liabilities increase is less than 30% compared to previous reported balance sheet number""",
        """debt-servicing: 
            High: Calculate the ratio of debt servicing to net cash accrual and provide a high-risk flag if the calculated ratio is greater than 1.5; 
            Low: Debt servicing to net cash accrual ratio is 1.5 or less"""
    ]

    # Bucket 2: Profitability & Performance (Quantitative)
    bucket_2_quant = [
        """revenue_decline: 
            High: Revenue falls for the company or any business line of the company by more than 25% compared to previous reported quarter number; 
            Low: Revenue falls by less than 25% compared to previous reported quarter number""",
        """profit_before_tax_decline: 
            High: Profitability or profit before tax (PBT) falls by more than 25% compared to previous reported quarter number for the company or any business line of the company; 
            Low: Profitability or profit before tax (PBT) falls by less than 25% compared to previous reported quarter number""",
        """profit_after_tax_decline: 
            High: Profit after tax (PAT) falls by more than 25% compared to previous reported quarter number for the company or any business line of the company; 
            Low: Profit after tax (PAT) falls by less than 25% compared to previous reported quarter number""",
        """EBIDTA_decline: 
            High: EBITDA falls by more than 25% compared to previous reported quarter number for the company or any business line of the company; 
            Low: EBITDA falls by less than 25% compared to previous reported quarter number"""
    ]
    
    # Bucket 3: Margins & Operational Efficiency (Quantitative)
    bucket_3_quant = [
        """margin_decline: 
            High: Operating margin falling more than 25% compared to previous reported quarter number for the company or any business line of the company; 
            Low: Operating margin falling less than 25% compared to previous reported quarter number""",
        """gross_margin: 
            High: Gross margin falling more than 300 bps (basis points) i.e. if the gross margin is falling by more than 3% for the company or any business line of the company; 
            Low: Gross margin falling by less than 300 bps (basis points) i.e.3%""",
        """cash_and_cash_equivalents_balance: 
            High: Cash and cash equivalents balance for the company or any business line of the company falling more than 25% compared to previous reported balance sheet number; 
            Low: Cash and cash equivalents balance falling less than 25% compared to previous reported balance sheet number""",
        """others_4: 
            High: If there is number comparison provided in transcript which is beyond severity logic, and if decline/moderation is more than 25%; 
            Low: If decline/moderation is 25% or less, or no concerning comparisons are provided"""
    ]
    
    # Bucket 4: Working Capital & Asset Management (Quantitative)
    bucket_4_quant = [
        """receivable_days: 
            High: Receivable days OR debtor days for the company or any business line of the company are increased more than 30% compared to previous reported balance sheet number; 
            Low: Receivable days or debtor's days are increased but less than 30% compared to previous reported balance sheet number""",
        """payable_days: 
            High: Payable days or creditors days for the company or any business line of the company increase by more than 30% compared to previous reported balance sheet number; 
            Low: Payable days or creditors days increase is less than 30% compared to previous reported balance sheet number""",
        """receivables: 
            High: Receivables or debtors for the company or any business line of the company are increased more than 30% compared to previous reported balance sheet number; 
            Low: Receivables or debtors are increase is less than 30% compared to previous reported balance sheet number""",
        """payables: 
            High: Payables or creditors for the company or any business line of the company increase by more than 30% compared to previous reported balance sheet number; 
            Low: Payables or creditors is less than 30% compared to previous reported balance sheet number"""
    ]
    
    # Bucket 5: Asset Quality & Impairments (Quantitative)
    bucket_5_quant = [
        """asset_decline: 
            High: Asset value for the company or any business line of the company falls by more than 30% compared to the previous reported balance sheet number; 
            Low: Asset value falls by less than 30% compared to previous reported balance sheet number""",
        """impairment: 
            High: Impairment or devaluation more than 25% of previous reported net worth from balance sheet for the company or any business line of the company; 
            Low: Impairment or devaluation less than 25% of previous reported net worth from balance sheet""",
        """one-time_expenses: 
            High: One-time expenses or losses more than 25% of current quarter's EBITDA for the company or any business line of the company; 
            Low: One-time expenses or losses less than 25% of current quarter's EBITDA""",
        """provisioning: 
            High: Provisioning or write-offs more than 25% of current quarter's EBITDA for the company or any business line of the company; 
            Low: Provisioning or write-offs less than 25% of current quarter's EBITDA"""
    ]
    
    # Bucket 6: Other Quantitative Risks (Quantitative)
    bucket_6_quant = [
        """others_1: 
            High: If there are any other material issues with quantified impact more than 25% of current quarter EBITDA; 
            Low: If there are no other material issues with quantified impact more than 25% of current quarter EBITDA""",
        """others_5: 
            High: If there is increase in amount for 'Receivables-due for more than 6 months or one year'; 
            Low: No significant increase in long-term receivables""",
        """others_6: 
            High: If transcript mentions indemnity assets, such as "indemnity claims" or "indemnity receivable" more than 25% of net worth; 
            Low: Indemnity assets less than 25% of net worth or not mentioned"""
    ]
    # ENHANCED QUALITATIVE BUCKETS WITH FEW-SHOT EXAMPLES (2 buckets)
    # Bucket 7: Management & Regulatory Issues (Qualitative) - ENHANCED
    bucket_7_qual = [
        """management_issues: 
            High: If found any management or strategy related issues or concerns in a question or a conclusion of any discussion. This can include any delays, failures, dysfunction, instability etc. due to management / strategy / governance.
            <examples>
            high risk sample 1: Management not able to improve margin and hence not being able to expand as planned in Middle East
            high risk sample 2: Leadership transition has caused delays in executing the new product roadmap, impacting revenue targets
            high risk sample 3: Strategic shift away from core markets has resulted in declining customer retention
            </examples>

            Low: If found no issues related to management or strategy or no concerns or a conclusion of any discussion related to management and strategy""",
            
        """regulatory_compliance: 
            High: If found any regulatory issues as a concern or a conclusion of any discussion. This includes warnings from regulators, delays in obtaining permits, approvals, licenses, or clearances, and any regulatory challenges affecting operations, transactions, or strategic initiatives.
            <examples>
            high risk sample 1: Approval processes for international customers' qualification of electrolyte salts have experienced delays due to the need to achieve stable and uniform production quality meeting stringent standards
            high risk sample 2: Q4 faced multiple regulatory challenges related to store launches, point-of-sale etc
            high risk sample 3: The sale of the steel plant is delayed pending regulatory clearances expected by Q1 FY '25, with lender NOCs for the vertical split also pending
            </examples>

            Low: If there is no clear concern for the company basis the discussion on the regulatory issues""",
            
        """market_competition: 
            High: Identify any signs of competitive intensity, new entrants, pricing pressure (including dumping or price changes), or decline in market share. Also include indirect competitive risks such as supply-side constraints, delays, or resource shortages that impact market performance. Additionally, capture macroeconomic factors—such as adverse currency movements, inflation, or trade barriers—that materially affect the company's competitive position or pricing power.
            
            <examples>
            high risk sample 1: Local competition has intensified with increased branding and promotional activities by regional players
            high risk sample 2: Advanced intermediates face challenges due to aggressive low-priced imports from China, impacting domestic production and pricing
            high risk sample 3: Adverse movement in exchange rate: Further there was a translation loss due to adverse movement in exchange rate between the USD and the INR and the AUD, INR compared to March 2022
            high risk sample 4: Pricing pressure and competition, especially from China, have led to lower realizations despite volume growth of about 20%, resulting in only 9% revenue growth
            </examples>

            Low: Low competitive intensity or no new entrants, or no decline in market share""",
            
        ]
    
    # Bucket 8: Qualitative Risk Indicators (Qualitative) - ENHANCED
    bucket_8_qual = [
        """others_3: 
            High: Identify any mention of severe financial deterioration using strong adjectives. This includes phrases such as "significant decline" in key metrics (revenue, EBITDA, PAT/PBT, margins), "high erosion of net worth," "massive losses," or similar expressions that imply material financial stress. Include such mentions even if the exact quantum is not provided. Also capture other high-impact adjectives like "sharp drop," "steep fall," "substantial deterioration," "severe contraction," or "critical pressure" when used in the context of financial performance.
            <examples>
            high risk sample 1: We reported massive losses in Q3, primarily driven by one-time restructuring costs and adverse market conditions
            high risk sample 2: There has been a high erosion of net worth following the impairment of legacy assets and continued losses in the international business
            high risk sample 3: The company experienced a significant decline in EBITDA due to underperformance in its core segment
            </examples>

            Low: No mentions of significant decline, erosion, massive losses, huge decline, net loss, or adverse adjectives""",

        """operational_disruptions: 
            High: Identify any operational challenges or supply chain issues mentioned as a concern or conclusion. This includes disruptions due to labor shortages, productivity losses, client-side delays, weather-related impacts, or any other factors that materially affect execution, delivery timelines.
            <examples>
            high risk sample 1: In the last two quarters, we have had a bit of operational challenges, first with service and then with registrations
            high risk sample 2: Labor costs increased due to shortage of labor supply and in Australia, labor cost, site and site overheads increased due to loss of productivity on account of extreme weather conditions
            high risk sample 3: Clients delaying final handover
            high risk sample 4: The lack of equipment availability has led to continuous projected delays, thereby decreasing the 2022 forecast by 600 megawatts to 8.1 gigawatts, the lowest annual total since 2018
            </examples>

            Low: If there is no clear concern for the company basis the discussion on the operational or supply chain issues"""
        ]
    
    return [bucket_1_quant, bucket_2_quant, bucket_3_quant, bucket_4_quant, 
            bucket_5_quant, bucket_6_quant, bucket_7_qual, bucket_8_qual]

def create_previous_data_buckets(previous_year_data: str):
    """Organize previous year data into 8 buckets matching the criteria buckets"""
    
    # Parse the previous year data to extract relevant metrics for each bucket
    lines = previous_year_data.strip().split('\n')
    data_dict = {}
    
    for line in lines:
        if line.strip():
            parts = line.split('\t')
            if len(parts) >= 3:
                metric = parts[0].strip()
                value = '\t'.join(parts[1:]).strip()
                data_dict[metric.lower()] = f"{metric}\t{value}"
    
    # Bucket 1: Core Debt & Leverage (Quantitative)
    bucket_1_data = ""
    for key in ['debt as per previous reported balance sheet number', 'current quarter ebitda', 'ebitda as per previous reported quarter number', 'short term borrowings as per the previous reported balance sheet number','previous quarter net cash accrual(NCA)']:
        if key in data_dict:
            bucket_1_data += data_dict[key] + "\n"
    
    # Bucket 2: Profitability & Performance (Quantitative)
    bucket_2_data = ""
    for key in ['revenue as per previous reported quarter number', 'profit before tax as per previous reported quarter number', 'profit after tax as per previous reported quarter number', 'ebitda as per previous reported quarter number']:
        if key in data_dict:
            bucket_2_data += data_dict[key] + "\n"
    
    # Bucket 3: Margins & Operational Efficiency (Quantitative)
    bucket_3_data = ""
    for key in ['operating margin as per previous quarter number', 'current quarter ebitda', 'cash balance as per previous reported balance sheet number']:
        if key in data_dict:
            bucket_3_data += data_dict[key] + "\n"
    
    # Bucket 4: Working Capital & Asset Management (Quantitative)
    bucket_4_data = ""
    for key in ['receivable days as per previous reported balance sheet number', 'payable days as per previous reported balance sheet number', 'receivables as per previous reported balance sheet number', 'payables as per previous reported balance sheet number']:
        if key in data_dict:
            bucket_4_data += data_dict[key] + "\n"
    
    # Bucket 5: Asset Quality & Impairments (Quantitative)
    bucket_5_data = ""
    for key in ['asset value as per previous reported balance sheet number', 'previous reported net worth from balance sheet', 'current quarter ebitda']:
        if key in data_dict:
            bucket_5_data += data_dict[key] + "\n"
    
    # Bucket 6: Other Quantitative Risks (Quantitative)
    bucket_6_data = ""
    for key in ['current quarter ebitda', 'previous reported net worth from balance sheet', 'revenue as per previous reported quarter number']:
        if key in data_dict:
            bucket_6_data += data_dict[key] + "\n"
    
    # Bucket 7: Management & Regulatory Issues (Qualitative) - No specific financial data needed
    bucket_7_data = "No specific financial metrics required for qualitative analysis"
    
    # Bucket 8: Qualitative Risk Indicators (Qualitative) - No specific financial data needed
    bucket_8_data = "No specific financial metrics required for qualitative analysis"
    
    return [bucket_1_data, bucket_2_data, bucket_3_data, bucket_4_data, 
            bucket_5_data, bucket_6_data, bucket_7_data, bucket_8_data]

def classify_all_flags_with_enhanced_buckets(all_flags_with_context: List[str], previous_year_data: str, llm: AzureOpenAILLM) -> Dict[str, List[Dict[str, str]]]:
    """
    Enhanced classification using 8 total LLM calls for all flags combined - one call per bucket
    """
    
    criteria_buckets = create_criteria_buckets()
    data_buckets = create_previous_data_buckets(previous_year_data)
    
    bucket_names = [
        "Core Debt & Leverage (Quantitative)",
        "Profitability & Performance (Quantitative)", 
        "Margins & Operational Efficiency (Quantitative)",
        "Working Capital & Asset Management (Quantitative)",
        "Asset Quality & Impairments (Quantitative)",
        "Other Quantitative Risks (Quantitative)",
        "Management & Regulatory Issues (Qualitative)",
        "Qualitative Risk Indicators (Qualitative)"
    ]
    
    # Prepare all flags text for analysis with clear numbering
    all_flags_text = ""
    for i, flag in enumerate(all_flags_with_context, 1):
        all_flags_text += f"\n--- FLAG_{i} ---\n{flag}\n"
    
    bucket_results = {}
    
    for i, (criteria_bucket, data_bucket, bucket_name) in enumerate(zip(criteria_buckets, data_buckets, bucket_names)):
        criteria_list = "\n\n".join(criteria_bucket)
        
        # Different prompts for quantitative vs qualitative buckets
        if "Quantitative" in bucket_name:
            prompt = f"""You are an experienced financial analyst. Your goal is to classify given red flags gathered from earnings call transcript document into High/Low Risk based on QUANTITATIVE thresholds.
 
Red Flags to be analyzed:-
{all_flags_text}
 
High/Low Risk identification criteria (QUANTITATIVE - focus on numbers and percentages):-
{criteria_list}
 
Financial Metrics of the company needed for analysis:-
{data_bucket}
 
<instructions>
1. Review each flag against the above given QUANTITATIVE criteria and the financial metrics.
2. Classify ONLY the red flags that match the criteria in this bucket.
3. For each matching flag, determine if it's High or Low risk based on the EXACT numerical thresholds in the criteria.
4. Use the exact flag numbering format: FLAG_1, FLAG_2, etc.
5. Focus on specific numbers, percentages, ratios mentioned in the flags.
6. Echo back the numbers you gave from the red flags and the financial metrics for accuracy check.
7. Use accurate formula for number comparisons.
8. Important: If a single flag matches multiple criteria, create separate output sections for each matching criterion. Do not combine them.
9. If no flags match the criteria in this bucket, respond with "No flags match the criteria in this bucket."
</instructions>
 
Output format - For each matching flag:
Flag_Number: FLAG_X (where X is the flag number)
Flag_Name: The red flag name
Matched_Criteria: [exact criteria name from the criteria list-only one criteria per output]
Risk_Level: [High or Low]
Reasoning: [brief explanation with specific numbers/evidence from the flag and financial metrics]
Relevant_Financial: [extract all the relevant financial metrics if high risk is identified else NA]

<example>
Example: If FLAG_1 matches both "debt_increase" and "debt_ebitda" criteria, provide TWO separate outputs:

Flag_Number: FLAG_1
Flag_Name: Company debt concerns
Matched_Criteria: debt_increase
Risk_Level: High
Reasoning: Debt increased by 40% exceeding 30% threshold
Relevant_Financials: Previous debt: 446Cr, Current debt: 624Cr

Flag_Number: FLAG_1  
Flag_Name: Company debt concerns
Matched_Criteria: debt_ebitda
Risk_Level: High
Reasoning: Debt/EBITDA ratio is 4.2x, exceeding 3x threshold
Relevant_Financials: Debt: 624Cr, EBITDA: 150Cr
</example>

<review>
1. Only analyze flags that specifically match the QUANTITATIVE criteria in this bucket.
2. Use exact flag numbering: FLAG_1, FLAG_2, FLAG_3, etc.
3. Ensure risk level determination follows the exact numerical thresholds in the criteria.
4. If a flag doesn't match any criteria in this bucket, don't include it in the output.
</review>
"""
        else:  # Qualitative bucket
            prompt = f"""You are an experienced financial analyst. Your goal is to classify given red flags gathered from earnings call transcript document into High/Low Risk based on QUALITATIVE indicators.
 
Red Flags to be analyzed:-
{all_flags_text}
 
High/Low Risk identification criteria (QUALITATIVE - focus on concerns, issues, and strategic matters):-
{criteria_list}
 
<instructions>
1. Review each flag against the above given QUALITATIVE criteria.
2. Classify ONLY the red flags that match the criteria in this bucket.
3. For each matching flag, determine if it's High or Low risk based on the presence/absence of concerns mentioned in the criteria.
4. Use the exact flag numbering format: FLAG_1, FLAG_2, etc.
5. Focus on management issues, regulatory concerns, operational problems, and strategic uncertainties.
6. Refer to the sample examples provided in criteria_list to help identify high risk flags accurately.
7. IMPORTANT: If a single flag matches multiple criteria, create SEPARATE outputs for each criteria match. Do not combine them.
8. If no flags match the criteria in this bucket, respond with "No flags match the criteria in this bucket."
</instructions>
 
Output format - For each matching flag and each matching criteria:
Flag_Number: FLAG_X (where X is the flag number)
Flag_Name: The Red Flag name
Matched_Criteria: [exact single criteria name from the criteria list - only one criteria per output]
Risk_Level: [High or Low]
Reasoning: [brief explanation for the specific criteria fulfilled with evidence from the flag about the qualitative concern]
Relevant_Financials: [extract all the relevant financial metrics if high risk is identified else NA]

<review>
1. Create separate outputs for each criteria match - do not combine multiple criteria in one output.
2. Use exact flag numbering: FLAG_1, FLAG_2, FLAG_3, etc.
3. Ensure each output has only ONE criteria in Matched_Criteria field.
4. If a flag matches 2 criteria, provide 2 separate outputs.
5. Extract relevant financial data for high risk qualitative flags where available.
</review>
"""

        try:
            logger.info(f"Analyzing all flags against {bucket_name}...")
            response = llm._call(prompt, temperature=0.0)
            bucket_results[bucket_name] = response
            
        except Exception as e:
            logger.error(f"Error analyzing {bucket_name}: {e}")
            bucket_results[bucket_name] = f"Error in {bucket_name}: {str(e)}"
    
    return bucket_results

def parse_bucket_results_to_classifications_enhanced(bucket_results: Dict[str, str], all_flags_with_context: List[str]) -> List[Dict[str, str]]:
    """
    Parse bucket results with separate outputs for each criteria match
    """
    flag_classifications = []
    
    for bucket_name, bucket_response in bucket_results.items():
        if isinstance(bucket_response, str) and "No flags match" not in bucket_response and "Error" not in bucket_response:
            
            # Split response into individual flag entries
            sections = re.split(r'\n\s*(?=Flag_Number:\s*FLAG_\d+)', bucket_response.strip())
            
            for section in sections:
                if not section.strip():
                    continue
                
                # Initialize variables for each section
                flag_number = None
                flag_name = None
                matched_criteria = None
                risk_level = None
                reasoning = None
                relevant_financials = None
                
                # Parse each line in the section
                lines = section.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    
                    if line.startswith('Flag_Number:'):
                        flag_number_text = line.replace('Flag_Number:', '').strip()
                        flag_match = re.search(r'FLAG_(\d+)', flag_number_text)
                        if flag_match:
                            flag_number = int(flag_match.group(1))
                            
                    elif line.startswith('Flag_Name:'):
                        flag_name = line.replace('Flag_Name:', '').strip()
                        
                    elif line.startswith('Matched_Criteria:'):
                        matched_criteria = line.replace('Matched_Criteria:', '').strip()
                        # Clean up any brackets
                        matched_criteria = re.sub(r'^\[|\]$', '', matched_criteria).strip()
                        
                    elif line.startswith('Risk_Level:'):
                        risk_level_text = line.replace('Risk_Level:', '').strip()
                        # Clean up any brackets
                        risk_level_text = re.sub(r'^\[|\]$', '', risk_level_text).strip()
                        if 'High' in risk_level_text:
                            risk_level = 'High'
                        elif 'Low' in risk_level_text:
                            risk_level = 'Low'
                            
                    elif line.startswith('Reasoning:'):
                        reasoning = line.replace('Reasoning:', '').strip()
                        # Clean up any brackets
                        reasoning = re.sub(r'^\[|\]$', '', reasoning).strip()
                        
                    elif line.startswith('Relevant_Financials:'):
                        relevant_financials = line.replace('Relevant_Financials:', '').strip()
                        # Clean up any brackets
                        relevant_financials = re.sub(r'^\[|\]$', '', relevant_financials).strip()
                
                # Create classification entry if we have all required fields
                if (flag_number is not None and flag_name and matched_criteria and 
                    risk_level and reasoning and 1 <= flag_number <= len(all_flags_with_context)):
                    
                    flag_index = flag_number - 1
                    original_flag = all_flags_with_context[flag_index]
                    
                    flag_classifications.append({
                        'original_flag_number': flag_number,
                        'flag': f"{flag_name} [Criteria: {matched_criteria}]",
                        'flag_with_context': original_flag,
                        'matched_criteria': matched_criteria,
                        'risk_level': risk_level,
                        'reasoning': reasoning,
                        'relevant_financials': relevant_financials if relevant_financials and relevant_financials.lower() != 'na' else 'NA',
                        'bucket': bucket_name  # This comes from the loop variable
                    })

                    logger.info(f"Parsed: FLAG_{flag_number} - {matched_criteria} - {risk_level} risk in {bucket_name}")

                else:
                    # Debug incomplete parsing
                    if flag_number is not None:
                        logger.warning(f"Debug: Incomplete parsing for FLAG_{flag_number} in {bucket_name}")
                        logger.warning(f"  Flag Name: {flag_name}")
    
    # Add unmatched flags as Low risk
    matched_flag_numbers = set(c['original_flag_number'] for c in flag_classifications)
    for i, flag_with_context in enumerate(all_flags_with_context, 1):
        if i not in matched_flag_numbers:
            # Extract flag description
            flag_lines = flag_with_context.strip().split('\n')
            flag_description = flag_lines[0] if flag_lines else flag_with_context
            flag_description = re.sub(r'^\d+\.\s*', '', flag_description).strip()
            
            flag_classifications.append({
                'original_flag_number': i,
                'flag': f"{flag_description} [No Match]",
                'flag_with_context': flag_with_context,
                'matched_criteria': 'None',
                'risk_level': 'Low',
                'reasoning': 'No matching criteria found across all buckets',
                'relevant_financials': 'NA',
                'bucket': 'None'
            })
    
    return flag_classifications

def extract_flags_with_complete_context(second_response: str) -> List[str]:
    """
    Enhanced flag extraction that preserves complete context including original quotes and page references
    """
    flags_with_context = []
    lines = second_response.split('\n')
    current_flag = ""
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Check if this is the start of a new flag
        if re.match(r'^\d+\.\s+', line):
            # Save previous flag if it exists
            if current_flag.strip():
                flags_with_context.append(current_flag.strip())
            
            # Start new flag
            current_flag = line
            
            # Look ahead to capture original quotes and page references
            j = i + 1
            while j < len(lines) and not re.match(r'^\d+\.\s+', lines[j].strip()):
                next_line = lines[j].strip()
                if next_line:  # Only add non-empty lines
                    current_flag += "\n" + next_line
                j += 1
        
    # Don't forget the last flag
    if current_flag.strip():
        flags_with_context.append(current_flag.strip())
    
    # Clean and validate flags
    cleaned_flags = []
    for flag in flags_with_context:
        # Remove any prefixes but keep the complete context
        flag = re.sub(r'^The potential red flag you observed - ', '', flag)
        flag = flag.strip()
        
        if flag and len(flag) > 10:  # Minimum length check
            cleaned_flags.append(flag)
    
    return cleaned_flags

# ==============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ==============================================================================

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    sections = fourth_response.split('###')
   
    for section in sections:
        if not section.strip():
            continue
           
        lines = section.split('\n')
        category_name = ""
        bullets = []
       
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line[1:].strip()
                if bullet_text:
                    bullets.append(bullet_text)
       
        if category_name and bullets:
            categories_summary[category_name] = bullets
   
    return categories_summary

def generate_strict_high_risk_summary(classification_results: List[Dict[str, str]], previous_year_data: str, llm: AzureOpenAILLM) -> List[str]:
    """Generate VERY concise 1-2 line summaries for high risk flags using classification data"""
    
    # Filter only high risk flags
    high_risk_classifications = [result for result in classification_results if result['risk_level'] == 'High']
    
    if not high_risk_classifications:
        return []
    
    # Create consolidated output from classification results
    output_from_all_buckets_where_high_risk_identified = ""

    for i, classification in enumerate(high_risk_classifications, 1):
        output_from_all_buckets_where_high_risk_identified += f"""
--- HIGH RISK CLASSIFICATION {i} ---
Original Flag Number: {classification.get('original_flag_number', 'Unknown')}
Flag: {classification.get('flag', 'Unknown flag')}
Matched Criteria: {classification.get('matched_criteria', 'Unknown criteria')}
Risk Level: {classification.get('risk_level', 'Unknown')}
Reasoning: {classification.get('reasoning', 'No reasoning provided')}
Relevant Financials: {classification.get('relevant_financials', 'NA')}

"""
    
    # Single LLM call with new prompt format
    prompt = f"""<role>
You are an experienced financial analyst working in ratings company. Your goal is to review the high risk red flag identified for accuracy and generate summary of high-risk financial red flag identified from given context.
The context is delimited by ####.
</role>
 
<instructions>
1. Analyze the financials, red flag identified and the contexts, the criteria which led to high risk identification.
2. Ensure the accuracy of the identification of the red flag to be high risk.
3. Create a very concise 1-2 line summary for each high-risk flag.
4. Include exact numbers, percentages, ratios, and dates whenever mentioned which led to identification of high risk flag.
5. Be factual and direct - no speculation or interpretation.
6. Provide multiple outputs when multiple points are meeting the criteria for high risk.
7. Ensure subsequent statements are cautious and do not downplay the risk.
8. Avoid neutral/positive statements that contradict the warning.
9. If applicable, specify whether the flag is for: specific business unit/division, consolidated financials, standalone financials, or geographical region. Maintain professional financial terminology.
10. Generate the summary from reasoning with original quotes.
</instructions>
 
<context>
####
{output_from_all_buckets_where_high_risk_identified}
####
 
</context>
 
<output_format>
For each high risk classification, provide:
Classification_Number: [1, 2, 3, etc.]
high_risk_flag: yes if it is actually high risk after review, no otherwise.
high_risk_flag_summary: [if high risk, provide factual summary]
</output_format>
 
<review>
1. Ensure summary is exactly 1-2 lines and preserves all quantitative information
2. Confirm that all summaries are based solely on information from the input document context
3. Check that each summary maintains a cautious tone without downplaying risks
4. Ensure proper business unit/division specification where applicable
5. Verify that the summary uses professional financial terminology
6. Check that no speculative or interpretive language is used
7. Ensure all relevant exact numbers, percentages and dates from the context are preserved
8. Verify that the output follows the output format specified above
9. No explanation needed.
</review>"""

    try:
        response = llm._call(prompt, temperature=0.1)
        
        # Parse the response to extract summaries
        concise_summaries = []
        lines = response.strip().split('\n')
        
        current_classification = {}
        for line in lines:
            line = line.strip()
            
            if line.startswith('Classification_Number:'):
                # Save previous classification if it exists and is confirmed high risk
                if (current_classification.get('high_risk_flag') == 'yes' and 
                    current_classification.get('high_risk_flag_summary')):
                    summary = current_classification['high_risk_flag_summary']
                    # Clean up summary
                    clean_summary = re.sub(r'^\[|\]$', '', summary).strip()
                    if clean_summary and not clean_summary.endswith('.'):
                        clean_summary += '.'
                    if clean_summary:
                        concise_summaries.append(clean_summary)
                
                # Start new classification
                current_classification = {}
                
            elif line.startswith('high_risk_flag:'):
                flag_value = line.split(':', 1)[1].strip().lower()
                current_classification['high_risk_flag'] = 'yes' if 'yes' in flag_value else 'no'
                
            elif line.startswith('high_risk_flag_summary:'):
                summary = line.split(':', 1)[1].strip()
                current_classification['high_risk_flag_summary'] = summary
        
        # Process the last classification
        if (current_classification.get('high_risk_flag') == 'yes' and 
            current_classification.get('high_risk_flag_summary')):
            summary = current_classification['high_risk_flag_summary']
            clean_summary = re.sub(r'^\[|\]$', '', summary).strip()
            if clean_summary and not clean_summary.endswith('.'):
                clean_summary += '.'
            if clean_summary:
                concise_summaries.append(clean_summary)
        
        if concise_summaries:
            logger.info(f"Initial concise summaries count: {len(concise_summaries)}")
            deduplicated_summaries = generate_deduplicated_high_risk_summary(concise_summaries, llm)
            logger.info(f"Final deduplicated summaries count: {len(deduplicated_summaries)}")
            return deduplicated_summaries
        else:
            return []      
    except Exception as e:
        logger.error(f"Error generating high risk summaries: {e}")

def generate_deduplicated_high_risk_summary(concise_summaries: List[str], llm: AzureOpenAILLM) -> List[str]:
    """
    Generate deduplicated high risk summary in bullet points from concise summaries
    """
    if not concise_summaries or len(concise_summaries) == 0:
        return []
    
    # Format input summaries with bullet points
    summaries_text = ""
    for summary in concise_summaries:
        summaries_text += f"• {summary}\n"
    
    prompt = f"""You are an experienced financial analyst to identify and eliminate duplicate high risk red flags.
You excel at recognizing when multiple high risk flags describe the same underlying financial issue, even when worded differently, and consolidating them into single.

Rules:
- Merge flags about the same financial issue
- One financial value cannot be part of multiple red flags
- Preserve all numbers and percentages
- Preserve all qualitative issues
- Number the final deduplicated flags
- Flags should have 1-2 sentences
- Only mention original quotes in filtered list of red flags, no explanatory tone
- Filter, consolidate aggressively

Input Red Flags:
{summaries_text}

OUTPUT FORMAT:
1. [First deduplicated flag]
2. [Second deduplicated flag]
etc.

Review:-
1. Only output the flags, no explanation needed.
2. Ensure same financial value is NOT repeat present in multiple flag.
3. Ensure aggressive deduplication with above rules so number of red flags are significantly less.
"""

    try:
        response = llm._call(prompt, temperature=0.1)
        
        # Parse the response into bullet points
        deduplicated_bullets = []
        lines = response.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            # Match numbered lines like "1. [content]" or "1. content"
            if re.match(r'^\d+\.\s+', line):
                # Remove the number prefix and clean up
                bullet = re.sub(r'^\d+\.\s+', '', line).strip()
                # Remove brackets if present
                bullet = re.sub(r'^\[|\]$', '', bullet).strip()
                if bullet:
                    deduplicated_bullets.append(bullet)
        
        return deduplicated_bullets
        
    except Exception as e:
        logger.error(f"Error in high risk deduplication: {e}")
        return concise_summaries


def create_word_document(pdf_name: str, folder_name: str, company_info: str, risk_counts: Dict[str, int],
                        deduplicated_summaries: List[str], summary_by_categories: Dict[str, List[str]],
                        output_folder: str) -> str:
    """Create a formatted Word document with concise high risk summaries"""
   
    try:
        doc = docx.Document()
       
        # Document title
        title = doc.add_heading(company_info, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Flag Distribution section
        flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
        flag_dist_heading.runs[0].bold = True
       
        # Create flag distribution table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
       
        # Use the passed deduplicated count
        high_count = risk_counts.get('High', 0)
        low_count = risk_counts.get('Low', 0)
        
        # Safely set table cells
        table.cell(0, 0).text = 'High Risk'
        table.cell(0, 1).text = str(high_count)
           
        doc.add_paragraph('')
       
        # High Risk Flags section with already generated deduplicated summaries
        if deduplicated_summaries:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
           
            # Use the already generated summaries - no function call needed
            for summary in deduplicated_summaries:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(summary)
        else:
            high_risk_heading = doc.add_heading('High Risk Summary:', level=2)
            if len(high_risk_heading.runs) > 0:
                high_risk_heading.runs[0].bold = True
            doc.add_paragraph('No high risk flags identified.')
       
        # Horizontal line
        doc.add_paragraph('_' * 50)
       
        # Summary section (4th iteration results)
        summary_heading = doc.add_heading('Summary', level=1)
        if len(summary_heading.runs) > 0:
            summary_heading.runs[0].bold = True
       
        # Add categorized summary
        if summary_by_categories and len(summary_by_categories) > 0:
            for category, bullets in summary_by_categories.items():
                if bullets and len(bullets) > 0:
                    cat_heading = doc.add_heading(str(category), level=2)
                    if len(cat_heading.runs) > 0:
                        cat_heading.runs[0].bold = True
                   
                    for bullet in bullets:
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(str(bullet))
                   
                    doc.add_paragraph('')
        else:
            doc.add_paragraph('No categorized summary available.')
       
        # Save document
        doc_filename = f"{pdf_name}_Report.docx"
        doc_path = os.path.join(output_folder, doc_filename)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        with open(temp_file.name, "wb") as f:
            doc.save(f)
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            s3.upload_fileobj(doc_buffer, BUCKET_NAME,
                            f"genai_summarization_output/{folder_name}/{doc_filename}")
       
        return doc_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")

system_prompt_step_1 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags. 
7. Include every separate statical value indicating the decline or increase in financial stress.
8. Include every separate idea indicating financial stress with original quotes and context.
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. Ensure the original quotes are comprehensive and all inclusive for the red flags identified. 
5. No explanation needed.
</review>

####
{keywords_part1}
####

"""

system_prompt_step_2 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags. 
7. Include every separate statical value indicating the decline or increase in financial stress.
8. Include every separate idea indicating financial stress with original quotes and context.
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. No explanation needed.
</review>

####
{keywords_part2}
####

"""
system_prompt_step_3 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags. 
7. Include every separate statical value indicating the decline or increase in financial stress.
8. Include every separate idea indicating financial stress with original quotes and context.
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. Ensure the original quotes are comprehensive and all inclusive for the red flags identified. 
5. No explanation needed.
</review>

####
{keywords_part3}
####
"""
system_prompt_step_4 = f"""
<role> 
You are an expert financial research analyst. Your goal is to analyze Earnings call transcript about a company and identify potential causes for concern: red flags based on given list of criteria. 
</role> 

<instructions> 
1. List of criterias are delimited by ####.
2. Earnings call transcript document is delimited by %%%%.
3. Criteria is provided with format <Criteria Name>:<its description>.
4. Analyze the Earnings Call Transcript document and identify the red flags according to the given list of criteria.
5. A criteria may appear multiple times in the document; you need to evaluate each instance and flag it as a new point only if it appears in a different paragraph.                                        
6. Only identify a criteria if it is associated with a negative cause for concern, and refrain from highlighting positive or neutral flags. 
7. Include every separate statical value indicating the decline or increase in financial stress.
8. Include every separate idea indicating financial stress with original quotes and context.
</instructions>

For each identified negative red flag, strictly adhere to the following output format: 
<output format> 
1. <The criteria name identified> - <Provide the entire original quote and text that led to the identification of the red flag, along with the page number where the statement was found.> 
   Context - <all the relevant contexts summary from the document that led to the identification of the red flag>
2. <next criteria identified name> - <original quote>
   Context - <all relevant context summary>
</output format>

<review>
1. Please ensure if all negative cause for concern red flags are provided in the response.
2. Please analyze the document again to ensure no red flags are missed.
3. Kindly ensure the response follows the output format given above.
4. Ensure the original quotes are comprehensive and all inclusive for the red flags identified. 
5. No explanation needed.
</review>

####
{keywords_part4}
####
"""

# ==============================================================================
# MAIN PROCESSING PIPELINE WITH SPLIT FIRST ITERATION
# ==============================================================================

def process_pdf_enhanced_pipeline_with_split_iteration(pdf_path: str, folder_name: str, previous_year_data: str, 
                               output_folder: str = "results", 
                               api_key: str = None, azure_endpoint: str = None, 
                               api_version: str = None, deployment_name: str = "gpt-4.1"):
    """
    Enhanced processing pipeline with split first iteration 
    """
   
    os.makedirs(output_folder, exist_ok=True)
    pdf_name = Path(pdf_path).stem
   
    try:
        # Initialize LLM and load PDF
        # Change:
        llm_client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=azure_endpoint, 
            api_version=api_version,
        )
        
        llm = AzureOpenAILLM(
            api_key=api_key,
            azure_endpoint=azure_endpoint, 
            api_version=api_version,
            deployment_name=deployment_name
        )

        docs = mergeDocs(pdf_path, split_pages=False)
        context = docs[0]["context"]
        
        # Make a chat completions call
        response1 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_1},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )

        # Make a chat completions call
        response2 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_2},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )

        # Make a chat completions call
        response3 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_3},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )

        # Make a chat completions call
        response4 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt_step_4},
                {"role": "user", "content": f"%%%%{context}%%%%"}
            ]
        )

        first_response = response1.choices[0].message.content + "\n" + response2.choices[0].message.content+"\n" + response3.choices[0].message.content+"\n" + response4.choices[0].message.content
        first_part_1= response1.choices[0].message.content + "|" + response2.choices[0].message.content
        first_part_2= response3.choices[0].message.content + "|" + response4.choices[0].message.content

       # ITERATION 2: Enhanced Deduplication - Modified for direct client approach
        logger.info("Running 2nd iteration 'A' - Enhanced Deduplication...")
        second_system_prompt_A = """<role>
        You are an experienced financial analyst for analyzing earnings call transcripts. Your goal is to identify and eliminate duplicate red flags while maintaining comprehensive analysis integrity.
        You excel at recognizing when multiple red flags describe the same underlying financial issue, even when worded differently, and consolidating them into single, comprehensive red flags while preserving all supporting evidence.
        Input document is delimited by ####.
</role>
<instruction>
        deduplication rules:
        1. merge red flags that refer to the same financial metric, issue, or concern
        2. combine red flags about the same business area/division/segment  
        3. consolidate similar operational or strategic concerns
        4. eliminate redundant mentions of the same data point or statistic
        5. preserve all original quotes, speaker attributions, and page references from merged items
        6. maintain sequential numbering (1, 2, 3, etc.) after deduplication
        7. do not lose any substantive financial concerns or statistic refering to declining metrics - only remove true duplicates
        8. be aggressive in removing duplicates while preserving all important context and evidence
        9. any quarter to quarter or year on year financial metric decline by more than 25% needs to be present in the corresponding red flag
        </instruction>

<output format>
        1. <the criteria name identified> - <provide all the entire original quotes and text that led to the identification of the red flag, along with the page number where the statement was found.>
        context - <all the relevant contexts summary from the document that led to the identification of the red flag>
        2. <next criteria identified name> - <original quotes>
        context - <all relevant context summary>
</output format>

<review>
        1. Ensure that all duplicate red flags referring to the same underlying financial issue are properly merged.
        2. Verify that no substantive financial concerns or statistic are lost during the deduplication process.
        3. Confirm that all original quotes and page references are preserved in the consolidated flags.
        4. Check that the response follows the exact output format specified above.
        5. Verify that merged flags contain comprehensive evidence from all related duplicates.
        6. Confirm the response starts immediately with "1." without any introduction.
        7. Double-check that speaker attributions are maintained in the original quotes.
        8. Ensure all financial stress points are covered with original quotes in relevant red flags.
        9. Analyze the input document again and ensure all financial stress/concerns/issues and statistics are covered with original quotes and merged in relevant red flags.
</review>"""
        second_user_content_A = f"""<context>
        Earnings call transcripts red flags for deduplication:-
        {first_part_1}
</context>
        Provide deduplicated analysis with merged duplicates and preserved evidence:"""

        # Use direct client approach
        response1 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": second_system_prompt_A},
                {"role": "user", "content": f"####{second_user_content_A}####"}
            ]
        )

        logger.info("Running 2nd 'B' iteration - Enhanced Deduplication...")
        second_system_prompt_B = """<role>
        You are an experienced financial analyst for analyzing earnings call transcripts. Your goal is to identify and eliminate duplicate red flags while maintaining comprehensive analysis integrity.
        You excel at recognizing when multiple red flags describe the same underlying financial issue, even when worded differently, and consolidating them into single, comprehensive red flags while preserving all supporting evidence.
        Input document is delimited by ####.
</role>
<instruction>
        deduplication rules:
        1. merge red flags that refer to the same financial metric, issue, or concern
        2. combine red flags about the same business area/division/segment  
        3. consolidate similar operational or strategic concerns
        4. eliminate redundant mentions of the same data point or statistic
        5. preserve all original quotes, speaker attributions, and page references from merged items
        6. maintain sequential numbering (1, 2, 3, etc.) after deduplication
        7. do not lose any substantive financial concerns or statistic refering to declining metrics - only remove true duplicates
        8. be aggressive in removing duplicates while preserving all important context and evidence
        9. any quarter to quarter or year on year financial metric decline by more than 25% needs to be present in the corresponding red flag
        </instruction>

<output format>
        1. <the criteria name identified> - <provide all the entire original quotes and text that led to the identification of the red flag, along with the page number where the statement was found.>
        context - <all the relevant contexts summary from the document that led to the identification of the red flag>
        2. <next criteria identified name> - <original quotes>
        context - <all relevant context summary>
</output format>

<review>
        1. Ensure that all duplicate red flags referring to the same underlying financial issue are properly merged.
        2. Verify that no substantive financial concerns or statistic are lost during the deduplication process.
        3. Confirm that all original quotes and page references are preserved in the consolidated flags.
        4. Check that the response follows the exact output format specified above.
        5. Verify that merged flags contain comprehensive evidence from all related duplicates.
        6. Confirm the response starts immediately with "1." without any introduction.
        7. Double-check that speaker attributions are maintained in the original quotes.
        8. Ensure all financial stress points are covered with original quotes in relevant red flags.
        9. Analyze the input document again and ensure all financial stress/concerns/issues and statistics are covered with original quotes and merged in relevant red flags.
</review>"""
        second_user_content_B = f"""<context>
        Earnings call transcripts red flags for deduplication:-
        {first_part_2}
</context>
        Provide deduplicated analysis with merged duplicates and preserved evidence:"""

        # Use direct client approach
        response2 = llm_client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": second_system_prompt_B},
                {"role": "user", "content": f"####{second_user_content_B}####"}
            ]
        )

        second_response = response1.choices[0].message.content+response2.choices[0].message.content
                
        # ITERATION 3: Categorization (UNCHANGED)
        logger.info("Running 3rd iteration - Categorization...")
        third_prompt = f"""<role>
You are a senior financial analyst expert in financial risk categorization with deep knowledge of balance sheet analysis, P&L assessment, and corporate risk frameworks.
</role>

<system_prompt>
You excel at organizing financial risks into standardized categories, ensuring comprehensive coverage of all financial risk areas, and maintaining accuracy in risk classification.
</system_prompt>

<instruction>
Categorize the identified red flags into the following 7 standardized categories based on their financial nature and business impact.

MANDATORY CATEGORIES:
1. Balance Sheet Issues: Assets, liabilities, equity, debt, and overall financial position concerns
2. P&L (Income Statement) Issues: Revenue, expenses, profits, and financial performance concerns  
3. Liquidity Issues: Short-term obligations, cash flow, debt repayment, working capital concerns
4. Management and Strategy Issues: Leadership, governance, decision-making, strategy, and vision concerns
5. Regulatory Issues: Compliance, laws, regulations, and regulatory body concerns
6. Industry and Market Issues: Market position, industry trends, competitive landscape concerns
7. Operational Issues: Internal processes, systems, infrastructure, and operational efficiency concerns

CATEGORIZATION RULES:
- Assign each red flag to the MOST relevant category only
- Do not create new categories - use only the 7 listed above
- Preserve all Original Quotes exactly as provided
- Maintain sequential organization within each category

OUTPUT FORMAT:
### Balance Sheet Issues
- [Red flag 1 with original quote and page reference]
- [Red flag 2 with original quote and page reference]

### P&L (Income Statement) Issues
- [Red flag 1 with original quote and page reference]

Continue this format for all applicable categories.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

DEDUPLICATED ANALYSIS TO CATEGORIZE:
{second_response}c
</context>

Provide categorized analysis:"""
        
        third_response = llm._call(third_prompt)
        
        # ITERATION 4: Summary Generation (UNCHANGED)
        logger.info("Running 4th iteration - Summary Generation...")
        fourth_prompt = f"""<role>
You are an expert financial summarization specialist with expertise in creating concise, factual, and comprehensive summaries that preserve critical quantitative data and key insights.
</role>

<system_prompt>
You excel at distilling complex financial analysis into clear, actionable summaries while maintaining objectivity, preserving all quantitative details, and ensuring no critical information is lost.
</system_prompt>

<instruction>
Create a comprehensive summary of each category of red flags in bullet point format following these strict guidelines.

SUMMARY REQUIREMENTS:
1. Retain ALL quantitative information (numbers, percentages, ratios, dates)
2. Maintain completely neutral, factual tone - no opinions or interpretations
3. Include every red flag from each category - no omissions
4. Base content solely on the provided categorized analysis
5. Preserve specific data points and statistics wherever mentioned
6. Each bullet point should capture key details of individual red flags
7. Balance thoroughness with conciseness
8. Use professional financial terminology
9. Ensure category-specific content alignment

OUTPUT FORMAT:
### Balance Sheet Issues
* [Summary of red flag 1 with specific data points and factual information]
* [Summary of red flag 2 with specific data points and factual information]

### P&L (Income Statement) Issues  
* [Summary of red flag 1 with specific data points and factual information]

Continue this format for all 7 categories that contain red flags.

CRITICAL: Each bullet point represents a concise summary of individual red flags with preserved quantitative details.
</instruction>

<context>
ORIGINAL DOCUMENT:
{context}

CATEGORIZED ANALYSIS TO SUMMARIZE:
{third_response}
</context>

Provide factual category summaries:"""
        
        fourth_response = llm._call(fourth_prompt)
        
        # ITERATION 5: Efficient Bucket-Based Classification (UNCHANGED)
        logger.info("Running 5th iteration - Efficient Bucket-Based Classification...")
        
        try:
            flags_with_context = extract_flags_with_complete_context(second_response)
            logger.info(f"\nFlags with context extracted: {len(flags_with_context)}")
            
        except Exception as e:
            logger.error(f"Error parsing flags with context: {e}")
            flags_with_context = ["Error in flag parsing"]

        classification_results = []
        high_risk_flags = []
        low_risk_flags = []

        if len(flags_with_context) > 0 and flags_with_context[0] != "Error in flag parsing":
            try:
                logger.info(f"Analyzing all {len(flags_with_context)} flags using 8 bucket calls.")
                
                bucket_results = classify_all_flags_with_enhanced_buckets(flags_with_context, previous_year_data, llm)
                classification_results = parse_bucket_results_to_classifications_enhanced(bucket_results, flags_with_context)

                for result in classification_results:
                    if (result['risk_level'].lower() == 'high' and 
                        result['matched_criteria'] != 'None'):
                        high_risk_flags.append(result['flag'])
                    else:
                        low_risk_flags.append(result['flag'])
                        
            except Exception as e:
                logger.error(f"Error in efficient bucket classification: {e}")
                for flag_with_context in flags_with_context:
                    flag_description = flag_with_context.split('\n')[0]
                    flag_description = re.sub(r'^\d+\.\s+', '', flag_description).strip()
                    
                    classification_results.append({
                        'flag': flag_description,
                        'flag_with_context': flag_with_context,
                        'matched_criteria': 'None',
                        'risk_level': 'Low',
                        'reasoning': f'Classification failed: {str(e)}',
                        'bucket': 'Error'
                    })
                    low_risk_flags.append(flag_description)

        # **Generate deduplicated summaries for accurate count**
        logger.info("\nGenerating deduplicated high risk summary for final count...")
        deduplicated_summaries = []
        try:
            high_risk_classifications = [result for result in classification_results if result['risk_level'] == 'High']
            
            if high_risk_classifications:
                deduplicated_summaries = generate_strict_high_risk_summary(classification_results, previous_year_data, llm)
            
            # Update risk counts with deduplicated numbers
            risk_counts_deduplicated = {
                'High_Raw': len(high_risk_flags),  # Original count
                'High': len(deduplicated_summaries),  # Deduplicated count
                'Low': len(low_risk_flags),
                'Total': len(flags_with_context) if flags_with_context and flags_with_context[0] != "Error in flag parsing" else 0
            }
            
        except Exception as e:
            logger.error(f"Error generating deduplicated summaries: {e}")
            # Fallback to original counts
            risk_counts_deduplicated = {
                'High': len(high_risk_flags),
                'Low': len(low_risk_flags),
                'Total': len(flags_with_context) if flags_with_context and flags_with_context[0] != "Error in flag parsing" else 0
            }
        
        if len(deduplicated_summaries) > 0:
            logger.info(f"\n--- HIGH RISK FLAGS (after deduplication) ---")
            for i, summary in enumerate(deduplicated_summaries, 1):
                logger.info(f"  {i}. {summary}")
        else:
            logger.info("  No high risk flags identified after deduplication")
        
        # Word Document Creation
        logger.info("Creating Word document...")
        try:
            company_info = extract_company_info_from_pdf(pdf_path, llm)
            summary_by_categories = parse_summary_by_categories(fourth_response)
        
            word_doc_path = create_word_document(
                pdf_name=pdf_name,
                folder_name=folder_name,
                company_info=company_info,
                risk_counts=risk_counts_deduplicated,  # Use deduplicated counts
                deduplicated_summaries=deduplicated_summaries,  # Pass the already generated summaries
                summary_by_categories=summary_by_categories,
                output_folder=output_folder
            )
                
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            word_doc_path = None
       
        # Save all results to CSV files (MODIFIED)
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        results_summary = pd.DataFrame({
            "pdf_name": [pdf_name] * 5,
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Red Flags",
                "Enhanced Deduplication",
                "Categorization",
                "Summary Generation", 
                "Enhanced Context-Based Classification"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                f"Enhanced Context-Based Classification: {risk_counts_deduplicated['High']} High Risk (after deduplication), {risk_counts_deduplicated['Low']} Low Risk flags from {risk_counts_deduplicated['Total']} total flags"
            ],
            "timestamp": [timestamp] * 5
        })
       
        results_file = os.path.join(output_folder, f"{pdf_name}_split_iteration_pipeline_results.csv")
        results_summary.to_csv(results_file, index=False)
        
        if len(classification_results) > 0:
            classification_df = pd.DataFrame(classification_results)
            classification_file = os.path.join(output_folder, f"{pdf_name}_split_iteration_classification.csv")
            classification_df.to_csv(classification_file, index=False)

        return results_summary
       
    except Exception as e:
        logger.error(f"Error processing {pdf_name}: {str(e)}")
        return None

# ==============================================================================
# MAIN FUNCTION
# ==============================================================================

def main_ect(pdf_path, folder_name, mf): 

    API_CONFIG = {
        "api_key": "84998c",
        "azure_endpoint": "https://crisil-pp-gpt.openai.azure.com",
        "api_version": "2025-01-01-preview",
        "deployment_name": "gpt-4.1"
    }
    
    PATHS_CONFIG = {
        "output_folder": r"sterlin_dec_results_split_iteration_5"
    }

    os.makedirs(PATHS_CONFIG["output_folder"], exist_ok=True)
        
    # Use the new split iteration function
    result = process_pdf_enhanced_pipeline_with_split_iteration(
        pdf_path=pdf_path,
        folder_name=folder_name,
        previous_year_data=mf,
        output_folder=PATHS_CONFIG["output_folder"],
        api_key=API_CONFIG["api_key"],
        azure_endpoint=API_CONFIG["azure_endpoint"],
        api_version=API_CONFIG["api_version"],
        deployment_name=API_CONFIG["deployment_name"]
    )
