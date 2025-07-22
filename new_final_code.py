import os
import time
import pandas as pd
import fitz  
import requests
import warnings
import hashlib
import logging
import json
import ast
from typing import Dict, List, Any
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re
 
warnings.filterwarnings('ignore')
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
 
def getFilehash(file_path: str):
    """Generate SHA3-256 hash of a file"""
    with open(file_path, 'rb') as f:
        return hashlib.sha3_256(f.read()).hexdigest()
 
class HostedLLM:
    """Custom LLM class for hosted Llama model"""
   
    def __init__(self, endpoint: str):
        self.endpoint = endpoint
   
    def _call(self, prompt: str) -> str:
        """Make API call to hosted LLM"""
        try:
            prompt_template = f"""<|begin_of_text|><|start_header_id|>user<|end_header_id|>
{prompt}<|eot_id|><|start_header_id|>assistant<|end_header_id|>
"""
           
            payload = json.dumps({
                "provider": "tgi",
                "deployment": "Llama 3.3 v1",
                "spec_version": 1,
                "input_text": prompt_template,
                "params": {"temperature": 0.1}
            })
           
            headers = {
                'token': '0e53d2cf9f724a94a6ca0a9c880fdee7',
                'Content-Type': 'application/json'
            }
           
            response = requests.post(
                url="https://llmgateway.crisil.local/api/v1/llm",
                headers=headers,
                data=payload,
                verify=False
            )
           
            print(f"Response Status: {response.status_code}")
            print(f"Response Text: {response.text[:500]}...")  # First 500 chars
           
            if response.status_code != 200:
                return f"LLM Call Failed: HTTP {response.status_code} - {response.text}"
           
            response_v = ast.literal_eval(response.text)
            resp_o = response_v['output']
            output = str(resp_o).replace(prompt_template, "")
            return output.strip()
           
        except requests.exceptions.RequestException as e:
            return f"LLM Call Failed - Network Error: {e}"
        except json.JSONDecodeError as e:
            return f"LLM Call Failed - JSON Error: {e}"
        except KeyError as e:
            return f"LLM Call Failed - Missing Key: {e}"
        except Exception as e:
            return f"LLM Call Failed - General Error: {e}"
 
class PDFExtractor:
    """Class for extracting text from PDF files"""
   
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract text from each page of a PDF file"""
        start_time = time.time()
        try:
            doc = fitz.open(pdf_path)
            pages = []
           
            for page_num, page in enumerate(doc):
                text = page.get_text()
                pages.append({
                    "page_num": page_num + 1,
                    "text": text
                })
           
            # Explicitly close the document to free memory
            doc.close()
            logger.info(f"PDF text extraction took {time.time() - start_time:.2f} seconds")
            return pages
           
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
 
def mergeDocs(pdf_path: str, split_pages: bool = False) -> List[Dict[str, Any]]:
    """Merge PDF documents into a single context"""
    extractor = PDFExtractor()
    pages = extractor.extract_text_from_pdf(pdf_path)
   
    if split_pages:
        return [{"context": page["text"], "page_num": page["page_num"]} for page in pages]
    else:
        # Merge all pages into single context
        all_text = "\n".join([page["text"] for page in pages])
        return [{"context": all_text}]
 
class LlamaQueryPipeline:
    """Main pipeline class for querying PDF content with Llama"""
   
    def __init__(self, pdf_path: str, queries_csv_path: str = None, llm_endpoint: str = "https://llmgateway.crisil.local/api/v1/llm", previous_results_path: str = None):
        """Initialize the pipeline"""
        self.llm = HostedLLM(endpoint=llm_endpoint)
        self.docs = mergeDocs(pdf_path, split_pages=False)
       
        # Load queries from Excel or CSV (only if provided)
        if queries_csv_path:
            if queries_csv_path.endswith('.xlsx'):
                queries_df = pd.read_excel(queries_csv_path)
            else:
                queries_df = pd.read_csv(queries_csv_path)
            self.queries = queries_df["prompt"].tolist()
        else:
            self.queries = []
           
        self.pdf_path = pdf_path
        self.pdf_name = Path(pdf_path).stem  # Get filename without extension
       
        # Load previous results if provided
        self.previous_results = None
        if previous_results_path and os.path.exists(previous_results_path):
            self.previous_results = pd.read_csv(previous_results_path)
   
    def query_llama_with_chaining(self, new_queries_csv_path: str, iteration_number: int = 2) -> pd.DataFrame:
        """Query the Llama API using previous results for chaining"""
        if self.previous_results is None:
            raise ValueError("No previous results loaded. Please provide previous_results_path in __init__")
       
        # Load new queries
        if new_queries_csv_path.endswith('.xlsx'):
            new_queries_df = pd.read_excel(new_queries_csv_path)
        else:
            new_queries_df = pd.read_csv(new_queries_csv_path)
       
        new_queries = new_queries_df["prompt"].tolist()
       
        sys_prompt = f"""You must answer the question strictly based on the below given context.
 
Context:
{self.docs[0]["context"]}
 
"""
       
        results = []
       
        # Process each new query with corresponding previous response
        for i, new_query in enumerate(new_queries):
            start = time.perf_counter()
           
            try:
                # Get the corresponding previous response (if available)
                if i < len(self.previous_results):
                    # Check if it's 3rd iteration (chained results) or 2nd iteration (original results)
                    if 'chained_response' in self.previous_results.columns:
                        previous_response = self.previous_results.iloc[i]['chained_response']
                    else:
                        previous_response = self.previous_results.iloc[i]['response']
                   
                    # Create chained prompt: new query + previous response
                    chained_prompt = f"""Previous Analysis: {previous_response}
 
Based on the above analysis and the original context, please answer: {new_query}
 
Answer:"""
                else:
                    # If no previous response available, use regular prompt
                    chained_prompt = f"""Question: {new_query}
Answer:"""
               
                full_prompt = f"{sys_prompt}\n{chained_prompt}"
               
                # Get response from Llama
                response_text = self.llm._call(full_prompt)
                end = time.perf_counter()
               
                # Calculate token approximations
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
               
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
               
            except Exception as e:
                end = time.perf_counter()
               
                usage = {
                    "iteration": iteration_number,
                    "query_id": i + 1,
                    "original_query": new_queries_df.iloc[i]["prompt"] if i < len(new_queries_df) else new_query,
                    "previous_response": previous_response if i < len(self.previous_results) else "",
                    "new_query": new_query,
                    "chained_response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
           
            results.append(usage)
       
        return pd.DataFrame(results)
 
    def query_llama(self, maintain_conversation: bool = True, enable_chaining: bool = False) -> pd.DataFrame:
        """Query the Llama API for a list of queries using the provided context"""
        sys_prompt = f"""You are a financial analyst expert specializing in identifying red flags from earnings call transcripts and financial documents.

COMPLETE DOCUMENT TO ANALYZE:
{self.docs[0]["context"]}

Your task is to analyze the ENTIRE document above and identify ALL potential red flags. 

CRITICAL OUTPUT FORMAT REQUIREMENTS:
- Number each red flag sequentially (1, 2, 3, etc.)
- Start each entry with: "The potential red flag you observed - [brief description]"
- Follow with "Original Quote:" and then the exact quote with speaker names
- Include page references where available: (Page X)
- Ensure comprehensive analysis of the entire document
- Do not miss any sections or concerning statements

EXAMPLE FORMAT:
1. The potential red flag you observed - Debt reduction is lower than expected  
Original Quote:  
"Vikrant Kashyap: Have you -- are you able to reduce any debt in quarter one in Middle East and India?  
Ramesh Kalyanaraman: So India, we are not reduced, but the cash balance has been increased to around INR75 crores, but we have not reduced any debt in Q1 in India and Middle East because Middle East we have not converted any showroom in Q1." (Page 9)  

2. The potential red flag you observed - Margin pressure/Competition intensifying/Cost inflation  
Original Quote:  
"Pulkit Singhal: Thank you for the opportunity and congrats on the good set of performance. Just the first question is really on the margins, which seems to have been elusive. I mean looking at the company for the last two years, we seem to be executing quite well in terms of store expansions and revenue growth and clearly delivering higher than expected there. But margin expansion just has been completely elusive.  
And I find it surprising your comment that while growing at 30% growth and 12% SSSG, I mean, which is quite healthy, you're still talking about high competitive intensity kind of quite contradictory that with such high growth rates, we have to invest so high. So can you talk about this a bit more? I mean we don't expect with lower revenue growth rates that you would not have to invest in the business. And it's only during a higher revenue growth that you expect margin expansion.  
Ramesh Kalyanaraman: Yes. So you're right. We are -- meaning somewhere we have missed out on the operating leverage for advertisements that's why I told you that even Q1, it was a miss. And regarding competition, I will tell you where in new markets, where we assume that we will not spend too much because the brand is already aware and the location is the only thing which has to be communicated.  
When you see the local players, regional players or the micro market players there becoming extremely active because of our showroom launch then we will have to increase the noise level there. Otherwise, we will lose our market share. And existing local players they increase their activity around our launch time. So that is where we also put more money so that we don't end up losing the market share or we don't end up taking out lesser from the competition." (Page 12-13)  

3. The potential red flag you observed - Debt high/Increase in borrowing cost  
Original Quote:  
"Vikrant Kashyap: My question is how are you going to address this? Because if you continue to grow at a higher level, but bottom line is not expanding related to the top line, it will going to impact your overall performance? So, what are the steps you are taking to improve the bottom line in the businesses  
Sanjay Raghuraman: Finance costs will be taken care because we told you when we convert stores, that money is going to reduce our debt. Okay. And again, FOCO, when you do FOCO showrooms, the margins will come down. And surely, that will have an impact on the gross margin. Okay. And interest, if you look at actually, the interest rates have been going up last year. So next year, that will be the base, right? So then again, we will not have this kind of issue is what we feel. So interest rates have been going up over the past year, one year, in that region.  
And we are also beginning to repay loans now because of conversion. So all put together, interest part will be taken care but other area where FOCO showrooms will surely reduce our margin. We cannot have the own store margin. So that should be the way we should look at it." (Page 9)

Continue this exact format for ALL red flags identified throughout the document.

"""
       
        prompt_template = """Question: {query}

Analyze the complete document and provide ALL red flags in the exact numbered format specified above. Be thorough and comprehensive - cover the entire document.

Answer:"""
       
        conversation_history = ""
        results = []
        previous_output = ""  # For prompt chaining
       
        for i, query in enumerate(self.queries, 1):
            start = time.perf_counter()
           
            try:
                if enable_chaining and i > 1 and previous_output:
                    chained_query = f"{query}\n\nPrevious context: {previous_output}"
                else:
                    chained_query = query
               
                if maintain_conversation and conversation_history:
                    full_prompt = f"{sys_prompt}\n{conversation_history}\n{prompt_template.format(query=chained_query)}"
                else:
                    full_prompt = f"{sys_prompt}\n{prompt_template.format(query=chained_query)}"
               
                response_text = self.llm._call(full_prompt)
                end = time.perf_counter()
               
                input_tokens = len(full_prompt.split())
                completion_tokens = len(response_text.split()) if response_text else 0
               
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": chained_query if enable_chaining else query,
                    "response": response_text,
                    "completion_tokens": completion_tokens,
                    "input_tokens": input_tokens,
                    "response_time": f"{end - start:.2f}"
                }
               
                # Update conversation history for next iteration
                if maintain_conversation:
                    conversation_history += f"\nQuestion: {chained_query}\nAnswer: {response_text}\n"
               
                # Store output for next chaining iteration
                if enable_chaining:
                    previous_output = response_text
               
            except Exception as e:
                end = time.perf_counter()
               
                usage = {
                    "query_id": i,
                    "query": query,
                    "chained_query": query,
                    "response": f"Error: {str(e)}",
                    "completion_tokens": None,
                    "input_tokens": None,
                    "response_time": f"{end - start:.2f}"
                }
               
                if enable_chaining:
                    previous_output = ""
           
            results.append(usage)
       
        return pd.DataFrame(results)
   
    def save_results(self, results_df: pd.DataFrame, output_path: str = None):
        """Save results to CSV file"""
        if output_path is None:
            # Generate filename based on PDF name and timestamp
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_path = f"llama_query_results_{self.pdf_name}_{timestamp}.csv"
       
        results_df.to_csv(output_path, index=False)
        return output_path

def extract_company_info_from_pdf(pdf_path: str, llm: HostedLLM) -> str:
    """Extract company name, quarter, and financial year from first page of PDF"""
    try:
        doc = fitz.open(pdf_path)
        first_page_text = doc[0].get_text()
        doc.close()
        
        # Limit text to first 2000 characters to avoid token limits
        first_page_text = first_page_text[:2000]
        
        prompt = f"""
You are a financial document analyst. Extract the company name, quarter, and financial year from the following text which is from the first page of an earnings call transcript or financial document.

Text from first page:
{first_page_text}

Please identify:
1. Company Name (full company name including Ltd/Limited/Inc etc.)
2. Quarter (Q1/Q2/Q3/Q4)
3. Financial Year (FY23/FY24/FY25 etc.)

Format your response as: [Company Name]-[Quarter][Financial Year]
Example: Reliance Industries Limited-Q4FY25

If you cannot find clear information, make the best estimate based on available data.

Response:"""
        
        response = llm._call(prompt)
        # Clean the response to get just the formatted string
        response_lines = response.strip().split('\n')
        for line in response_lines:
            if '-Q' in line and 'FY' in line:
                return line.strip()
        
        # Fallback - return first non-empty line
        return response_lines[0].strip() if response_lines else "Unknown Company-Q1FY25"
        
    except Exception as e:
        logger.error(f"Error extracting company info: {e}")
        return "Unknown Company-Q1FY25"

def parse_risk_classifications(combined_fifth_response: str) -> Dict[str, List[str]]:
    """Parse the risk classification response to extract categorized flags"""
    risk_flags = {
        'High': [],
        'Medium': [],
        'Low': []
    }
    
    print("DEBUG: Parsing risk classifications...")
    print(f"DEBUG: Response length: {len(combined_fifth_response)}")
    
    # Split by categories
    categories = combined_fifth_response.split('###')
    
    for category_text in categories:
        if not category_text.strip():
            continue
            
        lines = category_text.split('\n')
        current_flag = ""
        
        for line in lines:
            line = line.strip()
            
            # Check for bullet points (handle * format) - more flexible matching
            if line.startswith('*') and not 'Risk Classification' in line:
                current_flag = line[1:].strip()  # Remove bullet point
                print(f"DEBUG: Found flag: {current_flag}")
                continue
            
            # Check for risk classifications - exact pattern matching
            if current_flag and line.startswith('  - '):
                line_clean = line.replace('  - ', '').strip().lower()
                
                # Look for exact "high: yes" pattern
                if line_clean == 'high: yes':
                    risk_flags['High'].append(current_flag)
                    print(f"DEBUG: Added HIGH risk: {current_flag}")
                    # Don't reset current_flag here, continue checking other risk levels
                
                # Look for exact "medium: yes" pattern  
                elif line_clean == 'medium: yes':
                    risk_flags['Medium'].append(current_flag)
                    print(f"DEBUG: Added MEDIUM risk: {current_flag}")
                    # Don't reset current_flag here, continue checking other risk levels
                
                # Look for exact "low: yes" pattern
                elif line_clean == 'low: yes':
                    risk_flags['Low'].append(current_flag)
                    print(f"DEBUG: Added LOW risk: {current_flag}")
                    # Don't reset current_flag here, continue checking other risk levels
            
            # Reset current_flag when we encounter next bullet or category
            elif line.startswith('*') or line.startswith('###'):
                current_flag = ""
    
    print(f"DEBUG: Final counts - High: {len(risk_flags['High'])}, Medium: {len(risk_flags['Medium'])}, Low: {len(risk_flags['Low'])}")
    print(f"DEBUG: High risk flags: {risk_flags['High']}")
    
    return risk_flags

def parse_summary_by_categories(fourth_response: str) -> Dict[str, List[str]]:
    """Parse the 4th iteration summary response by categories"""
    categories_summary = {}
    
    # Split by ### headers
    sections = fourth_response.split('###')
    
    for section in sections:
        if not section.strip():
            continue
            
        lines = section.split('\n')
        category_name = ""
        bullets = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('*') and not line.startswith('-'):
                # This is likely the category name
                category_name = line.strip()
            elif line.startswith('*') or line.startswith('-'):
                # This is a bullet point
                bullet_text = line[1:].strip()  # Remove bullet symbol
                if bullet_text:
                    bullets.append(bullet_text)
        
        if category_name and bullets:
            categories_summary[category_name] = bullets
    
    return categories_summary

def create_word_document(pdf_name: str, company_info: str, risk_flags: Dict[str, List[str]], 
                        summary_by_categories: Dict[str, List[str]], output_folder: str) -> str:
    """Create a formatted Word document with the analysis results"""
    
    # Create new document
    doc = Document()
    
    # Set document title
    title = doc.add_heading(company_info, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Flag Distribution section
    flag_dist_heading = doc.add_heading('Flag Distribution:', level=2)
    flag_dist_heading.runs[0].bold = True
    
    # Create table for flag distribution
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Add table headers and data
    high_count = len(risk_flags['High'])
    medium_count = len(risk_flags['Medium'])
    low_count = len(risk_flags['Low'])
    total_count = high_count + medium_count + low_count
    
    table.cell(0, 0).text = 'High Risk'
    table.cell(0, 1).text = str(high_count)
    table.cell(1, 0).text = 'Medium Risk'
    table.cell(1, 1).text = str(medium_count)
    table.cell(2, 0).text = 'Low Risk'
    table.cell(2, 1).text = str(low_count)
    table.cell(3, 0).text = 'Total Flags'
    table.cell(3, 1).text = str(total_count)
    
    # Make table headers bold
    for i in range(4):
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # Add space
    doc.add_paragraph('')
    
    # Debug: Print what we're about to add to the document
    print(f"DEBUG WORD: About to add {high_count} high risk flags to document")
    print(f"DEBUG WORD: High risk flags: {risk_flags['High']}")
    
    # Add High Risk Flags section only
    if risk_flags['High']:
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
        
        for flag in risk_flags['High']:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(flag)
            print(f"DEBUG WORD: Added flag to document: {flag}")
    else:
        # If no high risk flags, add a note
        high_risk_heading = doc.add_heading('High Risk Flags:', level=2)
        high_risk_heading.runs[0].bold = True
        no_flags_para = doc.add_paragraph('No high risk flags identified.')
        print("DEBUG WORD: No high risk flags found - added 'No high risk flags identified' message")
    
    # Add horizontal line
    doc.add_paragraph('_' * 50)
    
    # Add Summary section
    summary_heading = doc.add_heading('Summary', level=1)
    summary_heading.runs[0].bold = True
    
    # Add categorized summary
    for category, bullets in summary_by_categories.items():
        if bullets:  # Only add if there are bullets
            # Add category as subheading
            cat_heading = doc.add_heading(category, level=2)
            cat_heading.runs[0].bold = True
            
            # Add bullet points for this category
            for bullet in bullets:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(bullet)
            
            # Add space between categories
            doc.add_paragraph('')
    
    # Save document
    doc_filename = f"{pdf_name}_Report.docx"
    doc_path = os.path.join(output_folder, doc_filename)
    doc.save(doc_path)
    
    return doc_path

def process_single_pdf_five_iterations(pdf_path: str, queries_csv_path: str, previous_year_data: str, output_folder: str = "results"):
    """
    Process a single PDF through the 5-iteration pipeline including risk classification
    """
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Get PDF name without extension
    pdf_name = Path(pdf_path).stem
    
    print(f"\nProcessing: {pdf_name}")
    print("=" * 50)
    
    try:
        # ITERATION 1: Initial red flag identification
        print("Running 1st iteration - Initial Analysis...")
        pipeline_1st = LlamaQueryPipeline(
            pdf_path=pdf_path,
            queries_csv_path=queries_csv_path
        )
        
        # Run 1st iteration
        first_results_df = pipeline_1st.query_llama(maintain_conversation=True, enable_chaining=False)
        
        # Get first response for chaining
        first_response = first_results_df.iloc[0]['response']
        
        # ITERATION 2: Deduplication and cleanup
        print("Running 2nd iteration - Deduplication...")
        second_prompt = """Remove the duplicates from the above context. Also if the Original Quote and Keyword identifies is same remove them. 
        
    Ensure that:
    1. No duplicate red flags are present
    2. Similar or redundant information is consolidated
    3. All unique red flags are retained
    4. The output maintains the structure and clarity of the original analysis"""
        
        second_full_prompt = f"""You must answer the question strictly based on the below given context.

    Context:
    {pipeline_1st.docs[0]["context"]}

    Previous Analysis: {first_response}

    Based on the above analysis and the original context, please answer: {second_prompt}

    Answer:"""
        
        second_response = pipeline_1st.llm._call(second_full_prompt)
        
        # ITERATION 3: Categorization of red flags
        print("Running 3rd iteration - Categorization...")
        third_prompt = """You are an expert in financial analysis tasked at categorizing the below identified red flags related to a company's financial health and operations. You need to categorize the red flags into following categories based on their original quotes and the identified keyword.

- Balance Sheet Issues: Red flags related to the company's assets, liabilities, equity, debt and overall financial position.
- P&L (Income Statement) Issues: Red flags related to the company's revenues, expenses, profits, and overall financial performance.
- Liquidity Issues: Concerns related to the company's ability to meet its short-term obligations, such as cash flow problems, debt repayment issues, or insufficient working capital.
- Management and Strategy related Issues: Concerns related to leadership, governance, decision-making processes, overall strategy, vision, and direction. 
- Regulatory Issues: Concerns related Compliance with laws, regulations. 
- Industry and Market Issues: Concerns related Position within the industry, market trends, and competitive landscape. 
- Operational Issues: Concerns related Internal processes, systems, and infrastructure.

While categorizing the red flags strictly adhere to the following guidelines:
1. Please review the below red flags and assign each one to the most relevant category.
2. Do not loose information from the Original Quotes keep them as it is.
3. If a red flag could fit into multiple categories, please assign it to the one that seems most closely related, do not leave any flag unclassified or fit it into multiple categories.
4. While classifying, classify it in a such a way that the flags come under the categories along with their content. Strictly do not create a new category stick to what is mentioned above like an "Additional Red Flags", classify the flags in the above mentioned category only.
5. Do not repeat a category more than once in the output.

**Output Format**:
### Balance Sheet Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### P&L (Income Statement) Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### Liquidity Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### Management and Strategy related Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### Regulatory Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### Industry and Market Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

### Operational Issues
- [Red flag 1 from previous analysis with original quote]
- [Red flag 2 from previous analysis with original quote]

Continue this format for all categories, ensuring every red flag from the previous analysis is categorized properly."""
        
        third_full_prompt = f"""You must answer the question strictly based on the below given context.

    Context:
    {pipeline_1st.docs[0]["context"]}

    Previous Analysis: {second_response}

    Based on the above analysis and the original context, please answer: {third_prompt}

    Answer:"""
        
        third_response = pipeline_1st.llm._call(third_full_prompt)
        
        # ITERATION 4: Detailed summary generation
        print("Running 4th iteration - Summary Generation...")
        fourth_prompt = """Based on the categorized red flags from the previous analysis, provide a comprehensive and detailed summary of each category of red flags in bullet point format. Follow these guidelines:

    1. **Retain all information**: Ensure that no details are omitted or lost during the summarization process
    2. **Maintain a neutral tone**: Present the summary in a factual and objective manner, avoiding any emotional or biased language
    4. **Include all red flags**: Incorporate every red flag within the category into the summary, without exception
    5. **Balance detail and concision**: Provide a summary that is both thorough and concise, avoiding unnecessary elaboration while still conveying all essential information
    6. **Incorporate quantitative data**: Wherever possible, include quantitative data and statistics to support the summary and provide additional context
    7. **Category-specific content**: Ensure that the summary is generated based solely on the content present within each category
    8. **Summary should be factual**: Avoid any subjective interpretations or opinions
    9. **Use bullet points**: Each red flag should be summarized as a separate bullet point with key details and data points

    Format the output exactly like this example:
    ### Balance Sheet Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]
    * [Summary of red flag 3 with specific data points and factual information]

    ### P&L (Income Statement) Issues  
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]
    * [Summary of red flag 3 with specific data points and factual information]

    ### Liquidity Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]

    ### Management and Strategy related Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]

    ### Regulatory Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]

    ### Industry and Market Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]

    ### Operational Issues
    * [Summary of red flag 1 with specific data points and factual information]
    * [Summary of red flag 2 with specific data points and factual information]

    Continue this format for all 7 categories. Each bullet point should be a concise summary that captures the key details of each red flag within that category, including relevant quantitative data where available."""
        
        fourth_full_prompt = f"""You must answer the question strictly based on the below given context.

    Context:
    {pipeline_1st.docs[0]["context"]}

    Previous Analysis: {third_response}

    Based on the above analysis and the original context, please answer: {fourth_prompt}

    Answer:"""
        
        fourth_response = pipeline_1st.llm._call(fourth_full_prompt)
        
        # ITERATION 5: Risk Classification
        print("Running 5th iteration - Risk Classification...")
        
        # Define risk criteria for each category
        risk_criteria = {
            "Balance Sheet Issues": {
                "debt_increase": "High: Debt increase by >=40% compared to previous reported b/s number; Medium: Debt increase between 25 to 40% compared to previous reported b/s number; Low: Debt increase is less than 25% compared to previous reported b/s number",
                "provisioning": "High: provisioning or write-offs more than 25% of current quarter's EBIDTA; Medium: provisioning or write-offs between 10 to 25% of current quarter's EBIDTA; Low: provisioning or write-offs less than 10% of current quarter's EBIDTA",
                "asset_decline": "High: Asset value falls by >=40% compared to previous reported b/s number; Medium: Asset value falls between 25% to 40% compared to previous reported b/s number; Low: Asset value falls by less than 25% compared to previous reported b/s number",
                "receivable_days": "High: receivable days increase by >=40% compared to previous reported b/s number; Medium: receivable days increase between 25 to 40% compared to previous reported b/s number; Low: receivable days increase is less than 25% compared to previous reported b/s number",
                "payable_days": "High: payable days increase by >=40% compared to previous reported b/s number; Medium: payable days increase between 25 to 40% compared to previous reported b/s number; Low: payable days increase is less than 25% compared to previous reported b/s number",
                "debt_ebitda": "High: Debt/EBITDA > 4x; Medium: Debt/EBITDA 2-4x; Low: Debt/EBITDA < 2x"
            },
            "P&L (Income Statement) Issues": {
                "revenue_decline": "High: revenue or profitability falls by >=40% compared to previous reported quarter number; Medium: revenue or profitability falls between 25% to 40% compared to previous reported quarter number; Low: revenue or profitability falls by less than 25% compared to previous reported quarter number",
                "onetime_expenses": "High: one-time expenses or losses more than 25% of current quarter's EBIDTA; Medium: one-time expenses or losses between 10 to 25% of current quarter's EBIDTA; Low: one-time expenses or losses less than 10% of current quarter's EBIDTA",
                "margin_decline": "High: gross margin or operating margin falling more than 25% compared to previous reported quarter number; Medium: gross margin or operating margin falling between 10 to 25% compared to previous reported quarter number; Low: gross margin or operating margin falling less than 10% compared to previous reported quarter number"
            },
            "Liquidity Issues": {
                "cash_balance": "High: cash balance falling more than 25% compared to previous reported b/s number; Medium: cash balance falling between 10 to 25% compared to previous reported b/s number; Low: cash balance falling less than 10% compared to previous reported b/s number",
                "short_term_debt": "High: Short-term debt or current liabilities increase by >=40% compared to previous reported b/s number; Medium: Short-term debt or current liabilities increase between 25 to 40% compared to previous reported b/s number; Low: Short-term debt or current liabilities increase is less than 25% compared to previous reported b/s number"
            },
            "Management and Strategy related Issues": {
                "management_issues": "High: High management turnover or key personnel departures, Poor track record of execution or delivery, High employee attrition rates; Medium: Some management turnover or key personnel departures, Some concerns around execution or delivery, Moderate employee attrition rates; Low: Low management turnover or key personnel departures, Strong track record of execution or delivery, Low employee attrition rates"
            },
            "Regulatory Issues": {
                "regulatory_compliance": "High: Material non-compliance with regulations or laws, Repeated regulatory issues or warnings; Medium: Some non-compliance with regulations or laws, Some regulatory issues or warnings; Low: No material non-compliance with regulations or laws, No regulatory issues or warnings"
            },
            "Industry and Market Issues": {
                "market_competition": "High: High competitive intensity or new entrants, Material decline in market share; Medium: Some competitive intensity or new entrants, Some decline in market share; Low: Low competitive intensity or new entrants, Stable or increasing market share"
            },
            "Operational Issues": {
                "operational_disruptions": "High: Material disruptions to operations or supply chain, Significant IT or system failures; Medium: Some disruptions to operations or supply chain, Some IT or system failures; Low: No Material disruptions to operations or supply chain, No significant IT or system failures"
            }
        }
        
        # Process each category for risk classification
        fifth_results = []
        categories = ["Balance Sheet Issues", "P&L (Income Statement) Issues", "Liquidity Issues", 
                     "Management and Strategy related Issues", "Regulatory Issues", 
                     "Industry and Market Issues", "Operational Issues"]
        
        for category in categories:
            print(f"Classifying risk for: {category}")
            
            # Create category-specific prompt
            category_criteria = risk_criteria.get(category, {})
            criteria_text = ""
            for criteria_name, criteria_desc in category_criteria.items():
                criteria_text += f"{criteria_name}: {criteria_desc}\n"
            
            fifth_prompt = f"""
            
You are a financial risk analyst. Based on the summary provided and the previous year's financial data, classify each red flag in the {category} category into High, Medium, Low, or Not_Applicable risk levels.

Criteria for {category}:
{criteria_text}

Summary from previous iteration:
{fourth_response}

Previous Year Financial Data:
{previous_year_data}

For each bullet point under {category}, analyze it against the criteria and classify the risk level.

Output Format:
### {category} Risk Classification

For each red flag found in {category}:
* [Red flag summary]
  - High: yes/no
  - Medium: yes/no 
  - Low: yes/no
  - Not_Applicable: yes/no

If no red flags exist in this category, state: "No red flags present in this category."

IMPORTANT: Provide ONLY yes/no answers without any explanations or additional text after yes/no.

Analysis should be based on:
1. Quantitative thresholds mentioned in the criteria
2. Comparison with previous year data provided
3. Qualitative assessment of the described issues

"""
            
            fifth_full_prompt = f"""You must answer the question strictly based on the below given context.

Context:
{pipeline_1st.docs[0]["context"]}

Previous Analysis: {fourth_response}

Based on the above analysis and the original context, please answer: {fifth_prompt}

Answer:"""
            
            category_response = pipeline_1st.llm._call(fifth_full_prompt)
            fifth_results.append({
                "category": category,
                "risk_classification": category_response
            })
        
        # Combine all risk classifications
        combined_fifth_response = ""
        for result in fifth_results:
            combined_fifth_response += f"\n{result['risk_classification']}\n"
        
        # Extract company information from first page
        print("Extracting company information...")
        company_info = extract_company_info_from_pdf(pdf_path, pipeline_1st.llm)
        print(f"Identified company: {company_info}")
        
        # Parse risk classifications
        print("Parsing risk classifications...")
        risk_flags = parse_risk_classifications(combined_fifth_response)
        
        # Parse summary by categories
        print("Parsing summary by categories...")
        summary_by_categories = parse_summary_by_categories(fourth_response)
        
        # Create Word document
        print("Creating Word document...")
        word_doc_path = create_word_document(
            pdf_name=pdf_name,
            company_info=company_info,
            risk_flags=risk_flags,
            summary_by_categories=summary_by_categories,
            output_folder=output_folder
        )
        
        # Save all results together
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        all_results = pd.DataFrame({
            "pdf_name": [pdf_name, pdf_name, pdf_name, pdf_name, pdf_name],
            "iteration": [1, 2, 3, 4, 5],
            "stage": [
                "Initial Analysis",
                "Deduplication", 
                "Categorization",
                "Summary Generation",
                "Risk Classification"
            ],
            "prompt": [
                first_results_df.iloc[0]['query'],  # Original query from 1st iteration
                second_prompt,
                third_prompt,
                fourth_prompt,
                "Risk Classification for all categories"
            ],
            "response": [
                first_response,
                second_response,
                third_response,
                fourth_response,
                combined_fifth_response
            ],
            "timestamp": [timestamp, timestamp, timestamp, timestamp, timestamp]
        })
        
        # Save complete results
        complete_output_file = os.path.join(output_folder, f"{pdf_name}_complete_5iteration_pipeline_results.csv")
        all_results.to_csv(complete_output_file, index=False)
        
        # Save individual risk classification results
        risk_df = pd.DataFrame(fifth_results)
        risk_output_file = os.path.join(output_folder, f"{pdf_name}_risk_classification.csv")
        risk_df.to_csv(risk_output_file, index=False)
        
        print(f"Complete 5-iteration pipeline finished for {pdf_name}!")
        print(f"CSV Results saved to: {complete_output_file}")
        print(f"Risk classification saved to: {risk_output_file}")
        print(f"Word document saved to: {word_doc_path}")
        
        return all_results
        
    except Exception as e:
        print(f"Error processing {pdf_name}: {str(e)}")
        # Save error log
        error_df = pd.DataFrame({
            "pdf_name": [pdf_name],
            "error": [str(e)],
            "timestamp": [time.strftime("%Y%m%d_%H%M%S")]
        })
        error_file = os.path.join(output_folder, f"{pdf_name}_error_log.csv")
        error_df.to_csv(error_file, index=False)
        return None

def run_multiple_pdfs_five_iterations_pipeline(pdf_folder_path: str, queries_csv_path: str, previous_year_data: str, output_folder: str = "results"):
    """
    Process multiple PDFs from a folder through the 5-iteration pipeline including risk classification
    
    Args:
        pdf_folder_path: Path to folder containing PDF files
        queries_csv_path: Path to CSV/Excel file containing queries
        previous_year_data: String containing previous year financial data
        output_folder: Path to output folder for results
    """
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Get all PDF files from the folder
    pdf_files = glob.glob(os.path.join(pdf_folder_path, "*.pdf"))
    
    if not pdf_files:
        print(f"No PDF files found in {pdf_folder_path}")
        return
    
    print(f"Found {len(pdf_files)} PDF files to process:")
    for pdf_file in pdf_files:
        print(f"  - {os.path.basename(pdf_file)}")
    
    print(f"\nStarting batch processing with 5 iterations...")
    print(f"Output folder: {output_folder}")
    print("=" * 60)
    
    # Process each PDF
    successful_processing = []
    failed_processing = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {os.path.basename(pdf_file)}")
        
        try:
            result = process_single_pdf_five_iterations(
                pdf_path=pdf_file,
                queries_csv_path=queries_csv_path,
                previous_year_data=previous_year_data,
                output_folder=output_folder
            )
            
            if result is not None:
                successful_processing.append(os.path.basename(pdf_file))
            else:
                failed_processing.append(os.path.basename(pdf_file))
                
        except Exception as e:
            print(f"Failed to process {os.path.basename(pdf_file)}: {str(e)}")
            failed_processing.append(os.path.basename(pdf_file))
    
    if successful_processing:
        print(f"\nSuccessfully processed:")
        for file in successful_processing:
            print(f"  ✓ {file}")
    
    if failed_processing:
        print(f"\nFailed to process:")
        for file in failed_processing:
            print(f"  ✗ {file}")

def main_batch_processing_five_iterations():
    
    pdf_folder_path = r"pdf_KEC"
    queries_csv_path = r"EWS_prompts_v2.xlsx"       
    output_folder = r"KEC_with_word"
    
    # Previous year data - you can modify this for each company
    previous_year_data = """
Parameter	Mar-23	Unit
Previous reported Debt	3,194	Cr
Current quarter ebidta	244	Cr
Previous reported asset value	18,668	Cr
Previous reported receivable days	260	days
Previous reported payable days	236	days
Previous reported revenue	17,282	Cr
Previous reported profitability	943	Cr
Previous reported operating margin	5.50%	 
Previous reported cash balance	344	Cr
Previous reported current liabilities	14,209	Cr
"""

    run_multiple_pdfs_five_iterations_pipeline(
        pdf_folder_path=pdf_folder_path,
        queries_csv_path=queries_csv_path,
        previous_year_data=previous_year_data,
        output_folder=output_folder
    )

if __name__ == "__main__":
    main_batch_processing_five_iterations()
